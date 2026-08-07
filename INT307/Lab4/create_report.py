from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('INT307 Web Application Security', level=0)
subtitle = doc.add_heading('Lab 4: File Upload Vulnerabilities in DVWA', level=1)

# Student Info
doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Program: ICDFA Internship')
doc.add_paragraph('Date: August 2026')

doc.add_paragraph('')

# Environment
doc.add_heading('Lab Environment', level=1)
env = doc.add_paragraph()
env.add_run('Attacker machine (Kali Linux): ').bold = True
env.add_run('legion@pastor, 192.168.5.139\n')
env2 = doc.add_paragraph()
env2.add_run('Target (OWASP BWA VM): ').bold = True
env2.add_run('192.168.5.130 (DVWA v1.8, PHP 5.3.2-1ubuntu4.30)\n')
env3 = doc.add_paragraph()
env3.add_run('Tools used: ').bold = True
env3.add_run('Weevely 4.0.1, Burp Suite Community Edition v2026.3.2')

doc.add_paragraph('')

# Exercise 1
doc.add_heading('Exercise 1: Exploring the File Upload Page', level=1)
doc.add_paragraph('[INSERT SCREENSHOT: DVWA Upload page showing form fields]')
doc.add_heading('Observations', level=2)
doc.add_paragraph(
    'The Upload page presents a minimal form: a single file input labeled '
    '"Choose an image to upload," a Browse button, and an Upload button. '
    'No visible client-side restriction on file type or size was displayed '
    'on the page itself.'
)
doc.add_heading('Test Upload (test.txt)', level=2)
doc.add_paragraph('[INSERT SCREENSHOT: Successful upload of test.txt]')
doc.add_paragraph(
    'A harmless text file (test.txt) was uploaded successfully with no '
    'restriction, confirming the form label ("image") is not enforced. '
    'The success message revealed the storage path: '
    '../../hackable/uploads/test.txt, with the original filename preserved '
    'unmodified.'
)

doc.add_paragraph('')

# Exercise 2
doc.add_heading('Exercise 2: Generate a PHP Shell Using Weevely', level=1)
doc.add_paragraph(
    'Weevely 4.0.1 was already installed as a system package on Kali '
    '(/usr/bin/weevely), so no manual clone of the legacy weevely3 '
    'repository was required.'
)
doc.add_paragraph('[INSERT SCREENSHOT: weevely generate command and output]')
doc.add_heading('Shell Generation', level=2)
doc.add_paragraph(
    'Command used: weevely generate YourPassword123 shell.php\n'
    'Result: shell.php generated successfully, 696 bytes.'
)
doc.add_heading('Code Review', level=2)
doc.add_paragraph(
    'Inspecting shell.php with cat revealed obfuscated/encoded binary-like '
    'content rather than readable PHP source. Weevely 4.x obfuscates its '
    'generated payload by design, making static signature detection more '
    'difficult for WAFs or antivirus tools.'
)

doc.add_paragraph('')

# Exercise 3
doc.add_heading('Exercise 3: Upload the Malicious Shell at Different Security Levels', level=1)

doc.add_heading('Low Security Level', level=2)
doc.add_paragraph('[INSERT SCREENSHOT: Successful shell.php upload at Low level]')
doc.add_paragraph(
    'shell.php was uploaded successfully with no validation whatsoever. '
    'Upload path confirmed: ../../hackable/uploads/shell.php. This '
    'demonstrates a critical vulnerability: no file type, extension, or '
    'content checking is performed at Low security level.'
)

doc.add_heading('Medium Security Level', level=2)
doc.add_paragraph('[INSERT SCREENSHOT: Direct upload blocked - "Your image was not uploaded"]')
doc.add_paragraph(
    'A direct upload attempt of shell.php was rejected. The intercepted '
    'request in Burp Suite showed the file part sent with '
    'Content-Type: application/x-php.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Burp intercepted request before modification]')
doc.add_paragraph('[INSERT SCREENSHOT: Burp intercepted request with Content-Type changed to image/jpeg]')
doc.add_paragraph('[INSERT SCREENSHOT: Successful upload after Content-Type bypass]')
doc.add_paragraph(
    'Changing only the Content-Type header value to image/jpeg in Burp '
    'Suite, then forwarding the request, resulted in a successful upload. '
    'This confirms Medium security level validates only the client-supplied '
    'Content-Type header - a value fully controlled by the attacker and '
    'never verified against the file\'s actual content.'
)

doc.add_heading('High Security Level', level=2)
doc.add_paragraph('[INSERT SCREENSHOT: Direct upload blocked at High level]')
doc.add_paragraph(
    'A direct upload attempt was blocked, as expected. The same '
    'Content-Type spoofing technique that succeeded at Medium level was '
    'attempted via Burp Suite and failed.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Burp request with filename changed to shell.phtml and Content-Type spoofed]')
doc.add_paragraph(
    'An extended bypass attempt was also made, combining a filename change '
    '(shell.php to shell.phtml) with the Content-Type spoof. This also '
    'failed, with the server again returning "Your image was not '
    'uploaded." This indicates High security level performs validation '
    'beyond headers and filenames - most likely inspecting actual file '
    'content or structure (e.g., verifying genuine image data), which '
    'cannot be bypassed through header or filename manipulation alone.'
)

doc.add_paragraph('')

# Exercise 4
doc.add_heading('Exercise 4: Accessing the Uploaded Shell', level=1)
doc.add_paragraph(
    'Security level was reset to Low to test shell access. Browsing '
    'directly to http://192.168.5.130/dvwa/hackable/uploads/shell.php '
    'returned a blank page with no visible prompt - this is expected '
    'behavior for Weevely-generated shells, which only respond to '
    'specifically-crafted authenticated POST requests rather than showing '
    'an interactive password form in the browser.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Blank browser response from shell.php]')
doc.add_paragraph(
    'Connecting via the Weevely client directly established a session '
    'successfully:\n'
    'weevely http://192.168.5.130/dvwa/hackable/uploads/shell.php '
    'YourPassword123'
)
doc.add_paragraph('[INSERT SCREENSHOT: Successful Weevely connection and weevely> prompt]')
doc.add_heading('Command Execution Issue', level=2)
doc.add_paragraph(
    'Attempting basic commands (whoami, :system_info) at the weevely> '
    'prompt failed with HTTP 500 errors. Investigation via DVWA\'s PHP '
    'Info page revealed the target is running PHP 5.3.2-1ubuntu4.30 '
    '(released 2010). The disable_functions directive was confirmed empty, '
    'ruling out server-side function restrictions as the cause.'
)
doc.add_paragraph('[INSERT SCREENSHOT: PHP Info page showing PHP version]')
doc.add_paragraph(
    'Root cause: Weevely 4.0.1 generates payloads targeting modern PHP '
    'syntax and behavior. Executed against PHP 5.3.2, the payload triggers '
    'a server-side error rather than completing command execution. This is '
    'a tooling/version compatibility issue rather than a defensive control - '
    'it illustrates that exploit tooling must match the target software '
    'stack, and legacy backends can break modern offensive tools even when '
    'the initial file upload vulnerability is fully exploitable.'
)

doc.add_paragraph('')

# Exercise 5
doc.add_heading('Exercise 5: Understanding Security Measures', level=1)
doc.add_heading('Summary of Outcomes Across Security Levels', level=2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
hdr[0].text = 'Security Level'
hdr[1].text = 'Direct Upload'
hdr[2].text = 'Content-Type Spoof (Burp)'
hdr[3].text = 'Extension + Content-Type Spoof (Burp)'

rows_data = [
    ('Low', 'Allowed', 'N/A - not needed', 'N/A'),
    ('Medium', 'Blocked', 'Bypassed', 'N/A - not needed'),
    ('High', 'Blocked', 'Blocked', 'Blocked'),
]
for i, row in enumerate(rows_data, start=1):
    cells = table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph('')

doc.add_heading('How Burp Suite Was Utilized', level=2)
doc.add_paragraph(
    'Burp Suite\'s Proxy Intercept feature captured the raw multipart '
    'upload request before it reached the server, allowing direct inline '
    'editing of the Content-Type header and filename field within the '
    'request body - values that are normally set automatically by the '
    'browser and not editable through the DVWA form itself. This '
    'demonstrated that any client-supplied header or filename is inherently '
    'untrustworthy from the server\'s perspective, since Burp shows how '
    'trivially these values can be rewritten in transit.'
)

doc.add_heading('Reflection: Effectiveness of Security Measures', level=2)
doc.add_paragraph(
    'Low security level performed no validation at all, making full '
    'compromise trivial. Medium security level validated only the '
    'client-supplied Content-Type header, which is easily spoofed since it '
    'is attacker-controlled and never verified against the actual file '
    'content. High security level appears to validate genuine file content '
    '(likely checking for real image structure, such as through '
    'getimagesize() or magic byte inspection), which could not be bypassed '
    'through header or filename manipulation alone.'
)
doc.add_paragraph(
    'The overall lesson from this lab is that effective file upload '
    'security requires content-based validation - inspecting the actual '
    'bytes and structure of an uploaded file - combined with server-side '
    'extension whitelisting, storing uploads outside the web root or in '
    'non-executable directories, and randomizing or rejecting suspicious '
    'filenames. Relying solely on client-supplied metadata such as headers '
    'or filenames, as seen at the Medium level, provides only superficial '
    'protection that is trivially bypassed with tools like Burp Suite.'
)

doc.add_paragraph('')

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab demonstrated the practical exploitation of file upload '
    'vulnerabilities in DVWA across three security levels. At Low level, '
    'a malicious PHP shell was uploaded without restriction. At Medium '
    'level, Content-Type header validation was bypassed using Burp Suite '
    'by spoofing the header to image/jpeg. At High level, both header and '
    'filename-based bypass attempts failed, indicating deeper content '
    'validation. Weevely successfully generated an obfuscated PHP shell '
    'and established a session connection, though command execution was '
    'blocked by a PHP version incompatibility (target running PHP 5.3.2 '
    'from 2010) rather than a security control - a distinction worth '
    'noting when interpreting exploit tool failures in real engagements. '
    'This lab reinforced that layered, content-aware validation is '
    'essential to defending against file upload attacks, and that '
    'client-supplied metadata should never be trusted for security '
    'decisions.'
)

doc.save('/home/legion/INT307_Lab4_Report.docx')
print("Report saved to /home/legion/INT307_Lab4_Report.docx")
