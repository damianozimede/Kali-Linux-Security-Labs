#!/usr/bin/env python3
"""
INT307 - Web Application Security
Lab 3: Command Injection in DVWA
Report generator (python-docx)

Usage:
    python3 ~/Kali-Linux-Security-Labs/INT307/Lab3/create_report.py

Screenshots are NOT embedded automatically - this script inserts plain-text
placeholders. Open the generated .docx in LibreOffice/Word afterwards and
paste in screenshots (captured via Windows Snipping Tool) where indicated.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILENAME = "INT307_Lab3_Command_Injection_Report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_title(doc, text):
    h = doc.add_heading(text, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_subtitle_centered(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bold_para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p


def add_screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[INSERT SCREENSHOT: {description}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)


def add_simple_table(doc, headers, rows, col_widths_in=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for p_ in hdr_cells[i].paragraphs:
            for r in p_.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], 'DCE6F1')

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")
    return table


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---------------- TITLE PAGE ----------------
    add_title(doc, "INT307: Web Application Security")
    add_subtitle_centered(doc, "Lab 3: Command Injection in DVWA", size=16)
    doc.add_paragraph("")
    add_subtitle_centered(doc, "Student: Damian Ozimede Patrick", size=12)
    add_subtitle_centered(doc, "GitHub: github.com/damianozimede", size=12, bold=False)
    add_subtitle_centered(
        doc,
        "Environment: Kali Linux (legion@pastor, 192.168.5.139) - OWASP BWA VM (192.168.5.130)",
        size=12, bold=False
    )
    add_subtitle_centered(doc, "Date: July 2026", size=12, bold=False)
    add_page_break(doc)

    # ---------------- OBJECTIVE / ENVIRONMENT ----------------
    add_h1(doc, "Objective")
    add_para(
        doc,
        "To identify and exploit Command Injection vulnerabilities within the Damn "
        "Vulnerable Web Application (DVWA), testing across Low, Medium, and High security "
        "levels, and to understand the sanitization approaches that defeat each technique."
    )

    add_h1(doc, "Environment")
    add_simple_table(
        doc,
        ["Item", "Value"],
        [
            ["Attacker machine", "Kali Linux - legion@pastor (192.168.5.139)"],
            ["Target VM", "OWASP BWA - 192.168.5.130"],
            ["Target application", "DVWA - Command Execution module"],
            ["Target URL", "http://192.168.5.130/dvwa/vulnerabilities/exec/"],
            ["DVWA login", "admin / admin"],
        ],
        col_widths_in=[2.3, 4.2]
    )

    add_h2(doc, "Note on scope")
    add_para(
        doc,
        "This DVWA installation's Security Level dropdown only offers Low, Medium, and "
        "High - there is no \"Impossible\" option available, unlike newer DVWA versions. "
        "Exercise 4 (Impossible Command Injection) was therefore adapted: rather than "
        "testing a security level that does not exist in this environment, the High-level "
        "filter was probed further using additional shell metacharacter techniques beyond "
        "the pipe character already tested in Exercise 3."
    )

    # ================= EXERCISE 1 =================
    add_h1(doc, "Exercise 1: Low-Level Command Injection")
    add_bold_para(doc, "Baseline test:")
    add_code_block(doc, "127.0.0.1")
    add_para(
        doc,
        "Submitted at Security Level: Low. The application returned standard ping output "
        "(3 packets transmitted, 0% packet loss), confirming the input is passed to a "
        "system ping command and the raw output is returned to the browser."
    )
    add_bold_para(doc, "Injection payload:")
    add_code_block(doc, "127.0.0.1 ; whoami ; cat /etc/passwd")
    add_para(
        doc,
        "Output Observed: All three commands executed sequentially. The ping output "
        "appeared as expected, followed by \"www-data\" (confirming commands run as the "
        "web server's process user), followed by the complete contents of /etc/passwd - "
        "revealing every system account, including root, service accounts (mysql, "
        "postgres, sshd), and the regular user account (user)."
    )
    add_screenshot_placeholder(doc, "Command injection Low - whoami and /etc/passwd disclosed via ; chaining")
    add_para(
        doc,
        "Implications: The semicolon (;) is a shell command separator, and this "
        "application passes user input directly into a shell command with no "
        "sanitization at Low security. This allows an attacker to execute arbitrary "
        "operating system commands with the privileges of the web server process. Impact "
        "includes full user enumeration and, since arbitrary command execution is "
        "possible, could extend to reading/writing any file the www-data user can access, "
        "installing a reverse shell, or pivoting further into the host system. This is a "
        "critical vulnerability, as it grants an attacker a foothold on the underlying "
        "server, not just the web application."
    )

    # ================= EXERCISE 2 =================
    add_h1(doc, "Exercise 2: Medium-Level Command Injection")
    add_bold_para(doc, "Payload:")
    add_code_block(doc, "127.0.0.1 | whoami")
    add_para(
        doc,
        "Output Observed: Submitted at Security Level: Medium. The response showed only "
        "\"www-data\" - no ping output was displayed at all, suggesting the application "
        "shows only the final command's output rather than all commands in the chain. "
        "This confirms the pipe (|) character was not filtered or blocked at Medium "
        "security."
    )
    add_screenshot_placeholder(doc, "Command injection Medium - whoami executed via | pipe, ping output suppressed")
    add_para(
        doc,
        "Implications: This indicates the Medium-level filter likely blacklists specific "
        "characters used at Low (such as the semicolon) but does not account for other "
        "valid shell command-chaining operators like the pipe. This is a textbook example "
        "of an incomplete blacklist: blocking one injection technique while leaving "
        "functionally equivalent alternatives open. An attacker who finds one payload "
        "blocked should systematically try alternate shell metacharacters (|, &&, ||, "
        "backticks, $()) rather than assuming the vulnerability is closed."
    )

    # ================= EXERCISE 3 =================
    add_h1(doc, "Exercise 3: High-Level Command Injection")
    add_bold_para(doc, "Payloads tested:")
    add_code_block(
        doc,
        "127.0.0.1 |whoami   (as specified by the lab, no space)\n"
        "127.0.0.1 | whoami  (same payload that succeeded at Medium)"
    )
    add_para(
        doc,
        "Output Observed: Both variants returned \"ERROR: You have entered an invalid "
        "IP\" at High security level. Neither payload executed, regardless of whitespace "
        "around the pipe character."
    )
    add_screenshot_placeholder(doc, "Command injection High - both pipe variants blocked, invalid IP error")
    add_para(
        doc,
        "Implications: Unlike the Medium-level filter, which appeared to miss the pipe "
        "character while blocking semicolons, the High-level filter blocks the pipe "
        "character regardless of surrounding whitespace. This indicates the High-level "
        "filter uses a more complete blacklist and/or input validation approach - likely "
        "checking that input matches a strict IP address format, or blacklisting a wider "
        "set of shell metacharacters, rather than targeting only specific known payloads. "
        "This is a meaningful security improvement over Medium, though a blacklist-based "
        "approach can still potentially be bypassed by other, untested metacharacters - a "
        "genuinely secure implementation would use strict input validation (e.g. "
        "regex-matching a valid IPv4/IPv6 format) combined with parameterized command "
        "execution (avoiding shell string concatenation entirely), rather than "
        "blacklisting."
    )

    # ================= EXERCISE 4 (ADAPTED) =================
    add_h1(doc, "Exercise 4: Further Probing High-Level Filtering (Adapted)")
    add_para(
        doc,
        "As noted above, this DVWA installation does not offer an \"Impossible\" security "
        "level, so this exercise was adapted to further probe the High-level filter's "
        "robustness using multiple alternate command-injection techniques, testing "
        "whether any shell metacharacter beyond the pipe could bypass it."
    )
    add_bold_para(doc, "Payloads tested (all at High security level):")
    add_code_block(
        doc,
        "127.0.0.1 && whoami\n"
        "127.0.0.1 `whoami`      (backtick command substitution)\n"
        "127.0.0.1 $(whoami)     (command substitution)"
    )
    add_para(
        doc,
        "Output Observed: All three payloads returned \"ERROR: You have entered an "
        "invalid IP\", identical to the pipe-based payloads tested in Exercise 3."
    )
    add_screenshot_placeholder(doc, "Command injection High - && and backtick/$() substitution payloads all blocked")
    add_para(
        doc,
        "Implications: The High-level filter consistently blocks every shell "
        "metacharacter tested - semicolon, pipe, logical AND, and both forms of command "
        "substitution - regardless of syntax variation. This strongly suggests the filter "
        "is not simply blacklisting individual known-bad payloads, but is likely "
        "validating that the input matches a strict expected format (e.g. a valid IPv4 "
        "address via regex), which is a fundamentally more robust approach than "
        "blacklisting specific characters. While a small number of injection techniques "
        "were tested and none succeeded, a blacklist-only implementation can rarely be "
        "proven completely secure through testing alone - true security would "
        "additionally require avoiding shell string concatenation entirely (e.g. using a "
        "language-native ping library or an execve-style call with argument arrays rather "
        "than a shell string), which cannot be confirmed without reviewing the "
        "application's source code."
    )

    # ================= SUMMARY TABLE =================
    add_h1(doc, "Summary")
    add_simple_table(
        doc,
        ["Security Level", "Payload(s) Tested", "Result"],
        [
            ["Low", "127.0.0.1 ; whoami ; cat /etc/passwd", "Full execution - www-data + /etc/passwd disclosed"],
            ["Medium", "127.0.0.1 | whoami", "Executed - pipe not filtered"],
            ["High", "127.0.0.1 |whoami / | whoami", "Blocked - \"invalid IP\" error"],
            ["High (adapted)", "&&, backticks, $()", "Blocked - all variants rejected"],
        ],
        col_widths_in=[1.5, 3.0, 2.5]
    )

    # ================= CONCLUSION =================
    add_h1(doc, "Conclusion")
    add_para(
        doc,
        "This lab demonstrated Command Injection vulnerabilities in DVWA across Low, "
        "Medium, and High security levels. At Low, the application applied no input "
        "sanitization whatsoever, allowing full command chaining via the semicolon "
        "operator and resulting in complete disclosure of system user accounts. At "
        "Medium, a partial blacklist blocked the semicolon but failed to account for the "
        "pipe character, allowing continued command execution - a clear example of an "
        "incomplete blacklist defense. At High, the filter successfully blocked every "
        "shell metacharacter tested, including pipes, logical AND, and both forms of "
        "command substitution (backticks and $()), suggesting a stricter validation "
        "approach, likely based on matching a valid IP address format rather than simply "
        "blacklisting known-bad characters."
    )
    add_bold_para(doc, "Key recommendations:")
    add_bullets(
        doc,
        [
            "Never pass user input directly into a shell command string. Use "
            "language-native libraries (e.g. a ping library) or execute commands via "
            "argument arrays (execve-style) rather than shell string concatenation, which "
            "eliminates the injection vector entirely regardless of which characters an "
            "attacker uses.",
            "If shell execution cannot be avoided, validate input strictly against an "
            "expected format (e.g. a regex matching valid IPv4/IPv6 addresses) before use, "
            "rejecting anything that does not match exactly.",
            "Avoid blacklist-based character filtering as a sole defense - as seen at "
            "Medium level, blacklists are easily incomplete and can be bypassed by "
            "metacharacters the developer did not anticipate.",
            "Run web application processes with the minimum privileges necessary (least "
            "privilege), so that even if command injection occurs, the impact is limited.",
            "Apply defense in depth: combine input validation, least-privilege process "
            "accounts, and monitoring/logging of unusual command execution patterns.",
        ]
    )
    add_para(
        doc,
        "Ethical note: all testing in this lab was performed against a deliberately "
        "vulnerable, isolated training environment (DVWA on the OWASP BWA VM) with no "
        "real user data, for educational purposes only."
    )

    doc.save(OUTPUT_FILENAME)
    print(f"Report saved: {OUTPUT_FILENAME}")


if __name__ == "__main__":
    build_report()
