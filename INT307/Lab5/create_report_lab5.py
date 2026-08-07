from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
doc.add_heading('INT307 Web Application Security', level=0)
doc.add_heading('Lab 5: Security Misconfiguration in DVWA', level=1)

# Student Info
doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Program: ICDFA Internship')
doc.add_paragraph('Date: August 2026')

doc.add_paragraph('')

# Environment
doc.add_heading('Lab Environment', level=1)
env = doc.add_paragraph()
env.add_run('Attacker machine (Kali Linux): ').bold = True
env.add_run('legion@pastor, 192.168.5.139\n')
env2 = doc.add_paragraph()
env2.add_run('Target (OWASP BWA VM): ').bold = True
env2.add_run('192.168.5.130 (DVWA v1.8, Apache/2.2.14 Ubuntu, PHP 5.3.2-1ubuntu4.30)\n')
env3 = doc.add_paragraph()
env3.add_run('Tools used: ').bold = True
env3.add_run('Firefox Developer Tools, Burp Suite Community Edition v2026.3.2')

doc.add_paragraph('')

# Exercise 1
doc.add_heading('Exercise 1: Default Credentials and Permissions', level=1)
doc.add_paragraph('[INSERT SCREENSHOT: Login failed using admin/password]')
doc.add_paragraph(
    'An attempt to log in using the commonly documented DVWA default '
    'credentials (admin / password) failed. Testing revealed the actual '
    'working credentials configured on this instance are admin / admin.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Successful login with admin/admin]')
doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'The credentials in active use (admin / admin) are arguably weaker '
    'than the documented default (admin / password), since the username '
    'and password are identical and require no reference to '
    'documentation to guess. This demonstrates that default credential '
    'risk is not limited to "textbook" defaults - any predictable, weak, '
    'or unchanged credential set exposes the application to trivial '
    'unauthorized access. Default and weak credentials remain one of the '
    'most common initial access vectors in real-world breaches because '
    'they require no technical exploitation, only a guess.'
)

doc.add_paragraph('')

# Exercise 2
doc.add_heading('Exercise 2: Directory Listing', level=1)
doc.add_paragraph('[INSERT SCREENSHOT: Directory listing of /dvwa/hackable/uploads/]')
doc.add_paragraph(
    'Navigating directly to http://192.168.5.130/dvwa/hackable/uploads/ '
    'returned a full Apache-generated directory index rather than a 403 '
    'Forbidden or 404 error. The listing exposed the presence, filenames, '
    'and upload timestamps of every file in the folder, including '
    'dvwa_email.png, shell.php, and test.txt - the exact malicious PHP '
    'web shell and test file uploaded during the previous file upload lab.'
)
doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'Directory listing being enabled is a serious misconfiguration in its '
    'own right, and it significantly compounds file upload '
    'vulnerabilities. An attacker does not need to guess or predict an '
    'uploaded filename in advance; they can simply browse the directory '
    'and click straight through to any shell, backup file, or sensitive '
    'document sitting in an exposed folder. This finding directly links '
    'Exercise 2 to the file upload vulnerability explored in Exercise 5, '
    'demonstrating how multiple small misconfigurations combine into a '
    'much more serious exploitation chain.'
)

doc.add_paragraph('')

# Exercise 3
doc.add_heading('Exercise 3: Error Handling Misconfiguration', level=1)
doc.add_paragraph('[INSERT SCREENSHOT: 404 Not Found for invalid_page]')
doc.add_paragraph(
    'Accessing a non-existent page (/dvwa/vulnerabilities/invalid_page) '
    'returned a generic Apache 404 "Not Found" message with no stack '
    'trace, internal file paths, or server version banner - a properly '
    'handled web-server-level error.'
)
doc.add_paragraph(
    'To test error handling at the application layer, a malformed '
    'parameter was submitted to the SQL Injection page '
    '(?id=test\'&Submit=Submit).'
)
doc.add_paragraph('[INSERT SCREENSHOT: Raw MySQL syntax error exposed]')
doc.add_paragraph(
    'This returned a raw, unhandled MySQL error: "You have an error in '
    'your SQL syntax; check the manual that corresponds to your MySQL '
    'server version for the right syntax to use near \'\'test\'\'\' at '
    'line 1."'
)
doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'This is a clear inconsistency in error handling across layers of '
    'the stack. The web server itself (Apache) handles routing errors '
    'like 404s cleanly, with no information disclosure. However, at the '
    'application layer, unsanitized input reaching the database directly '
    'triggers a verbose, unhandled MySQL error message. This confirms '
    'the backend database technology (MySQL) and reveals that raw '
    'database errors are surfaced directly to end users rather than '
    'being caught and replaced with a generic error page. This gives an '
    'attacker a significant head start in crafting SQL injection '
    'payloads, since they receive direct feedback confirming both the '
    'database type and that their input is reaching the query layer '
    'unfiltered. Effective error handling must be applied consistently '
    'across every layer of an application, not just at the web server.'
)

doc.add_paragraph('')

# Exercise 4
doc.add_heading('Exercise 4: Security Headers', level=1)
doc.add_paragraph('[INSERT SCREENSHOT: Response headers - general headers]')
doc.add_paragraph('[INSERT SCREENSHOT: Response headers - Server and X-Powered-By]')
doc.add_paragraph(
    'HTTP response headers were inspected using Firefox Developer Tools '
    '(Network tab) while loading a DVWA page.'
)

doc.add_heading('Headers Present', level=2)
doc.add_paragraph(
    'Standard headers observed included Cache-Control, Connection, '
    'Content-Encoding, Content-Length, Content-Type, Date, Expires, '
    'Keep-Alive, Pragma, and Vary - none of which are security-specific.'
)
doc.add_paragraph(
    'Notably, the Server header disclosed extensive version information: '
    'Apache/2.2.14 (Ubuntu) mod_mono/2.4.3 PHP/5.3.2-1ubuntu4.30 with '
    'Suhosin-Patch, along with mod_python, mod_ssl, OpenSSL, and Perl '
    'version details. An X-Powered-By header additionally confirmed '
    'PHP/5.3.2-1ubuntu4.30.'
)

doc.add_heading('Security Headers Checked', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Security Header'
hdr[1].text = 'Present?'
rows_data = [
    ('X-Content-Type-Options', 'Missing'),
    ('X-Frame-Options', 'Missing'),
    ('Content-Security-Policy', 'Missing'),
    ('Strict-Transport-Security', 'Missing'),
]
for i, row in enumerate(rows_data, start=1):
    cells = table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph('')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'None of the four requested security headers are implemented on '
    'this application. Their absence has concrete consequences: missing '
    'X-Frame-Options or a Content-Security-Policy frame-ancestors '
    'directive leaves the application vulnerable to clickjacking via '
    'iframe embedding; missing X-Content-Type-Options: nosniff permits '
    'MIME-sniffing attacks where a browser may misinterpret a file\'s '
    'content type; and the missing Strict-Transport-Security header is '
    'consistent with the application running over plain HTTP without '
    'TLS, meaning secure transport has not been configured at all. '
    'Compounding this, the server actively over-shares detailed version '
    'information via the Server and X-Powered-By headers, effectively '
    'handing an attacker a precise software fingerprint that can be '
    'matched directly against known CVEs for those exact component '
    'versions. Best practice is to implement all relevant security '
    'headers and to suppress or minimize version-disclosing headers to '
    'reduce the attack surface available through reconnaissance alone.'
)

doc.add_paragraph('')

# Exercise 5
doc.add_heading('Exercise 5: Insecure File Upload Configuration', level=1)
doc.add_paragraph(
    'This exercise was addressed by revisiting and reframing findings '
    'from the prior File Upload Vulnerabilities lab, viewed here '
    'specifically through a misconfiguration lens rather than an '
    'exploitation lens.'
)
doc.add_heading('Access and Test Summary', level=2)
doc.add_paragraph(
    'DVWA\'s Upload feature was previously tested extensively. At Low '
    'security level, a Weevely-generated PHP web shell (shell.php) was '
    'uploaded with no restriction on file type, content, or size, and '
    'stored directly under the publicly web-accessible path '
    '/dvwa/hackable/uploads/.'
)
doc.add_heading('Reflection: Was the File Uploaded?', level=2)
doc.add_paragraph(
    'Yes, at Low security level the upload succeeded outright with no '
    'validation. At Medium level, upload was blocked based only on the '
    'client-supplied Content-Type header, and this check was bypassed '
    'using Burp Suite by changing the header from application/x-php to '
    'image/jpeg - the server never verified the actual file content. At '
    'High level, both header spoofing and a combined filename/header '
    'bypass attempt (renaming to shell.phtml plus spoofing the header) '
    'failed, indicating deeper content-based validation at that tier.'
)
doc.add_heading('Misconfiguration Analysis', level=2)
doc.add_paragraph(
    '1. Uploads are stored inside the web root, in a location where PHP '
    'execution is enabled - the hackable/uploads/ directory is both '
    'publicly reachable via HTTP and configured to execute PHP files '
    'directly, rather than being isolated outside the web root or in a '
    'non-executable directory.'
)
doc.add_paragraph(
    '2. Directory listing is enabled on the uploads folder (confirmed '
    'in Exercise 2), which significantly compounds the file upload risk: '
    'an attacker does not need to guess or predict the uploaded '
    'filename, since the entire directory is browsable and clickable.'
)
doc.add_paragraph(
    '3. Validation is inconsistent across configurable security levels - '
    'Low applies zero validation, while Medium applies only a '
    'client-trusted header check that is trivially defeated with basic '
    'tooling such as Burp Suite. This indicates upload security was not '
    'treated as a first-class design concern.'
)
doc.add_heading('Implications', level=2)
doc.add_paragraph(
    'Combined, these misconfigurations allow an attacker to achieve '
    'remote code execution on the server via a straightforward '
    'upload-then-execute chain, with directory listing removing even the '
    'minor obstacle of needing to know the uploaded filename in advance.'
)
doc.add_heading('Mitigation Recommendations', level=2)
doc.add_paragraph(
    'Store all uploaded files outside the web root, or in a directory '
    'with script execution explicitly disabled at the web server level. '
    'Validate file content (magic bytes/structure) rather than trusting '
    'headers or extensions. Disable directory listing globally (Options '
    '-Indexes in Apache configuration). Rename uploaded files to random, '
    'non-guessable values server-side. Apply consistent, content-based '
    'validation regardless of any configurable "security level" - '
    'production systems should not offer a low-security mode at all.'
)

doc.add_paragraph('')

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab explored a range of security misconfigurations across DVWA, '
    'demonstrating that vulnerabilities are not limited to application '
    'logic flaws but frequently arise from how an application and its '
    'underlying server are configured and maintained. Weak default '
    'credentials (admin/admin), directory listing exposing sensitive '
    'uploaded files, inconsistent error handling between the web server '
    'and application layers, a complete absence of standard security '
    'headers combined with excessive version disclosure, and insecure '
    'file upload storage all compounded one another to create a much '
    'more serious overall attack surface than any single issue alone. '
    'This reinforces a core lesson of secure application design: '
    'misconfigurations are cumulative, and addressing them requires a '
    'defense-in-depth approach across every layer of the stack, from '
    'server hardening and credential management to HTTP header policy '
    'and file handling logic.'
)

doc.save('/home/legion/INT307_Lab5_Report.docx')
print("Report saved to /home/legion/INT307_Lab5_Report.docx")
