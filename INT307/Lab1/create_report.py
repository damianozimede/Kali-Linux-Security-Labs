#!/usr/bin/env python3
"""
INT307 - Web Application Security
Lab 1: Manual and Automated SQL Injection Testing
Report generator (python-docx)

Usage:
    python3 ~/Kali-Linux-Security-Labs/INT307/Lab1/create_report.py

Screenshots are NOT embedded automatically - this script inserts plain-text
placeholders. Open the generated .docx in Word/LibreOffice afterwards and
paste in screenshots (captured via Windows Snipping Tool) where indicated.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILENAME = "INT307_Lab1_SQL_Injection_Report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_title(doc, text):
    h = doc.add_heading(text, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_subtitle_centered(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bold_para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    # Shade the paragraph light grey to visually set it apart as code
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p


def add_screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[INSERT SCREENSHOT: {description}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)


def add_simple_table(doc, headers, rows, col_widths_in=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for p_ in hdr_cells[i].paragraphs:
            for r in p_.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], 'DCE6F1')

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacing after table
    return table


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_report():
    doc = Document()

    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---------------- TITLE PAGE ----------------
    add_title(doc, "INT307: Web Application Security")
    add_subtitle_centered(doc, "Lab 1: Manual and Automated SQL Injection Testing", size=16)
    doc.add_paragraph("")
    add_subtitle_centered(doc, "Student: Damian Ozimede Patrick", size=12)
    add_subtitle_centered(doc, "GitHub: github.com/damianozimede", size=12, bold=False)
    add_subtitle_centered(
        doc,
        "Environment: Kali Linux (legion@pastor, 192.168.5.139) - OWASP BWA VM (192.168.5.130)",
        size=12, bold=False
    )
    add_subtitle_centered(doc, "Date: July 2026", size=12, bold=False)
    add_page_break(doc)

    # ---------------- OBJECTIVE / ENVIRONMENT ----------------
    add_h1(doc, "Objective")
    add_para(
        doc,
        "To understand and practice SQL injection techniques using both manual testing "
        "methods and automated tools (sqlmap) against a deliberately vulnerable web "
        "application (OWASP Mutillidae II), hosted on the OWASP Broken Web Applications "
        "(BWA) virtual machine."
    )

    add_h1(doc, "Environment")
    add_simple_table(
        doc,
        ["Item", "Value"],
        [
            ["Attacker machine", "Kali Linux - legion@pastor (192.168.5.139)"],
            ["Target VM", "OWASP BWA - 192.168.5.130"],
            ["Target application", "OWASP Mutillidae II v2.6.24, Security Level 0 (Hosed)"],
            ["Backend DBMS", "MySQL 5.1.41-3ubuntu12.6-log"],
            ["Web/App server", "Apache 2.2.14 (Ubuntu), PHP 5.3.2-1ubuntu4.30"],
            ["Target OS", "Ubuntu 10.04 (Lucid Lynx)"],
            ["Target database (Mutillidae)", "nowasp"],
        ],
        col_widths_in=[2.3, 4.2]
    )

    # ================= EXERCISE 1 =================
    add_h1(doc, "Exercise 1: Identify Vulnerable Parameters")
    add_para(doc, "Objective: Identify input fields that may be vulnerable to SQL injection.")
    add_para(
        doc,
        "The Mutillidae application was accessed at http://192.168.5.130/mutillidae/. "
        "Navigating to OWASP 2013 -> A1 Injection (SQL) -> SQLi - Extract Data -> "
        "User Info (SQL) led to the User Lookup (SQL) page."
    )
    add_bold_para(doc, "Fields identified:")
    add_bullets(doc, ["Name (username) - text input", "Password - text input"])
    add_para(
        doc,
        "Both fields are submitted via HTTP GET to user-info.php, meaning both parameters "
        "are visible directly in the URL. This makes the page a strong candidate for manual "
        "SQL injection testing, since payloads can be tested simply by editing the URL."
    )
    add_screenshot_placeholder(doc, "Mutillidae User Lookup (SQL) page showing Name and Password fields")

    # ================= EXERCISE 2 =================
    add_h1(doc, "Exercise 2: Basic SQL Injection Testing")
    add_para(doc, "Objective: Perform basic SQL injection attacks to evaluate application response.")
    add_bold_para(doc, "Payload 1:")
    add_code_block(doc, "' OR '1'='1")
    add_para(
        doc,
        "Result: \"Authentication Error: Bad user name or password\" - "
        "Results for \"' OR '1'='1\". 0 records found."
    )
    add_para(
        doc,
        "This payload did not succeed. Note: the lab's suggested payload syntax did not "
        "include a comment terminator, so the remainder of the original query "
        "(AND password='...') was likely left intact, breaking the intended logic without "
        "producing a usable true condition."
    )
    add_bold_para(doc, "Payload 2 (adjusted):")
    add_code_block(doc, "' OR 1=1-- -")
    add_para(
        doc,
        "Result: 24 records found, returning full username/password/signature data for "
        "every account in the accounts table, including admin/admin."
    )
    add_screenshot_placeholder(doc, "Successful basic SQL injection - 24 records returned via ' OR 1=1-- -")
    add_para(
        doc,
        "Finding: The Name field is vulnerable to SQL injection. The payload breaks out of "
        "the intended query logic, comments out the remainder of the original query, and "
        "forces the WHERE clause to always evaluate true, exposing the entire accounts "
        "table without valid credentials."
    )

    # ================= EXERCISE 3 =================
    add_h1(doc, "Exercise 3: Error-Based SQL Injection")
    add_para(doc, "Objective: Use error messages to extract database information.")
    add_bold_para(doc, "Payload:")
    add_code_block(doc, "' AND 1=CONVERT(int, (SELECT @@version))-- -")
    add_para(
        doc,
        "Note: this payload uses Microsoft SQL Server syntax (CONVERT), while Mutillidae's "
        "backend is MySQL. The syntax mismatch itself triggered a verbose MySQL error."
    )
    add_para(doc, "Result: MySQL error 1064 (syntax error), which disclosed the full underlying query structure:")
    add_code_block(
        doc,
        "SELECT * FROM accounts WHERE username='' AND 1=CONVERT(int, "
        "(SELECT @@version))-- -' AND password=''"
    )
    add_para(
        doc,
        "The error also revealed: MySQL version 5.1.73 (client library), the accounts "
        "table name, and internal server file paths "
        "(e.g. /owaspbwa/mutillidae-git/classes/MySQLHandler.php) along with a full PHP "
        "stack trace."
    )
    add_screenshot_placeholder(doc, "Verbose MySQL error disclosing query structure, DB version, and server file paths")
    add_para(
        doc,
        "Finding: Error-based SQL injection confirmed. Verbose error handling discloses "
        "sensitive information about database structure and server internals, which an "
        "attacker can use to refine further attacks without needing valid data output."
    )

    # ================= EXERCISE 4 =================
    add_h1(doc, "Exercise 4: Boolean-Based SQL Injection")
    add_para(doc, "Objective: Manipulate queries to obtain true/false responses.")
    add_para(
        doc,
        "Initial attempts using AND-based payloads with a bare trailing space after -- "
        "(as suggested by the lab) behaved inconsistently, likely due to trailing "
        "whitespace being trimmed by the application before reaching the SQL query. The "
        "-- - comment format (used successfully in Exercise 2) was used instead to ensure "
        "reliable query termination."
    )
    add_simple_table(
        doc,
        ["Payload", "Condition", "Result"],
        [
            ["' OR 1=1-- -", "TRUE", "24 records found"],
            ["' OR 1=2-- -", "FALSE", "0 records found, no error"],
        ],
        col_widths_in=[2.5, 1.5, 2.5]
    )
    add_screenshot_placeholder(doc, "Boolean FALSE condition (' OR 1=2-- -) returning 0 records with no SQL error")
    add_para(
        doc,
        "Finding: Boolean-based blind SQL injection confirmed. The application's response "
        "(record count / presence of an authentication error) reliably distinguishes TRUE "
        "from FALSE injected conditions, even without relying on verbose error messages. "
        "This technique remains viable in environments where detailed error output has "
        "been suppressed."
    )

    # ================= EXERCISE 5 =================
    add_h1(doc, "Exercise 5: Union-Based SQL Injection")
    add_para(doc, "Objective: Retrieve data from other tables using the UNION operator.")
    add_para(
        doc,
        "The lab's suggested 4-column payload failed with MySQL error 1222 (\"different "
        "number of columns\"). A 3-column retry also failed. Rather than guess further, "
        "the column count was determined systematically using the ORDER BY technique:"
    )
    add_code_block(
        doc,
        "' ORDER BY 1-- -   through   ' ORDER BY 7-- -   (succeeded)\n"
        "' ORDER BY 8-- -   (failed - column count error)"
    )
    add_para(doc, "This confirmed the accounts table (accessed via SELECT *) has exactly 7 columns.")
    add_bold_para(doc, "Final payload:")
    add_code_block(doc, "' UNION SELECT 1,user(),database(),version(),5,6,7-- -")
    add_para(doc, "Result: 1 record returned -")
    add_bullets(
        doc,
        [
            "Current DB user: mutillidae@localhost",
            "Current database: nowasp",
            "MySQL version: 5.1.41-3ubuntu12.6-log",
        ]
    )
    add_screenshot_placeholder(doc, "UNION-based injection revealing DB user, database name, and MySQL version")
    add_para(
        doc,
        "Finding: Union-based SQL injection confirmed. Once the correct column count is "
        "established (via ORDER BY probing), arbitrary data - including data unrelated to "
        "the original query's table - can be retrieved by aligning it with the injected "
        "UNION SELECT's column positions."
    )

    # ================= EXERCISE 6 =================
    add_h1(doc, "Exercise 6: Retrieving Database Information")
    add_para(doc, "Objective: Extract tables and columns from the database.")
    add_bold_para(doc, "Table enumeration payload:")
    add_code_block(
        doc,
        "' UNION SELECT 1,table_name,3,4,5,6,7 FROM information_schema.tables "
        "WHERE table_schema=database()-- -"
    )
    add_para(doc, "Result: 12 tables found in the nowasp database:")
    add_bullets(
        doc,
        [
            "accounts, balloon_tips, blogs_table, captured_data, credit_cards, help_texts, "
            "hitlog, level_1_help_include_files, page_help, page_hints, pen_test_tools, "
            "youtubevideos"
        ]
    )
    add_para(doc, "The credit_cards table stood out as a high-value target for further testing.")
    add_screenshot_placeholder(doc, "information_schema.tables enumeration - 12 tables listed")
    add_bold_para(doc, "Column enumeration payload (accounts table):")
    add_code_block(
        doc,
        "' UNION SELECT 1,column_name,3,4,5,6,7 FROM information_schema.columns "
        "WHERE table_name='accounts'-- -"
    )
    add_para(
        doc,
        "Result: 9 columns found - cid, username, password, mysignature, is_admin, "
        "firstname, lastname, userid, account."
    )
    add_screenshot_placeholder(doc, "information_schema.columns enumeration - 9 columns of the accounts table")
    add_para(
        doc,
        "Finding: The presence of an is_admin column is significant - it is a privilege "
        "flag that could be targeted for further exploitation. Combined with the table "
        "enumeration, this demonstrates that UNION-based SQL injection can be used to map "
        "an entire database schema - every table and column - without any authenticated "
        "access."
    )

    # ================= EXERCISE 7 =================
    add_h1(doc, "Exercise 7: SQLMap Basic Commands")
    add_para(doc, "Objective: Use sqlmap to automatically confirm and characterize the vulnerability.")
    add_code_block(
        doc,
        'sqlmap -u "http://192.168.5.130/mutillidae/index.php?page=user-info.php&'
        'username=admin&password=admin&user-info-php-submit-button=View+Account+Details" '
        '-p username --batch --dbs'
    )
    add_para(
        doc,
        "(The -p username flag scopes testing to the parameter already confirmed "
        "vulnerable manually, and --batch accepts sqlmap's default answers automatically, "
        "significantly reducing scan time.)"
    )
    add_para(doc, "Result: sqlmap confirmed the username parameter injectable via four distinct techniques:")
    add_bullets(
        doc,
        [
            "Boolean-based blind (OR, WHERE/HAVING clause, MySQL comment)",
            "Error-based (EXTRACTVALUE, MySQL >= 5.1)",
            "Time-based blind (query SLEEP)",
            "UNION query (7 columns) - matching manual findings exactly",
        ]
    )
    add_para(
        doc,
        "It also fingerprinted the backend: MySQL >= 5.1, Apache 2.2.14, PHP 5.3.2, "
        "Ubuntu 10.04 (Lucid Lynx) - and enumerated 34 databases present on the shared "
        "BWA server (hosting many vulnerable applications beyond Mutillidae, e.g. dvwa, "
        "joomla, wordpress)."
    )
    add_screenshot_placeholder(doc, "sqlmap injection point detection (4 techniques) and DBMS/OS fingerprint")
    add_screenshot_placeholder(doc, "sqlmap database enumeration results (--dbs, 34 databases)")
    add_para(
        doc,
        "Finding: Automated tooling independently validated and extended manual findings "
        "within seconds, confirming multiple injection techniques and full server-level "
        "database enumeration."
    )

    # ================= EXERCISE 8 =================
    add_h1(doc, "Exercise 8: Enumerate Users and Passwords")
    add_para(doc, "Objective: Extract MySQL database management system user and password information.")
    add_code_block(doc, 'sqlmap -u "..." -p username --batch -D nowasp --users --passwords')
    add_para(
        doc,
        "Result: 38 MySQL DBMS user accounts enumerated, including root@127.0.0.1, "
        "root@localhost, and root@brokenwebapps. sqlmap then retrieved password hashes "
        "for all accounts and performed a dictionary-based cracking attack, successfully "
        "cracking the majority of passwords - notably, one of the root account's password "
        "hashes was cracked to the clear-text password \"user\"."
    )
    add_screenshot_placeholder(doc, "sqlmap enumerated MySQL database users (--users), 38 total")
    add_screenshot_placeholder(doc, "sqlmap cracked password hashes for MySQL users (--passwords), including root account")
    add_para(
        doc,
        "Finding: Weak, dictionary-crackable passwords on privileged database accounts can "
        "be fully compromised via automated SQL injection tooling, potentially granting an "
        "attacker full control over the MySQL server and every database it hosts."
    )

    # ================= EXERCISE 9 =================
    add_h1(doc, "Exercise 9: Dumping Data")
    add_para(doc, "Objective: Retrieve all entries from a specific table.")
    add_code_block(doc, 'sqlmap -u "..." -p username --batch -D nowasp -T accounts --dump')
    add_para(
        doc,
        "Result: All 24 rows and 7 columns of the accounts table were extracted "
        "automatically in a single command, producing a clean formatted table and a CSV "
        "export. This dump confirmed the actual values of the is_admin column: three "
        "accounts (admin, adrian, ABaker) are flagged as administrators."
    )
    add_screenshot_placeholder(doc, "sqlmap full dump of nowasp.accounts table, 24 entries, showing is_admin flags")
    add_para(
        doc,
        "Finding: Automated dumping drastically speeds up full data exfiltration compared "
        "to manual, column-by-column UNION queries, and can reveal information (such as "
        "which specific accounts hold admin privileges) that manual extraction had not yet "
        "fully surfaced."
    )

    # ================= EXERCISE 10 =================
    add_h1(doc, "Exercise 10: Specify Columns and Tables")
    add_para(doc, "Objective: Enumerate columns in a specific table and extract only the data of interest.")
    add_code_block(doc, 'sqlmap -u "..." -p username --batch -D nowasp -T credit_cards --columns')
    add_para(
        doc,
        "Result: 4 columns identified in the credit_cards table - ccid (int), "
        "ccnumber (text), ccv (text), expiration (date)."
    )
    add_screenshot_placeholder(doc, "sqlmap column enumeration for nowasp.credit_cards table")
    add_code_block(
        doc,
        'sqlmap -u "..." -p username --batch -D nowasp -T credit_cards '
        '-C ccnumber,ccv,expiration --dump'
    )
    add_para(
        doc,
        "Result: 5 complete credit card records extracted - full card numbers, CVV codes, "
        "and expiration dates."
    )
    add_screenshot_placeholder(doc, "sqlmap targeted column dump - credit_cards (ccnumber, ccv, expiration)")
    add_para(
        doc,
        "Finding: An attacker does not need to dump entire tables; SQL injection allows "
        "surgical extraction of exactly the data of interest by specifying table and "
        "column names directly, making attacks faster and harder to detect. Full payment "
        "card data extractable via a single unauthenticated web form field represents the "
        "most severe risk identified in this lab."
    )

    # ================= ADDITIONAL: AUTH BYPASS =================
    add_h1(doc, "Additional Exercise: Bypassing Authentication")
    add_para(doc, "Objective: Log in as an admin user without knowing the password.")
    add_para(
        doc,
        "On the Mutillidae Login page (not User Lookup), the following payload was "
        "submitted in the username field, with the password field left blank:"
    )
    add_code_block(doc, "admin' -- -")
    add_para(
        doc,
        "Result: Successfully authenticated as the admin user. The application header "
        "confirmed: \"Logged In Admin: admin (g0t r00t?)\" - matching the exact signature "
        "seen in earlier data dumps for that account."
    )
    add_screenshot_placeholder(doc, "Successful authentication bypass - logged in as admin via SQL injection")
    add_para(
        doc,
        "Finding: The trailing -- - comments out the password validation portion of the "
        "backend query, so the query effectively becomes a lookup by username only. This "
        "demonstrates that SQL injection can be used not just to extract data, but to gain "
        "full authenticated access to privileged accounts without any valid credentials."
    )

    # ================= ADDITIONAL: LOG INJECTION =================
    add_h1(doc, "Additional Exercise: Log Injection Testing")
    add_para(doc, "Objective: Determine whether application logs can be manipulated via SQL injection input.")
    add_para(
        doc,
        "The Mutillidae View Log page was reviewed after SQL injection testing. The raw "
        "payloads used during testing - including the authentication bypass payload and "
        "complex sqlmap UNION strings - appeared in the log verbatim and unsanitized, e.g.:"
    )
    add_code_block(doc, "User ' OR 1=1-- - attempting to authenticate")
    add_screenshot_placeholder(doc, "Mutillidae View Log showing raw, unsanitized SQL injection payloads recorded verbatim")
    add_para(
        doc,
        "Finding: User-supplied input is logged without sanitization or encoding. This "
        "confirms a log injection vulnerability: an attacker could inject falsified log "
        "entries, characters that break log-parsing tools, or (if the log viewer renders "
        "entries unsafely) script content, creating a secondary XSS risk. It also means "
        "application logs cannot be fully trusted as an audit trail, since their content "
        "is directly attacker-influenced."
    )

    # ================= ADDITIONAL: BURP SUITE =================
    add_h1(doc, "Additional Exercise: Using Burp Suite for Manual Testing")
    add_para(doc, "Objective: Intercept and modify an HTTP request to Mutillidae using Burp Suite.")
    add_para(
        doc,
        "Burp Suite's Proxy tool was used to intercept a live GET request to "
        "user-info.php. The username parameter was manually edited within Burp's request "
        "editor, from a normal value to the injection payload:"
    )
    add_code_block(doc, "' OR 1=1 -- -")
    add_para(
        doc,
        "The modified request was forwarded to the server. Burp's \"Edited request\" view "
        "confirmed the payload was correctly URL-encoded and included in the outgoing "
        "request, and the server responded with HTTP 200 OK."
    )
    add_screenshot_placeholder(doc, "Burp Suite Proxy - edited request showing modified username parameter with SQL injection payload")
    add_para(
        doc,
        "Finding: SQL injection testing does not require browser input at all. Using an "
        "intercepting proxy, an attacker can directly manipulate any request parameter "
        "in-flight, bypassing any client-side validation (such as JavaScript length or "
        "character checks) entirely. (Note: this particular test was performed while an "
        "authenticated admin session was already active from the prior exercise, so the "
        "results table did not render identically to the earlier manual test; logging out "
        "before testing would be expected to restore the full 24-record output seen "
        "previously.)"
    )

    # ================= CONCLUSION =================
    add_h1(doc, "Conclusion")
    add_para(
        doc,
        "This lab provided hands-on experience with both manual and automated SQL "
        "injection techniques against OWASP Mutillidae II. Manual testing confirmed the "
        "User Lookup (SQL) page is vulnerable to basic, error-based, boolean-based, and "
        "union-based SQL injection, allowing full enumeration of the application's "
        "database schema (12 tables, including a sensitive credit_cards table) without "
        "authentication. The Login page was also shown to be vulnerable to authentication "
        "bypass, allowing admin-level access without valid credentials."
    )
    add_para(
        doc,
        "Automated testing with sqlmap independently confirmed all manual findings, "
        "additionally revealing the underlying MySQL server hosts 34 databases and 38 "
        "database user accounts, many with weak, dictionary-crackable passwords - "
        "including the root account. sqlmap was also used to perform full and targeted "
        "data extraction, culminating in the recovery of complete, unencrypted credit "
        "card records (card number, CVV, and expiration date) directly from the database."
    )
    add_para(
        doc,
        "Additional testing showed that the application logs user input verbatim without "
        "sanitization (a log injection risk), and that request parameters can be "
        "manipulated directly via an intercepting proxy such as Burp Suite, independent "
        "of any client-side input restrictions."
    )
    add_bold_para(doc, "Key recommendations:")
    add_bullets(
        doc,
        [
            "Use parameterized queries / prepared statements for all database access - "
            "never concatenate user input directly into SQL strings.",
            "Disable verbose database error messages in production; log them server-side only.",
            "Enforce strong, unique passwords for all database accounts, especially "
            "privileged accounts such as root.",
            "Sanitize or encode all user input before writing it to application logs.",
            "Apply the principle of least privilege to database accounts used by web applications.",
        ]
    )
    add_para(
        doc,
        "Ethical note: all testing in this lab was performed against a deliberately "
        "vulnerable, isolated training environment (OWASP BWA VM) with no real user data, "
        "for educational purposes only."
    )

    doc.save(OUTPUT_FILENAME)
    print(f"Report saved: {OUTPUT_FILENAME}")


if __name__ == "__main__":
    build_report()
