#!/usr/bin/env python3
"""
INT307 - Web Application Security
Lab 2: XSS Vulnerabilities in DVWA
Report generator (python-docx)

Usage:
    python3 ~/Kali-Linux-Security-Labs/INT307/Lab2/create_report.py

Screenshots are NOT embedded automatically - this script inserts plain-text
placeholders. Open the generated .docx in LibreOffice/Word afterwards and
paste in screenshots (captured via Windows Snipping Tool) where indicated.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILENAME = "INT307_Lab2_XSS_DVWA_Report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_title(doc, text):
    h = doc.add_heading(text, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_subtitle_centered(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bold_para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p


def add_screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[INSERT SCREENSHOT: {description}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)


def add_simple_table(doc, headers, rows, col_widths_in=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for p_ in hdr_cells[i].paragraphs:
            for r in p_.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], 'DCE6F1')

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")
    return table


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---------------- TITLE PAGE ----------------
    add_title(doc, "INT307: Web Application Security")
    add_subtitle_centered(doc, "Lab 2: XSS Vulnerabilities in DVWA", size=16)
    doc.add_paragraph("")
    add_subtitle_centered(doc, "Student: Damian Ozimede Patrick", size=12)
    add_subtitle_centered(doc, "GitHub: github.com/damianozimede", size=12, bold=False)
    add_subtitle_centered(
        doc,
        "Environment: Kali Linux (legion@pastor, 192.168.5.139) - OWASP BWA VM (192.168.5.130)",
        size=12, bold=False
    )
    add_subtitle_centered(doc, "Date: July 2026", size=12, bold=False)
    add_page_break(doc)

    # ---------------- OBJECTIVE / ENVIRONMENT ----------------
    add_h1(doc, "Objective")
    add_para(
        doc,
        "To identify and exploit Cross-Site Scripting (XSS) vulnerabilities within the "
        "Damn Vulnerable Web Application (DVWA), covering Reflected and Stored XSS across "
        "Low, Medium, and High security levels, and to document mitigation strategies."
    )

    add_h1(doc, "Environment")
    add_simple_table(
        doc,
        ["Item", "Value"],
        [
            ["Attacker machine", "Kali Linux - legion@pastor (192.168.5.139)"],
            ["Target VM", "OWASP BWA - 192.168.5.130"],
            ["Target application", "DVWA (Damn Vulnerable Web Application)"],
            ["DVWA login", "admin / admin"],
            ["Database setup", "Performed via /dvwa/setup.php - Create/Reset Database"],
        ],
        col_widths_in=[2.3, 4.2]
    )

    add_h2(doc, "Note on scope")
    add_para(
        doc,
        "This DVWA installation does not include a separate \"XSS (DOM)\" module - only "
        "\"XSS reflected\" and \"XSS stored\" are present in the application menu. DOM-based "
        "XSS testing was therefore not performed, as the vulnerable page does not exist in "
        "this environment. Testing proceeded with the two available XSS categories: "
        "Reflected and Stored."
    )

    # ================= STEP 1-2: SETUP & IDENTIFY =================
    add_h1(doc, "Step 1-2: Setup and Identify Vulnerable Input Fields")
    add_para(
        doc,
        "DVWA was accessed at http://192.168.5.130/dvwa/. The database was initialized via "
        "the Setup page (Create / Reset Database), which created the users and guestbook "
        "tables successfully. Login was completed using admin/admin."
    )
    add_para(
        doc,
        "Exploring the application sidebar identified two relevant vulnerable pages:"
    )
    add_bullets(
        doc,
        [
            "XSS reflected - a single \"What's your name?\" input field, reflected back "
            "into the page immediately upon submission via a GET parameter (name).",
            "XSS stored - a guestbook form with Name and Message fields, where submitted "
            "entries are saved to the database and displayed to all visitors.",
        ]
    )
    add_screenshot_placeholder(doc, "DVWA database setup page showing successful setup")

    # ================= XSS REFLECTED =================
    add_h1(doc, "XSS (Reflected) Testing")

    add_h2(doc, "Low")
    add_bold_para(doc, "Payload:")
    add_code_block(doc, "<svg onload=alert('BugBot19 was here')>")
    add_para(
        doc,
        "Submitted into the \"What's your name?\" field at Security Level: Low. Page source "
        "confirmed the payload was reflected into the HTML completely unescaped, and the "
        "browser executed it, producing a JavaScript alert dialog showing \"BugBot19 was "
        "here\"."
    )
    add_screenshot_placeholder(doc, "Reflected XSS alert dialog triggered via <svg onload> payload at Low security")
    add_para(
        doc,
        "Finding: user input is echoed back into the page without any sanitization or "
        "HTML encoding at Low security."
    )

    add_h2(doc, "Medium")
    add_bold_para(doc, "Initial payload (failed):")
    add_code_block(doc, "<script>alert('BugBot19 was here')</script>")
    add_para(
        doc,
        "This did not execute; the page displayed the literal text "
        "\"Hello alert('BugBot19 was here')\". Inspecting the page source revealed the "
        "filter had removed only the opening <script> tag, leaving the closing </script> "
        "tag intact - the signature of a naive, case-sensitive string-replacement filter."
    )
    add_bold_para(doc, "Bypass payload (succeeded):")
    add_code_block(doc, "<ScRipt>alert('BugBot19 was here')</ScRipt>")
    add_para(
        doc,
        "Since the filter's match is case-sensitive, mixed-case tags are not recognized as "
        "<script> and pass through unmodified. This executed successfully, producing the "
        "expected alert dialog."
    )
    add_screenshot_placeholder(doc, "Reflected XSS at Medium - filter bypassed via case variation <ScRipt>, alert dialog shown")
    add_para(
        doc,
        "Finding: blacklist-based filtering is fundamentally weak - case variation alone "
        "defeats a naive string-match filter. It was also observed that the <svg onload=...> "
        "payload used at Low was fully stripped at Medium, suggesting the filter's scope "
        "differs by tag type rather than applying uniform sanitization."
    )

    add_h2(doc, "High")
    add_bold_para(doc, "Payload:")
    add_code_block(doc, "<ScRipt>alert('BugBot19 was here')</ScRipt>")
    add_para(
        doc,
        "The same case-variation bypass from Medium was tested first at High, and it "
        "succeeded immediately without modification, producing the alert dialog."
    )
    add_screenshot_placeholder(doc, "Reflected XSS at High - same case-variation bypass still succeeds, alert dialog shown")
    add_para(
        doc,
        "Finding: DVWA's High-level filter for XSS Reflected does not add case-insensitive "
        "matching over the Medium-level filter. This highlights a key weakness in "
        "blacklist-based defenses: increasing a \"security level\" label does not "
        "guarantee that known bypass techniques are closed - the underlying filtering "
        "logic must be reviewed for completeness, not assumed to improve with a difficulty "
        "setting."
    )

    # ================= XSS STORED =================
    add_h1(doc, "XSS (Stored) Testing")

    add_h2(doc, "Low")
    add_bold_para(doc, "Payload (Message field):")
    add_code_block(doc, "<script>alert(document.domain)</script>")
    add_para(
        doc,
        "Submitted with Name: test2 at Security Level: Low. The alert fired immediately "
        "upon submission, displaying \"192.168.5.130\" (the value of document.domain), "
        "confirming the script executed with access to page context, not just a static "
        "string."
    )
    add_para(
        doc,
        "Critically, the payload was also persisted to the guestbook - the entry shows "
        "Name: test2 with an empty Message field (the script tag consumed the visible "
        "content). This means the malicious script will re-execute for any user who later "
        "views this page."
    )
    add_screenshot_placeholder(doc, "Stored XSS alert firing on submission, showing document.domain value")
    add_screenshot_placeholder(doc, "Stored XSS payload persisted in guestbook, confirming stored/persistent nature")
    add_para(
        doc,
        "Finding: stored XSS is more severe than reflected XSS - it requires no crafted "
        "link or social engineering, and persists as a standing threat to every subsequent "
        "visitor to the page."
    )

    add_h2(doc, "Medium")
    add_bold_para(doc, "Payload (Message field):")
    add_code_block(doc, "<img src=x onerror=alert(document.cookie)>")
    add_para(
        doc,
        "Submitted with Name: test3 at Security Level: Medium. The payload executed "
        "successfully via the onerror event handler - since src=\"x\" is an invalid image "
        "path, onerror triggers automatically, executing the injected JavaScript. This "
        "bypasses filters that specifically target <script> tags, since no script tag is "
        "present at all."
    )
    add_para(
        doc,
        "The payload was persisted to the guestbook (Name: test3, empty Message), "
        "confirming stored XSS remains exploitable at Medium via non-script-tag vectors."
    )
    add_screenshot_placeholder(doc, "Stored XSS Medium - alert via <img onerror> payload")
    add_screenshot_placeholder(doc, "Stored XSS Medium - payload persisted in guestbook as test3")

    add_h2(doc, "High")
    add_bold_para(doc, "Payload attempts:")
    add_code_block(
        doc,
        "<body onload=alert('BugBot19')>\n"
        "<img src=x onerror=alert('BugBot19')>"
    )
    add_para(
        doc,
        "Both payloads were submitted at Security Level: High (Name: test4, test5). "
        "Neither executed - both were rendered as literal text, with angle brackets "
        "converted to HTML entities (&lt; and &gt;) before being stored/displayed, e.g. "
        "&lt;img src=x onerror=alert('BugBot19')&gt;."
    )
    add_screenshot_placeholder(doc, "Stored XSS High - payload neutralized via HTML entity encoding, shown as literal text")
    add_para(
        doc,
        "Finding: this is a fundamentally stronger defense than the Low/Medium filters, "
        "since it does not rely on blacklisting specific tags or keywords. Instead, it "
        "neutralizes any HTML markup by encoding it for safe display, meaning no tag-based "
        "XSS payload can execute at this level. This demonstrates the correct mitigation "
        "approach: output encoding (contextual escaping) is far more robust than "
        "blacklist-based input filtering, which can always be bypassed with tag or case "
        "variations, as shown in the Reflected XSS exercises."
    )

    # ================= CUSTOM PAYLOAD =================
    add_h1(doc, "Create Your Own Payload")
    add_para(
        doc,
        "Rather than using only the lab's suggested payloads verbatim, an original bypass "
        "was developed for the Reflected XSS page at Medium and High security levels. "
        "Initial testing showed the standard payload "
        "<script>alert('BugBot19 was here')</script> failed at Medium - inspection of the "
        "page source revealed the filter performs a simple, case-sensitive string "
        "replacement, removing the literal substring <script> but leaving the closing "
        "</script> tag untouched."
    )
    add_para(
        doc,
        "Based on this observation, a custom payload using mixed-case tags was developed:"
    )
    add_code_block(doc, "<ScRipt>alert('BugBot19 was here')</ScRipt>")
    add_para(
        doc,
        "Since the filter's string-match is case-sensitive, this payload is not recognized "
        "as <script> and passes through unmodified. It was tested and confirmed successful "
        "at both Medium and High security levels, proving that DVWA's XSS Reflected filter "
        "relies on a weak, case-sensitive blacklist rather than genuine sanitization - and "
        "that the \"High\" setting does not add case-insensitivity over \"Medium.\""
    )
    add_para(
        doc,
        "This demonstrates a general principle in web security testing: understanding why "
        "a filter fails (by inspecting source/behavior) is far more effective than blindly "
        "trying payloads from a list."
    )

    # ================= CONCLUSION =================
    add_h1(doc, "Conclusion and Mitigation Strategies")
    add_para(
        doc,
        "This lab demonstrated both Reflected and Stored Cross-Site Scripting "
        "vulnerabilities in DVWA across Low, Medium, and High security levels. Reflected "
        "XSS proved exploitable at all three levels - the Low level applied no filtering "
        "at all, while Medium and High relied on a case-sensitive blacklist that was "
        "bypassed using mixed-case HTML tags, showing no meaningful improvement between "
        "the two settings. Stored XSS followed a similar pattern at Low and Medium, where "
        "script tags and event-handler-based payloads (img onerror) executed and "
        "persisted in the guestbook database, posing a standing risk to every future "
        "visitor. Only at High security level was Stored XSS properly mitigated, via "
        "output encoding that converts angle brackets to HTML entities regardless of tag "
        "or keyword used."
    )
    add_bold_para(doc, "Key recommendations:")
    add_bullets(
        doc,
        [
            "Apply context-aware output encoding (HTML entity encoding) to all "
            "user-supplied data before rendering it in a page - this is effective "
            "regardless of which tags or keywords an attacker uses, unlike blacklist "
            "filtering.",
            "Avoid blacklist/keyword-based input filtering as a primary defense; it is "
            "trivially bypassed via case variation, alternate tags, or encoding tricks.",
            "Apply a strict Content Security Policy (CSP) to restrict inline script "
            "execution as a defense-in-depth measure.",
            "Sanitize and validate input on the server side, never relying solely on "
            "client-side JavaScript validation.",
            "Treat all stored user input (e.g. guestbook/comment fields) as untrusted at "
            "the point of display, not only at the point of submission.",
        ]
    )
    add_para(
        doc,
        "Ethical note: all testing in this lab was performed against a deliberately "
        "vulnerable, isolated training environment (DVWA on the OWASP BWA VM) with no "
        "real user data, for educational purposes only."
    )

    doc.save(OUTPUT_FILENAME)
    print(f"Report saved: {OUTPUT_FILENAME}")


if __name__ == "__main__":
    build_report()
