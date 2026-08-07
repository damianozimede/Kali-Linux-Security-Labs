from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
doc.add_heading('INT307 Web Application Security', level=0)
doc.add_heading('Lab 6: Insecure Direct Object References (IDOR) in DVWA', level=1)

# Student Info
doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Program: ICDFA Internship')
doc.add_paragraph('Date: August 2026')

doc.add_paragraph('')

# Environment
doc.add_heading('Lab Environment', level=1)
env = doc.add_paragraph()
env.add_run('Attacker machine (Kali Linux): ').bold = True
env.add_run('legion@pastor, 192.168.5.139\n')
env2 = doc.add_paragraph()
env2.add_run('Target (OWASP BWA VM): ').bold = True
env2.add_run('192.168.5.130 (DVWA v1.8)\n')
env3 = doc.add_paragraph()
env3.add_run('Note on lab adaptation: ').bold = True
env3.add_run(
    'DVWA v1.8 does not include a dedicated IDOR module '
    '(/vulnerabilities/idor/ returns a 404 Not Found). This is a known '
    'gap, as a standalone IDOR page was only added in later DVWA '
    'versions. The lab was adapted to use the SQL Injection page '
    '(/vulnerabilities/sqli/?id=), whose "id" parameter directly '
    'references a database user record with no ownership or '
    'authorization check - a textbook IDOR pattern, independent of any '
    'SQL injection technique.'
)

doc.add_paragraph('')

# Exercise 1
doc.add_heading('Exercise 1: Understanding IDOR', level=1)
doc.add_heading('Definition', level=2)
doc.add_paragraph(
    'Insecure Direct Object Reference (IDOR) occurs when an application '
    'exposes a direct reference to an internal object - such as a '
    'database record, file, or user account - typically via a URL '
    'parameter or form field, without verifying that the currently '
    'logged-in user is actually authorized to access that specific '
    'object.'
)
doc.add_heading('How IDOR Vulnerabilities Arise', level=2)
doc.add_paragraph(
    'The vulnerability arises because the application trusts the '
    'object reference itself (e.g., id=1) as sufficient proof of '
    'authorization, rather than checking whether the requesting user '
    'has permission to view or modify that particular record. This '
    'typically happens when developers implement authentication '
    '(proving who a user is) correctly, but skip or under-implement '
    'authorization (verifying what that user is allowed to access) at '
    'the individual-record level.'
)
doc.add_heading('Real-World Examples', level=2)
doc.add_paragraph(
    'IDOR has been responsible for real-world breaches, including '
    'exposed invoice or document IDs that let any authenticated user '
    'view other customers\' invoices by changing a number in the URL, '
    'social media platforms where changing a numeric user ID in an API '
    'request exposed private profile data, and banking or fintech '
    'applications where account or transaction IDs could be incremented '
    'to view other users\' financial records.'
)

doc.add_paragraph('')

# Exercise 2
doc.add_heading('Exercise 2: Identify IDOR Vulnerabilities in DVWA', level=1)
doc.add_paragraph(
    'Using the adapted SQL Injection page as the IDOR test target, the '
    'User ID field was submitted with a value of 1.'
)
doc.add_paragraph('[INSERT SCREENSHOT: SQLi page result for id=1, returning admin/admin]')
doc.add_paragraph(
    'The result returned the logged-in user\'s own record (ID: 1, First '
    'name: admin, Surname: admin) - a legitimate, expected result for '
    'an authorized lookup of one\'s own data.'
)
doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'The URL parameter directly references a database record by its '
    'numeric ID (?id=1), with no session-based ownership check tying '
    'the requested ID to the identity of the logged-in user. This '
    'structure is the precondition for an IDOR vulnerability: any '
    'change to this parameter will be executed by the server without '
    'question.'
)

doc.add_paragraph('')

# Exercise 3
doc.add_heading('Exercise 3: Exploit IDOR to Access Unauthorized Data', level=1)
doc.add_paragraph(
    'The id parameter was changed sequentially to access other user '
    'records, without any special tooling or injection technique - '
    'simply incrementing the value in the URL.'
)
doc.add_paragraph('[INSERT SCREENSHOT: id=2 result - Gordon Brown]')
doc.add_paragraph('[INSERT SCREENSHOT: id=3 result - Hack Me]')
doc.add_paragraph('[INSERT SCREENSHOT: id=4 result - Pablo Picasso]')

doc.add_heading('Results', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'User ID'
hdr[1].text = 'First Name'
hdr[2].text = 'Surname'
rows_data = [
    ('1', 'admin', 'admin'),
    ('2', 'Gordon', 'Brown'),
    ('3', 'Hack', 'Me'),
    ('4', 'Pablo', 'Picasso'),
]
for i, row in enumerate(rows_data, start=1):
    cells = table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph('')

doc.add_heading('Analysis', level=2)
doc.add_paragraph(
    'IDs 1 through 4 all returned distinct, complete user records with '
    'no restriction of any kind. No session-to-record ownership '
    'validation exists anywhere in the request flow - the application '
    'trusts the id parameter alone as sufficient authorization to '
    'return data. Because these IDs are small sequential integers, an '
    'attacker could trivially script a loop (id=1 through id=1000, for '
    'example) to enumerate and exfiltrate the entire user table within '
    'seconds, without needing any SQL injection payload at all - this '
    'is pure IDOR.'
)
doc.add_heading('Reflection: Sensitive Data and Risk', level=2)
doc.add_paragraph(
    'The data exposed in this lab was limited to first and last names, '
    'but the same flawed access-control pattern, applied to a '
    'real-world application, typically extends to far more sensitive '
    'fields such as email addresses, phone numbers, order history, '
    'financial records, or private messages, depending on what the '
    'vulnerable endpoint returns. IDOR is considered a low-effort, '
    'high-impact vulnerability class precisely because exploitation '
    'requires no special tools or skills beyond changing a number in a '
    'URL, making it one of the most commonly found - and most commonly '
    'overlooked - issues in real-world web applications.'
)

doc.add_paragraph('')

# Exercise 4
doc.add_heading('Exercise 4: Preventing IDOR Vulnerabilities', level=1)
doc.add_heading('Mitigation Strategies', level=2)
doc.add_paragraph(
    '1. Access control / authorization checks on every object '
    'reference - the core fix. Every request that includes an object '
    'identifier must be checked server-side to confirm the currently '
    'authenticated user actually owns or is permitted to view that '
    'specific record, not just that they are logged in generally.'
)
doc.add_paragraph(
    '2. Indirect object references - instead of exposing raw, '
    'predictable database primary keys (sequential integers) directly '
    'in the URL, map them to random, per-session, non-guessable tokens '
    '(such as a UUID). This does not replace proper authorization '
    'checks, but removes the ability to simply guess or enumerate '
    'adjacent records.'
)
doc.add_paragraph(
    '3. Enforce authorization at every layer, not just the UI - many '
    'applications hide a "view other users" link in the interface for '
    'regular users but forget that the underlying endpoint is still '
    'directly reachable via URL manipulation. Authorization must be '
    'enforced server-side, on every request, regardless of what the UI '
    'displays.'
)
doc.add_paragraph(
    '4. Logging and monitoring access attempts - logging every '
    'object-access request, including denied ones, allows detection of '
    'enumeration patterns, such as one account rapidly requesting '
    'sequential IDs, which is a strong signal of an IDOR attack in '
    'progress. Rate-limiting and alerting on such patterns adds a '
    'detection layer even where a flaw exists.'
)

doc.add_heading('Reflection: Applying This to DVWA', level=2)
doc.add_paragraph(
    'For the IDOR pattern demonstrated on DVWA\'s SQL Injection page, '
    'the application would need a server-side check comparing the '
    'requested id against the logged-in session\'s own user ID (or an '
    'explicit role-based permission check for administrative access to '
    'other records) before executing the query at all, rather than '
    'processing any id value supplied in the URL unconditionally. In a '
    'production application, this is typically implemented as reusable '
    'authorization middleware applied consistently across all endpoints '
    'that accept object identifiers, rather than relying on individual '
    'developers to remember to add the check to each new feature '
    'separately.'
)

doc.add_paragraph('')

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab demonstrated an Insecure Direct Object Reference '
    'vulnerability using DVWA\'s SQL Injection page as a stand-in for a '
    'dedicated IDOR module not present in this DVWA version. Simply '
    'changing a numeric ID parameter in the URL allowed complete, '
    'unauthorized access to every user record in the application\'s '
    'database, with no authentication bypass, injection payload, or '
    'specialized tooling required. This lab reinforced that '
    'authorization checks must be enforced independently for every '
    'object a user requests, at the server level, rather than relying '
    'on authentication alone or on the assumption that predictable '
    'identifiers will not be manipulated.'
)

doc.save('/home/legion/INT307_Lab6_Report.docx')
print("Report saved to /home/legion/INT307_Lab6_Report.docx")
