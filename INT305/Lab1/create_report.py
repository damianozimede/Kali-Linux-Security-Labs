from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT305: Network Security and Protocols', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 1: User Management', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Student: Damian Ozimede Patrick')
doc.add_paragraph('Course: INT305 - Network Security and Protocols')
doc.add_paragraph('GitHub: github.com/damianozimede/Kali-Linux-Security-Labs')
doc.add_paragraph('Platform: Kali Linux | Username: legion | Date: June 2026')

# 1. Objective
doc.add_heading('1. Objective', 2)
doc.add_paragraph(
    'The objective of this lab was to learn how to create, modify, and manage user accounts '
    'securely in Linux. This included understanding the importance of user account management '
    'for maintaining system security, and practising best practices in user administration such '
    'as user creation, group assignment, password management, account locking and unlocking, '
    'and account deletion.'
)

# 2. Environment
doc.add_heading('2. Environment', 2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Operating System: Kali Linux (legion@pastor)')
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('Machine IP: 192.168.5.139')
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run('All tasks performed directly in the Kali Linux terminal')

# 3. Creating User Accounts
doc.add_heading('3. Creating User Accounts', 2)
doc.add_paragraph(
    'The first task was to create a new user named student1 using the adduser command, '
    'which automatically handles home directory creation and prompts for user details:'
)
doc.add_paragraph('Command: sudo adduser student1')
doc.add_paragraph(
    'The command prompted for a password and optional user information fields which were '
    'left blank. The system confirmed the password was updated successfully and the user '
    'information was saved.'
)
doc.add_paragraph('[SCREENSHOT 1: Terminal output of sudo adduser student1]')

doc.add_paragraph('The user account was verified using:')
doc.add_paragraph('Command: id student1')
doc.add_paragraph(
    'Output: uid=1003(student1) gid=1003(student1) groups=1003(student1),100(users) '
    'confirming the account was created successfully.'
)
doc.add_paragraph('[SCREENSHOT 2: Terminal output of id student1]')

# 4. Modifying User Accounts
doc.add_heading('4. Modifying User Accounts', 2)
doc.add_heading('4.1 Changing the Password', 3)
doc.add_paragraph(
    'The password for student1 was changed using the passwd command to ensure secure access:'
)
doc.add_paragraph('Command: sudo passwd student1')
doc.add_paragraph(
    'The system prompted for a new password and confirmation. It confirmed the password '
    'was updated successfully.'
)
doc.add_paragraph('[SCREENSHOT 3: Terminal output of sudo passwd student1]')

doc.add_heading('4.2 Creating a Group and Adding student1', 3)
doc.add_paragraph(
    'A group named students was created and student1 was added to it to manage permissions effectively:'
)
doc.add_paragraph('Command: sudo groupadd students')
doc.add_paragraph('Command: sudo usermod -aG students student1')
doc.add_paragraph('Command: id student1')
doc.add_paragraph(
    'The updated output showed uid=1003(student1) gid=1003(student1) '
    'groups=1003(student1),100(users),1004(students), confirming student1 was successfully '
    'added to the students group.'
)
doc.add_paragraph('[SCREENSHOT 4: id student1 output showing students group membership]')

# 5. Locking and Unlocking
doc.add_heading('5. Locking and Unlocking User Accounts', 2)
doc.add_heading('5.1 Checking for Active Processes', 3)
doc.add_paragraph('Active processes under student1 were checked before locking:')
doc.add_paragraph('Command: ps -u student1')
doc.add_paragraph(
    'The output returned an empty process list confirming there were no active processes '
    'running under student1, as expected since the account had just been created.'
)
doc.add_paragraph('[SCREENSHOT 5: ps -u student1 showing empty process list]')

doc.add_heading('5.2 Locking the Account', 3)
doc.add_paragraph(
    'The student1 account was locked using usermod -L to prevent login without deleting the account:'
)
doc.add_paragraph('Command: sudo usermod -L student1')
doc.add_paragraph('Command: sudo passwd -S student1')
doc.add_paragraph(
    'Output: student1 L 2026-06-19 0 99999 7 -1. The L in the second field confirms '
    'the account is locked.'
)
doc.add_paragraph('[SCREENSHOT 6: sudo passwd -S student1 showing L status]')

doc.add_heading('5.3 Unlocking the Account', 3)
doc.add_paragraph('The account was unlocked using usermod -U:')
doc.add_paragraph('Command: sudo usermod -U student1')
doc.add_paragraph('Command: sudo passwd -S student1')
doc.add_paragraph(
    'Output: student1 P 2026-06-19 0 99999 7 -1. The P confirms the password is active '
    'and the account is accessible again.'
)
doc.add_paragraph('[SCREENSHOT 7: sudo passwd -S student1 showing P status]')

# 6. Deleting User Accounts
doc.add_heading('6. Deleting User Accounts', 2)
doc.add_heading('6.1 Deleting User but Keeping Home Directory', 3)
doc.add_paragraph(
    'The student1 account was deleted using deluser without flags, preserving the home directory:'
)
doc.add_paragraph('Command: sudo deluser student1')
doc.add_paragraph('Command: id student1')
doc.add_paragraph('Command: ls /home/student1')
doc.add_paragraph(
    'id returned no such user confirming the account was deleted. ls returned Permission denied '
    'rather than No such file or directory, confirming the home directory still existed on the system.'
)
doc.add_paragraph('[SCREENSHOT 8: id student1 showing no such user and ls /home/student1 showing Permission denied]')

doc.add_heading('6.2 Complete Removal Including Home Directory', 3)
doc.add_paragraph(
    'The student1 account was recreated then completely removed including the home directory:'
)
doc.add_paragraph('Command: sudo deluser --remove-home student1')
doc.add_paragraph(
    'Both id student1 and ls /home/student1 confirmed complete removal, returning '
    'no such user and No such file or directory respectively.'
)
doc.add_paragraph('[SCREENSHOT 9: Complete removal confirmation showing both user and home directory gone]')

# 7. Analysis
doc.add_heading('7. Analysis and Findings', 2)
doc.add_paragraph(
    'This lab highlighted several important aspects of user management in Linux. The adduser '
    'command is the recommended approach for creating users in Debian-based systems as it handles '
    'home directory creation and guides the administrator interactively through the setup process.'
)
doc.add_paragraph(
    'Group management through groupadd and usermod is an effective way to control permissions '
    'across multiple users. By adding student1 to the students group, permissions could be applied '
    'at the group level, which is far more scalable and secure in environments with many users.'
)
doc.add_paragraph(
    'The ability to lock and unlock accounts using usermod -L and usermod -U is a valuable security '
    'feature. Rather than deleting an account when a user needs to be temporarily suspended, locking '
    'preserves all data and settings while preventing login.'
)
doc.add_paragraph(
    'The difference between deluser and deluser --remove-home is an important distinction. Keeping '
    'the home directory after deletion is useful for forensic analysis, while complete removal is '
    'appropriate when the user is permanently leaving the organisation.'
)

# 8. Conclusion
doc.add_heading('8. Conclusion', 2)
doc.add_paragraph(
    'This lab provided practical hands-on experience with Linux user account management. All core '
    'tasks were completed successfully including creating a user, modifying account details, assigning '
    'group membership, changing passwords, locking and unlocking accounts, and performing both partial '
    'and complete account deletion. These skills are fundamental to maintaining system security in any '
    'Linux environment, as improper user management is one of the most common sources of unauthorised '
    'access and privilege escalation vulnerabilities.'
)

# Appendix
doc.add_heading('Appendix: Commands Used', 2)
commands = [
    'sudo adduser student1',
    'id student1',
    'sudo passwd student1',
    'sudo groupadd students',
    'sudo usermod -aG students student1',
    'ps -u student1',
    'sudo usermod -L student1',
    'sudo passwd -S student1',
    'sudo usermod -U student1',
    'sudo deluser student1',
    'ls /home/student1',
    'sudo deluser --remove-home student1',
]
for cmd in commands:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd)

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT305/Lab1/INT305_Lab1_Report.docx')
print("Report saved successfully!")
