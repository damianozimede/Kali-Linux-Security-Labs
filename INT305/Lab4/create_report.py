#!/usr/bin/env python3
"""
INT305 Lab 4 - Report Generator
Author: Damian Ozimede Patrick
Course: INT305 - Secure User Access Management in Linux
GitHub: github.com/damianozimede/Kali-Linux-Security-Labs
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_PATH = os.path.expanduser(
    "~/Kali-Linux-Security-Labs/INT305/Lab4/INT305_Lab4_Report.docx"
)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"Note: {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_report():
    doc = Document()

    # ── Title Page ──────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("INT305: Secure User Access Management in Linux")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Lab 4: Advanced User Access Control and Security Auditing")
    r2.bold = True
    r2.font.size = Pt(16)

    doc.add_paragraph()

    for line in [
        "Submitted by: Damian Ozimede Patrick",
        "Course: INT305 | Lab 4",
        "GitHub: github.com/damianozimede/Kali-Linux-Security-Labs",
        "System: Kali Linux | Username: legion | IP: 192.168.5.139",
        "Date: 02 July 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(11)

    doc.add_page_break()

    # ── PART 1 ──────────────────────────────────────────────────────
    add_heading(doc, "Part 1: Password Policy Configuration")

    add_heading(doc, "Exercise 1.1: Configure Password Aging", level=2)
    add_para(doc, "Password aging was configured for testuser using the chage command to enforce regular password changes.")
    add_para(doc, "Commands used:", bold=True)
    for cmd in [
        "sudo chage -l testuser",
        "sudo chage -M 90 testuser",
        "sudo chage -m 7 testuser",
        "sudo chage -W 14 testuser",
    ]:
        add_code(doc, cmd)

    add_para(doc, "System-wide policy configured in /etc/login.defs:", bold=True)
    for line in ["PASS_MAX_DAYS   90", "PASS_MIN_DAYS   7", "PASS_WARN_AGE   14"]:
        add_code(doc, line)

    add_heading(doc, "Question 1.1: Why is it important to set a minimum number of days between password changes?", level=3)
    add_para(doc, "Setting a minimum number of days between password changes prevents users from immediately reverting to their old password after being forced to change it. Without this control, a user could change their password and then change it straight back, effectively bypassing the rotation policy. This ensures each new password is used for a meaningful period before another change is permitted, maintaining the security intent of password aging.")

    add_heading(doc, "Question 1.2: What are the security implications of setting PASS_MAX_DAYS too high or too low?", level=3)
    add_para(doc, "Setting PASS_MAX_DAYS too high means compromised credentials remain valid longer, giving attackers an extended window to exploit a stolen password. Setting it too low creates password fatigue — users forced to change passwords too frequently resort to weak, predictable patterns or write passwords down, ultimately reducing security. The recommended balance is 60–90 days: frequent enough to limit exposure without driving insecure workarounds.")

    add_heading(doc, "Exercise 1.2: Implement Password Complexity Requirements", level=2)
    add_para(doc, "The libpam-pwquality package was installed and /etc/security/pwquality.conf was configured:")
    for line in [
        "minlen = 12",
        "dcredit = -1   # at least one digit",
        "ucredit = -1   # at least one uppercase",
        "lcredit = -1   # at least one lowercase",
        "ocredit = -1   # at least one special character",
        "difok = 3      # 3 chars must differ from old password",
        "maxrepeat = 2  # no more than 2 consecutive identical chars",
        "usercheck = 1  # reject passwords containing username",
    ]:
        add_code(doc, line)

    add_heading(doc, "Question 1.3: What happens when you try to set a password that doesn't meet complexity requirements?", level=3)
    add_para(doc, "The system displays a 'BAD PASSWORD' warning describing exactly which requirement fails. Regular users have the change rejected and must try again. With sudo/root privileges the warning appears but the password is still accepted — this is by design, allowing administrators to override policy when necessary.")

    add_heading(doc, "Question 1.4: How do password complexity requirements help prevent brute-force attacks?", level=3)
    add_para(doc, "Complexity requirements significantly expand the search space an attacker must work through. A password requiring uppercase, lowercase, digits, and special characters across 12+ characters has an astronomically larger number of possible combinations than a simple dictionary word or short PIN, making automated brute-force and dictionary attacks computationally impractical — especially when combined with account lockout policies.")

    doc.add_page_break()

    # ── PART 2 ──────────────────────────────────────────────────────
    add_heading(doc, "Part 2: Account Security Auditing")

    add_heading(doc, "Exercise 2.1: Identify Inactive User Accounts", level=2)
    add_note(doc, "The lab instructs use of sudo lastlog and sudo lastb. Both are unavailable on this Kali version as the legacy wtmp/btmp and /var/log/lastlog systems have been replaced by journald and lastlog2 (SQLite-backed). The lastlog2 package was installed and initialized manually. Equivalent commands are documented below.")
    add_para(doc, "Commands used:", bold=True)
    for cmd in [
        "sudo lastlog2 --active   # equivalent to: sudo lastlog | grep -v 'Never'",
        "sudo lastlog2 -b 90      # equivalent to: sudo lastlog -b 90",
        "sudo awk -F: '($2 == \"\") {print $1}' /etc/shadow",
        "awk -F: '($3 == \"0\") {print $1}' /etc/passwd",
        "sudo passwd -S -a",
    ]:
        add_code(doc, cmd)

    add_para(doc, "Key findings:", bold=True)
    add_bullet(doc, "lastlog2 --active: Only 'legion' had a recorded login — consistent with a freshly provisioned lab VM.")
    add_bullet(doc, "lastlog2 -b 90: No output — no accounts inactive for 90+ days.")
    add_bullet(doc, "Empty passwords: None found — all accounts have passwords set.")
    add_bullet(doc, "UID 0: Only 'root' — no rogue root-equivalent accounts.")
    add_bullet(doc, "passwd -S -a: Most service accounts correctly locked (L). testuser, testuser3, testuser4 have aging applied (7/90/14). Other user accounts still on system defaults — flagged in recommendations.")

    add_heading(doc, "Question 2.1: Why are inactive accounts a security risk?", level=3)
    add_para(doc, "Inactive accounts are dangerous primarily because nobody monitors them. If credentials for a dormant account are leaked, there is no legitimate user around to notice unauthorized access. They also tend to retain stale permissions never revoked after a role change, creating unnecessary attack surface. They are frequently targeted in credential-stuffing attacks since a successful login won't trigger the kind of alert an active user would raise immediately.")

    add_heading(doc, "Question 2.2: What actions should be taken for accounts inactive for 90+ days?", level=3)
    add_para(doc, "First verify with the account owner or manager whether the account is still needed. If confirmed unneeded, lock it immediately using passwd -l or usermod -L rather than deleting — to preserve audit trails and file ownership records. If it must remain active, force a password reset and verify it meets current policy. Organizations should set automatic disable thresholds (90 days triggers lock, 180 days triggers deletion review) so this is systematic rather than manual. All actions should be logged as part of the access review process.")

    add_heading(doc, "Exercise 2.2: Audit User Permissions and Group Memberships", level=2)
    add_para(doc, "Commands used:", bold=True)
    for cmd in [
        "groups legion",
        "getent group sudo",
        "sudo cat /etc/sudoers",
        "grep -v '/nologin\\|/false' /etc/passwd",
        "sudo find / -perm -4000 -type f -ls 2>/dev/null",
    ]:
        add_code(doc, cmd)

    add_para(doc, "Key findings:", bold=True)
    add_bullet(doc, "legion belongs to: legion, adm, dialout, cdrom, floppy, sudo, audio, dip, video, plugdev, users, netdev, scanner, bluetooth, lpadmin, wireshark, kaboxer.")
    add_bullet(doc, "Only legion is in the sudo group — single administrative account confirmed.")
    add_bullet(doc, "sudoers: root and %sudo have full access; mary has NOPASSWD for apt; paul has NOPASSWD for systemctl. Principle of least privilege correctly applied.")
    add_bullet(doc, "Service accounts postgres, arpwatch, sbuild have interactive login shells — a hardening gap. Should be changed to /usr/sbin/nologin.")
    add_bullet(doc, "SUID scan returned expected system binaries only. No suspicious or unknown SUID files found.")

    add_heading(doc, "Question 2.3: Why are SUID files a potential security concern?", level=3)
    add_para(doc, "SUID files execute with the file owner's permissions — usually root — regardless of who runs them. While many SUID binaries are legitimate (passwd needs SUID-root to write to /etc/shadow), any vulnerability in an SUID binary can be exploited to escalate from a regular user to root. This is a classic privilege escalation vector: attackers scan for unusual or outdated SUID binaries and cross-reference against known exploits. All SUID files should be periodically audited and the bit should only be set where strictly necessary.")

    add_heading(doc, "Question 2.4: What is the principle of least privilege?", level=3)
    add_para(doc, "The principle of least privilege means granting users only the minimum access rights needed for their specific role — nothing more. In user management this is demonstrated by mary and paul's sudoers entries: instead of full sudo group membership, they are restricted to exact commands relevant to their role. This limits the blast radius of a compromised account, reduces accidental misconfiguration risk, and makes auditing simpler since each account's purpose and scope are clearly defined.")

    doc.add_page_break()

    # ── PART 3 ──────────────────────────────────────────────────────
    add_heading(doc, "Part 3: User Activity Monitoring")

    add_heading(doc, "Exercise 3.1: Analyze Authentication Logs", level=2)
    add_note(doc, "sudo lastb is not available on this Kali version. The equivalent command used was: sudo journalctl -u ssh --since '7 days ago' | grep -i 'failed|invalid'")
    add_para(doc, "Commands used:", bold=True)
    for cmd in [
        "sudo last -a",
        "sudo journalctl -u ssh --since '7 days ago' | grep -i 'failed\\|invalid'",
        "sudo tail -20 /var/log/auth.log",
        "w",
        "who -a",
        "last -n 20",
    ]:
        add_code(doc, cmd)

    add_para(doc, "Key findings:", bold=True)
    add_bullet(doc, "last -a: legion logged in from 192.168.5.139 multiple times on Jun 25. No unexpected sessions.")
    add_bullet(doc, "journalctl revealed repeated failed root authentication attempts on Jun 25 between 09:32-09:33, including Google Authenticator 'Invalid verification code' errors.")
    add_bullet(doc, "Investigation confirmed this was self-generated: MFA configured in Lab 3 was still active and the verification code requirement was not anticipated during a routine login attempt.")
    add_bullet(doc, "From log evidence alone, this pattern is indistinguishable from a brute-force attack — demonstrating why monitoring matters even when failures are benign.")
    add_bullet(doc, "auth.log did not capture the same failures found in journalctl, confirming journalctl is the more complete authentication log source on this system.")

    add_heading(doc, "Question 3.1: What patterns might indicate a brute-force attack?", level=3)
    add_para(doc, "Indicators include: repeated failed login attempts in quick succession from the same IP or targeting the same account; many failed attempts across different usernames from one source (credential stuffing); failed logins at unusual hours; and sudden spikes in authentication failures compared to baseline. Targeting of privileged accounts like root is particularly significant.")

    add_heading(doc, "Question 3.2: Why is it important to monitor root login attempts?", level=3)
    add_para(doc, "Root has unrestricted access to the entire system. A compromised root account represents complete system compromise — full access to all files, ability to install backdoors, delete logs, and control all other accounts. Monitoring root login attempts provides early warning of privilege-escalation attempts and allows security teams to respond before damage is done.")

    add_heading(doc, "Exercise 3.2: Implement Account Lockout Policies", level=2)
    add_note(doc, "pam_tally2.so has been removed from current Kali PAM packages and replaced by pam_faillock.so. The faillock equivalent was used throughout.")
    add_para(doc, "Configuration added to /etc/pam.d/common-auth:", bold=True)
    for line in [
        "auth    required       pam_faillock.so preauth silent deny=5 unlock_time=1800",
        "auth    [success=1 default=ignore]    pam_unix.so nullok",
        "auth    [default=die]  pam_faillock.so authfail deny=5 unlock_time=1800",
    ]:
        add_code(doc, line)

    add_para(doc, "Deprecated commands and modern equivalents:", bold=True)
    add_bullet(doc, "sudo pam_tally2 --user=username  →  sudo faillock --user username")
    add_bullet(doc, "sudo pam_tally2 --user=username --reset  →  sudo faillock --user username --reset")

    add_para(doc, "Lockout incident:", bold=True)
    add_para(doc, "Testing confirmed the lockout policy works — incorrect passwords correctly triggered a system-wide account lockout. The 30-minute auto-unlock did not clear within the observed timeframe, suggesting a possible issue with this faillock build. Recovery required GRUB-level root shell access (init=/bin/bash) and manual faillock reset. A secondary recovery account was subsequently created as a safeguard.")

    add_heading(doc, "Question 3.3: Trade-offs between security and usability for account lockout?", level=3)
    add_para(doc, "Lockout policies prevent brute-force attacks but create real usability risks: legitimate users can lock themselves out through honest mistakes, and if recovery is not straightforward this means real downtime — as demonstrated in this lab. Attackers can also weaponize lockout for denial-of-service. Good policy design tunes thresholds carefully, ensures auto-unlock is reliable, and maintains a tested out-of-band recovery path so a lockout never results in complete access loss.")

    doc.add_page_break()

    # ── PART 4 ──────────────────────────────────────────────────────
    add_heading(doc, "Part 4: Sudo Access Management")

    add_heading(doc, "Exercise 4.1: Configure Fine-Grained Sudo Permissions", level=2)
    add_para(doc, "Configuration added to /etc/sudoers via visudo:", bold=True)
    for line in [
        "legion ALL=(ALL) NOPASSWD: /usr/bin/systemctl restart apache2, /usr/bin/systemctl status apache2",
        "legion ALL=(webuser) /usr/bin/vim /var/www/html/*",
        "%developers ALL=(ALL:ALL) ALL",
        "Defaults logfile=\"/var/log/sudo.log\"",
        "Defaults log_input, log_output",
    ]:
        add_code(doc, line)

    add_para(doc, "Test result: sudo systemctl status apache2 ran without a password prompt — NOPASSWD rule confirmed active. Sudo log at /var/log/sudo.log recorded both test commands with full timestamp, TTY, working directory, and exact command.")

    add_heading(doc, "Question 4.1: Why use visudo instead of directly editing /etc/sudoers?", level=3)
    add_para(doc, "visudo validates sudoers file syntax before saving. A typo when editing directly with nano or vim creates a broken file immediately with no validation — and a syntax error in /etc/sudoers can make sudo completely non-functional, locking the administrator out of all privileged access. visudo detects errors before writing and allows correction.")

    add_heading(doc, "Question 4.2: Security implications of NOPASSWD in sudo configurations?", level=3)
    add_para(doc, "NOPASSWD removes password authentication for specific listed commands. If a session is left unattended, anyone at that terminal can run those privileged commands without authentication. In the event of account compromise, an attacker gains instant access without knowing the user's password. NOPASSWD on specific low-risk commands in controlled contexts is manageable; NOPASSWD: ALL is effectively passwordless root access and should never be used unless fully justified.")

    doc.add_page_break()

    # ── PART 5 ──────────────────────────────────────────────────────
    add_heading(doc, "Part 5: Security Compliance and Best Practices")

    add_heading(doc, "Exercise 5.1: Security Audit Script", level=2)
    add_note(doc, "The lab's original script used sudo lastb for failed logins, replaced with sudo journalctl -u ssh --since '7 days ago' | grep -i 'failed|invalid' since lastb is unavailable on this Kali version.")
    add_para(doc, "Script security_audit.sh was created and run successfully. All 7 checks produced expected output. Key results:")
    add_bullet(doc, "UID 0: Only root — no rogue accounts.")
    add_bullet(doc, "Empty passwords: None found.")
    add_bullet(doc, "Sudo group: legion and recovery (recovery account added as lab safety net).")
    add_bullet(doc, "Password aging: PASS_MAX_DAYS 90, PASS_MIN_DAYS 7, PASS_WARN_AGE 14 — confirmed active.")
    add_bullet(doc, "SUID files: 10 listed — expected system binaries only.")
    add_bullet(doc, "Failed logins: MFA-related failures from Jun 25 — self-generated, documented above.")
    add_bullet(doc, "Logged-in users: legion (graphical) and recovery (tty2) — both expected.")

    add_heading(doc, "Question 5.1: What additional checks would you add?", level=3)
    add_bullet(doc, "Accounts inactive for 90+ days: sudo lastlog2 -b 90")
    add_bullet(doc, "Password expiration status: sudo passwd -S -a")
    add_bullet(doc, "Service accounts with unnecessary login shells (postgres, arpwatch, sbuild found in this audit)")
    add_bullet(doc, "World-writable files: sudo find / -perm -0002 -type f 2>/dev/null")
    add_bullet(doc, "Listening network services: ss -tulnp")
    add_bullet(doc, "Unowned files: sudo find / -nouser -o -nogroup 2>/dev/null")
    add_bullet(doc, "SSH hardening check: grep -E 'PermitRootLogin|PasswordAuthentication' /etc/ssh/sshd_config")

    add_heading(doc, "Exercise 5.2: User Access Report", level=2)
    add_para(doc, "user_access_report.sh was created and run successfully. Security recommendations documented in security_recommendations.txt:")
    add_bullet(doc, "Service accounts (postgres, arpwatch, sbuild) have unnecessary login shells — should be changed to /usr/sbin/nologin.")
    add_bullet(doc, "User accounts (john, alice, bob, charlie, mary, paul, student1) not subject to password aging — chage should be applied to all active accounts.")
    add_bullet(doc, "Account lockout (pam_faillock) configured and verified — but unlock_time behaviour requires further investigation.")
    add_bullet(doc, "Secondary recovery account created and verified as administrative lockout safeguard.")
    add_bullet(doc, "SSH root login and Google Authenticator MFA configuration should be reviewed and documented.")

    add_heading(doc, "Question 5.2: How often should user access reviews be conducted?", level=3)
    add_para(doc, "High-security environments (financial, healthcare, government) should review monthly (every 30 days). Standard business environments typically review quarterly (every 90 days). At minimum all environments should review annually. Beyond scheduled reviews, access should also be reviewed immediately following trigger events: staff departures, role changes, security incidents, or significant system changes. Regulatory frameworks including ISO 27001, PCI-DSS, and HIPAA each specify their own mandated intervals.")

    doc.add_page_break()

    # ── PART 6 ──────────────────────────────────────────────────────
    add_heading(doc, "Part 6: Practical Challenge — Secure a Multi-User System")

    add_heading(doc, "Step 1: Create Groups", level=2)
    for cmd in ["sudo groupadd developers", "sudo groupadd analysts", "sudo groupadd admins"]:
        add_code(doc, cmd)
    add_para(doc, "Result: developers:x:1005 | analysts:x:2009 | admins:x:2010 — all three confirmed.")

    add_heading(doc, "Step 2: Create Users", level=2)
    for cmd in [
        "sudo useradd -m -s /bin/bash -G developers dev1",
        "sudo useradd -m -s /bin/bash -G developers dev2",
        "sudo useradd -m -s /bin/bash -G analysts analyst1",
        "sudo useradd -m -s /bin/bash -G analysts analyst2",
        "sudo useradd -m -s /bin/bash -G admins admin1",
        "sudo useradd -m -s /bin/bash -G admins admin2",
    ]:
        add_code(doc, cmd)
    add_para(doc, "Result: developers:dev1,dev2 | analysts:analyst1,analyst2 | admins:admin1,admin2")

    add_heading(doc, "Step 3: Password Policies", level=2)
    add_para(doc, "Passwords set for all six users. pwquality.conf enforces minimum 12 characters with complexity requirements system-wide. Password aging applied (60-day expiry, 10-day minimum, 14-day warning):")
    for cmd in [
        "sudo chage -M 60 -m 10 -W 14 dev1",
        "sudo chage -M 60 -m 10 -W 14 dev2",
        "sudo chage -M 60 -m 10 -W 14 analyst1",
        "sudo chage -M 60 -m 10 -W 14 analyst2",
        "sudo chage -M 60 -m 10 -W 14 admin1",
        "sudo chage -M 60 -m 10 -W 14 admin2",
    ]:
        add_code(doc, cmd)
    add_para(doc, "Verification (dev1): Password expires Aug 30 2026 | Min days: 10 | Max days: 60 | Warning: 14 days")

    add_heading(doc, "Step 4: Sudo Access Configuration", level=2)
    add_code(doc, "%admins ALL=(ALL:ALL) ALL")
    add_code(doc, "%developers ALL=(ALL) NOPASSWD: /usr/bin/systemctl restart apache2, /usr/bin/systemctl restart nginx")
    add_para(doc, "Security rationale:", bold=True)
    add_bullet(doc, "Admins: full sudo justified — administrative duties require unrestricted system access.")
    add_bullet(doc, "Developers: restricted to specific web service commands only — principle of least privilege applied.")
    add_bullet(doc, "Analysts: no sudo access — data analysis role requires no system-level privileges.")

    add_heading(doc, "Step 5: Account Lockout Policy", level=2)
    add_code(doc, "auth    required       pam_faillock.so preauth silent deny=3 unlock_time=900")
    add_code(doc, "auth    [success=1 default=ignore]    pam_unix.so nullok")
    add_code(doc, "auth    [default=die]  pam_faillock.so authfail deny=3 unlock_time=900")
    add_bullet(doc, "deny=3: blocks automated brute-force tools while allowing for genuine user mistakes.")
    add_bullet(doc, "unlock_time=900: 15-minute auto-unlock balances security with usability.")
    add_bullet(doc, "Policy verified working — lockout triggered and confirmed active across all PAM services.")

    add_heading(doc, "Step 6: Weekly Audit Script", level=2)
    add_para(doc, "weekly_audit.sh created covering 9 security checks. Generates a timestamped report file on each run. Test run (02 July 2026): All 9 checks completed successfully. Report saved to weekly_audit_20260702.txt.")

    doc.add_page_break()

    # ── SUMMARY ─────────────────────────────────────────────────────
    add_heading(doc, "Lab Summary")
    add_bullet(doc, "Part 1: Password aging and complexity configured for testuser and applied system-wide.")
    add_bullet(doc, "Part 2: User accounts audited — key findings: service accounts with unnecessary shells; users missing aging policy.")
    add_bullet(doc, "Part 3: Auth logs analyzed using journalctl. Account lockout configured using pam_faillock and verified working.")
    add_bullet(doc, "Part 4: Fine-grained sudo configured via visudo with logging enabled.")
    add_bullet(doc, "Part 5: Automated security audit and user access report scripts created and run.")
    add_bullet(doc, "Part 6: Complete multi-user security configuration implemented from scratch.")
    doc.add_paragraph()
    add_para(doc, "Throughout this lab several deprecated tools were encountered (lastlog, lastb, pam_tally2) and replaced with modern equivalents (lastlog2, journalctl, pam_faillock), reflecting the current state of security tooling on Kali Linux 7.0.12.")

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
