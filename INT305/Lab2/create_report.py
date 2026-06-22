from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT305: Network Security and Protocols', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 2: File and Directory Permissions', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Student: Damian Ozimede Patrick')
doc.add_paragraph('Course: INT305 - Network Security and Protocols')
doc.add_paragraph('GitHub: github.com/damianozimede/Kali-Linux-Security-Labs')
doc.add_paragraph('Platform: Kali Linux | Username: legion | Date: June 2026')

# 1. Objective
doc.add_heading('1. Objective', 2)
doc.add_paragraph(
    'The objective of this lab was to understand how to set and manage permissions for files '
    'and directories in Linux, and to learn the significance of permissions in securing files '
    'and directories against unauthorised access. The lab covered checking existing permissions, '
    'modifying them using both numeric and symbolic modes, managing directory permissions, '
    'and changing file and directory ownership.'
)

# 2. Environment
doc.add_heading('2. Environment', 2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('Operating System: Kali Linux (legion@pastor)')
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('Machine IP: 192.168.5.139')
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run('User: legion (administrator), student1 (test user)')
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run('Group: students')

# 3. Understanding Permissions
doc.add_heading('3. Understanding File Permissions', 2)
doc.add_paragraph(
    'Linux file permissions control who can read, write, and execute files. Each file has three '
    'sets of permissions assigned to three categories of users: the owner, the group, and others. '
    'Each set can have up to three permissions: read (r), write (w), and execute (x). These are '
    'represented in the ls -l output as a string of ten characters, for example -rwxr-xr-x.'
)
doc.add_paragraph('A test file was created and its default permissions were checked:')
doc.add_paragraph('Command: touch testfile.txt && ls -l testfile.txt')
doc.add_paragraph(
    'Output: -rw-rw-r-- 1 legion legion 0 Jun 22 08:40 testfile.txt. The first character (-) '
    'means it is a regular file. The next three (rw-) are owner permissions showing read and '
    'write but no execute. The middle three (rw-) are group permissions also showing read and '
    'write. The final three (r--) are others permissions showing read only.'
)
doc.add_paragraph('[SCREENSHOT 1: ls -l testfile.txt showing default -rw-rw-r-- permissions]')

# 4. Numeric Mode
doc.add_heading('4. Changing Permissions - Numeric Mode', 2)
doc.add_paragraph(
    'Permissions were changed using numeric (octal) mode. Each permission type is assigned a '
    'value: read = 4, write = 2, execute = 1. These are added together for each user category '
    'to produce a three-digit number. For example, 755 means the owner gets 7 (rwx), '
    'the group gets 5 (r-x), and others get 5 (r-x).'
)
doc.add_paragraph('Command: chmod 755 testfile.txt && ls -l testfile.txt')
doc.add_paragraph(
    'Output: -rwxr-xr-x 1 legion legion 0 Jun 22 08:40 testfile.txt, confirming the owner '
    'now has full read, write and execute access while group and others have read and execute only.'
)
doc.add_paragraph('[SCREENSHOT 2: ls -l testfile.txt showing -rwxr-xr-x after chmod 755]')

# 5. Symbolic Mode
doc.add_heading('5. Changing Permissions - Symbolic Mode', 2)
doc.add_paragraph(
    'Permissions were also changed using symbolic mode, which allows targeted adjustments to '
    'specific permissions without affecting others. The syntax uses letters to represent user '
    'categories (u = user/owner, g = group, o = others) and operators (+ to add, - to remove, '
    '= to set exactly).'
)
doc.add_paragraph('Command: chmod u+x testfile.txt && ls -l testfile.txt')
doc.add_paragraph(
    'Since the file already had execute permission for the owner from chmod 755, the output '
    'remained -rwxr-xr-x. The key difference between the two modes is that numeric mode sets '
    'all permissions at once for all three categories, while symbolic mode makes targeted changes '
    'to specific permissions without disturbing the rest. For example, if the file had been '
    'rw-r--r--, running chmod u+x would produce rwxr--r-- without touching group or others.'
)
doc.add_paragraph('[SCREENSHOT 3: ls -l testfile.txt showing permissions after chmod u+x]')

# 6. Directory Permissions
doc.add_heading('6. Managing Directory Permissions', 2)
doc.add_heading('6.1 Creating and Setting Directory Permissions', 3)
doc.add_paragraph(
    'A directory named project was created and its permissions were set to 755 to allow '
    'the owner full access while restricting group and others to read and execute:'
)
doc.add_paragraph('Command: mkdir project && chmod 755 project && ls -ld project')
doc.add_paragraph(
    'Output: drwxr-xr-x 2 legion legion 4096 Jun 22 08:50 project. The d at the beginning '
    'confirms this is a directory. Permissions 755 give the owner full rwx access while group '
    'and others can read and navigate into the directory but cannot create or delete files within it.'
)
doc.add_paragraph('[SCREENSHOT 4: ls -ld project showing drwxr-xr-x with legion as owner]')

doc.add_heading('6.2 Changing Directory Ownership', 3)
doc.add_paragraph(
    'The student1 user was recreated since it had been deleted in Lab 1, then the ownership '
    'of the project directory was changed to student1 with the students group:'
)
doc.add_paragraph('Command: sudo chown student1:students project && ls -ld project')
doc.add_paragraph(
    'Output: drwxr-xr-x 2 student1 students 4096 Jun 22 08:50 project, confirming the '
    'ownership was successfully transferred. The chown command with the format owner:group '
    'sets both the user and group ownership in a single operation.'
)
doc.add_paragraph('[SCREENSHOT 5: ls -ld project showing student1 students as owner and group]')

# 7. Exercise
doc.add_heading('7. Exercise - Creating my_project', 2)
doc.add_paragraph(
    'The exercise required creating a directory named my_project, setting its permissions '
    'to 755, and changing its ownership to student1:students:'
)
doc.add_paragraph('Command: mkdir my_project && chmod 755 my_project && sudo chown student1:students my_project && ls -ld my_project')
doc.add_paragraph(
    'Output: drwxr-xr-x 2 student1 students 4096 Jun 22 08:58 my_project, confirming all '
    'three requirements were met. The directory was created, permissions set to 755, and '
    'ownership assigned to student1 with the students group.'
)
doc.add_paragraph('[SCREENSHOT 6: ls -ld my_project showing drwxr-xr-x with student1 students ownership]')

# 8. Analysis
doc.add_heading('8. Analysis and Security Significance', 2)
doc.add_paragraph(
    'File and directory permissions are one of the foundational mechanisms for securing a Linux '
    'system. Incorrect permissions can expose sensitive files to unauthorised users, allow '
    'malicious users to modify system files, or permit execution of untrusted scripts. Setting '
    'a file to 777 (rwxrwxrwx) gives every user on the system full access, which is a serious '
    'security risk in a multi-user environment.'
)
doc.add_paragraph(
    'The principle of least privilege applies directly to file permissions. Users and processes '
    'should only have the minimum permissions necessary to perform their tasks. Using 755 for '
    'directories and 644 for files is a common and secure default that gives the owner '
    'appropriate control while limiting what others can do.'
)
doc.add_paragraph(
    'Ownership management through chown is equally important. Assigning files to the correct '
    'user and group ensures that permission rules are applied to the right people. Group-based '
    'ownership allows teams to share access to project directories without opening them up to '
    'the entire system.'
)

# 9. Conclusion
doc.add_heading('9. Conclusion', 2)
doc.add_paragraph(
    'This lab provided practical experience with Linux file and directory permissions. The key '
    'commands chmod and chown were used to control access to files and directories, and the '
    'difference between numeric and symbolic permission modes was demonstrated. The exercise '
    'reinforced these concepts by creating a directory with specific permissions and ownership, '
    'mirroring real-world scenarios where access control is critical to system security. '
    'Understanding and correctly applying file permissions is an essential skill for any system '
    'administrator or security professional working in a Linux environment.'
)

# Appendix
doc.add_heading('Appendix: Commands Used', 2)
commands = [
    'touch testfile.txt',
    'ls -l testfile.txt',
    'chmod 755 testfile.txt',
    'chmod u+x testfile.txt',
    'mkdir project',
    'chmod 755 project',
    'ls -ld project',
    'sudo adduser student1',
    'sudo chown student1:students project',
    'mkdir my_project',
    'chmod 755 my_project',
    'sudo chown student1:students my_project',
    'ls -ld my_project',
]
for cmd in commands:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd)

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT305/Lab2/INT305_Lab2_Report.docx')
print("Report saved successfully!")
