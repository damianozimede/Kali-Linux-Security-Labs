from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT303: Networking Fundamentals', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 3: TCP/IP Protocol Stack and Packet Inspection', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Student: Damian Patrick Ozimede')
doc.add_paragraph('Course: INT303 — Networking Fundamentals')
doc.add_paragraph('Lab: Lab 3')
doc.add_paragraph('Date: June 2026')

doc.add_paragraph('')

# Exercise 1
doc.add_heading('Exercise 1: Understanding the TCP/IP Model', 2)

doc.add_heading('Question 1: Explain the differences between the TCP/IP and OSI models.', 3)
doc.add_paragraph('The TCP/IP model and the OSI model are both frameworks for understanding network communication but differ in several important ways.\n\nNumber of Layers: The OSI model has 7 layers while the TCP/IP model has only 4 layers. The TCP/IP model combines several OSI layers into single layers for simplicity.\n\nPurpose and Origin: The OSI model was developed by the International Organisation for Standardisation as a theoretical reference model and was never fully implemented as a working protocol suite. The TCP/IP model was developed by the United States Department of Defense as a practical working protocol suite that became the foundation of the modern internet.\n\nLayer Structure: The OSI model separates the Application layer into three distinct layers — Application, Presentation, and Session. The TCP/IP model combines all three into a single Application layer. Similarly the OSI model separates the Physical and Data Link layers while TCP/IP combines them into a single Link layer.\n\nFlexibility: The OSI model is protocol-independent and describes how any networking protocol should work. The TCP/IP model is built specifically around the TCP and IP protocols making it less flexible but more directly applicable to real-world networking.\n\nAdoption: The TCP/IP model is the standard used in practice for all internet and modern network communications. The OSI model is primarily used as a teaching and reference tool.')

doc.add_heading('Question 2: Which layers of the TCP/IP model correspond to specific OSI layers?', 3)

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'TCP/IP Layer'
hdr[1].text = 'Corresponding OSI Layers'
hdr[2].text = 'Key Protocols'
rows = [
    ('Application', 'Layer 7 – Application', 'HTTP, HTTPS, FTP, SMTP, DNS, SSH'),
    ('Application', 'Layer 6 – Presentation', 'SSL/TLS, JPEG, ASCII'),
    ('Application', 'Layer 5 – Session', 'NetBIOS, RPC'),
    ('Transport', 'Layer 4 – Transport', 'TCP, UDP'),
    ('Internet', 'Layer 3 – Network', 'IP, ICMP, ARP'),
    ('Link', 'Layer 2 – Data Link', 'Ethernet, MAC, ARP'),
    ('Link', 'Layer 1 – Physical', 'Cables, Wi-Fi, fibre optics'),
]
for i, (tcpip, osi, protocols) in enumerate(rows):
    row = table.rows[i+1].cells
    row[0].text = tcpip
    row[1].text = osi
    row[2].text = protocols

doc.add_paragraph('\nThe TCP/IP Application layer combines OSI Layers 5, 6, and 7. The Transport layer corresponds directly to OSI Layer 4. The Internet layer corresponds to OSI Layer 3. The Link layer combines OSI Layers 1 and 2.')

doc.add_paragraph('')

# Exercise 2
doc.add_heading('Exercise 2: Capturing and Analyzing TCP Packets', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('sudo tcpdump -i eth0 tcp and host 192.168.5.130')

doc.add_heading('Question 1: What happens during the TCP handshake (SYN, SYN-ACK, ACK)?', 3)
doc.add_paragraph('The TCP three-way handshake was clearly visible in the first three packets of the capture:\n\nPacket 1 — SYN (19:06:13.402906): The Kali machine (192.168.5.139) initiated the connection by sending a SYN packet to the OWASP VM (192.168.5.130) on port 80. This packet contained Flags [S], Sequence number 1209652905, Window size 64240, and MSS 1460. It says "I want to connect and my starting sequence number is 1209652905".\n\nPacket 2 — SYN-ACK (19:06:13.405092): The OWASP VM responded with Flags [S.], Sequence number 1777608837, and Acknowledgement number 1209652906, confirming receipt and readiness.\n\nPacket 3 — ACK (19:06:13.405216): The Kali machine completed the handshake with Flags [.] and Acknowledgement number 1. The connection was now fully established.\n\nThe entire three-way handshake completed in just 0.002310 seconds.')

doc.add_heading('Question 2: Identify and describe the flags used in TCP communication.', 3)
doc.add_paragraph('[S] — SYN: Initiates a TCP connection. Sets the starting sequence number. Only seen during connection establishment.\n\n[S.] — SYN-ACK: Server accepts the connection, acknowledges the client SYN, and sends its own sequence number.\n\n[.] — ACK: The most common flag. Confirms successful receipt of data up to a certain sequence number.\n\n[P.] — PSH-ACK: Instructs the receiver to pass data immediately to the application without buffering. Seen when HTTP GET request was sent and when final data chunks were delivered.\n\n[F.] — FIN-ACK: Signals the sender has finished sending data. Used in the connection termination sequence.')

doc.add_heading('Question 3: How does TCP maintain reliability during transmission?', 3)
doc.add_paragraph('TCP maintains reliability through four key mechanisms:\n\nSequence Numbers — Every packet is numbered so data can be reassembled in the correct order. Sequence numbers incremented consistently throughout the capture (1:2897, 2897:4345, 4345:5793 etc.).\n\nAcknowledgements — Every packet received was immediately acknowledged. If an ACK is not received within a timeout period TCP automatically retransmits the missing data.\n\nFlow Control (Window Size) — The window size field tells the sender how much data the receiver can handle at once, preventing the sender from overwhelming the receiver.\n\nClean Connection Termination — The FIN/ACK exchange ensured all data was fully delivered before the connection closed, preventing data loss.')

doc.add_paragraph('')

# Exercise 3
doc.add_heading('Exercise 3: Investigating IP Packets (Network Layer)', 2)
doc.add_heading('Tool Used:', 3)
doc.add_paragraph('Wireshark on eth0\nFilter: ip.addr == 192.168.5.130')

doc.add_heading('Question 1: What fields can you see in the IP packet header?', 3)
doc.add_paragraph('The expanded Internet Protocol Version 4 section in Wireshark revealed:\n\nVersion: 4 (IPv4)\nHeader Length: 20 bytes (5 x 32-bit words)\nDifferentiated Services Field: 0x00 (DSCP: CS0, ECN: Not-ECT)\nTotal Length: 60 bytes\nIdentification: 0x029c (668)\nFlags: 0x2 — Don\'t Fragment\nFragment Offset: 0\nTime to Live (TTL): 64\nProtocol: TCP (6)\nHeader Checksum: 0xabc2\nSource Address: 192.168.5.139\nDestination Address: 192.168.5.130')

doc.add_heading('Question 2: What is the significance of each of these fields?', 3)
doc.add_paragraph('Version — Identifies this as IPv4.\n\nHeader Length — Tells the receiver where the header ends and data begins. 20 bytes is the standard minimum IP header size.\n\nDifferentiated Services Field — Used for Quality of Service (QoS). A value of 0x00 means normal priority.\n\nTotal Length — The complete size of the packet including header and payload.\n\nIdentification — A unique number for each packet used to reassemble fragmented packets at the destination.\n\nFlags (Don\'t Fragment) — Tells routers not to fragment this packet. If too large for a network segment it will be dropped and an ICMP error sent back.\n\nFragment Offset — Indicates where in the original packet this fragment belongs. Value 0 means not fragmented.\n\nTTL — Set to 64, decrements by 1 at each router hop. Prevents packets circulating endlessly in routing loops.\n\nProtocol — Value 6 identifies TCP as the transport layer protocol. Other values include 17 for UDP and 1 for ICMP.\n\nHeader Checksum — Verifies the integrity of the IP header. Mismatched checksums cause the packet to be discarded.\n\nSource and Destination Addresses — Identify the sending host (192.168.5.139) and receiving host (192.168.5.130) at the Network layer.')

doc.add_heading('Question 3: How does IP routing work in this scenario? Are there any hops?', 3)
doc.add_paragraph('Both the Kali machine (192.168.5.139) and the OWASP VM (192.168.5.130) are on the same local subnet (192.168.5.0/24). Because they share the same subnet, packets do not need to pass through any routers.\n\nWhen the Kali machine sends a packet to 192.168.5.130, it checks its routing table, determines the destination is on the same local network, uses ARP to resolve the MAC address, and delivers the packet directly over the local Ethernet network.\n\nThis is confirmed by the TTL value of 64 — it did not decrement below 64, confirming zero router hops. This is consistent with the traceroute result from Lab 1 where the OWASP VM was reached in just one effective hop on the local network.')

doc.add_paragraph('')

# Exercise 4
doc.add_heading('Exercise 4: Application Layer Analysis (HTTP/SSH Traffic)', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('Wireshark filter: tcp.port == 80 and ip.addr == 192.168.5.130\nsudo tcpdump -i eth0 port 22 and host 192.168.5.130\nssh -oHostKeyAlgorithms=+ssh-rsa root@192.168.5.130')

doc.add_heading('Question 1: Analyze the HTTP packets. What information is available in the HTTP request and response?', 3)
doc.add_paragraph('The Wireshark HTTP capture showed the full HTTP conversation:\n\nHTTP Request:\nSource: 192.168.5.139, Destination: 192.168.5.130\nMethod: GET / HTTP/1.1\nLength: 143 bytes\nThe request asked for the root page (/) of the OWASP web server using HTTP version 1.1\n\nHTTP Response:\nSource: 192.168.5.130, Destination: 192.168.5.139\nStatus: HTTP/1.1 200 OK — request was successful\nResponse delivered in multiple TCP segments totalling 28,515 bytes of HTML content\nTCP PDU reassembly confirmed the webpage was broken into chunks and reassembled at the destination\n\nHTTP is a plaintext protocol — the request method, URL, and response status were all clearly visible without any encryption.')

doc.add_heading('Question 2: For SSH traffic, what is the significance of encrypted packets? Can you analyze the payload?', 3)
doc.add_paragraph('The SSH tcpdump capture recorded 24 packets. The SSH version exchange was visible in plain text at the start:\nKali machine: SSH-2.0-OpenSSH_10.2p1 Debian-6\nOWASP VM: SSH-2.0-OpenSSH_5.3p1 Debian-3ubuntu4\n\nAfter this initial exchange, all subsequent traffic was completely encrypted. The packet payloads showed only raw encrypted binary data — no readable content, usernames, passwords, or commands were visible.\n\nThis is the fundamental significance of SSH encryption — even though packets can be captured, their contents cannot be read without the private encryption keys. This contrasts sharply with HTTP where all data is transmitted in plain text and fully visible in packet captures.')

doc.add_heading('Question 3: How does the application layer play a role in data exchange?', 3)
doc.add_paragraph('The Application layer is the topmost layer and plays a direct role in how meaningful data is formatted, requested, and delivered:\n\nFor HTTP: The Application layer defined the rules for how the Kali machine requested a web page (GET / HTTP/1.1) and how the OWASP VM responded (HTTP/1.1 200 OK). Without HTTP the TCP connection would just be a raw data stream with no structure.\n\nFor SSH: The Application layer handled the SSH protocol negotiation, version exchange, encryption algorithm selection, and authentication.\n\nIn both cases the Application layer relied on the layers below it — Transport (TCP), Internet (IP), and Link (Ethernet) — to physically move data. The Application layer concerns itself only with what the data means and how it should be formatted.')

doc.add_paragraph('')

# Exercise 5
doc.add_heading('Exercise 5: Error Handling in TCP/IP Transmission', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('ping 192.168.5.130\nsudo tcpdump -i eth0 host 192.168.5.130\nsudo ifconfig eth0 down\nsudo ifconfig eth0 up')

doc.add_heading('Question 1: What happens when packets are dropped or delayed?', 3)
doc.add_paragraph('The ping sequence numbers jumped from icmp_seq=13 directly to icmp_seq=32, meaning 18 consecutive packets were completely lost during the outage. The tcpdump capture confirmed this — after seq 13 at 18:20:03, the next packet did not appear until 18:20:22, a gap of approximately 19 seconds.\n\nThe final ping statistics confirmed: 38 packets transmitted, only 20 received, giving a 47.3684% packet loss. When packets are dropped using ICMP, the data is simply lost with no automatic recovery since ICMP has no built-in retransmission mechanism.')

doc.add_heading('Question 2: How does TCP ensure data reliability in the presence of errors?', 3)
doc.add_paragraph('TCP handles errors through several built-in reliability mechanisms:\n\nRetransmission — If a TCP segment is sent but no ACK is received within a timeout period, TCP automatically retransmits the missing segment without any intervention from the application.\n\nDuplicate ACKs — If packets arrive out of order, the receiver sends duplicate ACKs for the last successfully received packet, signalling to the sender that something was missed.\n\nTimeout and Backoff — If multiple retransmissions fail, TCP increases the waiting time between retries (exponential backoff) to avoid overwhelming an already congested network.')

doc.add_heading('Question 3: How do retransmissions and sequence numbers work in TCP?', 3)
doc.add_paragraph('Every byte of data sent over a TCP connection is assigned a sequence number. The receiver confirms exactly which bytes have been received by sending ACK packets containing the next expected sequence number.\n\nIf a segment is lost, the receiver stops sending ACKs beyond the missing segment. The sender notices ACKs have stopped and after its retransmission timeout expires, resends the missing segment. Once received, the receiver acknowledges all buffered data and normal flow resumes.\n\nThis system ensures that even in the presence of network disruptions, TCP can detect exactly what was lost and recover precisely — no data is duplicated and no data is permanently lost as long as the connection eventually recovers.')

doc.add_paragraph('')

# Exercise 6
doc.add_heading('Exercise 6: ICMP and Ping Inspection (Network Layer)', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('sudo tcpdump -i eth0 icmp and host 192.168.5.130\nping -c 4 192.168.5.130')

doc.add_heading('Question 1: What are the key fields in an ICMP packet?', 3)
doc.add_paragraph('The tcpdump capture showed 8 ICMP packets — 4 Echo Requests and 4 Echo Replies. Key fields:\n\nType — Identifies the purpose: Type 8 = Echo Request, Type 0 = Echo Reply\nCode — Additional context for the Type field. For Echo Request/Reply the code is always 0\nChecksum — Verifies the integrity of the ICMP packet\nIdentifier — Unique ID for each ping session (54147 in this capture) allowing replies to be matched to requests\nSequence Number — Increments with each packet (1, 2, 3, 4) to track order and identify missing replies\nLength — Each packet was 64 bytes, the default ping size on Linux')

doc.add_heading('Question 2: How does ICMP assist in diagnosing network connectivity issues?', 3)
doc.add_paragraph('Testing reachability — Ping uses ICMP Echo Request and Reply to test whether a host is alive and reachable.\n\nMeasuring latency — Round trip times (min 1.081 ms, avg 1.700 ms, max 3.098 ms) measure how long packets take to travel to the destination and back. High latency indicates congestion or routing problems.\n\nDetecting packet loss — Ping statistics immediately show packet loss percentage, pinpointing exactly when and where network failures occur.\n\nTraceroute — ICMP is also used by traceroute to map network paths by exploiting the TTL field.')

doc.add_heading('Question 3: What is the significance of TTL in ICMP and general IP packets?', 3)
doc.add_paragraph('TTL had a value of 64 in all captured ICMP packets, which is the default for Linux systems.\n\nPreventing routing loops — Every time a packet passes through a router the TTL decrements by 1. If TTL reaches 0 the router discards the packet and sends an ICMP Time Exceeded message back. This prevents packets circulating indefinitely.\n\nInferring network topology — TTL reveals OS information and hop count. TTL=64 is typical of Linux (Windows uses 128, Cisco uses 255). The OWASP VM replies arriving with TTL=64 confirmed it is Linux-based and that packets travelled through zero router hops.')

doc.add_paragraph('')

# Exercise 7
doc.add_heading('Exercise 7: Analyzing UDP Packets (Transport Layer)', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('sudo tcpdump -i eth0 udp and host 192.168.5.130\nsudo nmap -sU --top-ports 20 192.168.5.130')

doc.add_heading('Question 1: Compare UDP with TCP. What are the major differences?', 3)

table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Feature'
hdr2[1].text = 'TCP'
hdr2[2].text = 'UDP'
udp_rows = [
    ('Connection', 'Three-way handshake required', 'No connection — just sends'),
    ('Header Size', '20 bytes minimum', 'Only 8 bytes'),
    ('Sequencing', 'Sequence numbers track order', 'No sequence numbers'),
    ('Acknowledgement', 'Every packet acknowledged', 'No acknowledgements'),
    ('Flags', 'SYN, ACK, FIN, PSH, RST', 'None'),
    ('Error Recovery', 'Retransmission on failure', 'No retransmission'),
]
for i, (feature, tcp, udp) in enumerate(udp_rows):
    row = table2.rows[i+1].cells
    row[0].text = feature
    row[1].text = tcp
    row[2].text = udp

doc.add_paragraph('\nIn the tcpdump capture, UDP packets were fired off rapidly in bursts with no handshake and no waiting for replies. The nmap scan sent probes to 20 different ports almost simultaneously — something TCP could never do without first establishing individual connections.')

doc.add_heading('Question 2: Why does UDP not ensure reliability, and when is it preferred?', 3)
doc.add_paragraph('UDP does not ensure reliability by design — it was built for speed and simplicity with no connection establishment, acknowledgements, sequence numbers, or retransmission.\n\nUDP is preferred when:\nSpeed matters more than perfection — Online gaming, video streaming, and VoIP all use UDP. A dropped frame is far less disruptive than TCP retransmission delays.\nSmall single-request transactions — DNS uses UDP because queries and responses fit in a single packet. TCP handshakes for every DNS query would be unnecessarily slow.\nBroadcasting — NetBIOS and DHCP use UDP broadcasts to reach multiple hosts simultaneously. TCP cannot broadcast.\nNetwork scanning — UDP allows rapid probing of multiple ports without TCP connection overhead.')

doc.add_heading('Question 3: How does UDP manage data transmission without acknowledgements?', 3)
doc.add_paragraph('UDP simply takes data from the application, adds an 8-byte header (source port, destination port, length, checksum), and sends it onto the network.\n\nThe tcpdump capture demonstrated this clearly. Nmap sent probes to ports including netbios-ns, tftp, snmp, ntp, isakmp, dhcp, syslog, and others — all without waiting for responses. Only netbios-ns responded (337 and 229 byte replies) while the rest went unanswered. Nmap noted these as closed or open|filtered and moved on.\n\nAny reliability needed must be handled by the application layer above UDP. Video streaming applications implement their own buffering, DNS has its own retry logic. This keeps UDP lightweight and fast while pushing reliability complexity to applications that need it.')

doc.add_paragraph('')

# Exercise 8
doc.add_heading('Exercise 8: OS Detection via Nmap', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('sudo nmap -O 192.168.5.130')

doc.add_heading('Question 1: How does nmap perform OS detection?', 3)
doc.add_paragraph('Nmap performs OS detection by sending specially crafted TCP, UDP, and ICMP packets and analysing the responses. Every OS implements the TCP/IP stack slightly differently, creating a unique fingerprint compared against nmap\'s database.\n\nKey techniques include:\nTCP ISN Sampling — Analyses the pattern of Initial Sequence Numbers generated by the target\nTCP Options Analysis — Examines which TCP options the target supports and their order\nIP TTL and Window Size — Checks default TTL and TCP window size\nICMP Response Analysis — Analyses how the target responds to specific ICMP probes\n\nThe scan successfully identified the OWASP VM as running Linux 2.6.X, specifically Linux 2.6.17-2.6.36, consistent with the older Ubuntu-based OWASP BWA virtual machine.')

doc.add_heading('Question 2: What packet characteristics help identify the OS?', 3)
doc.add_paragraph('TTL Value — The OWASP VM responded with TTL=64, the default for Linux. Windows uses TTL=128, Cisco routers use TTL=255.\n\nTCP Window Size — The OWASP VM consistently used window size 5792, characteristic of older Linux kernel 2.6.x versions.\n\nOpen Ports — The combination of open ports (22, 80, 139, 143, 443, 445, 5001, 8080, 8081) provided additional OS context.\n\nMAC Address — 00:0C:29:59:23:DD confirmed a VMware virtual machine (00:0C:29 prefix belongs to VMware).\n\nNetwork Distance — 1 hop confirmed direct local network reachability.')

doc.add_heading('Question 3: Why is OS detection important in network analysis and vulnerability assessment?', 3)
doc.add_paragraph('Vulnerability identification — Knowing the exact OS and version allows identification of known vulnerabilities. Linux 2.6.17-2.6.36 is a very old kernel with numerous documented CVEs that can be targeted with specific exploits.\n\nAttack surface mapping — Different operating systems have different default services, ports, and configurations, helping analysts understand available attack vectors.\n\nNetwork inventory — OS detection helps administrators maintain an accurate inventory and identify unauthorised or outdated systems quickly.\n\nPatch management — Knowing which OS versions are running helps prioritise patching efforts, focusing on the most vulnerable systems first.\n\nIncident response — Quickly identifying the OS of affected systems helps responders understand likely attack vectors and potential impact.')

doc.add_paragraph('')

# Exercise 9
doc.add_heading('Exercise 9: Analyzing ARP Traffic (Link Layer)', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('sudo arp-scan -I eth0 192.168.5.130')

doc.add_heading('Question 1: What information can you gather from ARP packets?', 3)
doc.add_paragraph('The arp-scan output revealed:\nIP Address: 192.168.5.130\nMAC Address: 00:0c:29:59:23:dd\nInterface used: eth0\nKali MAC Address: 00:0c:29:db:e8:53\nKali IP Address: 192.168.5.139\nScan duration: 0.117 seconds\n\nThe vendor lookup returned Unknown due to permission issues with vendor database files, however from the nmap scan we know the prefix 00:0c:29 belongs to VMware, confirming the OWASP VM is a virtual machine.\n\nARP packets reveal two critical pieces of information — the IP address and physical MAC address of every device on the local network.')

doc.add_heading('Question 2: How does the ARP protocol function at the Link layer?', 3)
doc.add_paragraph('ARP operates at Layer 2 — the Data Link Layer and bridges Layer 3 (IP addresses) and Layer 2 (MAC addresses):\n\nStep 1 — ARP Request (Broadcast): The Kali machine sent a broadcast to FF:FF:FF:FF:FF:FF asking "Who has 192.168.5.130? Tell 192.168.5.139". Every device on the network received this broadcast.\n\nStep 2 — ARP Reply (Unicast): The OWASP VM responded directly: "192.168.5.130 is at 00:0c:29:59:23:dd".\n\nStep 3 — ARP Cache Update: The Kali machine stored this IP-to-MAC mapping in its ARP cache to avoid broadcasting again for every packet.\n\nThe arp-scan confirmed this process completed in just 0.117 seconds.')

doc.add_heading('Question 3: What role does ARP play in communication between your system and the OWASP VM?', 3)
doc.add_paragraph('ARP plays an essential role — without it the two machines could not communicate on the local network even if both IP addresses were known.\n\nIP packets contain source and destination IP addresses, but Ethernet frames require MAC addresses for local delivery. Every time the Kali machine sends data to the OWASP VM:\n1. The application generates data addressed to IP 192.168.5.130\n2. The IP layer wraps it in an IP packet\n3. The Data Link layer needs the MAC address of 192.168.5.130\n4. ARP provides this mapping (00:0c:29:59:23:dd)\n5. The Ethernet frame is created and transmitted\n\nThis was visible throughout all previous exercises — ARP exchanges always preceded the first TCP or ICMP communication with the OWASP VM.')

doc.add_paragraph('')

# Exercise 10
doc.add_heading('Exercise 10: Troubleshooting Network Connectivity Using TCP/IP Knowledge', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('ip route show\nsudo ip route del default\nping -c 4 8.8.8.8\nping -c 4 192.168.5.130\nsudo ip route add default via 192.168.5.2')

doc.add_heading('Question 1: What is the issue you simulated, and how did it affect network communication?', 3)
doc.add_paragraph('The network issue simulated was a missing default gateway — one of the most common and impactful network misconfigurations.\n\nBefore the simulation, the routing table showed:\ndefault via 192.168.5.2 dev eth0 — the default gateway for all external traffic\n192.168.5.0/24 dev eth0 — the local subnet route\n\nAfter removing the default gateway with sudo ip route del default:\nExternal connectivity failed completely — Pinging 8.8.8.8 returned "Network is unreachable" immediately. Without a default gateway the Kali machine had no way of forwarding packets outside the 192.168.5.0/24 subnet.\n\nLocal connectivity remained fully intact — Pinging 192.168.5.130 continued working with 0% packet loss. The OWASP VM is on the same local subnet so no gateway is needed for direct delivery.')

doc.add_heading('Question 2: How did you diagnose and resolve the issue?', 3)
doc.add_paragraph('Diagnosis: The ip route show command displayed the complete routing table confirming the default gateway was 192.168.5.2. The immediate "Network is unreachable" error from ping to 8.8.8.8 confirmed the diagnosis — the kernel checked the routing table, found no matching route, and rejected the packet before transmission. This is a Layer 3 failure.\n\nThe fact that local pings to 192.168.5.130 continued working confirmed the network interface (Layers 1 and 2) was healthy and only the routing configuration (Layer 3) was affected.\n\nResolution: The issue was resolved by restoring the default gateway:\nsudo ip route add default via 192.168.5.2\n\nThis re-added the default route telling the Kali machine to send all external traffic to the gateway at 192.168.5.2 for forwarding. External connectivity was immediately restored.')

doc.add_heading('Question 3: What tools and techniques would you recommend for real-world network troubleshooting?', 3)
doc.add_paragraph('ip route show / ip addr show — Always start here. Verify the interface is up, has the correct IP address, subnet mask, and default gateway. Most connectivity issues stem from misconfigurations at this level.\n\nping — The first test for any connectivity issue. Ping the local gateway first, then a known external address. If the gateway responds but external addresses do not, the issue is upstream.\n\ntraceroute — When ping fails to an external address, traceroute identifies exactly where packets stop, helping pinpoint whether the issue is at the local router, ISP, or further upstream.\n\ntcpdump / Wireshark — When higher-level tools do not reveal the problem, packet capture gives a ground-truth view of exactly what is happening. Missing ARP replies, unexpected RST flags, or ICMP errors can all be spotted here.\n\nnetstat / ss — Verify which ports are listening, what connections are established, and whether services are running correctly.\n\nnmap — Verify which services are reachable on a remote host and confirm firewall rules.\n\narp -a / arp-scan — Verify Layer 2 connectivity and confirm MAC address resolution is working correctly.\n\nThe recommended approach is always to work bottom-up through the TCP/IP layers — start at the Link layer (is the interface up?), then Internet layer (is routing correct?), then Transport layer (are the right ports open?), then Application layer (is the service responding?). This systematic approach quickly narrows down any network issue.')

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT303/Lab3/INT303_Lab3_Report.docx')
print("Report saved successfully!")
