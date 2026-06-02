from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT303: Networking Fundamentals', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 1: Understanding Network Layers and TCP/IP Model', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Student Info
doc.add_paragraph('Student: Damian Patrick Ozimede')
doc.add_paragraph('Course: INT303 — Networking Fundamentals')
doc.add_paragraph('Lab: Lab 1')
doc.add_paragraph('Date: June 2026')

doc.add_paragraph('')

# Exercise 1
doc.add_heading('Exercise 1: Understanding OSI and TCP/IP Models', 2)

doc.add_heading('Question 1: List the OSI model layers and describe the function of each.', 3)

doc.add_paragraph('Layer 7 – Application Layer\nThis is the layer closest to the end user. It provides network services directly to applications such as web browsers, email clients, and file transfer programs. Protocols at this layer include HTTP, HTTPS, FTP, SMTP, and DNS.')
doc.add_paragraph('Layer 6 – Presentation Layer\nThis layer is responsible for translating, encrypting, and compressing data so that it can be understood by the Application layer. It acts as a translator between the network and the application, handling data formatting such as JPEG, MP4, SSL/TLS encryption, and ASCII encoding.')
doc.add_paragraph('Layer 5 – Session Layer\nThis layer manages and controls the connections (sessions) between two communicating devices. It establishes, maintains, and terminates sessions. Examples include NetBIOS and RPC (Remote Procedure Call).')
doc.add_paragraph('Layer 4 – Transport Layer\nThis layer is responsible for end-to-end communication, error detection, flow control, and data segmentation. It ensures complete data transfer between hosts. The two main protocols here are TCP (reliable, connection-oriented) and UDP (faster, connectionless).')
doc.add_paragraph('Layer 3 – Network Layer\nThis layer handles logical addressing and routing of data packets across networks. It determines the best path for data to travel from source to destination. The key protocol here is IP (Internet Protocol). Routers operate at this layer.')
doc.add_paragraph('Layer 2 – Data Link Layer\nThis layer handles physical addressing (MAC addresses) and ensures reliable data transfer between two directly connected devices. It packages raw bits from the Physical layer into frames. Switches and network interface cards (NICs) operate at this layer. ARP also operates here.')
doc.add_paragraph('Layer 1 – Physical Layer\nThis is the lowest layer and deals with the actual physical transmission of raw binary data (0s and 1s) over a physical medium such as cables, fibre optics, or radio waves. It defines hardware specifications like voltage levels, cable types, and pin layouts.')

doc.add_heading('Question 2: Match each OSI layer with its corresponding TCP/IP layer.', 3)

table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'OSI Layer'
hdr[1].text = 'TCP/IP Layer'
rows = [
    ('Layer 7 – Application', 'Application'),
    ('Layer 6 – Presentation', 'Application'),
    ('Layer 5 – Session', 'Application'),
    ('Layer 4 – Transport', 'Transport'),
    ('Layer 3 – Network', 'Internet'),
    ('Layer 2 – Data Link', 'Link (Network Access)'),
    ('Layer 1 – Physical', 'Link (Network Access)'),
]
for i, (osi, tcpip) in enumerate(rows):
    row = table.rows[i+1].cells
    row[0].text = osi
    row[1].text = tcpip

doc.add_paragraph('The TCP/IP model condenses the 7 OSI layers into 4 layers. The top three OSI layers (Application, Presentation, Session) are all merged into a single Application layer in TCP/IP. The Transport and Network layers map directly across. The bottom two OSI layers (Data Link and Physical) are combined into the Link layer in TCP/IP.')

doc.add_paragraph('')

# Exercise 2
doc.add_heading('Exercise 2: Pinging the OWASP Broken Web Application', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('ping -c 4 192.168.5.130')

doc.add_heading('Question 1: What happens when you ping the OWASP application? Describe the process.', 3)
doc.add_paragraph('When the ping command was run against the OWASP Broken Web Application at 192.168.5.130, the machine responded successfully to all 4 packets sent. The terminal showed that 64 bytes of data were received for each packet, confirming that the OWASP VM is online and reachable on the network.')
doc.add_paragraph('The results showed the following:\nicmp_seq=1 — response time of 0.837 ms\nicmp_seq=2 — response time of 1.23 ms\nicmp_seq=3 — response time of 0.685 ms\nicmp_seq=4 — response time of 1.48 ms')
doc.add_paragraph('A total of 4 packets were transmitted and 4 were received, giving a 0% packet loss. The average round-trip time was 1.055 ms, which indicates a very fast and stable connection between the Kali machine (192.168.5.139) and the OWASP VM (192.168.5.130). The TTL (Time To Live) value was 64, which is typical for a Linux-based system, meaning the packet did not pass through any routers before reaching its destination.')

doc.add_heading('Question 2: Which OSI layer does the ping command operate at? Explain.', 3)
doc.add_paragraph('The ping command operates at Layer 3 — the Network Layer of the OSI model. Ping uses the ICMP (Internet Control Message Protocol), which is a Network layer protocol. When ping is executed, it sends ICMP Echo Request packets to the target IP address (192.168.5.130) and waits for ICMP Echo Reply packets in return. The Network layer is responsible for logical addressing and routing, and ICMP is used here specifically to test whether a host is reachable across an IP network. The fact that we used an IP address (not a MAC address or hostname port) further confirms that this operates at the Network layer.')

doc.add_paragraph('')

# Exercise 3
doc.add_heading('Exercise 3: Tracing the Path to the OWASP Application', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('traceroute 192.168.5.130')

doc.add_heading('Question 1: How many hops did it take to reach the OWASP VM?', 3)
doc.add_paragraph('It took 2 hops to reach the OWASP Broken Web Application at 192.168.5.130.\n\nHop 1 showed three asterisks (* * *), which means that the first device along the path did not respond to the traceroute probe packets. This is common behaviour when a device such as a router or gateway is configured to block or ignore ICMP/UDP probe packets for security reasons.\n\nHop 2 was the OWASP VM itself at 192.168.5.130, which responded successfully with three round-trip times of 4.441 ms, 4.135 ms, and 3.782 ms, confirming the destination was reached.')

doc.add_heading('Question 2: Describe the significance of each hop and what role traceroute plays in network troubleshooting.', 3)
doc.add_paragraph('Each hop in a traceroute result represents a device (such as a router, switch, or gateway) that the packet passes through on its way to the destination. The response times shown for each hop help identify where delays or failures are occurring along the network path.\n\nHop 1 (* * *) — This device did not respond to the traceroute probes. This could be a virtual network gateway or router between the Kali machine and the OWASP VM that is configured to silently drop probe packets. While it appears as a non-response, it did not stop the packet from reaching the destination, meaning it is still forwarding traffic correctly.\n\nHop 2 (192.168.5.130) — This is the OWASP Broken Web Application VM. It responded to all three probes with consistent times around 4 ms, confirming successful delivery.\n\nTraceroute plays a vital role in network troubleshooting because it maps the entire path a packet takes from source to destination. It helps network engineers and security professionals identify bottlenecks, unresponsive devices, routing loops, and points of failure within a network. In a security context, it can also reveal the network topology and intermediate devices between two hosts.')

doc.add_paragraph('')

# Exercise 4
doc.add_heading('Exercise 4: Viewing Active Connections to OWASP VM', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('netstat -an & curl http://192.168.5.130 > /dev/null 2>&1')

doc.add_heading('Question 1: What connections do you see? Identify the source and destination IP addresses.', 3)
doc.add_paragraph('The netstat output displayed all active internet connections and UNIX domain sockets on the Kali machine. The two most notable entries in the active internet connections section were:\n\nTCP — 127.0.0.1:9050 listening — This is the Tor proxy service running locally on the Kali machine, listening for connections on the loopback interface.\n\nUDP — 192.168.5.139:68 to 192.168.5.254:67 — This shows the Kali machine (192.168.5.139) communicating with the network gateway (192.168.5.254) via DHCP. Port 68 is the DHCP client port and port 67 is the DHCP server port, meaning the Kali machine has an active DHCP lease from the gateway.\n\nWhen curl connected to the OWASP VM at 192.168.5.130 over HTTP (port 80), the TCP connection was established and completed so quickly that netstat captured it in the brief window between the two commands. The source IP was the Kali machine at 192.168.5.139 and the destination was the OWASP VM at 192.168.5.130.')

doc.add_heading('Question 2: Explain how the Transport Layer (TCP/UDP) is involved in this communication.', 3)
doc.add_paragraph('The Transport Layer (Layer 4 of the OSI model) plays a critical role in the communication between the Kali machine and the OWASP VM. Two Transport layer protocols were observed in this exercise:\n\nTCP (Transmission Control Protocol) was used for the HTTP connection to the OWASP VM on port 80. TCP is a connection-oriented protocol, meaning it establishes a reliable connection through a three-way handshake (SYN, SYN-ACK, ACK) before any data is exchanged. It guarantees that all data packets are delivered in the correct order and retransmits any lost packets. This is why it is used for web traffic — reliability is essential when loading web pages.\n\nUDP (User Datagram Protocol) was observed in the DHCP communication between the Kali machine and the gateway. UDP is a connectionless protocol that does not guarantee delivery or order of packets. It is faster than TCP because it skips the handshake process, making it suitable for services like DHCP where speed matters more than guaranteed delivery.\n\nTogether, these two Transport layer protocols ensure that different types of network communication are handled appropriately based on whether reliability or speed is the priority.')

doc.add_paragraph('')

# Exercise 5
doc.add_heading('Exercise 5: TCP vs. UDP', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('nmap -sT 192.168.5.130\nsudo nmap -sU --top-ports 100 192.168.5.130')

doc.add_heading('Question 1: What are the key differences between TCP and UDP in terms of reliability and speed?', 3)
doc.add_paragraph('TCP (Transmission Control Protocol) and UDP (User Datagram Protocol) are both Transport Layer protocols but they work in very different ways.\n\nTCP is a connection-oriented protocol, meaning it establishes a formal connection between two devices before any data is sent. This is done through a process called the three-way handshake, where the sender sends a SYN packet, the receiver responds with a SYN-ACK, and the sender confirms with an ACK. TCP also guarantees that all packets are delivered in the correct order and retransmits any that are lost. This makes TCP highly reliable but slightly slower due to the overhead of all this checking and confirming. TCP is used for web browsing (HTTP/HTTPS), email, and file transfers where data accuracy is critical.\n\nUDP on the other hand is a connectionless protocol. It simply sends packets to the destination without establishing a connection first, without confirming delivery, and without caring about the order of arrival. This makes UDP significantly faster than TCP but unreliable in terms of guaranteed delivery. UDP is used for services where speed is more important than perfection, such as online gaming, video streaming, voice calls, DNS lookups, and DHCP.')

doc.add_heading('Question 2: Based on the scan results, list which services are using TCP and which are using UDP.', 3)

table2 = doc.add_table(rows=10, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Port'
hdr2[1].text = 'Service'
hdr2[2].text = 'Purpose'
tcp_rows = [
    ('22/tcp', 'SSH', 'Secure remote login'),
    ('80/tcp', 'HTTP', 'Web traffic'),
    ('139/tcp', 'NetBIOS-SSN', 'Windows file sharing sessions'),
    ('143/tcp', 'IMAP', 'Email receiving'),
    ('443/tcp', 'HTTPS', 'Secure web traffic'),
    ('445/tcp', 'Microsoft-DS', 'SMB file sharing'),
    ('5001/tcp', 'Commplex-link', 'Application communication'),
    ('8080/tcp', 'HTTP-Proxy', 'Alternative web port'),
    ('8081/tcp', 'Blackice-Icecap', 'Alternative web port'),
]
for i, (port, service, purpose) in enumerate(tcp_rows):
    row = table2.rows[i+1].cells
    row[0].text = port
    row[1].text = service
    row[2].text = purpose

doc.add_paragraph('')
doc.add_paragraph('UDP Services:')
table3 = doc.add_table(rows=2, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Port'
hdr3[1].text = 'Service'
hdr3[2].text = 'Purpose'
row3 = table3.rows[1].cells
row3[0].text = '137/udp'
row3[1].text = 'NetBIOS-NS'
row3[2].text = 'NetBIOS name service for network device discovery'

doc.add_paragraph('\nThe scan results clearly show that the OWASP VM runs predominantly TCP-based services, which makes sense given that it is a web application testing platform. The only confirmed open UDP port was 137 (NetBIOS-NS), which is used for resolving NetBIOS names to IP addresses on the local network. The remaining 45 UDP ports returned an open/filtered status, meaning nmap could not definitively determine whether they were open or blocked by a firewall.')

doc.add_paragraph('')

# Exercise 6
doc.add_heading('Exercise 6: Discovering MAC Addresses with ARP', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('arp -a')

doc.add_heading('Question 1: What is the MAC address associated with the OWASP VM\'s IP?', 3)
doc.add_paragraph('The ARP table returned three entries for devices on the local network. The entry for the OWASP Broken Web Application at 192.168.5.130 showed the following:\n\nIP Address: 192.168.5.130\nMAC Address: 00:0C:29:59:23:DD\nInterface: eth0\n\nThe other two entries in the ARP table were:\n192.168.5.2 — MAC address 00:50:56:E8:EA:83 (a VMware virtual network device)\n192.168.5.254 — MAC address 00:50:56:EB:AC:66 (the network gateway/router)\n\nIt is worth noting that the MAC address 00:0C:29:59:23:DD of the OWASP VM matches exactly what was shown in the nmap scan results from Exercise 5, confirming it is a VMware virtual machine.')

doc.add_heading('Question 2: Explain the significance of ARP in the Data Link Layer and how it contributes to successful communication.', 3)
doc.add_paragraph('ARP (Address Resolution Protocol) operates at Layer 2 — the Data Link Layer of the OSI model and serves as a critical bridge between the Network Layer (Layer 3) and the Data Link Layer (Layer 2).\n\nWhen the Kali machine wants to communicate with the OWASP VM at 192.168.5.130, it knows the IP address but needs the MAC address to actually deliver the data frame on the local network. This is because while IP addresses are used for routing traffic across networks, MAC addresses are what identify devices physically on the same local network segment.\n\nARP solves this by broadcasting a request across the entire local network asking "Who has IP address 192.168.5.130?" The OWASP VM responds with its MAC address 00:0C:29:59:23:DD. The Kali machine then stores this mapping in its ARP cache so it does not need to ask again for every single packet.\n\nWithout ARP, devices on a local network would have no way of translating logical IP addresses into physical MAC addresses, making communication between devices on the same network impossible. ARP is therefore a fundamental and essential protocol for any local network communication to succeed.')

doc.add_paragraph('')

# Exercise 7
doc.add_heading('Exercise 7: Capturing Network Traffic with tshark', 2)
doc.add_heading('Command Used:', 3)
doc.add_paragraph('sudo tshark -i eth0 host 192.168.5.130')

doc.add_heading('Question 1: Analyze the captured traffic. What protocols are in use?', 3)
doc.add_paragraph('The tshark capture recorded a total of 40 packets exchanged between the Kali machine (192.168.5.139) and the OWASP Broken Web Application (192.168.5.130). Two protocols were clearly visible in the captured traffic:\n\nTCP (Transmission Control Protocol) — The majority of the captured packets were TCP packets. These handled the reliable transmission of data between the two machines, managing the connection establishment, data transfer, and connection termination.\n\nHTTP (HyperText Transfer Protocol) — HTTP was visible in packets 4 and 36. Packet 4 showed a GET / HTTP/1.1 request sent from the Kali machine to the OWASP VM, requesting the homepage. Packet 36 showed the OWASP VM responding with HTTP/1.1 200 OK (text/html), meaning the request was successful and the webpage content was returned.')

doc.add_heading('Question 2: Can you identify the handshake process or other significant events in the captured packets?', 3)
doc.add_paragraph('Yes! The capture clearly shows several significant events in the communication between the two machines:\n\n1. The TCP Three-Way Handshake (Packets 1-3)\n\nPacket 1 — SYN: The Kali machine (192.168.5.139) sent a SYN packet to the OWASP VM on port 80, saying "I want to connect"\nPacket 2 — SYN, ACK: The OWASP VM (192.168.5.130) responded with a SYN-ACK packet, saying "I received your request and I am ready"\nPacket 3 — ACK: The Kali machine replied with an ACK packet, saying "Confirmed, let\'s communicate"\n\nThe three-way handshake was completed in under 1 millisecond, establishing a reliable connection before any data was sent.\n\n2. The HTTP Request and Response (Packets 4-36)\n\nPacket 4 — The Kali machine sent an HTTP GET request asking for the homepage of the OWASP VM\nPacket 6 — The OWASP VM began responding with HTTP/1.1 200 OK, meaning the page was found and data transfer began\nPackets 7-35 — A series of TCP ACK packets were exchanged as the OWASP VM broke the webpage content into multiple segments and sent them across, with the Kali machine acknowledging each one\nPacket 36 — The final HTTP response was delivered containing the complete HTML content of the OWASP homepage\n\n3. The TCP Connection Termination (Packets 38-40)\n\nPacket 38 — FIN, ACK: The Kali machine sent a FIN packet saying "I am done, closing the connection"\nPacket 39 — FIN, ACK: The OWASP VM acknowledged and also sent its own FIN packet\nPacket 40 — ACK: The Kali machine sent a final ACK confirming the connection was fully closed\n\nThis entire conversation happened in just 0.013 seconds, demonstrating how efficiently TCP manages network communication.')

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT303/Lab1/INT303_Lab1_Report.docx')
print("Report saved successfully!")
