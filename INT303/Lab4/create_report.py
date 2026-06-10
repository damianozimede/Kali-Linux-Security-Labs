from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT303: Networking Fundamentals', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 4: Simulating Network Routing and VLAN Configuration in Linux', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Student: Damian Patrick Ozimede')
doc.add_paragraph('Course: INT303 — Networking Fundamentals')
doc.add_paragraph('Lab: Lab 4')
doc.add_paragraph('Date: June 2026')

doc.add_paragraph('')

# Exercise 1
doc.add_heading('Exercise 1: Setting Up Static Routing in Linux', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('ip route show\nsudo ip route add 192.168.4.0/24 via 192.168.5.2 dev eth0\nip route show\nping -c 4 192.168.5.130')

doc.add_heading('Question 1: How does static routing work on a Linux system?', 3)
doc.add_paragraph('Static routing works by manually adding fixed entries to the kernel routing table that tell the system exactly where to send packets destined for specific networks.\n\nThe routing table is checked every time the system needs to send a packet. The kernel looks for the most specific matching route. Before adding the static route the routing table had two entries:\ndefault via 192.168.5.2 — for all traffic with no specific route\n192.168.5.0/24 dev eth0 — for local network traffic\n\nAfter running sudo ip route add 192.168.4.0/24 via 192.168.5.2 dev eth0, a third entry appeared:\n192.168.4.0/24 via 192.168.5.2 dev eth0 — directing traffic for the 192.168.4.0/24 network through the gateway at 192.168.5.2\n\nThe ping to the OWASP VM continued working with 0% packet loss and average response time of 1.050 ms, confirming existing routes remained intact after adding the new static route.')

doc.add_heading('Question 2: What challenges could arise when setting up routes manually?', 3)
doc.add_paragraph('Routes are not persistent — Static routes added with ip route are lost on reboot. They must be saved to configuration files for permanence.\n\nHuman error — Typing wrong IP addresses, subnet masks, or gateway addresses causes traffic to be sent to the wrong destination or dropped entirely.\n\nScalability — In large networks with hundreds of subnets, manually maintaining static routes becomes extremely complex. Dynamic routing protocols like OSPF and BGP were developed to solve this automatically.\n\nNo automatic failover — If the gateway becomes unreachable, traffic continues being sent to it and dropped. Dynamic protocols can detect failures and automatically reroute traffic.\n\nRoute conflicts — Overlapping or conflicting routes cause unpredictable routing behaviour requiring careful management of route metrics.')

doc.add_paragraph('')

# Exercise 2
doc.add_heading('Exercise 2: VLAN Configuration Using Network Namespaces', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('sudo ip netns add vlan1\nsudo ip netns add vlan2\nsudo ip link add veth1 type veth peer name veth2\nsudo ip link set veth1 netns vlan1\nsudo ip link set veth2 netns vlan2\nsudo ip netns exec vlan1 ip addr add 192.168.10.1/24 dev veth1\nsudo ip netns exec vlan2 ip addr add 192.168.10.2/24 dev veth2\nsudo ip netns exec vlan1 ping -c 4 192.168.10.2')

doc.add_heading('Question 1: How do network namespaces simulate VLANs in Linux?', 3)
doc.add_paragraph('Network namespaces create completely isolated network environments within the same Linux system. Each namespace has its own network interfaces, routing table, and firewall rules — completely separate from the host and other namespaces.\n\nIn this exercise two namespaces (vlan1 and vlan2) were created, each assigned a virtual ethernet interface (veth1 and veth2) connected like a virtual cable. vlan1 was assigned IP 192.168.10.1 and vlan2 was assigned 192.168.10.2. The ping from vlan1 to vlan2 succeeded with 0% packet loss and sub-millisecond response times (avg 0.110 ms), confirming isolated but connected virtual network segments — simulating the behaviour of VLANs.')

doc.add_heading('Question 2: What are the benefits of using namespaces in network isolation?', 3)
doc.add_paragraph('Security isolation — Processes in one namespace cannot see or interfere with network traffic in another, preventing unauthorised access between segments.\n\nTesting and development — Network namespaces allow complex network topologies to be simulated on a single machine without physical hardware.\n\nContainer networking — Tools like Docker use network namespaces internally to give each container its own isolated network stack.\n\nTraffic control — Different firewall rules and routing policies can be applied per namespace, giving fine-grained control over traffic flow between segments.')

doc.add_paragraph('')

# Exercise 3
doc.add_heading('Exercise 3: IP Address Assignment and Subnetting in Linux', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('sudo ip netns exec vlan1 ip addr show\nsudo ip netns exec vlan2 ip addr show\nsudo ip netns exec vlan2 ip addr add 192.168.20.1/24 dev veth2\nsudo ip netns exec vlan1 ip route add 192.168.20.0/24 via 192.168.10.2\nsudo ip netns exec vlan1 ip route show\nsudo ip netns exec vlan1 ping -c 4 192.168.20.1')

doc.add_heading('Question 1: How does subnetting work in Linux environments?', 3)
doc.add_paragraph('Subnetting divides a large network into smaller segments. In this exercise two subnets were created — 192.168.10.0/24 and 192.168.20.0/24. vlan2 was given addresses on both subnets, acting as a router between them.\n\nThe routing table in vlan1 confirmed this:\n192.168.10.0/24 dev veth1 — directly connected subnet\n192.168.20.0/24 via 192.168.10.2 — reach the second subnet through vlan2\n\nThe ping from vlan1 to 192.168.20.1 succeeded with 0% packet loss and avg 0.122 ms, confirming vlan1 successfully reached the second subnet through vlan2.')

doc.add_heading('Question 2: What challenges arise when configuring subnets manually?', 3)
doc.add_paragraph('Overlapping subnets — Assigning the same or overlapping IP ranges to different interfaces causes routing conflicts and unpredictable traffic behaviour.\n\nWrong subnet masks — An incorrect subnet mask causes the system to misidentify which addresses are local and which need routing, breaking connectivity silently.\n\nMissing routes — Forgetting to add a static route between subnets means packets have no path and are dropped, causing immediate connectivity failures.\n\nNo persistence — All manually configured addresses and routes are lost on reboot unless saved to configuration files.')

doc.add_paragraph('')

# Exercise 4
doc.add_heading('Exercise 4: Testing Connectivity Using Ping and Traceroute', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('ping -c 4 192.168.5.130\ntraceroute 192.168.5.130\nsudo ip netns exec vlan1 ping -c 4 192.168.10.2\nsudo ip netns exec vlan1 traceroute 192.168.10.2')

doc.add_heading('Question 1: What can traceroute reveal about network issues?', 3)
doc.add_paragraph('The traceroute results showed:\nHost to OWASP VM — reached in 1 hop at 192.168.5.130 with response times of 2.856-3.104 ms, confirming direct local delivery\nvlan1 to vlan2 — reached in 1 hop at 192.168.10.2 with response times of 0.550-0.719 ms, confirming the virtual ethernet link is working\n\nTraceroute can reveal:\nWhere connectivity breaks — * * * at a hop pinpoints exactly where failure occurs\nUnexpected routing paths — reveals if traffic is taking a longer detour than expected\nHigh latency hops — a sudden spike in response time indicates congestion or a slow link\nRouting loops — if the same IP appears repeatedly, a routing loop exists')

doc.add_heading('Question 2: How does packet routing affect network performance?', 3)
doc.add_paragraph('Number of hops — Every additional hop adds latency. Both the OWASP VM and vlan2 were reachable in 1 hop giving low response times. In real-world internet connections packets often traverse 10-20 hops, each adding milliseconds of delay.\n\nRouting decisions — Every router must look up the destination in its routing table for each packet. Large routing tables or inefficient lookups slow down forwarding.\n\nCongestion — If a routing path passes through a congested link, all packets using that route experience increased latency and potential packet loss.\n\nDirect vs routed delivery — Direct local delivery is significantly faster than routed delivery. The vlan1 to vlan2 ping averaged just 0.269 ms compared to 0.937 ms for the OWASP VM due to the virtual ethernet having no physical overhead.')

doc.add_paragraph('')

# Exercise 5
doc.add_heading('Exercise 5: Configuring iptables for Routing and Firewall Rules', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('echo 1 | sudo tee /proc/sys/net/ipv4/ip_forward\nsudo iptables -L -n\nsudo iptables -A FORWARD -s 192.168.10.0/24 -d 192.168.20.0/24 -j ACCEPT\nsudo iptables -A FORWARD -s 192.168.10.0/24 -d 192.168.5.130 -j DROP\nsudo iptables -L -n\nsudo ip netns exec vlan1 ping -c 4 192.168.5.130')

doc.add_heading('Question 1: How can iptables be used to simulate routing and firewall functionality?', 3)
doc.add_paragraph('iptables controls how packets are handled as they flow through the Linux kernel through three main chains:\nINPUT — controls packets destined for the local machine\nFORWARD — controls packets being routed through the machine between networks\nOUTPUT — controls packets generated by the local machine\n\nIP forwarding was enabled first by writing 1 to /proc/sys/net/ipv4/ip_forward, turning the Kali machine into a router.\n\nTwo FORWARD rules were added:\nACCEPT — allowing traffic from 192.168.10.0/24 to 192.168.20.0/24, permitting inter-VLAN communication\nDROP — blocking all traffic from 192.168.10.0/24 to 192.168.5.130, isolating vlan1 from the OWASP VM\n\nThe ping from vlan1 to 192.168.5.130 returned "Network is unreachable", confirming the DROP rule worked correctly.')

doc.add_heading('Question 2: What are common mistakes when configuring iptables rules?', 3)
doc.add_paragraph('Wrong rule order — iptables processes rules top to bottom and stops at the first match. Adding an ACCEPT after a DROP for the same traffic means the DROP always wins.\n\nForgetting IP forwarding — Adding FORWARD rules without enabling IP forwarding means packets are never forwarded regardless of the rules.\n\nLocking yourself out — Accidentally adding an INPUT DROP rule for SSH can permanently lock an administrator out of a remote server.\n\nRules not persistent — iptables rules are lost on reboot unless saved using iptables-save or iptables-persistent.\n\nOverly broad rules — Rules that are too general can accidentally block legitimate traffic beyond the intended scope.')

doc.add_paragraph('')

# Exercise 6
doc.add_heading('Exercise 6: Monitoring Traffic Using tcpdump', 2)
doc.add_heading('Commands Used:', 3)
doc.add_paragraph('sudo ip netns exec vlan1 tcpdump -i veth1\nsudo tcpdump -i eth0 src 192.168.5.130')

doc.add_heading('Question 1: How can tcpdump be used to diagnose network issues?', 3)
doc.add_paragraph('Verifying connectivity — The veth1 capture confirmed vlan1 and vlan2 were communicating correctly, showing clean ICMP request/reply pairs for all 4 ping packets.\n\nConfirming ARP resolution — After the ping, ARP exchanges confirmed vlan1 and vlan2 refreshed each other\'s MAC addresses (6e:0d:f8:ac:cd:82 and d6:ee:47:34:7f:6f), confirming Layer 2 communication inside the namespace.\n\nAnalysing specific traffic sources — The filter src 192.168.5.130 showed only packets from the OWASP VM, making it easy to analyse server behaviour without noise from other hosts.\n\nIdentifying protocols — The eth0 capture showed ARP, HTTP, TCP, and UDP NetBIOS traffic originating from the OWASP VM, giving a clear picture of active services.')

doc.add_heading('Question 2: What types of traffic do you observe during the simulation?', 3)
doc.add_paragraph('veth1 namespace capture:\nICMP Echo Requests and Replies — 4 request/reply pairs between 192.168.10.1 and 192.168.10.2 with identifier 33210 and sequence numbers 1-4\nARP Requests and Replies — Both namespaces refreshed each other\'s MAC addresses after the ping completed\n\neth0 capture (src 192.168.5.130):\nARP Reply — OWASP VM confirmed its MAC address as 00:0c:29:59:23:dd\nTCP SYN-ACK — OWASP VM accepted the HTTP connection on port 56076\nHTTP responses — Multiple TCP segments delivering HTTP/1.1 200 OK in chunks of 1448, 2896, and 4344 bytes totalling 28,515 bytes\nTCP FIN-ACK — Clean connection termination after data delivery\nUDP NetBIOS broadcasts — OWASP VM periodically sent NetBIOS datagrams to 192.168.5.255 announcing its presence')

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT303/Lab4/INT303_Lab4_Report.docx')
print("Report saved successfully!")
