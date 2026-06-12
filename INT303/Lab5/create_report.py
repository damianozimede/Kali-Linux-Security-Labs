from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('INT303: Networking Fundamentals', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 5: IP Address Analysis and Network Report', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Student: Damian Patrick Ozimede')
doc.add_paragraph('Course: INT303 — Networking Fundamentals')
doc.add_paragraph('Lab: Lab 5')
doc.add_paragraph('Date: June 2026')

doc.add_paragraph('')

# Introduction
doc.add_heading('Introduction', 2)
doc.add_paragraph('This lab analyses 10 real-world websites to understand IP addressing, classification, subnetting, and network scalability. By pinging each website, recording their IP addresses, classifying them by class (A, B, or C), calculating device capacity, and determining subnet masks, this report demonstrates practical application of IP networking concepts. Advanced tools such as traceroute and whois were also used to investigate routing paths and verify the ownership of IP addresses through official registry data.')

doc.add_paragraph('')

# Website List and IPs
doc.add_heading('Website List and IP Addresses', 2)

table1 = doc.add_table(rows=11, cols=3)
table1.style = 'Table Grid'
hdr1 = table1.rows[0].cells
hdr1[0].text = 'Website'
hdr1[1].text = 'IP Address'
hdr1[2].text = 'Ping Response'
sites = [
    ('google.com', '142.251.216.110', 'No reply (ICMP blocked)'),
    ('apple.com', '17.253.144.10', '180 ms'),
    ('microsoft.com', '150.171.109.2', 'No reply (ICMP blocked)'),
    ('amazon.com', '98.87.170.74', 'No reply (ICMP blocked)'),
    ('facebook.com', '57.144.38.1', '58.8 ms'),
    ('twitter.com', '162.159.140.229', '176 ms'),
    ('wikipedia.org', '185.15.58.224', '207 ms'),
    ('github.com', '140.82.121.3', '160 ms'),
    ('netflix.com', '54.155.178.5', 'No reply (ICMP blocked)'),
    ('bbc.co.uk', '151.101.128.81', '205 ms'),
]
for i, (site, ip, ping) in enumerate(sites):
    row = table1.rows[i+1].cells
    row[0].text = site
    row[1].text = ip
    row[2].text = ping

doc.add_paragraph('')
doc.add_heading('Bonus: Multiple IP Addresses (CDN/Load Balancing)', 3)
doc.add_paragraph('The dig command was used to check for multiple IP addresses per domain:\n\ngoogle.com — 1 IP only\namazon.com — 3 IPs: 98.87.170.71, 98.87.170.74, 98.82.161.185\nfacebook.com — 1 IP only\nnetflix.com — 3 IPs: 54.155.246.232, 18.200.8.190, 54.73.148.110\nmicrosoft.com — 1 IP only\n\nAmazon and Netflix clearly use load balancing — multiple servers handle traffic simultaneously across different geographic regions.')

doc.add_paragraph('')

# IP Classification
doc.add_heading('IP Address Classification', 2)
doc.add_paragraph('IP addresses are classified based on their first octet:\nClass A: 1-126 (default mask /8)\nClass B: 128-191 (default mask /16)\nClass C: 192-223 (default mask /24)\n\nOnly the first number of the IP address is needed to determine the class.')

table2 = doc.add_table(rows=11, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Website'
hdr2[1].text = 'IP Address'
hdr2[2].text = 'First Octet'
hdr2[3].text = 'Class'
class_data = [
    ('google.com', '142.251.216.110', '142', 'B'),
    ('apple.com', '17.253.144.10', '17', 'A'),
    ('microsoft.com', '150.171.109.2', '150', 'B'),
    ('amazon.com', '98.87.170.74', '98', 'A'),
    ('facebook.com', '57.144.38.1', '57', 'A'),
    ('twitter.com', '162.159.140.229', '162', 'B'),
    ('wikipedia.org', '185.15.58.224', '185', 'B'),
    ('github.com', '140.82.121.3', '140', 'B'),
    ('netflix.com', '54.155.178.5', '54', 'A'),
    ('bbc.co.uk', '151.101.128.81', '151', 'B'),
]
for i, (site, ip, octet, cls) in enumerate(class_data):
    row = table2.rows[i+1].cells
    row[0].text = site
    row[1].text = ip
    row[2].text = octet
    row[3].text = cls

doc.add_paragraph('')
doc.add_heading('Bonus: Reserved IP Ranges', 3)
doc.add_paragraph('All IP addresses fall into one of two categories:\n\nPublic IPs — Unique addresses used on the internet. All 10 websites use public IPs.\n\nPrivate IPs — Reserved ranges used only inside private networks, never directly on the internet:\n10.0.0.0 - 10.255.255.255 (10.0.0.0/8) — Class A private range\n172.16.0.0 - 172.31.255.255 (172.16.0.0/12) — Class B private range\n192.168.0.0 - 192.168.255.255 (192.168.0.0/16) — Class C private range\n\nNone of the 10 IP addresses fall into these reserved private ranges. All are public IPs assigned to large organisations.\n\nSignificance: Private IP ranges allow organisations to use the same internal addresses (e.g. 192.168.x.x) without conflicts, since these are never routed on the public internet. Routers use NAT (Network Address Translation) to translate private IPs to a single public IP when accessing the internet, conserving the limited pool of public IPv4 addresses.')

doc.add_paragraph('')

# Device Capacity
doc.add_heading('Number of Devices Supported', 2)
doc.add_paragraph('Device capacity is calculated based on the number of bits remaining for host addresses after the network portion:\nClass A: 32-8 = 24 host bits = 2^24 = 16,777,216 addresses\nClass B: 32-16 = 16 host bits = 2^16 = 65,536 addresses\nClass C: 32-24 = 8 host bits = 2^8 = 256 addresses\n\nUsable hosts = Total addresses - 2 (network address and broadcast address are reserved and cannot be assigned to devices).')

table3 = doc.add_table(rows=11, cols=4)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Website'
hdr3[1].text = 'Class'
hdr3[2].text = 'Total Addresses'
hdr3[3].text = 'Usable Hosts'
capacity_data = [
    ('google.com', 'B', '65,536', '65,534'),
    ('apple.com', 'A', '16,777,216', '16,777,214'),
    ('microsoft.com', 'B', '65,536', '65,534'),
    ('amazon.com', 'A', '16,777,216', '16,777,214'),
    ('facebook.com', 'A', '16,777,216', '16,777,214'),
    ('twitter.com', 'B', '65,536', '65,534'),
    ('wikipedia.org', 'B', '65,536', '65,534'),
    ('github.com', 'B', '65,536', '65,534'),
    ('netflix.com', 'A', '16,777,216', '16,777,214'),
    ('bbc.co.uk', 'B', '65,536', '65,534'),
]
for i, (site, cls, total, usable) in enumerate(capacity_data):
    row = table3.rows[i+1].cells
    row[0].text = site
    row[1].text = cls
    row[2].text = total
    row[3].text = usable

doc.add_paragraph('')
doc.add_heading('Bonus: Subnetting a Class B Network with /26', 3)
doc.add_paragraph('Taking google.com (142.251.216.110) as an example and subnetting with /26:\n/26 subnet mask = 255.255.255.192\nClass B default = /16, so /26 borrows 26-16 = 10 extra bits for subnets\nNumber of subnets = 2^10 = 1,024 subnets\nHosts per subnet = 2^(32-26) - 2 = 2^6 - 2 = 62 usable hosts\n\nThis shows how subnetting a large Class B network into /26 subnets creates 1,024 smaller segments each supporting 62 devices — useful for dividing large organisations into smaller manageable network segments such as departments or offices.\n\nThis calculation can also be applied to Class A. If a Class A network (/8) is subnetted to /26, it would create 2^(26-8) = 2^18 = 262,144 subnets, each still supporting 62 hosts. The hosts-per-subnet figure depends only on the final CIDR value, while the number of subnets depends on how many bits were borrowed from the original class.')

doc.add_paragraph('')

# Subnet Masks
doc.add_heading('Subnet Masks', 2)

table4 = doc.add_table(rows=11, cols=4)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Website'
hdr4[1].text = 'Class'
hdr4[2].text = 'Default Subnet Mask'
hdr4[3].text = 'CIDR'
mask_data = [
    ('google.com', 'B', '255.255.0.0', '/16'),
    ('apple.com', 'A', '255.0.0.0', '/8'),
    ('microsoft.com', 'B', '255.255.0.0', '/16'),
    ('amazon.com', 'A', '255.0.0.0', '/8'),
    ('facebook.com', 'A', '255.0.0.0', '/8'),
    ('twitter.com', 'B', '255.255.0.0', '/16'),
    ('wikipedia.org', 'B', '255.255.0.0', '/16'),
    ('github.com', 'B', '255.255.0.0', '/16'),
    ('netflix.com', 'A', '255.0.0.0', '/8'),
    ('bbc.co.uk', 'B', '255.255.0.0', '/16'),
]
for i, (site, cls, mask, cidr) in enumerate(mask_data):
    row = table4.rows[i+1].cells
    row[0].text = site
    row[1].text = cls
    row[2].text = mask
    row[3].text = cidr

doc.add_paragraph('')
doc.add_heading('Bonus: Custom Subnetting Scenarios (/25, /26, /27)', 3)
doc.add_paragraph('Using a Class B network (e.g. google.com 142.251.216.0) with custom subnet masks:')

table5 = doc.add_table(rows=4, cols=4)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
hdr5[0].text = 'Subnet Mask'
hdr5[1].text = 'CIDR'
hdr5[2].text = 'Hosts per Subnet'
hdr5[3].text = 'Number of Subnets'
custom_data = [
    ('255.255.255.128', '/25', '126', '512'),
    ('255.255.255.192', '/26', '62', '1,024'),
    ('255.255.255.224', '/27', '30', '2,048'),
]
for i, (mask, cidr, hosts, subnets) in enumerate(custom_data):
    row = table5.rows[i+1].cells
    row[0].text = mask
    row[1].text = cidr
    row[2].text = hosts
    row[3].text = subnets

doc.add_paragraph('\nHow the subnet mask values were derived:\nSubnet masks are built in binary, with each bit position in the final octet worth: 128, 64, 32, 16, 8, 4, 2, 1\n/25 turns on 1 bit beyond /24 = 128\n/26 turns on 2 bits = 128+64 = 192\n/27 turns on 3 bits = 128+64+32 = 224\n\nHosts per subnet formula: 2^(32-CIDR) - 2\nNumber of subnets formula (from Class B /16 base): 2^(CIDR-16)')

doc.add_heading('Custom Subnetting Scenarios: Discussion', 3)
doc.add_paragraph('Scenarios where custom subnetting is required: A large organisation with a single Class B network (65,534 hosts) rarely puts every device on one flat network. Instead they divide it by department or location — for example separate subnets for Finance, HR, IT, and Sales — each needing only 50-100 devices.\n\nHow altering the subnet mask affects device capacity and organization: Moving from /25 to /27, device capacity decreases (126 to 30 hosts) while the number of subnets increases (512 to 2,048). This is a direct tradeoff — borrowing more bits for subnets leaves fewer bits for hosts.\n\nBenefits of smaller subnets: Better security (a breach in one subnet does not directly expose others), reduced broadcast traffic (broadcasts stay within each smaller subnet improving performance), and easier troubleshooting (issues isolated to specific subnets).\n\nDrawback: More subnets means more complex routing tables and IP address management.\n\nThis mirrors the approach used in Lab 4 where separate subnets (192.168.10.0/24 and 192.168.20.0/24) were created with iptables controlling traffic between them — simulating real organisational network segmentation.')

doc.add_paragraph('')

# Advanced Tools
doc.add_heading('Advanced Tools and Analysis', 2)
doc.add_heading('Traceroute to apple.com', 3)
doc.add_paragraph('Command used: traceroute apple.com\n\nResults:\nHop 1 — 192.168.5.2 (5.1-5.3 ms) — Local gateway, a private IP address\nHops 2-30 — All * * * — No response from ISP/backbone routers\n\nAnalysis: Hop 1 is the local network gateway, a private IP (192.168.0.0/16 range), confirming the Kali machine exits the local network through this gateway with excellent response time (~5ms).\n\nHops 2-30 returned * * *, meaning intermediate routers did not respond to traceroute probes. This is normal behaviour — ISPs and backbone providers commonly drop ICMP/UDP traceroute probes for security and privacy reasons, preventing network topology mapping.\n\nDespite this, the ping to apple.com (17.253.144.10) succeeded with a response time of 180 ms, confirming traffic does reach Apple\'s servers successfully — the intermediate hops are simply invisible to traceroute, not actually broken.\n\nSignificance: This demonstrates that traceroute only shows hops where devices choose to respond. The actual path likely passes through multiple ISP routers and internet exchange points, but these are invisible due to security configurations.\n\nPrivate vs infrastructure: Hop 1 (192.168.5.2) is a private network device (local router). Apple\'s destination IP (17.253.144.10) belongs to Apple\'s own infrastructure rather than a third-party cloud provider.')

doc.add_heading('GeoIP Analysis (Verified via WHOIS)', 3)
doc.add_paragraph('The whois command was used to query the ARIN (American Registry for Internet Numbers) database for official documented proof of IP address ownership:\n\napple.com (17.253.144.10):\nNetRange: 17.0.0.0 - 17.255.255.255\nCIDR: 17.0.0.0/8\nOrganization: Apple Inc. (APPLEC-1-Z)\nThis confirms Apple owns the entire 17.0.0.0/8 address block — an extremely rare allocation size reserved only for very large organisations operating their own global private network.\n\ngithub.com (140.82.121.3):\nNetRange: 140.82.112.0 - 140.82.127.255\nCIDR: 140.82.112.0/20\nOrganization: GitHub, Inc. (GITHU)\nThis confirms GitHub directly owns this address block, consistent with the hostname "lb-140-82-121-3-fra" showing a load balancer in Frankfurt, Germany.\n\nbbc.co.uk (151.101.128.81):\nNetRange: 151.101.0.0 - 151.101.255.255\nCIDR: 151.101.0.0/16\nOrganization: Fastly, Inc. (SKYCA-3)\nThis confirms BBC\'s website is served through Fastly\'s CDN infrastructure rather than BBC\'s own servers.')

doc.add_heading('Data Centers, ISPs, or Specific Regions', 3)
doc.add_paragraph('Belonging to large tech company data centers (owned infrastructure):\napple.com (17.253.144.10) — Apple owns its own IP blocks and operates its own global data centers\nfacebook.com (57.144.38.1) — Hostname "edge-star-mini-shv-01-los4" reveals Meta\'s edge server in Los Angeles\ngithub.com (140.82.121.3) — Hostname "lb-140-82-121-3-fra" shows GitHub\'s load balancer in Frankfurt, Germany\nwikipedia.org (185.15.58.224) — Hostname "text-lb.drmrs.wikimedia.org" shows Wikimedia\'s server in Marseille, France (DRMRS = Marseille airport code)\n\nBelonging to cloud provider infrastructure (rented data centers):\namazon.com (98.87.170.74) and netflix.com (54.155.178.5, 18.200.8.190, 54.73.148.110) — These IPs fall within Amazon AWS address ranges, indicating both run on AWS cloud infrastructure across multiple regions\n\nBelonging to CDN providers:\nbbc.co.uk (151.101.128.81) — Belongs to Fastly, a CDN provider BBC uses to distribute content closer to users\ntwitter.com (162.159.140.229) — Belongs to Cloudflare, used for CDN and DDoS protection\n\nBelonging to standard global network backbones:\ngoogle.com (142.251.216.110) and microsoft.com (150.171.109.2) — These belong to Google\'s and Microsoft\'s own global network backbones, distributing traffic to regional data centers\n\nIn summary, all 10 websites are hosted on large-scale infrastructure — either owned directly by the company, rented from major cloud providers, or distributed through CDNs. This reflects how major websites rely on globally distributed infrastructure for fast and reliable access worldwide.')

doc.add_heading('Latency vs Geographic Distance', 3)

table6 = doc.add_table(rows=7, cols=3)
table6.style = 'Table Grid'
hdr6 = table6.rows[0].cells
hdr6[0].text = 'Website'
hdr6[1].text = 'Response Time'
hdr6[2].text = 'Likely Location'
latency_data = [
    ('facebook.com', '58.8 ms', 'Los Angeles, USA (local edge caching likely)'),
    ('github.com', '160 ms', 'Frankfurt, Germany'),
    ('twitter.com', '176 ms', 'Cloudflare CDN (European edge)'),
    ('apple.com', '180 ms', 'USA'),
    ('bbc.co.uk', '205 ms', 'Fastly CDN (UK/Europe)'),
    ('wikipedia.org', '207 ms', 'Marseille, France'),
]
for i, (site, time, location) in enumerate(latency_data):
    row = table6.rows[i+1].cells
    row[0].text = site
    row[1].text = time
    row[2].text = location

doc.add_paragraph('\nObservation: Facebook\'s significantly lower latency (58.8 ms) despite being hosted in Los Angeles suggests Meta operates a local caching/edge server within Nigeria or West Africa, while other services rely on more distant European or US servers, resulting in higher latencies of 160-207 ms.')

doc.add_paragraph('')

# Conclusion
doc.add_heading('Conclusion', 2)
doc.add_paragraph('This lab provided practical experience analysing real-world IP addresses across 10 major websites. By classifying each IP into Class A or B, calculating device capacities ranging from 65,534 to over 16 million hosts, and exploring subnet mask variations, this report demonstrates how IP address management directly impacts network scalability and organisation.\n\nThe traceroute exercise reinforced findings from earlier labs — ISPs commonly block traceroute probes for security reasons, though traffic still flows successfully. The whois lookups provided documented proof of IP ownership, confirming that major websites rely on a mix of owned infrastructure (Apple, GitHub, Wikimedia), cloud providers (AWS for Amazon and Netflix), and CDNs (Fastly for BBC, Cloudflare for Twitter).\n\nUnderstanding IP classification, subnetting, and address allocation is fundamental to network configuration, security auditing, and system administration. The custom subnetting exercise (/25, /26, /27) demonstrated the direct tradeoff between subnet size and the number of available subnets — a critical consideration when designing scalable and secure network architectures for real organisations.')

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT303/Lab5/INT303_Lab5_Report.docx')
print("Report saved successfully!")
