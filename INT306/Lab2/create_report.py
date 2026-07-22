"""
INT306 Lab 2 - Post-Quantum Cryptography (PQC) and Future-Proofing Security
Report generator using python-docx

Usage on Kali:
    source ~/int306-venv/bin/activate
    python3 ~/Kali-Linux-Security-Labs/INT306/Lab2/create_report.py

Screenshots are NOT inserted automatically - this script writes plain text
placeholders like [INSERT SCREENSHOT: description] for manual replacement
in Word after generation.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- Styles ----------
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_screenshot_placeholder(description):
    p = doc.add_paragraph()
    run = p.add_run(f"[INSERT SCREENSHOT: {description}]")
    run.italic = True
    run.font.color.rgb = None
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_answer_block(question, answer):
    q = doc.add_paragraph()
    qr = q.add_run("Question: ")
    qr.bold = True
    q.add_run(question)

    a = doc.add_paragraph()
    ar = a.add_run("Answer: ")
    ar.bold = True
    a.add_run(answer)


def add_simple_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], "D9D9D9")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("INT306: Cryptography\nLab 2 – Post-Quantum Cryptography (PQC) and Future-Proofing Security")
run.bold = True
run.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Damian Ozimede Patrick\nLeapSchool Africa / ICDFA Cybersecurity Internship Program")
sub_run.font.size = Pt(13)

doc.add_page_break()

# ============================================================
# ENVIRONMENT
# ============================================================
add_heading("1. Lab Environment", level=1)
add_simple_table(
    ["Component", "Detail"],
    [
        ["Kali Linux host", "legion@192.168.5.139"],
        ["Python environment", "~/int306-venv (cryptography, pycryptodome, flask, bcrypt)"],
        ["OpenSSL version", "3.6.3 (9 Jun 2026)"],
        ["Working directory", "~/Kali-Linux-Security-Labs/INT306/Lab2"],
        ["GitHub repo", "github.com/damianozimede/Kali-Linux-Security-Labs"],
    ]
)

# ============================================================
# PRE-REQUISITE SETUP
# ============================================================
add_heading("2. Pre-requisite Setup: Installing oqs-provider", level=1)
add_para(
    "Note: This setup phase is not explicitly listed as a numbered exercise in the lab "
    "document. It was required because a stock OpenSSL 3.x installation does not include "
    "ML-KEM, ML-DSA, or SLH-DSA support without an additional provider module. The following "
    "steps install liboqs and oqs-provider so that Exercise 2.2 onward can be completed.",
    italic=True
)

add_heading("2.1 Installing build dependencies", level=2)
add_code_block("sudo apt update && sudo apt install -y build-essential cmake ninja-build libssl-dev python3-pip git")
add_screenshot_placeholder("apt install output showing build tools installed")

add_heading("2.2 Building and installing liboqs", level=2)
add_code_block(
    "cd ~\n"
    "git clone --depth 1 https://github.com/open-quantum-safe/liboqs.git\n"
    "cd liboqs\n"
    "mkdir build && cd build\n"
    "cmake -GNinja -DCMAKE_INSTALL_PREFIX=/opt/liboqs ..\n"
    "ninja\n"
    "sudo ninja install"
)
add_screenshot_placeholder("liboqs build and install completing successfully")
add_para(
    "The installed headers confirmed support for the official NIST-standardized names "
    "(kem_ml_kem.h, sig_ml_dsa.h, sig_slh_dsa.h) alongside legacy naming (kem_kyber.h)."
)

add_heading("2.3 Building and installing oqs-provider", level=2)
add_code_block(
    "cd ~\n"
    "git clone --depth 1 https://github.com/open-quantum-safe/oqs-provider.git\n"
    "cd oqs-provider\n"
    "liboqs_DIR=/opt/liboqs cmake -GNinja -DOPENSSL_ROOT_DIR=/usr -S . -B build\n"
    "cmake --build build\n"
    "sudo cmake --install build"
)
add_screenshot_placeholder("oqs-provider build and install completing, oqsprovider.so installed")

add_heading("2.4 Registering the provider with OpenSSL (Kali-specific config)", level=2)
add_para(
    "Kali Linux uses a non-standard OpenSSL configuration structure: /etc/ssl/openssl.cnf "
    "includes /etc/ssl/kali.cnf, which defines a kali_wide_compatibility_providers section "
    "rather than the standard provider_sect. The provider was registered by editing this file."
)
add_code_block("sudo cp /etc/ssl/kali.cnf /etc/ssl/kali.cnf.bak")
add_para("A backup was taken before editing, in line with standard lab practice of reverting all environment changes afterward.")

add_code_block(
    "[kali_wide_compatibility_providers]\n"
    "default = default_sect\n"
    "legacy = legacy_sect\n"
    "oqsprovider = oqsprovider_sect\n\n"
    "[oqsprovider_sect]\n"
    "activate = 1"
)
add_screenshot_placeholder("nano showing the edited kali.cnf with oqsprovider section added")

add_code_block("openssl list -providers")
add_screenshot_placeholder("openssl list -providers showing oqsprovider active alongside default and legacy")

add_para(
    "A key finding at this stage: OpenSSL 3.6.3 already includes native support for the "
    "finalized NIST standards (ML-KEM and ML-DSA) directly in the 'default' provider. "
    "The oqs-provider module was found to add mainly hybrid combinations "
    "(e.g. x25519_mlkem768) and non-standardized alternative algorithms (FrodoKEM, BIKE, "
    "HQC, Falcon, etc.), rather than the core NIST algorithms themselves.",
)

add_para("Algorithm naming used in this lab (differs from the lab document's suggested names):")
add_simple_table(
    ["Lab document suggests", "Actual working name used"],
    [
        ["kyber768", "MLKEM768"],
        ["dilithium3", "MLDSA65"],
    ]
)

# ============================================================
# EXERCISE 2.1
# ============================================================
add_heading("3. Exercise 2.1: Setting Up the OQS Environment", level=1)

add_heading("3.1 Verify OpenSSL Version", level=2)
add_code_block("openssl version")
add_para("Output: OpenSSL 3.6.3 9 Jun 2026 — confirms version 3.0+ requirement is met.")
add_screenshot_placeholder("openssl version output")

add_heading("3.2 Provider architecture discussion", level=2)
add_answer_block(
    "Why is the 'Provider' architecture in OpenSSL 3.x significant for the transition to "
    "Post-Quantum Cryptography?",
    "In older versions of OpenSSL, cryptographic algorithms were built directly into the "
    "core codebase, meaning any new algorithm required modifying and recompiling OpenSSL "
    "itself. OpenSSL 3.x introduced a plugin-style 'Provider' architecture, allowing new "
    "algorithms to be added as self-contained modules that load into OpenSSL at runtime, "
    "without changing its core. This matters for PQC because the NIST standards (ML-KEM, "
    "ML-DSA, SLH-DSA) were only finalized in August 2024. The Provider architecture allowed "
    "the Open Quantum Safe project to make these algorithms usable immediately through "
    "oqs-provider, years before they would otherwise be natively bundled into an official "
    "OpenSSL release. It also supports 'crypto-agility' - the ability to swap cryptographic "
    "modules in and out without re-architecting an entire system - which is essential during "
    "the current transition period between classical and post-quantum cryptography."
)

# ============================================================
# EXERCISE 2.2
# ============================================================
add_heading("4. Exercise 2.2: Generating PQC Key Pairs", level=1)

add_heading("4.1 Generate ML-KEM-768 Key Pair (Kyber-768 equivalent)", level=2)
add_para(
    "As anticipated by the lab document's own note on algorithm naming variance, the "
    "suggested algorithm name 'kyber768' was not recognized by OpenSSL 3.6.3. The correct "
    "NIST-standardized name, MLKEM768, was used instead, since ML-KEM support is now built "
    "natively into OpenSSL's default provider."
)
add_code_block(
    "openssl genpkey -algorithm MLKEM768 -out pqc_kem_private.pem\n"
    "openssl pkey -in pqc_kem_private.pem -pubout -out pqc_kem_public.pem"
)
add_screenshot_placeholder("ML-KEM key pair generation commands")

add_heading("4.2 Generate ML-DSA-65 Key Pair (Dilithium3 equivalent)", level=2)
add_code_block(
    "openssl genpkey -algorithm MLDSA65 -out pqc_sig_private.pem\n"
    "openssl pkey -in pqc_sig_private.pem -pubout -out pqc_sig_public.pem"
)
add_screenshot_placeholder("ML-DSA key pair generation commands")

add_heading("4.3 Key Size Comparison", level=2)
add_code_block(
    "openssl genpkey -algorithm RSA -pkeyopt rsa_keygen_bits:2048 -out rsa2048_private.pem\n"
    "openssl genpkey -algorithm EC -pkeyopt ec_paramgen_curve:P-256 -out ecc256_private.pem\n"
    "ls -lh"
)
add_screenshot_placeholder("ls -lh output showing all generated keys with sizes")

add_simple_table(
    ["Key type", "Private key size"],
    [
        ["ECC P-256", "241 bytes"],
        ["RSA-2048", "~1.7 KB"],
        ["ML-KEM-768 (PQC)", "3.4 KB"],
        ["ML-DSA-65 (PQC)", "5.5 KB"],
    ]
)

add_answer_block(
    "Discuss the potential impact of these larger keys on network protocols like TLS and IKEv2.",
    "PQC keys are much larger than what current protocols were designed around. In this "
    "lab, ECC came to 241 bytes and RSA-2048 to about 1.7 KB, while ML-KEM came to 3.4 KB "
    "and ML-DSA to 5.5 KB - roughly 3 to 14 times larger. For TLS, this means the initial "
    "handshake (where keys and certificates are exchanged before any real data is sent) "
    "becomes noticeably bigger, adding a small amount of latency and bandwidth cost per "
    "connection - barely noticeable for a single connection, but adding up for a server "
    "handling thousands of handshakes per second. For IKEv2 (used in VPNs), the impact can "
    "be more serious: IKEv2 negotiation runs over UDP, which does not handle oversized "
    "messages as gracefully as TCP. If a PQC key or ciphertext is too large to fit in a "
    "single UDP packet, it must be fragmented, and many firewalls are configured to block "
    "fragmented UDP traffic as a security measure. This creates a real risk that VPN "
    "negotiations using PQC keys could fail on networks with strict firewall rules, unless "
    "fragmentation handling (such as RFC 7383) is explicitly supported and enabled. Overall, "
    "the added security PQC provides comes at the cost of size, which creates practical "
    "friction for protocols that were originally designed around small, fast key exchanges."
)

# ============================================================
# EXERCISE 2.3
# ============================================================
add_heading("5. Exercise 2.3: PQC Digital Signatures with ML-DSA", level=1)

add_code_block(
    'echo "This message is protected by NIST-standardized Post-Quantum Cryptography." > pqc_message.txt\n'
    "openssl dgst -sign pqc_sig_private.pem -out pqc_message.sig pqc_message.txt"
)
add_screenshot_placeholder("message creation and signing commands")
add_para("The resulting signature file (pqc_message.sig) was 3.3 KB - dramatically larger than a typical ECDSA P-256 signature (~70 bytes) or RSA-2048 signature (256 bytes).")

add_code_block("openssl dgst -verify pqc_sig_public.pem -signature pqc_message.sig pqc_message.txt")
add_para("Output: Verified OK")
add_screenshot_placeholder("signature verification showing Verified OK")

add_answer_block(
    "Explain the mathematical problem (Module Learning with Errors - MLWE) that provides "
    "the security foundation for ML-DSA. How does it differ from the factoring problem used in RSA?",
    "RSA's security relies on the fact that multiplying two very large prime numbers "
    "together is easy, but working backward from the product to find the original two "
    "primes is extremely difficult using classical computers - this is the factoring "
    "problem. Shor's algorithm breaks this because a quantum computer can perform that "
    "reverse step efficiently. ML-DSA is built on a completely different kind of problem "
    "called Module Learning With Errors (MLWE). Instead of primes, it uses structured "
    "mathematical grids called lattices. The core problem is: given a set of linear "
    "equations that have had small amounts of random 'noise' deliberately mixed in, "
    "recover the original hidden values despite that noise. 'Module' refers to a "
    "specific, more efficient structured version of this problem used in Dilithium/ML-DSA. "
    "The key difference is that no known quantum algorithm, including Shor's or Grover's, "
    "provides an efficient shortcut for solving noisy lattice problems the way Shor's "
    "algorithm does for factoring. This is why lattice-based problems like MLWE were "
    "selected as the foundation for NIST's post-quantum standards - they resist the "
    "specific mathematical shortcuts that quantum computers are known to provide."
)

# ============================================================
# EXERCISE 2.4
# ============================================================
add_heading("6. Exercise 2.4: Key Encapsulation (KEM) with ML-KEM", level=1)

add_para(
    "Note: the lab document's suggested command used '-derive' and '-peerout', which are "
    "legacy Diffie-Hellman-era options. OpenSSL 3.2+ introduced dedicated KEM operations "
    "('-encap'/'-decap'), which were required here instead."
)

add_heading("6.1 Encapsulation", level=2)
add_code_block("openssl pkeyutl -encap -inkey pqc_kem_public.pem -pubin -secret shared_secret.bin -out encapsulated_key.bin")
add_screenshot_placeholder("encapsulation command executed successfully")
add_para("Result: shared_secret.bin (32 bytes - a standard 256-bit symmetric key) and encapsulated_key.bin (1.1 KB ciphertext) were created.")

add_heading("6.2 Decapsulation", level=2)
add_code_block("openssl pkeyutl -decap -inkey pqc_kem_private.pem -secret recovered_secret.bin -in encapsulated_key.bin")
add_screenshot_placeholder("decapsulation command executed successfully")

add_heading("6.3 Verification", level=2)
add_code_block('diff shared_secret.bin recovered_secret.bin && echo "MATCH: Shared secret successfully recovered"')
add_para("Output: MATCH: Shared secret successfully recovered — confirming both sides derived the identical 32-byte shared secret.")
add_screenshot_placeholder("diff output confirming matching shared secrets")

add_answer_block(
    "Why is a Key Encapsulation Mechanism (KEM) preferred over traditional Diffie-Hellman "
    "(DH) for many post-quantum lattice-based algorithms?",
    "Classical Diffie-Hellman works because both sides can independently run the same "
    "mathematical operation using their own private key plus the other party's public key, "
    "and both arrive at the identical shared secret through symmetric math. Lattice-based "
    "problems, such as the one behind ML-KEM, do not have this same neat symmetric "
    "property. Instead, they naturally fit a one-directional design: one side generates a "
    "public key, the other side uses it to create ('encapsulate') a random secret plus a "
    "ciphertext, and only the private key holder can reverse ('decapsulate') that ciphertext "
    "back into the same secret. This was demonstrated directly in this lab - the original "
    "DH-style command using '-derive' and '-peerout' failed, and the correct KEM-native "
    "'-encap'/'-decap' commands were required instead. KEM is preferred for lattice-based "
    "PQC because it maps naturally onto how the underlying math works (easy to verify with "
    "the private key, hard to reverse without it), and avoids the subtle implementation "
    "mistakes that come from forcing lattice math into an older protocol shape it was not "
    "designed for."
)

# ============================================================
# SECTION 3
# ============================================================
add_heading("7. Hybrid Cryptography and Migration Strategies", level=1)

add_heading("7.1 Exercise 3.1: The Hybrid Approach", level=2)
add_para(
    "Cloudflare and Chrome have both worked to deploy hybrid post-quantum key exchange, but "
    "their experiences show this transition isn't as simple as flipping a switch."
)
add_para(
    "Cloudflare uses a hybrid combining the classical X25519 algorithm with ML-KEM (the "
    "finalized NIST standard, previously known as Kyber) to secure TLS 1.3 traffic across "
    "millions of domains automatically. They've since gone further, also adding ML-DSA (the "
    "post-quantum signature standard) alongside key exchange support."
)
add_para(
    "Chrome's rollout was less smooth. Google initially shipped hybrid PQC key exchange in "
    "Chrome version 124, but later rolled it back after it caused widespread TLS handshake "
    "failures in enterprise network environments - showing that even a well-tested hybrid "
    "approach can expose real-world compatibility problems that weren't caught in earlier "
    "testing."
)
add_para(
    "The core reason hybrid is used instead of switching to PQC alone: it acts as a safety "
    "net. If ML-KEM/Kyber's underlying math is ever found to have an unexpected weakness, "
    "the connection still stays as secure as the classical X25519 component alone. Neither "
    "algorithm has to be perfect on its own; only one of the two needs to hold."
)

add_answer_block(
    "What are the primary benefits and drawbacks of using hybrid cryptography instead of "
    "switching entirely to PQC immediately?",
    "Benefits: hybrid cryptography provides a safety net, since protection does not "
    "collapse if the new PQC math is later found to be weak; it allows gradual adoption "
    "across the internet without breaking existing systems overnight; and it has already "
    "been proven at scale by Cloudflare across millions of domains. Drawbacks: key and "
    "ciphertext sizes are dramatically larger - Kyber768 key shares are around 35 times "
    "larger than a plain X25519 exchange, adding measurable latency to every handshake. "
    "This overhead is barely noticeable on desktop or high-bandwidth connections, but "
    "causes real, measurable performance regressions on mobile devices. Even mature hybrid "
    "implementations can hit real interoperability problems, as shown by Chrome's rollback - "
    "the wider ecosystem (middleboxes, enterprise firewalls, older TLS stacks) is not fully "
    "ready yet. In short, hybrid cryptography trades some performance and complexity for "
    "insurance against the possibility that a brand-new cryptographic standard turns out to "
    "have a flaw nobody has found yet - a reasonable price during a transition period where "
    "trust in PQC algorithms is still being established."
)

add_heading("7.2 Exercise 3.2: Crypto-Agility and Inventory", level=2)
add_para("Cryptographic inventory (as CISO of a mid-sized organization) - systems requiring quantum-safe updates:")

inventory_items = [
    "VPNs / IPsec (IKEv2) - key exchange and authentication need PQC or hybrid support",
    "Web servers / TLS certificates - public-facing HTTPS, load balancers, reverse proxies",
    "SSH - used for all server administration access",
    "Internal Certificate Authorities (CAs) - would need PQC-capable CA software and reissued certificates",
    "Email encryption (S/MIME, PGP/GPG) - long-term encrypted archives are especially at risk",
    "Database encryption at rest - where RSA/ECC-wrapped keys protect stored data",
    "Code signing - software update mechanisms, firmware signing",
    "Backup and archival systems - long-lived encrypted data is exposed to Harvest Now, Decrypt Later attacks",
    "IoT / embedded devices - often hardest to update, may lack resources to run PQC algorithms given their larger key/signature sizes",
]
for item in inventory_items:
    doc.add_paragraph(item, style='List Bullet')

add_answer_block(
    "Discuss the concept of 'Harvest Now, Decrypt Later' (HNDL) attacks. Why does this "
    "threat make the transition to PQC urgent even before a cryptographically relevant "
    "quantum computer (CRQC) exists?",
    "A Harvest Now, Decrypt Later (HNDL) attack does not require a quantum computer to "
    "exist yet - it only requires an adversary to record and store encrypted traffic today, "
    "with the plan to decrypt it later once a sufficiently powerful quantum computer becomes "
    "available. Since RSA and ECC encrypted data can potentially be broken once a "
    "cryptographically relevant quantum computer (CRQC) exists, any data encrypted today "
    "with classical algorithms is already vulnerable, even if that break is still years "
    "away. This makes the threat urgent right now for one key reason: the danger isn't when "
    "the quantum computer arrives - it's when the data being protected today loses its need "
    "for confidentiality. For information with a short shelf life, HNDL isn't a serious "
    "concern. But for information that must remain secret for years or decades - medical "
    "records, government secrets, trade secrets, long-term financial data - the risk is real "
    "today, because an adversary intercepting and storing that traffic now could still "
    "successfully read it once quantum computers catch up. In short, the timeline that "
    "matters isn't 'when will quantum computers break RSA' - it's 'how long does this "
    "specific piece of data need to stay confidential.' If that answer is longer than the "
    "expected timeline to a CRQC, the data is already at risk, which is exactly why "
    "organizations are told to migrate to PQC now rather than waiting."
)

# ============================================================
# ENVIRONMENT REVERT
# ============================================================
add_heading("8. Environment Revert", level=1)
add_para(
    "In line with standard lab practice, the system-level configuration change made during "
    "setup was reverted after completing all exercises."
)
add_code_block(
    "sudo cp /etc/ssl/kali.cnf.bak /etc/ssl/kali.cnf\n"
    "openssl list -providers"
)
add_para("Confirmed output showed only 'default' and 'legacy' providers active, matching Kali's original configuration.")
add_screenshot_placeholder("openssl list -providers after revert, showing oqsprovider no longer active")

# ============================================================
# CONCLUSION
# ============================================================
add_heading("9. Conclusion", level=1)
add_para(
    "This lab provided hands-on experience with NIST's finalized Post-Quantum Cryptography "
    "standards - ML-KEM for key encapsulation and ML-DSA for digital signatures - using the "
    "Open Quantum Safe project on Kali Linux. A significant finding was that OpenSSL 3.6.3 "
    "already includes native support for these NIST-standardized algorithms directly in its "
    "default provider, with oqs-provider primarily adding hybrid combinations and "
    "non-standardized experimental algorithms. The lab also required correcting several "
    "outdated commands from the lab document - substituting current NIST algorithm names "
    "for legacy naming, and using modern KEM-native OpenSSL syntax in place of legacy "
    "Diffie-Hellman-era options. Beyond the hands-on exercises, this lab reinforced why "
    "the transition to post-quantum cryptography is both technically complex (due to larger "
    "key and signature sizes affecting protocols like TLS and IKEv2) and organizationally "
    "urgent (due to Harvest Now, Decrypt Later attacks), making crypto-agility an essential "
    "property for any modern security architecture."
)

# ============================================================
# SAVE
# ============================================================
doc.save("INT306_Lab2_Report.docx")
print("Report generated: INT306_Lab2_Report.docx")
