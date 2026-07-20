from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def add_placeholder(text):
    p = doc.add_paragraph()
    run = p.add_run(f"[INSERT SCREENSHOT: {text}]")
    run.italic = True
    run.font.color.rgb = None

# ===================== TITLE PAGE =====================
title = doc.add_heading('INT306: Cryptography - Lab 1 Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Levels 1-8 and Mastery Challenge')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Student: Damian Ozimede Patrick')
doc.add_paragraph('Kali Username: legion')
doc.add_page_break()

# ===================== SUMMARY SECTIONS =====================
doc.add_heading('1. Activities Performed', level=1)
doc.add_paragraph(
    "Completed all 8 levels of INT306 Lab 1 (Cryptography) plus the Mastery Challenge: "
    "classical ciphers (Caesar, Vigenere), symmetric encryption (AES-CBC, AES-GCM, ECB vs CBC "
    "comparison), hashing (SHA-256, MD5, HMAC, bcrypt, rainbow tables, MD5 collision demo), "
    "asymmetric cryptography (RSA manual calculation, OAEP encryption, Diffie-Hellman), digital "
    "signatures and PKI (signing, tamper detection, self-signed certificates, trust chains), "
    "real-world protocols (TLS handshake analysis in Wireshark, nmap cipher suite inspection, "
    "GPG encryption), broken crypto (a self-built padding oracle attack, hash length extension "
    "attack), and future cryptography (Zero-Knowledge Proofs, post-quantum algorithms). The "
    "Mastery Challenge combined DH key exchange, AES session encryption, SHA-3 signed audit "
    "trails, and a PQC migration plan into one simulated banking scenario."
)

doc.add_heading('2. Skills Learned', level=1)
doc.add_paragraph(
    "Practical use of OpenSSL for encryption, signing, and certificates; Python cryptographic "
    "programming using the cryptography and pycryptodome libraries; reading and interpreting "
    "Wireshark packet captures; using nmap for security scanning; installing and compiling "
    "security tools from source on Kali; and applying core cryptographic concepts (modular "
    "arithmetic, hashing properties, authenticated encryption) to real, hands-on exercises "
    "rather than just theory."
)

doc.add_heading('3. Challenges Faced', level=1)
doc.add_paragraph(
    "Early difficulty grasping number theory concepts (modular inverse, Euler's totient "
    "function) before seeing them applied practically in RSA. Environment setup issues "
    "(externally-managed-environment error, needing a Python virtual environment). Tool "
    "installation obstacles, including a dead GitHub repository for HashPump requiring a "
    "mirror, and a WebGoat version mismatch that lacked the expected padding oracle lesson, "
    "requiring a custom-built vulnerable server instead. A calculation error in the hash "
    "length extension attack (wrong key length) that had to be debugged and corrected."
)

doc.add_heading('4. Personal Reflection', level=1)
doc.add_paragraph(
    "This lab significantly deepened my understanding of how cryptography operates in "
    "practice, not just in theory. Building and successfully attacking my own vulnerable "
    "systems (the padding oracle server, the hash length extension forgery) made the "
    "underlying vulnerabilities feel real and memorable in a way that reading about them "
    "wouldn't have. Connecting concepts across exercises - like seeing the same avalanche "
    "effect explain both hash integrity and why Git detects tampering - helped the material "
    "feel connected rather than like isolated facts to memorize."
)
doc.add_page_break()

# ===================== LEVEL 1: CLASSICAL CRYPTOGRAPHY =====================
doc.add_heading('Level 1: Classical Cryptography', level=1)

doc.add_heading('Caesar Cipher - Manual Decryption', level=2)
doc.add_paragraph("Decrypted WTAAD with shift k=15 using D(x) = (x - 15) mod 26, giving the result HELLO.")
add_placeholder("manual working for WTAAD decryption")

doc.add_heading('Caesar Cipher - Brute Force', level=2)
doc.add_paragraph(
    "Attempted brute force decryption of ciphertext XAB across all 25 possible shifts (k=1 to "
    "k=25). No shift produced a valid English word. This is because a Caesar shift preserves "
    "the relative gaps between letters regardless of k, and XAB's letter-gap pattern does not "
    "match any real English word. This appears to be an error in the lab's provided ciphertext."
)
add_placeholder("brute force working for XAB")

doc.add_heading('Caesar Cipher - Python Implementation', level=2)
doc.add_paragraph(
    "Implemented a caesar_cipher() function supporting both encryption and decryption, "
    "handling uppercase/lowercase letters and leaving non-letter characters unchanged. "
    "Tested with plaintext CRYPTOGRAPHY and shift 7:"
)
doc.add_paragraph("Encrypted: JYFWAVNYHWOF")
doc.add_paragraph("Decrypted back: CRYPTOGRAPHY")
add_placeholder("Python function definition and test output")

doc.add_heading('Vigenere Cipher - Manual Decryption', level=2)
doc.add_paragraph(
    "Decrypted LXFOPVEFRNHR using key LEMON, applying P = (C - K) mod 26 to each letter pair, "
    "resulting in the plaintext ATTACK AT DAWN."
)
add_placeholder("manual Vigenere working")

doc.add_heading('Vigenere Cipher - Python Implementation', level=2)
doc.add_paragraph(
    "Implemented vigenere_decrypt(), which repeats the key automatically to match ciphertext "
    "length using modulo, and only advances the key index on actual letters. Tested with "
    "LXFOPVEFRNHR and key LEMON, correctly returning ATTACKATDAWN."
)
add_placeholder("Python function and test output")

doc.add_page_break()

# ===================== LEVEL 2: SYMMETRIC KEY CRYPTOGRAPHY =====================
doc.add_heading('Level 2: Symmetric Key Cryptography', level=1)

doc.add_heading('OpenSSL AES-256-CBC Encryption', level=2)
doc.add_paragraph(
    "Encrypted and decrypted a text file using AES-256-CBC via OpenSSL. Initial run showed a "
    "deprecation warning for the default key derivation method; re-ran using the -pbkdf2 flag, "
    "which resolved the warning and used a stronger, modern key derivation approach."
)
add_placeholder("OpenSSL AES-256-CBC encrypt/decrypt output")

doc.add_heading('AES-128-CBC - Wrong Password Test', level=2)
doc.add_paragraph(
    "Encrypted a file with AES-128-CBC and attempted decryption with an incorrect password. "
    "The result was a 'bad decrypt' error, since the wrong key produced garbage plaintext that "
    "failed the padding validity check. This demonstrates that CBC mode alone does not detect "
    "tampering or wrong keys gracefully - it fails via a padding error rather than a clear "
    "authentication failure."
)
add_placeholder("wrong password decrypt failure")

doc.add_heading('ECB vs CBC Visual Comparison', level=2)
doc.add_paragraph(
    "Created a test image with a repeating pattern of blue and yellow squares, then encrypted "
    "it using both AES-256-ECB and AES-256-CBC. Examining the raw bytes with xxd showed several "
    "identical 16-byte sequences repeating throughout the ECB-encrypted file, while the "
    "CBC-encrypted file showed no repeating patterns anywhere."
)
doc.add_paragraph(
    "This happens because AES encrypts data in fixed 16-byte blocks. In ECB mode, each block is "
    "encrypted independently, so identical plaintext blocks always produce identical ciphertext "
    "blocks. Because the test image had repeating squares, many underlying pixel blocks were "
    "identical, and ECB leaked this repetition directly into the ciphertext. CBC avoids this "
    "because each block is XORed with the previous ciphertext block before encryption, so even "
    "identical plaintext blocks produce different ciphertext depending on what came before them. "
    "This confirms why ECB is considered insecure for data with repeating structure, such as images."
)
add_placeholder("checkerboard test image")
add_placeholder("xxd output showing ECB repeating pattern vs CBC randomness")

doc.add_heading('Python AES-GCM', level=2)
doc.add_paragraph(
    "Encrypted the string 'Authenticated Data' using AES-GCM, including a randomly generated "
    "nonce (IV) and an authentication tag appended automatically to the ciphertext. Successfully "
    "decrypted it back to the original message. Then deliberately corrupted one byte of the "
    "ciphertext and attempted decryption again - this time it failed outright with an "
    "authentication error, rather than silently returning garbage as CBC would."
)
doc.add_paragraph(
    "This is the key difference between CBC and GCM: CBC has no built-in integrity check, so "
    "tampering produces garbage plaintext with no warning, while GCM's authentication tag "
    "detects any modification and refuses to decrypt, providing both confidentiality and "
    "integrity in one mechanism."
)
add_placeholder("AES-GCM encrypt/decrypt + tamper detection output")

doc.add_heading('Secure File Encryption Tool', level=2)
doc.add_paragraph(
    "Built a command-line Python tool (secure_file_tool.py) that derives a 256-bit key from a "
    "user password using PBKDF2 with a random salt, then encrypts/decrypts files using "
    "AES-256-GCM. Salt and nonce are stored alongside the ciphertext since they are not secret. "
    "Successfully encrypted a test file, decrypted it back to verify the contents matched "
    "exactly, then attempted decryption with an incorrect password - which failed cleanly due "
    "to GCM's authentication check, with no garbage output produced."
)
add_placeholder("encrypt, decrypt, and wrong-password test outputs")

doc.add_page_break()

# ===================== LEVEL 3: HASHING AND MESSAGE INTEGRITY =====================
doc.add_heading('Level 3: Hashing and Message Integrity', level=1)

doc.add_heading('Compute Hash Values', level=2)
doc.add_paragraph(
    "Computed MD5 and SHA-256 hashes for three strings using OpenSSL. MD5 produced a 128-bit "
    "hash shown as 32 hex characters, while SHA-256 produced a 256-bit hash shown as 64 hex "
    "characters - double the length. Both algorithms produce a fixed-length output regardless "
    "of input size."
)
add_placeholder("hash computation output for all three strings")

doc.add_heading('File Integrity Verification', level=2)
doc.add_paragraph(
    "Downloaded a small file and computed its SHA-256 checksum, then used sha256sum -c to "
    "verify the file against that expected hash, confirming an 'OK' result. This mirrors the "
    "real-world workflow of verifying downloaded software against a publisher's official "
    "checksum to detect corruption or tampering during transfer."
)
add_placeholder("download and checksum verification output")

doc.add_heading('Avalanche Effect', level=2)
doc.add_paragraph(
    "Computed SHA-256 hashes for two nearly identical sentences differing by a single letter "
    "(dog vs cog). The resulting hashes were completely different with no shared pattern, "
    "demonstrating the avalanche effect: a tiny input change causes a large, unpredictable "
    "change in hash output. This property is what makes hashing effective for detecting even "
    "minor file tampering or corruption."
)
add_placeholder("original.txt and modified.txt hash outputs")

doc.add_heading('HMAC Generation', level=2)
doc.add_paragraph(
    "Generated an HMAC-SHA256 for a sample financial transaction message using a secret key. "
    "Unlike a plain hash, which anyone can compute, an HMAC can only be reproduced by someone "
    "who knows the secret key, providing both integrity and authenticity rather than integrity "
    "alone."
)
add_placeholder("HMAC generation output")

doc.add_heading('Password Security (bcrypt)', level=2)
doc.add_paragraph(
    "Wrote a Python script using bcrypt to hash a password with an automatically generated "
    "salt, then verify a password attempt against the stored hash. Correct passwords verified "
    "as True, incorrect passwords as False. Bcrypt is preferred over plain SHA-256 for "
    "passwords because it is deliberately slow (resisting brute-force guessing), automatically "
    "generates and embeds a unique salt per password (preventing rainbow table attacks), and "
    "its cost factor can be increased over time as computers get faster - none of which SHA-256 "
    "provides."
)
add_placeholder("bcrypt hash and verification output")

doc.add_heading('Rainbow Table Concept', level=2)
doc.add_paragraph(
    "A rainbow table is a precomputed lookup table mapping common passwords to their hash "
    "values, allowing an attacker to instantly reverse a stolen hash back to its original "
    "password without guessing in real time. Salting defeats this attack because a random, "
    "unique salt is mixed into each password before hashing, meaning identical passwords "
    "produce different hashes for different users. This makes a precomputed table useless, "
    "since the attacker would need a separate table for every possible salt value."
)

doc.add_heading('Collision Demonstration', level=2)
doc.add_paragraph(
    "Downloaded two known MD5-colliding PDF files from the corkami/collisions GitHub "
    "repository. Both files produced the identical MD5 hash despite having different content, "
    "confirmed using diff to show the underlying bytes genuinely differed. This demonstrates "
    "that MD5 is not collision-resistant: an attacker could substitute a malicious file for a "
    "legitimate one while keeping the same MD5 hash, defeating integrity checks that rely on "
    "MD5 alone. For digital signatures, this means an attacker could potentially get a "
    "signature intended for one document accepted as valid for a different, malicious document "
    "with the same hash."
)
add_placeholder("md5sum output showing identical hashes")
add_placeholder("diff output showing different content")

doc.add_heading('Rainbow Table Creation', level=2)
doc.add_paragraph(
    "Built a Python dictionary mapping five common passwords to their MD5 hashes, then used it "
    "to instantly look up the original password for a given hash (successfully cracking "
    "'qwerty' from its MD5 hash). This simple table only works against passwords already in "
    "its list; a real attacker would need millions of entries for meaningful coverage, and "
    "salting (as discussed above) would defeat this approach entirely regardless of table size."
)
add_placeholder("rainbow table script output showing successful crack")

doc.add_heading('Real-World Application: Hash Functions in Git', level=2)
doc.add_paragraph(
    "Researched and wrote a report on how Git uses hash functions as the foundation of its "
    "content-addressable storage system, covering blobs, trees, and commits, the SHA-1 to "
    "SHA-256 migration, and the real-world SHAttered collision attack. Full write-up attached "
    "as a separate document (git_hash_report.docx)."
)

doc.add_heading('Implementing a Hash Table', level=2)
doc.add_paragraph(
    "Implemented a HashTable class in Python using chaining for collision resolution, with "
    "insert, search, and delete operations. Tested with five key-value pairs, confirming "
    "correct insertion, key updates, lookups, and deletion. The _hash method allows near-instant "
    "lookups by converting a key directly into a bucket index, rather than needing to scan every "
    "stored item - this is what gives hash tables their speed advantage over simple lists."
)
add_placeholder("hash table script output")

doc.add_page_break()

# ===================== LEVEL 4: ASYMMETRIC KEY CRYPTOGRAPHY =====================
doc.add_heading('Level 4: Asymmetric Key Cryptography', level=1)

doc.add_heading('RSA Key Pair Generation (OpenSSL)', level=2)
doc.add_paragraph("Generated a 2048-bit RSA private key and extracted the corresponding public key using OpenSSL.")
add_placeholder("key generation output")

doc.add_heading('Manual RSA Calculation', level=2)
doc.add_paragraph(
    "Given p=3 and q=11: calculated n = p*q = 33, and phi(n) = (p-1)(q-1) = 20. Chose e=3, "
    "verified gcd(3,20)=1. Found the modular inverse d=7, since (3*7) mod 20 = 1. This gave "
    "public key (n=33, e=3) and private key (n=33, d=7). Encrypted message M=5 using C = M^e "
    "mod n = 5^3 mod 33 = 26. Decrypted using M = C^d mod n = 26^7 mod 33 = 5, confirming a "
    "full, correct RSA round trip using small numbers."
)
add_placeholder("manual calculation working")

doc.add_heading('Python RSA Encryption with OAEP', level=2)
doc.add_paragraph(
    "Generated a 2048-bit RSA key pair using the cryptography library, then encrypted the "
    "message 'Top Secret' using the public key with OAEP padding, and successfully decrypted "
    "it back using the private key."
)
add_placeholder("RSA OAEP encrypt/decrypt output")

doc.add_heading('Diffie-Hellman Key Exchange Simulation', level=2)
doc.add_paragraph(
    "Implemented a Python simulation of Diffie-Hellman key exchange between Alice and Bob. "
    "Both parties generated their own private/public key pairs locally, exchanged only their "
    "public keys, and each independently derived an identical shared secret key using HKDF. "
    "The private keys were never transmitted; the shared secret arises mathematically because "
    "Alice computes (Bob's public key)^(her private key) and Bob computes (Alice's public "
    "key)^(his private key), both landing on the same value g^(ab) mod p due to how exponents "
    "combine. An eavesdropper observing only the public keys cannot reconstruct the shared "
    "secret without solving the Discrete Logarithm Problem, which is computationally infeasible."
)
add_placeholder("DH simulation output showing matching derived keys")

doc.add_page_break()

# ===================== LEVEL 5: DIGITAL SIGNATURES AND PKI =====================
doc.add_heading('Level 5: Digital Signatures and PKI', level=1)

doc.add_heading('Sign and Verify', level=2)
doc.add_paragraph(
    "Signed a text file using the RSA private key generated earlier, then verified the "
    "signature using the corresponding public key, confirming 'Verified OK'."
)
add_placeholder("sign and verify output")

doc.add_heading('Tamper Detection', level=2)
doc.add_paragraph(
    "Modified the signed file's contents after signing, then attempted verification again. "
    "This time verification failed, since the signature was generated for the original "
    "content's hash, and any change to the content changes that hash completely (the same "
    "avalanche effect observed in Level 3), breaking the signature match."
)
add_placeholder("verification failure output")

doc.add_heading('Self-Signed Certificate', level=2)
doc.add_paragraph(
    "Created a self-signed X.509 certificate for mysite.local, valid for 365 days, using a "
    "4096-bit RSA key. Verified the certificate details showed matching Issuer and Subject "
    "fields (both CN=mysite.local), confirming the certificate vouches for itself rather than "
    "being verified by an external Certificate Authority."
)
add_placeholder("certificate generation and details output")

doc.add_heading('Trust Chain Analysis', level=2)
doc.add_paragraph(
    "Examined google.com's live certificate chain using OpenSSL. The certificate's subject was "
    "*.google.com, issued by an intermediate Certificate Authority (Google Trust Services, CN=WE2) "
    "- unlike the self-signed certificate above, where issuer and subject were identical. Real-world "
    "HTTPS trust flows through a hierarchy: the website's certificate is signed by an intermediate "
    "CA, which is itself signed by a trusted Root CA pre-installed in browsers and operating systems."
)
add_placeholder("trust chain subject/issuer output")

doc.add_page_break()

# ===================== LEVEL 6: REAL-WORLD PROTOCOLS =====================
doc.add_heading('Level 6: Real-world Protocols and Advanced Topics', level=1)

doc.add_heading('Wireshark TLS Handshake Analysis', level=2)
doc.add_paragraph(
    "Captured live HTTPS traffic to google.com in Wireshark, filtered by 'tls'. Observed the "
    "Client Hello (including SNI=www.google.com) and Server Hello with Change Cipher Spec, "
    "followed almost immediately by Application Data packets. Since the connection used TLS "
    "1.3, the handshake was streamlined so that the certificate exchange and Finished messages "
    "are encrypted within the handshake itself rather than appearing as separately labeled "
    "plaintext packets, reflecting TLS 1.3's improved 1-RTT handshake and reduced plaintext "
    "exposure compared to older TLS versions."
)
doc.add_paragraph(
    "Drilling into the Client Hello packet showed 90 offered cipher suites, a client-generated "
    "random value used later in key derivation, a session ID for resumption, and various "
    "extensions including SNI."
)
add_placeholder("Wireshark packet list showing TLS handshake")
add_placeholder("Client Hello detail view")

doc.add_heading('Cipher Suite Inspection (nmap)', level=2)
doc.add_paragraph(
    "Scanned google.com's supported TLS cipher suites using nmap's ssl-enum-ciphers script. "
    "Nearly all offered ciphers scored grade A, except TLS_RSA_WITH_3DES_EDE_CBC_SHA, graded C "
    "and flagged as vulnerable to the SWEET32 attack due to its small 64-bit block size. "
    "TLSv1.0 and TLSv1.1 were also present, both considered deprecated protocol versions. "
    "Notably, TLSv1.3 offered TLS_AKE_WITH_AES_256_GCM_SHA384 and other modern AEAD ciphers "
    "using X25519MLKEM768 - a hybrid key exchange combining classical elliptic curve "
    "cryptography with the post-quantum ML-KEM algorithm, showing that post-quantum protection "
    "is already deployed in production by major providers."
)
add_placeholder("nmap ssl-enum-ciphers output")

doc.add_heading('GPG File Encryption', level=2)
doc.add_paragraph(
    "Generated a 3072-bit RSA GPG key pair, then encrypted a text file for myself as the "
    "recipient and successfully decrypted it back, confirming the original message was "
    "restored and identifying which key was used for decryption."
)
add_placeholder("GPG key generation, encryption, and decryption output")

doc.add_page_break()

# ===================== LEVEL 7: BROKEN CRYPTO & CRYPTANALYSIS =====================
doc.add_heading('Level 7: Broken Crypto and Cryptanalysis', level=1)

doc.add_heading('Padding Oracle Simulation', level=2)
doc.add_paragraph(
    "Since the available WebGoat version did not include a padding oracle lesson, built a "
    "custom vulnerable Flask server that encrypts a session cookie with AES-CBC and leaks "
    "whether submitted ciphertext has valid padding via HTTP status codes (200 for valid, 500 "
    "for invalid) - the actual flaw that enables this attack class. Used padbuster against the "
    "server, which correctly identified the error signature and systematically recovered the "
    "plaintext of the cookie's final block ('guest') purely by observing padding validity "
    "responses, without ever knowing the encryption key."
)
doc.add_paragraph(
    "This demonstrates that a server which distinguishes between padding errors and other "
    "errors leaks enough information for an attacker to decrypt data byte by byte through "
    "repeated queries. This is the vulnerability class behind real historical attacks such as "
    "POODLE against SSL 3.0."
)
add_placeholder("Flask server startup showing encrypted cookie")
add_placeholder("padbuster output showing recovered plaintext")

doc.add_heading('Hash Length Extension Attack', level=2)
doc.add_paragraph(
    "Simulated a scenario where a server signs data as SHA1(secret + message) instead of using "
    "HMAC. Starting only with a known signature, the known visible message, and the secret's "
    "length, used hashpump to forge a valid signature for an extended message with additional "
    "attacker-controlled data appended - without ever knowing the actual secret. An initial "
    "attempt failed due to miscounting the secret's byte length (used 16 instead of the correct "
    "15); once corrected, the forged signature was verified as mathematically valid by "
    "independently recomputing the hash with the real secret."
)
doc.add_paragraph(
    "This demonstrates why plain hash-then-secret constructions are unsafe for signing "
    "data: an attacker who intercepts one signed message can forge new signed messages with "
    "extra malicious parameters appended, and the server would accept them as genuine. HMAC "
    "avoids this vulnerability through its double-hashing construction, which breaks the "
    "length-extension property that plain hashing is vulnerable to."
)
add_placeholder("original hash computation")
add_placeholder("hashpump forged signature output")
add_placeholder("verification script showing Match: True")

doc.add_page_break()

# ===================== LEVEL 8: FUTURE OF CRYPTOGRAPHY =====================
doc.add_heading('Level 8: Future of Cryptography', level=1)

doc.add_heading('Zero-Knowledge Proofs - The Ali Baba Cave Analogy', level=2)
doc.add_paragraph(
    "Imagine a circular cave with one entrance that splits into two paths, A and B, which meet "
    "at a secret door deep inside. Peggy knows the password that opens this door, but she never "
    "tells Victor what it is. Victor wants to be convinced that Peggy really knows the password, "
    "without ever learning it himself."
)
doc.add_paragraph(
    "To do this, Peggy walks into the cave and takes either path A or B. Victor then stands at "
    "the entrance and randomly calls out which path he wants her to come back out from. If "
    "Peggy genuinely knows the password, she can always come out the requested path, using the "
    "secret door to cross over if she originally went down the other one. If she didn't know "
    "the password, she'd only have a 50% chance of guessing the right path correctly."
)
doc.add_paragraph(
    "They repeat this process many times. Since guessing correctly every single round by chance "
    "becomes extremely unlikely the more rounds they do, Victor becomes confident that Peggy "
    "genuinely knows the password - all without ever seeing the password itself, or the door "
    "being opened."
)

doc.add_heading('Post-Quantum Cryptography Research', level=2)
doc.add_paragraph(
    "NIST officially finalized its first three PQC standards in August 2024:"
)
doc.add_paragraph("1. ML-KEM (formerly CRYSTALS-Kyber) - FIPS 203, for general encryption/key exchange, replacing RSA and Diffie-Hellman. Already deployed in production, as observed in this lab's own nmap scan of google.com (X25519MLKEM768).", style='List Bullet')
doc.add_paragraph("2. ML-DSA (formerly CRYSTALS-Dilithium) - FIPS 204, the primary standard for digital signatures, replacing RSA/ECDSA signatures.", style='List Bullet')
doc.add_paragraph("3. SLH-DSA (formerly SPHINCS+) - FIPS 205, a hash-based signature scheme whose security rests solely on hash function properties rather than lattice math, serving as an independent backup approach.", style='List Bullet')

doc.add_page_break()

# ===================== MASTERY CHALLENGE =====================
doc.add_heading('Mastery Challenge: The Cryptographic Gauntlet', level=1)

doc.add_heading('Part 1: Key Management', level=2)
doc.add_paragraph(
    "For a banking system handling high-value transactions, RSA-4096 private keys cannot "
    "simply be stored as files on a server, the way private.pem was generated earlier in this "
    "lab using OpenSSL. In a real production system, private keys would instead be stored "
    "inside a Hardware Security Module (HSM) - a dedicated, tamper-resistant physical device "
    "designed so that the private key material never leaves the device, even during signing "
    "operations. The server sends data to the HSM to be signed, and only the signature comes "
    "back out; the key itself is never exposed to the operating system or any application code."
)
doc.add_paragraph(
    "For key rotation, keys would not be used indefinitely. A realistic policy would generate "
    "a new RSA-4096 key pair on a fixed schedule (for example, annually) or immediately if a "
    "compromise is suspected. Critically, old public keys must remain available for some time "
    "after rotation, so that signatures made with the old key can still be verified - otherwise, "
    "historical transactions could no longer be validated."
)
doc.add_paragraph(
    "Access control would restrict which systems and personnel are permitted to request a "
    "signing operation from the HSM at all, following the principle of least privilege - only "
    "the specific transaction-processing service should be authorized, not general staff or "
    "unrelated systems."
)
doc.add_paragraph(
    "Finally, for backup and recovery, the key material itself is too sensitive to be backed "
    "up as a single copy in one place. A common real-world approach is Shamir's Secret Sharing, "
    "which splits the key into multiple pieces distributed among different custodians, such "
    "that a minimum number of them must come together to reconstruct the key. This ensures no "
    "single compromised or malicious individual can recover the key alone."
)

doc.add_heading('Part 2: Secure Channel Simulation', level=2)
doc.add_paragraph(
    "Built a Python script simulating a TLS-like handshake between a bank and a client. "
    "Established a shared secret using Diffie-Hellman key exchange without transmitting either "
    "party's private key, derived a 256-bit AES session key from that shared secret using HKDF, "
    "then used the session key to encrypt a sample transaction with AES-256-GCM and successfully "
    "decrypted it back, confirming the full secure channel worked end-to-end."
)
add_placeholder("secure channel script output")

doc.add_heading('Part 3: Audit Trail', level=2)
doc.add_paragraph(
    "Extended the transaction handling with an audit trail: hashed the transaction using "
    "SHA-3-256, then signed that hash with a 4096-bit RSA private key (simulating the "
    "HSM-held signing key from Part 1) using PSS padding. Independently verified the signature "
    "against the public key, confirming the transaction record as authentic and untampered."
)
add_placeholder("audit trail script output")

doc.add_heading('Part 4: Quantum Readiness Migration Plan', level=2)
doc.add_paragraph(
    "The bank's current system, as built in this exercise, relies on RSA and Diffie-Hellman "
    "for key exchange and digital signatures - both vulnerable to Shor's algorithm on a "
    "sufficiently powerful quantum computer. A realistic migration plan would proceed in "
    "phases rather than a single cutover:"
)
doc.add_paragraph(
    "Phase 1 - Hybrid deployment: combine classical algorithms with post-quantum ones, as "
    "already demonstrated by Google's use of X25519MLKEM768 observed in this lab's nmap scan, "
    "ensuring security even if only one algorithm is eventually broken."
)
doc.add_paragraph(
    "Phase 2 - Migrate key exchange to ML-KEM (FIPS 203), replacing the Diffie-Hellman "
    "exchange used in Part 2."
)
doc.add_paragraph(
    "Phase 3 - Migrate digital signatures to ML-DSA (FIPS 204) for general use, with SLH-DSA "
    "(FIPS 205) considered for high-assurance signatures due to its independent, hash-based "
    "security foundation, replacing the RSA-4096 signatures used in Part 3."
)
doc.add_paragraph(
    "Phase 4 - Upgrade HSM firmware/hardware to support generating and storing post-quantum "
    "key types, since current HSMs are typically built around RSA/ECC operations."
)
doc.add_paragraph(
    "Ongoing - Prioritize early migration for particularly sensitive long-term data (audit "
    "records, archived transactions), due to the harvest-now-decrypt-later risk: encrypted "
    "data intercepted today could be stored and decrypted later once quantum computers mature."
)

doc.save('INT306_Lab1_Report.docx')
print("Report generated successfully.")
