from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)

def set_run(run, bold=False, size=12, color=None, font="Arial"):
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 56, 100)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, size=12, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, bold=bold, size=size)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=size)
    return p

def add_bold_label(doc, label, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r1 = p.add_run(label + " ")
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(text)
    set_run(r2, bold=False, size=12)
    return p

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("INT304: Network Security and Protocols")
set_run(r, bold=True, size=18, color=(31, 56, 100))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(48)
r = p.add_run("Lab 3: Network Security Protocols and Configuration")
set_run(r, bold=True, size=15, color=(46, 84, 150))

for line in ["Student: Damian Ozimede Patrick",
             "Course: INT304 — Network Security and Protocols",
             "Institution: International Cybersecurity & Digital Forensic Academy"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(line)
    set_run(r, size=12)

doc.add_page_break()

# Executive Summary
add_heading(doc, "Executive Summary", level=1)
add_para(doc, "This report documents the hands-on configuration and verification of five key network security protocols carried out on a Kali Linux virtual machine. The protocols implemented are SSL/TLS for securing web traffic, OpenVPN for encrypted VPN tunnelling, SSH for secure remote access, and 802.1X port-based network access control using FreeRADIUS. Each protocol was configured, tested, and verified using appropriate tools, and the results were captured through screenshots demonstrating successful implementation.")
add_para(doc, "The lab environment consisted of a single Kali Linux VM (IP: 192.168.5.139) acting as both server and client for all protocol configurations. Apache was used as the web server for SSL/TLS implementation, OpenVPN with Easy-RSA was used for VPN setup, OpenSSH was hardened with key-based authentication and port modification, and FreeRADIUS was deployed to simulate a RADIUS authentication server. All configurations were verified through live testing and packet capture using Wireshark.")

# Activity 1
add_heading(doc, "Activity 1: Overview of Network Security Protocols", level=1)
add_para(doc, "Before implementing the protocols, it is important to understand the role each one plays in securing network communications.")

add_heading(doc, "1.1 SSL/TLS (Secure Sockets Layer / Transport Layer Security)", level=2, color=(46, 84, 150))
add_para(doc, "SSL/TLS protocols secure data in transit between a client and a server by encrypting the communication channel. TLS is the modern successor to SSL and is used to protect HTTPS web traffic, email, and other application-layer protocols. It uses a certificate-based handshake to establish trust and negotiate encryption keys before any data is exchanged.")

add_heading(doc, "1.2 IPsec", level=2, color=(46, 84, 150))
add_para(doc, "IPsec (Internet Protocol Security) protects IP packets at the network layer by authenticating and encrypting each packet in a communication session. It is commonly used in VPN implementations to create secure tunnels between networks or between a remote user and a corporate network.")

add_heading(doc, "1.3 SSH (Secure Shell)", level=2, color=(46, 84, 150))
add_para(doc, "SSH provides a secure, encrypted channel for remote command-line access to Linux and Unix systems. It replaces older, insecure protocols such as Telnet and rlogin. SSH supports both password-based and public key authentication, with the latter being significantly more secure.")

add_heading(doc, "1.4 VPN (Virtual Private Network)", level=2, color=(46, 84, 150))
add_para(doc, "A VPN creates an encrypted tunnel between a client and a server over a public or untrusted network, allowing secure communication as if the devices were on the same private network. OpenVPN is one of the most widely used open-source VPN solutions, using SSL/TLS for key exchange and AES encryption for data protection.")

add_heading(doc, "1.5 802.1X", level=2, color=(46, 84, 150))
add_para(doc, "802.1X is a port-based network access control standard that requires devices to authenticate before being granted access to a network. It works in conjunction with a RADIUS authentication server to verify credentials and enforce access policies, making it highly effective at preventing unauthorised devices from joining a network.")

# Activity 2
add_heading(doc, "Activity 2: Configuring SSL/TLS on Apache Web Server", level=1)

add_heading(doc, "2.1 Configuration Details", level=2, color=(46, 84, 150))
add_para(doc, "Apache2 was installed on the Kali Linux VM and configured to support HTTPS using a self-signed SSL/TLS certificate. The following steps were carried out:")
for item in [
    "Apache2 was installed using apt and the SSL module was enabled using the a2enmod ssl command.",
    "The default SSL site was enabled using a2ensite default-ssl.",
    "A self-signed certificate was generated using OpenSSL with a 2048-bit RSA key, valid for 365 days. The certificate was issued to MedCore Hospital and Diagnostics Centre with the Common Name set to the Kali VM IP address 192.168.5.139.",
    "The Apache SSL configuration file (/etc/apache2/sites-available/default-ssl.conf) was updated to reference the newly generated certificate and private key.",
    "HTTP to HTTPS redirection was configured in the default virtual host file (000-default.conf) by adding a permanent redirect directive pointing to https://192.168.5.139.",
    "Apache was restarted to apply all changes."
]:
    add_bullet(doc, item)

add_heading(doc, "2.2 Verification", level=2, color=(46, 84, 150))
add_para(doc, "The SSL/TLS configuration was verified using two methods:")
for item in [
    "Terminal verification: The curl -k https://192.168.5.139 command was used to retrieve the Apache default page over HTTPS, confirming that the server was responding on port 443 with TLS encryption.",
    "Browser verification: The Kali Linux browser was used to navigate to https://192.168.5.139. After accepting the self-signed certificate warning, the Apache2 Debian Default Page loaded successfully, with the browser displaying a padlock icon confirming the HTTPS connection was active."
]:
    add_bullet(doc, item)
add_para(doc, "Both verification methods confirmed that SSL/TLS was successfully configured and that HTTP traffic was being redirected to HTTPS as intended.")

# Activity 3
add_heading(doc, "Activity 3: Implementing OpenVPN", level=1)

add_heading(doc, "3.1 Configuration Details", level=2, color=(46, 84, 150))
add_para(doc, "OpenVPN and Easy-RSA were used to set up a VPN server on the Kali Linux VM. The following steps were carried out:")
for item in [
    "OpenVPN and Easy-RSA were confirmed to be installed on the Kali VM.",
    "A new PKI directory was initialised using Easy-RSA in the ~/openvpn-ca directory.",
    "A Certificate Authority (CA) was built with the Common Name MedCore-VPN using ./easyrsa build-ca nopass.",
    "A server certificate and key were generated and signed by the CA using ./easyrsa build-server-full server nopass.",
    "Diffie-Hellman parameters were generated using ./easyrsa gen-dh to enable secure key exchange.",
    "A TLS authentication key was generated using openvpn --genkey secret ta.key for additional security.",
    "All generated files (ca.crt, server.crt, server.key, dh.pem, ta.key) were copied to /etc/openvpn/.",
    "A server configuration file was created at /etc/openvpn/server.conf, configuring OpenVPN to listen on UDP port 1194, use AES-256-CBC encryption, and assign VPN clients IP addresses in the 10.8.0.0/24 subnet.",
    "A client certificate and key were generated using ./easyrsa build-client-full client1 nopass.",
    "A client configuration file (client1.ovpn) was created with full paths to all certificate and key files."
]:
    add_bullet(doc, item)

add_heading(doc, "3.2 Verification", level=2, color=(46, 84, 150))
add_para(doc, "The VPN implementation was verified through four methods:")
for item in [
    "Server status: The OpenVPN server was started using systemctl and verified as active and running, with the tun0 tunnel interface created and assigned the IP address 10.8.0.1/24.",
    "Client connection: The client connected successfully using sudo openvpn --config ~/client1.ovpn, showing VERIFY OK for the MedCore-VPN certificate, Initialization Sequence Completed, client assigned VPN IP 10.8.0.2/24, and data channel encrypted with AES-256-GCM.",
    "Wireshark packet capture: Filtering by udp.port == 1194 revealed OpenVPN and TLSv1.3 encrypted packets on the loopback interface, confirming all VPN traffic was encrypted.",
    "Ping test: A ping to the VPN gateway 10.8.0.1 returned 4 packets transmitted, 4 received, 0% packet loss, confirming full VPN tunnel connectivity."
]:
    add_bullet(doc, item)

# Activity 4
add_heading(doc, "Activity 4: Configuring SSH for Secure Remote Access", level=1)

add_heading(doc, "4.1 Configuration Details", level=2, color=(46, 84, 150))
add_para(doc, "The OpenSSH server was configured on the Kali Linux VM with three specific security enhancements:")
for item in [
    "The SSH service was started and confirmed as active and running on the default port 22.",
    "The SSH configuration file (/etc/ssh/sshd_config) was edited to change the default port from 22 to 2222, disable root login by setting PermitRootLogin to no, and enable public key authentication by setting PubkeyAuthentication to yes.",
    "SSH was restarted to apply the configuration changes.",
    "A 4096-bit RSA key pair was generated using ssh-keygen.",
    "The public key was copied to the SSH server's authorised keys using ssh-copy-id, enabling password-less public key authentication.",
    "The SSH port was reverted to the default after testing, while PermitRootLogin no and PubkeyAuthentication yes were retained as permanent security improvements."
]:
    add_bullet(doc, item)

add_heading(doc, "4.2 Verification", level=2, color=(46, 84, 150))
add_para(doc, "The SSH configuration was verified through two methods:")
for item in [
    "Port verification: sudo ss -tlnp | grep sshd confirmed sshd listening on port 2222 on both IPv4 (0.0.0.0:2222) and IPv6 ([::]:2222).",
    "Public key authentication test: ssh -i ~/.ssh/id_rsa -p 2222 legion@127.0.0.1 connected successfully without any password prompt, confirming public key authentication was working correctly."
]:
    add_bullet(doc, item)

# Activity 5
add_heading(doc, "Activity 5: 802.1X Network Access Control with FreeRADIUS", level=1)

add_heading(doc, "5.1 Configuration Details", level=2, color=(46, 84, 150))
add_para(doc, "FreeRADIUS was installed and configured on the Kali Linux VM to simulate a RADIUS authentication server for 802.1X network access control:")
for item in [
    "FreeRADIUS was installed using apt and started using systemctl.",
    "The FreeRADIUS users file (/etc/freeradius/3.0/users) was edited to add a test user with username testuser and password testpass123.",
    "FreeRADIUS was restarted to apply the new user configuration.",
    "The radtest utility was used to simulate client authentication requests to the RADIUS server on port 1812."
]:
    add_bullet(doc, item)

add_heading(doc, "5.2 Verification", level=2, color=(46, 84, 150))
add_para(doc, "Two authentication tests were performed:")
for item in [
    "Successful authentication: radtest testuser testpass123 127.0.0.1 0 testing123 returned Access-Accept, confirming a device with valid credentials is granted network access.",
    "Failed authentication: radtest testuser wrongpassword 127.0.0.1 0 testing123 returned Access-Reject, confirming a device with invalid credentials is denied network access."
]:
    add_bullet(doc, item)
add_para(doc, "These two tests together demonstrate the core function of 802.1X — only authenticated and authorised devices are permitted to access network resources, while all others are blocked.")

# Activity 6
add_heading(doc, "Activity 6: Verification and Testing Summary", level=1)
add_para(doc, "The following summarises all verification tests performed for each protocol implemented in this lab:")
for protocol, command, result in [
    ("SSL/TLS", "curl -k https://192.168.5.139", "Apache default page returned over HTTPS — confirmed"),
    ("OpenVPN Server", "sudo systemctl status openvpn@server", "Active and running, tun0 interface created — confirmed"),
    ("OpenVPN Client", "sudo openvpn --config ~/client1.ovpn", "Initialization Sequence Completed, AES-256-GCM — confirmed"),
    ("OpenVPN Traffic", "Wireshark: udp.port == 1194", "Encrypted OpenVPN and TLSv1.3 packets captured — confirmed"),
    ("OpenVPN Tunnel", "ping -c 4 10.8.0.1", "4/4 packets received, 0% packet loss — confirmed"),
    ("SSH Port", "sudo ss -tlnp | grep sshd", "sshd listening on port 2222 — confirmed"),
    ("SSH Key Auth", "ssh -i ~/.ssh/id_rsa -p 2222 legion@127.0.0.1", "Password-less login successful — confirmed"),
    ("802.1X Accept", "radtest testuser testpass123 127.0.0.1 0 testing123", "Access-Accept received — confirmed"),
    ("802.1X Reject", "radtest testuser wrongpassword 127.0.0.1 0 testing123", "Access-Reject received — confirmed")
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{protocol}: ")
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(f"Command: {command} → Result: {result}")
    set_run(r2, size=12)

# Challenges
add_heading(doc, "Challenges Faced and Resolutions", level=1)

add_heading(doc, "6.1 OpenVPN Configuration Complexity", level=2, color=(46, 84, 150))
add_para(doc, "The most significant challenge encountered during this lab was the OpenVPN configuration. OpenVPN involves multiple interdependent components working together — the Public Key Infrastructure (PKI), the Certificate Authority, individual server and client certificates, Diffie-Hellman parameters, TLS authentication keys, and the server and client configuration files. Understanding how all of these components relate to each other and the order in which they must be configured was initially overwhelming.")
add_para(doc, "An early error occurred when the client configuration file referenced the ta.key file using a relative path, causing OpenVPN to fail with a 'Cannot pre-load keyfile' error. This was resolved by updating the client configuration to use absolute file paths for all certificate and key references. Once corrected, the VPN client connected successfully and the Initialization Sequence Completed message was displayed.")

add_heading(doc, "6.2 Real-World 802.1X Challenges", level=2, color=(46, 84, 150))
add_para(doc, "In the lab environment, 802.1X was simulated on a single machine acting as both the RADIUS server and the client. In a real-world network environment, the challenges would be significantly greater. The primary challenge would be managing the volume of authentication traffic generated by hundreds of devices — laptops, mobile phones, printers, IoT devices, and guest devices — all simultaneously authenticating through physical managed switches and wireless access points.")
add_para(doc, "Additional real-world challenges would include ensuring the RADIUS server does not become a single point of failure, managing certificates across a large number of devices, handling legacy devices that do not support 802.1X, and troubleshooting authentication failures across a complex multi-switch network topology.")

add_heading(doc, "6.3 SSL/TLS Certificate Warning", level=2, color=(46, 84, 150))
add_para(doc, "When accessing the Apache web server via the browser, a security warning was displayed because the certificate was self-signed and not issued by a trusted Certificate Authority. In a production environment, this would be resolved by obtaining a certificate from a trusted CA such as Let's Encrypt, which provides free, automatically renewable SSL/TLS certificates. For the purposes of this lab, the self-signed certificate was sufficient to demonstrate the SSL/TLS configuration.")

# Conclusion
add_heading(doc, "Conclusion", level=1)
add_para(doc, "This lab provided hands-on experience configuring and verifying four key network security protocols — SSL/TLS, OpenVPN, SSH, and 802.1X — on a Kali Linux virtual machine. Each protocol was successfully implemented and verified through live testing, command-line tools, and Wireshark packet analysis.")
add_para(doc, "The lab demonstrated that network security is not achieved through any single protocol or tool, but through a layered combination of controls — encrypting data in transit with TLS, securing remote access with hardened SSH, protecting network communications with VPN tunnelling, and controlling device access with 802.1X authentication. Each layer addresses a different threat vector, and together they form a comprehensive defence-in-depth strategy.")
add_para(doc, "The challenges encountered, particularly during the OpenVPN configuration, reinforced the importance of understanding how PKI and certificate-based authentication systems work before attempting to deploy them. Troubleshooting the ta.key path error and resolving it systematically was a valuable practical lesson in reading error output carefully and applying targeted fixes.")

# References
add_heading(doc, "References", level=1)
refs = [
    "OpenVPN Project. (2024). OpenVPN Documentation. https://openvpn.net/community-resources/reference-manual-for-openvpn-2-6/",
    "Apache Software Foundation. (2024). Apache SSL/TLS Encryption. https://httpd.apache.org/docs/2.4/ssl/ssl_howto.html",
    "OpenSSH Project. (2024). OpenSSH Manual Pages. https://www.openssh.com/manual.html",
    "FreeRADIUS Project. (2024). FreeRADIUS Documentation. https://freeradius.org/documentation/",
    "IEEE. (2010). IEEE 802.1X-2010: Port-Based Network Access Control. https://standards.ieee.org/ieee/802.1X/4645/",
    "NIST. (2020). SP 800-77 Rev. 1: Guide to IPsec VPNs. https://csrc.nist.gov/publications/detail/sp/800-77/rev-1/final",
    "Let's Encrypt. (2024). Getting Started with Free SSL/TLS Certificates. https://letsencrypt.org/getting-started/",
    "Wireshark Foundation. (2024). Wireshark User's Guide. https://www.wireshark.org/docs/wsug_html/",
    "Easy-RSA Project. (2024). Easy-RSA Documentation. https://easy-rsa.readthedocs.io/en/latest/",
    "CISA. (2023). Secure Shell (SSH) Hardening Guidance. https://www.cisa.gov/news-events/cybersecurity-advisories"
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    set_run(r, size=11)

# Save
output_path = os.path.expanduser("~/Kali-Linux-Security-Labs/INT304/Lab3/INT304_Lab3_Report.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print("Report saved successfully!")
