from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

doc = Document()

# Title
title = doc.add_heading('INT304: Network Security and Protocols', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Lab 4: Intrusion Detection Systems (IDS) and Traffic Analysis', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Student: Damian Ozimede Patrick')
doc.add_paragraph('Course: INT304 - Network Security and Protocols')
doc.add_paragraph('GitHub: github.com/damianozimede/Kali-Linux-Security-Labs')
doc.add_paragraph('Kali IP: 192.168.5.139 | Target (OWASP BWA): 192.168.5.130 | Interface: eth0')
doc.add_paragraph('Date: June 2026')

# 1. Objective
doc.add_heading('1. Objective', 2)
doc.add_paragraph(
    'The goal of this lab was to set up and configure Snort as an Intrusion Detection System (IDS) '
    'on Kali Linux, generate simulated attack traffic against the OWASP Broken Web Applications VM, '
    'capture and analyse that traffic using Wireshark, and evaluate how effectively Snort was able '
    'to detect and raise alerts for the attack in real time.'
)

# 2. Environment Setup
doc.add_heading('2. Environment Setup', 2)
doc.add_paragraph('The lab was carried out using the following environment:')
env = doc.add_paragraph(style='List Bullet')
env.add_run('Attacker Machine: Kali Linux (legion@pastor), IP address 192.168.5.139, interface eth0')
env2 = doc.add_paragraph(style='List Bullet')
env2.add_run('Target Machine: OWASP Broken Web Applications VM, IP address 192.168.5.130')
env3 = doc.add_paragraph(style='List Bullet')
env3.add_run('IDS: Snort 3.12.2.0 installed on Kali Linux')
env4 = doc.add_paragraph(style='List Bullet')
env4.add_run('Traffic capture tool: Wireshark monitoring eth0')
env5 = doc.add_paragraph(style='List Bullet')
env5.add_run('Attack generation tool: hping3 used to simulate a SYN flood')

# 3. Installing Snort
doc.add_heading('3. Installing and Verifying Snort', 2)
doc.add_paragraph(
    'Snort was not pre-installed on the machine, so it was installed using the Kali package manager. '
    'Running sudo apt install snort pulled down Snort version 3.12.2.0 along with all its dependencies '
    'including libdaq3, snort-rules-default, rsyslog, and several supporting libraries. The installation '
    'completed without errors. Snort was confirmed working by running snort --version, which displayed '
    'the full version string and library details. This version is Snort 3, which uses a Lua-based '
    'configuration format rather than the older snort.conf format used in Snort 2.'
)
doc.add_paragraph('[SCREENSHOT 1: Terminal output of snort --version showing Snort++ 3.12.2.0]')

# 4. Configuring Snort
doc.add_heading('4. Configuring Snort', 2)
doc.add_heading('4.1 Setting HOME_NET', 3)
doc.add_paragraph(
    'In Snort 3, the main configuration file is /etc/snort/snort.lua. The HOME_NET variable tells '
    'Snort which network to treat as the protected internal network. By default it was set to any. '
    'This was changed to 192.168.5.130/32 so that Snort would specifically monitor traffic directed '
    'at the OWASP VM. Line 24 was updated to read:'
)
doc.add_paragraph("HOME_NET = '192.168.5.130/32'")
doc.add_paragraph('[SCREENSHOT 2: grep output confirming HOME_NET set to 192.168.5.130/32 on line 24]')

doc.add_heading('4.2 Adding a Custom SYN Flood Detection Rule', 3)
doc.add_paragraph(
    'Snort 3 does not include a SYN flood detection rule in its default rule set, so a custom rule '
    'was written and added directly into the ips block of snort.lua. The rule added was:'
)
doc.add_paragraph(
    'alert tcp any any -> 192.168.5.130 80 (msg:"SYN Flood Detected"; flags:S; flow:stateless; '
    'detection_filter:track by_src, count 100, seconds 1; sid:1000001; rev:1;)'
)
doc.add_paragraph(
    'This rule watches for TCP packets with only the SYN flag set heading to port 80 on the OWASP VM. '
    'The detection_filter keyword means the alert only fires once a source sends more than 100 SYN '
    'packets within one second, which is the characteristic behaviour of a SYN flood. The flow:stateless '
    'option was included because SYN flood packets do not complete a full TCP handshake so Snort cannot '
    'track them as a normal stateful flow.'
)

# 5. Starting Snort
doc.add_heading('5. Starting Snort in IDS Mode', 2)
doc.add_paragraph('A log directory was created first:')
doc.add_paragraph('sudo mkdir -p /tmp/snort_logs')
doc.add_paragraph('Snort was then started in passive IDS mode using:')
doc.add_paragraph('sudo snort -c /etc/snort/snort.lua -i eth0 -A alert_fast -l /tmp/snort_logs')
doc.add_paragraph(
    'Once Snort loaded and initialised all its modules it displayed Commencing packet processing '
    'followed by ++ [0] eth0, confirming it was actively monitoring the network interface. Snort '
    'loaded 219 rules in total including the custom SYN flood rule and was ready to inspect live traffic.'
)
doc.add_paragraph('[SCREENSHOT 3: Snort terminal showing Commencing packet processing and ++ [0] eth0]')

# 6. Generating Traffic
doc.add_heading('6. Generating Attack Traffic with hping3', 2)
doc.add_paragraph(
    'With Snort running in one terminal, hping3 was used in a second terminal to simulate a SYN flood '
    'attack against the OWASP VM on port 80:'
)
doc.add_paragraph('sudo hping3 -S -p 80 --flood 192.168.5.130')
doc.add_paragraph(
    'The -S flag sets the SYN bit on every packet, -p 80 targets the web server port, and --flood '
    'sends packets as fast as the machine can generate them with no delay. In the first run hping3 '
    'transmitted 599,341 packets and in the second run it sent 817,013 packets before being stopped. '
    'In both cases 100% packet loss was recorded on the return side, which is expected from a SYN '
    'flood since no real TCP connection is being completed.'
)
doc.add_paragraph('[SCREENSHOT 4: hping3 terminal showing packets transmitted and 100% packet loss]')

# 7. Wireshark
doc.add_heading('7. Monitoring Traffic with Wireshark', 2)
doc.add_paragraph(
    'Wireshark was opened and configured to capture on eth0 with a display filter of tcp. During '
    'the flood Wireshark captured over 196,000 TCP packets. The packet list showed a continuous '
    'stream of packets flowing from 192.168.5.139 to 192.168.5.130 on port 80. The Info column '
    'showed [SYN] for outgoing packets and a mixture of [SYN, ACK] and [RST] responses from the '
    'target. The SYN, ACK responses show the OWASP VM attempting to respond to connection requests, '
    'while the RST packets represent Kali resetting those half-open connections. This pattern is the '
    'hallmark of a SYN flood and is clearly visible in the capture.'
)
doc.add_paragraph('[SCREENSHOT 5: Wireshark showing mass TCP SYN packets from 192.168.5.139 to 192.168.5.130:80]')

# 8. Snort Alerts
doc.add_heading('8. Analysing Snort Alerts', 2)
doc.add_paragraph(
    'As soon as the SYN flood started, the Snort terminal began filling up with alerts. '
    'Each alert line appeared in the following format:'
)
doc.add_paragraph(
    '06/18-19:28:28.595715 [**] [1:1000001:1] "SYN Flood Detected" [**] [Priority: 0] '
    '{TCP} 192.168.5.139:2361 -> 192.168.5.130:80'
)
doc.add_paragraph('Breaking down what each part of the alert means:')
a1 = doc.add_paragraph(style='List Bullet')
a1.add_run('06/18-19:28:28 is the timestamp showing when the alert fired')
a2 = doc.add_paragraph(style='List Bullet')
a2.add_run('[1:1000001:1] is the rule identifier matching the custom rule SID 1000001')
a3 = doc.add_paragraph(style='List Bullet')
a3.add_run('"SYN Flood Detected" is the message defined in the rule msg field')
a4 = doc.add_paragraph(style='List Bullet')
a4.add_run('[Priority: 0] is the default alert priority since none was explicitly set')
a5 = doc.add_paragraph(style='List Bullet')
a5.add_run('{TCP} confirms the protocol being matched')
a6 = doc.add_paragraph(style='List Bullet')
a6.add_run('192.168.5.139:xxxx -> 192.168.5.130:80 shows Kali sending to the OWASP VM on port 80')
doc.add_paragraph(
    'Hundreds of these alerts were generated per second throughout the flood, confirming that Snort '
    'was successfully detecting the attack in real time. The alerts fired continuously for as long as '
    'hping3 was running and stopped immediately once the flood was halted.'
)
doc.add_paragraph('[SCREENSHOT 6: Snort console showing rapid SYN Flood Detected alerts firing in real time]')

# 9. Analysis
doc.add_heading('9. Analysis and Findings', 2)
doc.add_paragraph(
    'The lab demonstrated that Snort 3, when configured with a suitable detection rule, is capable '
    'of identifying a SYN flood attack in real time. The custom rule worked as intended, triggering '
    'only when the packet rate from a single source exceeded 100 SYN packets per second to port 80. '
    'This threshold prevented false positives from normal web traffic while still catching the flood '
    'the moment it began.'
)
doc.add_paragraph(
    'An important finding was that Snort 3 does not ship with a built-in SYN flood rule. The default '
    'rule set focuses on file identification and protocol inspection rather than volume-based '
    'denial-of-service attacks. In a real-world deployment a security team would need to write or '
    'import additional rules to cover this type of attack.'
)
doc.add_paragraph(
    'Wireshark provided visual confirmation of everything Snort was detecting. The volume of packets '
    'in the capture, all heading to the same destination port from the same source, made the attack '
    'pattern immediately obvious. Used together the two tools offered both a high-level view of what '
    'was happening and a detailed per-packet breakdown useful for forensic investigation.'
)
doc.add_paragraph(
    'One limitation noted was that Snort was running in passive pcap mode, meaning it could detect '
    'and alert on attacks but could not block them. For active prevention Snort would need to be '
    'configured in inline mode using the afpacket DAQ module to drop malicious packets before they '
    'reach the target.'
)

# 10. Conclusion
doc.add_heading('10. Conclusion', 2)
doc.add_paragraph(
    'This lab provided practical hands-on experience with setting up and using an Intrusion Detection '
    'System in a controlled environment. Snort 3 was successfully installed, configured with a custom '
    'SYN flood detection rule, and run in IDS mode on Kali Linux. The hping3 tool was used to generate '
    'a realistic SYN flood attack against the OWASP VM, and both Snort and Wireshark were used '
    'simultaneously to observe and confirm the attack. Snort raised alerts correctly every time the '
    'flood threshold was exceeded, demonstrating that even a simple custom rule can be highly effective '
    'at catching volume-based network attacks when the detection logic matches the traffic pattern.'
)

# Appendix
doc.add_heading('Appendix: Commands Used', 2)
commands = [
    'sudo apt install snort',
    'snort --version',
    'grep -n "HOME_NET" /etc/snort/snort.lua',
    'sudo nano /etc/snort/snort.lua',
    'sudo mkdir -p /tmp/snort_logs',
    'sudo snort -c /etc/snort/snort.lua -i eth0 -A alert_fast -l /tmp/snort_logs',
    'sudo hping3 -S -p 80 --flood 192.168.5.130',
    'wireshark &',
    'cat /tmp/snort_logs/alert_fast.txt',
]
for cmd in commands:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd)

# Save
doc.save('/home/legion/Kali-Linux-Security-Labs/INT304/Lab4/INT304_Lab4_Report.docx')
print("Report saved successfully!")
