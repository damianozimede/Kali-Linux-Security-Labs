from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)

def set_run(run, bold=False, size=12, color=None, font="Arial"):
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 56, 100)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, size=12, bold=False, color=None, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    return p

def add_bold_label(doc, label, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r1 = p.add_run(label + " ")
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(text)
    set_run(r2, bold=False, size=12)
    return p

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after  = Pt(12)
r = p.add_run("INT304: Network Security and Protocols")
set_run(r, bold=True, size=18, color=(31, 56, 100))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(48)
r = p.add_run("Lab 1: Applied Network Security Analysis")
set_run(r, bold=True, size=15, color=(46, 84, 150))

for line in ["Student: Damian Ozimede Patrick",
             "Course: INT304 — Network Security and Protocols",
             "Institution: International Cybersecurity & Digital Forensic Academy"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(line)
    set_run(r, size=12)

doc.add_page_break()

# Executive Summary
add_heading(doc, "Executive Summary", level=1)
add_para(doc, "This report presents a structured incident analysis of the May 2021 Colonial Pipeline ransomware attack, one of the most consequential cyber incidents in recent history. The attack, carried out by the cybercriminal group DarkSide, resulted in the shutdown of one of the United States' largest fuel pipeline systems, causing widespread fuel shortages across the Southeastern states and forcing the company to pay approximately 4.4 million USD in ransom.")
add_para(doc, "The analysis conducted in this report covers three key areas: the identification of the initial attack vector and the vulnerabilities that made the breach possible; the role that weak or absent network security protocols played in allowing the attack to succeed and escalate; and a three-point, research-backed mitigation strategy designed to prevent a similar incident from occurring in the future.")
add_para(doc, "The findings reveal that the root cause of the breach was a fundamental failure in identity verification — specifically the absence of Multi-Factor Authentication (MFA) on a VPN account whose credentials had been previously compromised. Compounding this was the insufficient segmentation between the company's IT and Operational Technology (OT) networks, which created enough uncertainty during the incident to force a precautionary shutdown of pipeline operations. The recommended mitigations centre on enforcing MFA, implementing robust network segmentation, and adopting proactive dark web credential monitoring paired with strong password governance.")

# Part 1
add_heading(doc, "Part 1: Attack Vector and Vulnerability Identification", level=1)

add_heading(doc, "1.1 The Initial Attack Vector", level=2, color=(46, 84, 150))
add_para(doc, "The DarkSide group gained unauthorised access to Colonial Pipeline's network through a compromised Virtual Private Network (VPN) account. Investigators later confirmed that the attackers used a single set of leaked credentials — a username and password — to authenticate into the VPN and establish a foothold within the company's internal network. The password associated with the account was found in a collection of previously leaked credentials available on the dark web, indicating that it had been exposed through an earlier, unrelated data breach on a separate platform.")
add_para(doc, "This means the attackers did not need to carry out any sophisticated technical exploit to get in. They simply obtained valid credentials from an external source and used them to log in like any legitimate user would. The VPN, by design, trusted the credentials presented to it. Once authenticated, DarkSide moved laterally through the business network, exfiltrated approximately 100GB of data, and ultimately deployed their ransomware payload.")

add_heading(doc, "1.2 The Key Vulnerability", level=2, color=(46, 84, 150))
add_para(doc, "The central vulnerability that enabled this breach was weak, single-factor authentication on the VPN account. The account in question was protected by nothing more than a username and password combination. No secondary verification mechanism was in place. This meant that possession of the credentials alone was sufficient for full network access, with no additional identity verification required.")
add_para(doc, "In a properly hardened environment, a stolen password would represent a limited risk — it would be only one piece of a multi-step authentication process. In Colonial Pipeline's case, it was the only piece required. This single point of failure allowed what should have been an inconclusive credential leak to become a full network intrusion.")
add_para(doc, "Secondary contributing factors included the apparent failure to monitor whether employee credentials had appeared in known breach databases, and the apparent lack of regular forced password rotation policies. Had any of these controls been in place, the leaked credentials may have been identified and changed before they were ever used.")

add_heading(doc, "1.3 The Threat Type", level=2, color=(46, 84, 150))
add_para(doc, "This attack is categorised as a Ransomware attack. Once DarkSide had gained access to Colonial Pipeline's business network and completed their data exfiltration, they deployed ransomware — malicious software that encrypts the victim's files and systems, rendering them inaccessible until a decryption key is provided. The decryption key was offered by DarkSide in exchange for a ransom payment, which Colonial ultimately paid in Bitcoin.")
add_para(doc, "The defining characteristics that justify this categorisation are two-fold: first, the deliberate encryption of the victim's data and systems by the attackers to cause disruption; and second, the explicit demand for financial payment in exchange for restoring access. This fits squarely within the definition of ransomware. It is worth noting that DarkSide operated as a Ransomware-as-a-Service (RaaS) group, meaning they developed and maintained their ransomware toolkit and made it available to other criminal affiliates, taking a percentage of any ransom received.")

# Part 2
add_heading(doc, "Part 2: Protocol Failure and Relevance", level=1)

add_heading(doc, "2.1 Authentication Protocol Failure", level=2, color=(46, 84, 150))
add_para(doc, "The most significant protocol failure in this incident was the absence of Multi-Factor Authentication (MFA) on the compromised VPN account. MFA is an authentication mechanism that requires a user to verify their identity through two or more independent factors before access is granted. These factors typically fall into three categories: something the user knows (such as a password), something the user has (such as a mobile authenticator app or a hardware token), and something the user is (such as a biometric like a fingerprint).")
add_para(doc, "In the Colonial Pipeline incident, authentication relied solely on the first category — a username and password. When those credentials were stolen and used by DarkSide, the VPN system had no means of detecting that the person logging in was not the legitimate account holder. There was no second factor to challenge the attacker.")
add_para(doc, "Had MFA been enforced, the attack would very likely have failed at this stage. Even with valid credentials in hand, the attackers would have been prompted for a second factor — a one-time passcode sent to the real account holder's phone, or an approval request through an authenticator application. Since DarkSide had no access to the account holder's physical device or authenticator, they would have been unable to complete the login process. MFA effectively breaks the assumption that a stolen password equals stolen access, which is precisely why it is considered a foundational security control by organisations such as NIST, CISA, and the UK's NCSC.")

add_heading(doc, "2.2 Network Segmentation", level=2, color=(46, 84, 150))
add_para(doc, "Network segmentation is the practice of dividing a computer network into distinct sub-networks or zones, each with controlled access points between them. In the context of critical infrastructure organisations like Colonial Pipeline, the most important segmentation is between the IT network — which handles business operations such as billing, email, and corporate systems — and the OT network, which controls physical industrial processes such as the pipeline itself.")
add_para(doc, "In this incident, the ransomware attack directly targeted Colonial's IT network. The OT network — which controlled the actual pipeline operations — was not directly compromised. However, Colonial made the decision to proactively shut down the OT network as a precautionary measure. The reason for this decision is telling: the company could not be sufficiently confident that the breach was fully contained to the IT side. That uncertainty itself is a symptom of inadequate segmentation.")
add_para(doc, "If robust network segmentation had been in place — with clear, enforced boundaries between the IT and OT environments, supported by firewalls, air gaps, or strict access control lists — Colonial's security team would have been able to verify quickly and with confidence that the OT network was isolated and unaffected. Instead, the ambiguity around the breach's reach forced them to shut down pipeline operations as a precaution, which was the direct cause of the fuel shortages that followed. The real-world impact of the attack was therefore not caused by the ransomware reaching the pipeline systems — it was caused by the uncertainty that poor segmentation created.")

add_heading(doc, "2.3 Encryption Protocols: SSL/TLS and IPSec", level=2, color=(46, 84, 150))
add_para(doc, "Encryption protocols such as SSL/TLS and IPSec were not relevant to preventing the initial breach in this incident. VPN connections are by their nature encrypted — the VPN technology employed by Colonial Pipeline would have used encryption to protect data transmitted between remote users and the corporate network. The communication channel itself was therefore already secured.")
add_para(doc, "The attack did not succeed by intercepting or decrypting network traffic. DarkSide did not perform a Man-in-the-Middle attack or eavesdrop on communications. They succeeded by authenticating as a legitimate user with valid credentials. From the system's perspective, the encrypted VPN session they established was indistinguishable from a session initiated by the real account holder.")
add_para(doc, "This highlights an important principle in network security: encryption protects the confidentiality and integrity of data in transit, but it does not verify the identity of the person initiating the connection. That is the role of authentication. In this case, the authentication mechanism was the weak link — not the encryption. Strengthening SSL/TLS or IPSec configurations would have had no effect on the outcome of this attack.")

# Part 3
add_heading(doc, "Part 3: Mitigation Strategy", level=1)
add_para(doc, "The following three-point mitigation strategy is proposed for Colonial Pipeline to reduce the risk of a similar incident occurring in the future. Each point addresses a specific weakness identified in the 2021 attack and is grounded in established industry best practices.")

add_heading(doc, "3.1 Enforce Multi-Factor Authentication on All Remote Access", level=2, color=(46, 84, 150))
add_bold_label(doc, "Security Control:", "Mandatory MFA for all VPN accounts, remote desktop access, and privileged user accounts across the organisation.")
add_bold_label(doc, "How It Addresses the Vulnerability:", "The 2021 attack succeeded because a stolen password was all that was required to gain full network access. MFA directly eliminates this single point of failure. Even if an attacker obtains valid credentials through a dark web leak, a phishing campaign, or any other method, they would still be unable to complete authentication without the second factor — which remains in the possession of the legitimate account holder.")
add_bold_label(doc, "Industry Justification:", "CISA consistently lists MFA as one of the most impactful security controls an organisation can implement. Microsoft's research indicates that MFA blocks over 99.9% of account compromise attacks. NIST Special Publication 800-63B also outlines MFA as a requirement for any system handling sensitive or critical data. The fact that this control was absent from a VPN account at a national infrastructure operator represents a critical governance failure that must not be repeated.")

add_heading(doc, "3.2 Implement Robust IT/OT Network Segmentation", level=2, color=(46, 84, 150))
add_bold_label(doc, "Security Control:", "Enforce strict, verifiable separation between IT and OT networks using firewalls, demilitarised zones (DMZs), and where operationally feasible, air-gapped systems for the most critical OT components.")
add_bold_label(doc, "How It Addresses the Vulnerability:", "As established in Part 2, the fuel shortage caused by the 2021 attack was a direct consequence of Colonial being unable to confirm with certainty whether the OT network had been affected. Proper segmentation would have provided that certainty. With clear, enforced boundaries between IT and OT environments, any breach on the IT side could be contained and investigated without triggering a precautionary shutdown of pipeline operations.")
add_bold_label(doc, "Industry Justification:", "The IEC 62443 standard, which governs industrial cybersecurity, explicitly mandates the segmentation of IT and OT networks as a core security requirement for critical infrastructure operators. NIST SP 800-82 similarly recommends network segmentation as a primary defence strategy. For a company operating critical national infrastructure, this is not optional — it is a baseline requirement.")

add_heading(doc, "3.3 Deploy Dark Web Credential Monitoring and Enforce a Strong Password Policy", level=2, color=(46, 84, 150))
add_bold_label(doc, "Security Control:", "Implement a continuous dark web monitoring solution to detect when employee credentials appear in breach databases, combined with a mandatory password rotation policy, prohibition on password reuse across platforms, and the use of an enterprise password manager.")
add_bold_label(doc, "How It Addresses the Vulnerability:", "The credentials used in the Colonial Pipeline attack were already circulating on the dark web before the attack occurred. A proactive credential monitoring solution would have flagged this exposure and triggered an immediate forced password reset, closing the window of opportunity before DarkSide could exploit it. Additionally, had the organisation enforced a no-password-reuse policy and regularly rotated credentials, the likelihood of a previously leaked password remaining valid and usable would have been significantly reduced.")
add_bold_label(doc, "Industry Justification:", "Services such as Have I Been Pwned, SpyCloud, and enterprise solutions from vendors like CrowdStrike and Recorded Future provide automated dark web monitoring tailored to corporate environments. NIST SP 800-63B also advises organisations to check user-chosen passwords against lists of known compromised credentials at the point of creation or reset.")

# Conclusion
add_heading(doc, "Conclusion", level=1)
add_para(doc, "The 2021 Colonial Pipeline ransomware attack is a defining case study in modern network security failure. What made it particularly significant was not the sophistication of the attack — in technical terms, it was relatively straightforward. What made it catastrophic was the absence of basic, well-understood security controls that should have been standard practice for an operator of critical national infrastructure.")
add_para(doc, "A single compromised password, obtained from the dark web, was enough to bring down fuel supplies across multiple U.S. states for nearly a week and cost the company millions of dollars. The three mitigations proposed in this report — mandatory MFA, proper IT/OT network segmentation, and dark web credential monitoring with strong password governance — are not experimental or untested. They are foundational controls that, had they been implemented before May 2021, would in all likelihood have prevented the attack from succeeding.")
add_para(doc, "This incident serves as a reminder that cybersecurity failures at critical infrastructure organisations carry consequences that extend well beyond the digital realm. The lessons drawn from Colonial Pipeline must be applied not only by energy companies, but by any organisation whose network security posture has the potential to affect public safety and national stability.")

# References
add_heading(doc, "References", level=1)
refs = [
    "Cimpanu, C. (2021). Colonial Pipeline ransomware attack: Everything you need to know. ZDNet. https://www.zdnet.com/article/colonial-pipeline-ransomware-attack/",
    "CISA. (2021). Alert (AA21-131A): DarkSide Ransomware: Best Practices for Preventing Business Disruption from Ransomware Attacks. https://www.cisa.gov/news-events/cybersecurity-advisories/aa21-131a",
    "CISA. (2022). Implementing Multi-Factor Authentication. https://www.cisa.gov/mfa",
    "Turton, W., & Mehrotra, K. (2021). Hackers Breached Colonial Pipeline Using Compromised Password. Bloomberg. https://www.bloomberg.com/news/articles/2021-06-04/hackers-breached-colonial-pipeline-using-compromised-password",
    "NIST. (2017). SP 800-63B: Digital Identity Guidelines — Authentication and Lifecycle Management. https://pages.nist.gov/800-63-3/sp800-63b.html",
    "NIST. (2015). SP 800-82 Rev. 2: Guide to Industrial Control Systems (ICS) Security. https://csrc.nist.gov/publications/detail/sp/800-82/rev-2/final",
    "IEC. (2021). IEC 62443: Security for Industrial Automation and Control Systems. https://www.iec.ch/iec62443",
    "Microsoft Security. (2019). One simple action you can take to prevent 99.9 percent of attacks on your accounts. https://www.microsoft.com/en-us/security/blog/2019/08/20/one-simple-action-you-can-take-to-prevent-99-9-percent-of-account-attacks/",
    "SpyCloud. (2023). Enterprise Dark Web Monitoring. https://spycloud.com",
    "Have I Been Pwned. (2024). About Have I Been Pwned. https://haveibeenpwned.com/About"
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    set_run(r, size=11)

# Save
output_path = os.path.expanduser("~/Kali-Linux-Security-Labs/INT304/Lab1/INT304_Lab1_Report.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print("Report saved successfully!")
