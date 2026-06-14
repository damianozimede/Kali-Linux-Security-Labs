from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)

def set_run(run, bold=False, size=12, color=None, font="Arial"):
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 56, 100)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, size=12, bold=False, color=None, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=size)
    return p

def add_bold_label(doc, label, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r1 = p.add_run(label + " ")
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(text)
    set_run(r2, bold=False, size=12)
    return p

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("INT304: Network Security and Protocols")
set_run(r, bold=True, size=18, color=(31, 56, 100))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(48)
r = p.add_run("Lab 2: Network Security Policies and Risk Management")
set_run(r, bold=True, size=15, color=(46, 84, 150))

for line in ["Student: Damian Ozimede Patrick",
             "Course: INT304 — Network Security and Protocols",
             "Institution: International Cybersecurity & Digital Forensic Academy"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(line)
    set_run(r, size=12)

doc.add_page_break()

# Executive Summary
add_heading(doc, "Executive Summary", level=1)
add_para(doc, "This report documents the outcome of a structured network security policy and risk management exercise conducted for a hypothetical healthcare organisation, MedCore Hospital and Diagnostics Centre — a mid-sized private hospital based in Lagos, Nigeria. The exercise covers four key activities: an introduction to the core components of an effective network security policy; a comprehensive risk assessment of MedCore's key digital assets; the development of security policies derived from the risk assessment findings; and a reflective discussion on how organisations can adapt and maintain their security posture over time.")
add_para(doc, "The risk assessment identified five critical asset categories — patient medical records, payment and billing data, networked medical equipment, staff credentials and HR data, and hospital management software — and evaluated the threats, vulnerabilities, and potential impacts associated with each. Four of the five asset categories were rated as High risk, underscoring the significant cybersecurity exposure faced by healthcare organisations. Security policies were subsequently developed to address the identified weaknesses, covering access control, data protection, incident response, and staff training and awareness.")
add_para(doc, "The findings and recommendations in this report are grounded in established industry frameworks including NIST, ISO 27001, NDPR, and PCI-DSS, and reflect best practices applicable to healthcare institutions operating in both local and international regulatory environments.")

# Activity 1
add_heading(doc, "Activity 1: Introduction to Network Security Policies", level=1)
add_para(doc, "An effective network security policy provides an organisation with a formal, documented framework for protecting its digital assets, defining acceptable use of its systems, and establishing clear procedures for responding to security incidents. For a healthcare organisation like MedCore Hospital and Diagnostics Centre, such a policy is not merely a best practice — it is a legal and operational necessity. The following four components form the foundation of MedCore's network security policy.")

add_heading(doc, "1.1 Purpose and Scope", level=2, color=(46, 84, 150))
add_para(doc, "The primary purpose of MedCore's network security policy is to protect patient data, hospital systems, and connected medical equipment from unauthorised access, disruption, theft, or misuse. As a healthcare provider, MedCore handles some of the most sensitive categories of personal data — patient medical histories, diagnostic results, prescription records, and payment information. A breach of any of these assets carries serious consequences for patient privacy, hospital operations, and the organisation's legal standing.")
add_para(doc, "The scope of this policy extends to all individuals who access MedCore's network and systems, including full-time medical and administrative staff, part-time and contract employees, visiting consultants, third-party vendors, and any other parties granted access to hospital infrastructure. It covers all devices connected to the hospital network, whether hospital-owned or personally owned, and applies to both on-site and remote access scenarios.")

add_heading(doc, "1.2 User Responsibilities", level=2, color=(46, 84, 150))
add_para(doc, "User responsibilities define the rules and regulations governing every individual who accesses MedCore's network. Every user — whether a senior consultant, a junior nurse, an administrative officer, or a contracted IT technician — is expected to comply with these responsibilities as a condition of their network access.")
add_para(doc, "Key user responsibilities include: never sharing login credentials with colleagues or third parties; only accessing systems and data relevant to their specific role; reporting any suspicious activity, unusual system behaviour, or potential security incident to the IT security team immediately; not connecting personal or unauthorised devices to the hospital network without prior approval; and completing all mandatory cybersecurity awareness training as required by the hospital.")

add_heading(doc, "1.3 Incident Response Plan", level=2, color=(46, 84, 150))
add_para(doc, "MedCore's incident response plan is a formally documented, hierarchical procedure that defines exactly what must be done in the event of a security breach, cyberattack, or any other incident that threatens the integrity or availability of hospital systems and data. The plan ensures that all stakeholders know their roles and responsibilities when an incident occurs, and that the hospital's response is coordinated, swift, and effective.")
add_para(doc, "The incident response plan follows a structured sequence: detection and identification of the threat; escalation to the IT security team and relevant management; immediate containment to prevent further spread; impact assessment to determine the extent of the damage; eradication of the threat from all affected systems; system recovery and restoration of normal operations; notification of all relevant parties including patients, regulators, and law enforcement where required; post-incident monitoring for signs of recurrence; full documentation of the incident; and a formal review to identify lessons learned and improve future responses.")

add_heading(doc, "1.4 Compliance", level=2, color=(46, 84, 150))
add_para(doc, "Compliance means that MedCore must adhere to all stipulated protocols and standards set by recognised medical and data protection bodies, both within Nigeria and internationally. Failure to comply with these standards exposes the hospital to significant legal, financial, and reputational consequences, including regulatory fines, sanctions, and potential loss of operating licences.")
add_para(doc, "The key compliance frameworks relevant to MedCore include the Nigeria Data Protection Regulation (NDPR), which governs the collection, storage, and processing of personal data in Nigeria; the Health Insurance Portability and Accountability Act (HIPAA), an internationally referenced standard for the protection of healthcare data; ISO/IEC 27001, the global standard for information security management systems; and PCI-DSS (Payment Card Industry Data Security Standard), which applies to the handling of payment and billing data.")

# Activity 2
add_heading(doc, "Activity 2: Risk Assessment for MedCore Hospital and Diagnostics Centre", level=1)
add_para(doc, "The following risk assessment evaluates the key digital assets of MedCore Hospital and Diagnostics Centre, identifying the threats and vulnerabilities associated with each, analysing the potential impact of a successful attack, and assigning a risk rating based on the likelihood and severity of that impact.")

add_heading(doc, "2.1 Asset 1: Patient Medical Records (EHR)", level=2, color=(46, 84, 150))
add_bold_label(doc, "Threats:", "Ransomware attacks targeting and encrypting patient records; unauthorised access by internal staff or external hackers; data theft for resale on the dark web.")
add_bold_label(doc, "Vulnerabilities:", "Weak or absent access controls on the EHR system; lack of encryption for stored patient records; absence of audit logs to track who accessed or modified records.")
add_bold_label(doc, "Impact:", "A breach would expose highly sensitive personal and medical information, resulting in serious privacy violations, potential blackmail of patients, legal action under the NDPR, and a significant loss of patient trust.")
add_bold_label(doc, "Risk Rating:", "HIGH")

add_heading(doc, "2.2 Asset 2: Payment and Billing Data", level=2, color=(46, 84, 150))
add_bold_label(doc, "Threats:", "Financial fraud and theft of payment information; ransomware encrypting billing systems; unauthorised access to financial records.")
add_bold_label(doc, "Vulnerabilities:", "Weak encryption of stored and transmitted payment data; outdated payment processing systems; poor access controls and insecure data storage practices.")
add_bold_label(doc, "Impact:", "Compromise would result in direct financial losses, fraudulent transactions, regulatory fines for non-compliance with PCI-DSS, significant reputational damage, and erosion of patient trust.")
add_bold_label(doc, "Risk Rating:", "HIGH")

add_heading(doc, "2.3 Asset 3: Medical Equipment Connected to the Network", level=2, color=(46, 84, 150))
add_bold_label(doc, "Threats:", "Malware infection disrupting device functionality; unauthorised remote control of medical devices; exploitation of unpatched firmware vulnerabilities.")
add_bold_label(doc, "Vulnerabilities:", "Default factory passwords left unchanged; outdated firmware with known security flaws; lack of network segmentation; poor device authentication mechanisms.")
add_bold_label(doc, "Impact:", "An attack could directly endanger patient safety, cause life-critical device malfunctions, delay treatments, and result in prolonged operational downtime.")
add_bold_label(doc, "Risk Rating:", "HIGH")

add_heading(doc, "2.4 Asset 4: Staff Login Credentials and HR Data", level=2, color=(46, 84, 150))
add_bold_label(doc, "Threats:", "Credential theft through phishing attacks; brute-force attacks on staff accounts; insider misuse of privileged access.")
add_bold_label(doc, "Vulnerabilities:", "Weak passwords and widespread password reuse; absence of Multi-Factor Authentication (MFA); poor access management and privilege oversight.")
add_bold_label(doc, "Impact:", "Compromised credentials could enable unauthorised access, facilitate identity theft, enable payroll fraud, and serve as an entry point for broader attacks across MedCore's network.")
add_bold_label(doc, "Risk Rating:", "MEDIUM")

add_heading(doc, "2.5 Asset 5: Hospital Management Software", level=2, color=(46, 84, 150))
add_bold_label(doc, "Threats:", "Exploitation of unpatched software vulnerabilities; ransomware infection; unauthorised access leading to data corruption or manipulation.")
add_bold_label(doc, "Vulnerabilities:", "Unpatched software with known security flaws; weak authentication mechanisms; insecure system configurations; lack of regular security testing.")
add_bold_label(doc, "Impact:", "Compromise would disrupt core operational functions, prevent access to patient records, cause significant delays, and result in financial and reputational damage.")
add_bold_label(doc, "Risk Rating:", "HIGH")

# Activity 3
add_heading(doc, "Activity 3: Security Policy Development", level=1)
add_para(doc, "The following security policies have been developed based directly on the vulnerabilities and risks identified in the risk assessment above. Each policy addresses a specific area of weakness and is designed to reduce MedCore's overall risk exposure.")

add_heading(doc, "3.1 Access Control Policy", level=2, color=(46, 84, 150))
add_bold_label(doc, "Policy Statement:", "Access to MedCore's network, systems, and data shall be strictly controlled and granted only to authorised individuals on a need-to-know, role-based basis.")
for item in [
    "Multi-Factor Authentication (MFA) must be enforced for all staff accessing MedCore systems, particularly for remote access and administrative accounts.",
    "Strong password requirements must be enforced, with a minimum length of twelve characters including uppercase, lowercase, numbers, and special characters.",
    "Role-Based Access Control (RBAC) must be implemented — each staff member may only access systems and data relevant to their specific role.",
    "All remote access must be conducted through a secure, encrypted VPN connection.",
    "User accounts must be reviewed quarterly, and access revoked immediately upon staff departure or role change."
]:
    add_bullet(doc, item)
add_bold_label(doc, "Strengths and Weaknesses:", "The layered approach combining MFA, RBAC, and strong passwords creates multiple barriers for attackers. The primary weakness is staff resistance to MFA and complex passwords, which must be addressed through the training programme in Section 3.4.")

add_heading(doc, "3.2 Data Protection Policy", level=2, color=(46, 84, 150))
add_bold_label(doc, "Policy Statement:", "All sensitive data held or processed by MedCore must be protected through appropriate technical and administrative controls, both at rest and in transit.")
for item in [
    "All patient records and sensitive data must be encrypted using AES-256 at rest and TLS 1.2 or higher in transit.",
    "Regular automated backups of all critical data must be performed daily and stored securely offsite or in the cloud.",
    "Data no longer required must be securely deleted using approved data sanitisation methods.",
    "Payment and billing data must be handled in strict accordance with PCI-DSS standards.",
    "Transfer of sensitive data to personal USB drives or unauthorised external devices is strictly prohibited."
]:
    add_bullet(doc, item)
add_bold_label(doc, "Strengths and Weaknesses:", "Encryption ensures stolen data remains unusable. The weakness is the operational overhead of maintaining encrypted backups and ensuring third-party vendor compliance, requiring ongoing monitoring and resources.")

add_heading(doc, "3.3 Incident Response Procedures", level=2, color=(46, 84, 150))
add_bold_label(doc, "Policy Statement:", "MedCore must maintain a formally documented, tested, and regularly updated incident response procedure to ensure a coordinated and effective response to any security incident.")
for item in [
    "Detection and Identification — The threat is detected and its nature and scope identified.",
    "Escalation — The incident is immediately reported to the IT Security team and hospital management.",
    "Containment — Affected systems are isolated to prevent the threat from spreading.",
    "Impact Assessment — The full extent of the breach is evaluated.",
    "Eradication — The threat is completely removed from all affected systems.",
    "Recovery — Systems are restored to normal operation using clean backups where necessary.",
    "Notification — Affected patients, regulatory bodies, and law enforcement are informed where appropriate.",
    "Post-Incident Monitoring — Systems are closely monitored for signs of recurring activity.",
    "Documentation — A full record of the incident and response actions is documented.",
    "Review and Improvement — A formal post-incident review identifies lessons learned and improves future responses."
]:
    add_bullet(doc, item)
add_bold_label(doc, "Strengths and Weaknesses:", "The procedure covers the full incident lifecycle. Its effectiveness depends entirely on staff being trained in advance — an untested plan provides false assurance. Regular tabletop exercises and simulations are essential.")

add_heading(doc, "3.4 Training and Awareness Programme", level=2, color=(46, 84, 150))
add_bold_label(doc, "Policy Statement:", "All MedCore staff must participate in regular cybersecurity training and awareness activities to minimise the risk of human error.")
for item in [
    "Mandatory cybersecurity awareness training for all new staff during onboarding, covering password hygiene, phishing recognition, and incident reporting.",
    "Regular phishing simulation exercises to test staff awareness and identify individuals requiring additional training.",
    "Quarterly security refresher training for all employees.",
    "Dedicated advanced training for IT staff covering emerging threats, vulnerability management, and incident response.",
    "Clear, accessible reporting channels so all staff know who to contact when suspicious activity is detected.",
    "Physical awareness measures including posters and reminders reinforcing key security practices throughout the hospital."
]:
    add_bullet(doc, item)
add_bold_label(doc, "Strengths and Weaknesses:", "This programme addresses the human element — the most commonly exploited vulnerability. The weakness is that training effectiveness diminishes without reinforcement, and busy clinical environments may deprioritise security training. Leadership commitment is essential.")

# Activity 4
add_heading(doc, "Activity 4: Reflection", level=1)

add_heading(doc, "4.1 Adapting Security Policies to Evolving Threats", level=2, color=(46, 84, 150))
add_para(doc, "Security policies must never be treated as static documents. The cybersecurity threat landscape is in constant evolution — new ransomware variants emerge regularly, attack techniques become more sophisticated, and threat actors continuously identify novel ways to exploit vulnerabilities. For an organisation like MedCore, whose operations are directly linked to patient safety, the consequences of falling behind the threat curve are severe.")
add_para(doc, "To remain effective, MedCore must continuously monitor the threat landscape by subscribing to threat intelligence feeds and following advisories from bodies such as CISA, Nigeria's ngCERT, and the WHO's cybersecurity guidance for healthcare providers. Policy updates should be triggered not only on a fixed schedule but also reactively — whenever a significant new threat emerges in the healthcare sector. As the organisation grows or adopts new technologies, the scope and content of its security policies must be revised accordingly.")

add_heading(doc, "4.2 Regular Review and Update of Policies", level=2, color=(46, 84, 150))
add_para(doc, "MedCore should conduct a formal policy review at minimum every six months, or immediately following any security incident. Each review should involve IT security personnel, hospital management, compliance officers, and where possible an independent external auditor.")
add_para(doc, "The review process should assess whether existing controls remain effective, whether new assets require policy coverage, whether training materials reflect the current threat environment, and whether the organisation remains compliant with the latest NDPR and other regulatory requirements. All policy revisions must be formally documented, communicated to relevant staff, and signed off by senior management. Version control should be maintained so that the history of policy changes remains traceable and auditable.")

# Conclusion
add_heading(doc, "Conclusion", level=1)
add_para(doc, "This report has demonstrated that effective network security in a healthcare environment requires a structured, multi-layered approach that goes well beyond the installation of technical tools. For MedCore Hospital and Diagnostics Centre, the risk assessment revealed a consistently high level of exposure across nearly all asset categories, driven primarily by weaknesses in authentication, encryption, device management, and staff awareness.")
add_para(doc, "The four security policies developed in this report — covering access control, data protection, incident response, and training and awareness — collectively address these weaknesses and provide MedCore with a solid foundation for a robust security posture. However, policies alone are insufficient. Their effectiveness depends on consistent enforcement, regular review, and a genuine organisational commitment to security at every level.")
add_para(doc, "As cyber threats continue to evolve in sophistication and frequency, MedCore must treat cybersecurity not as a one-time project but as an ongoing operational priority. The cost of prevention will always be significantly less than the cost of a breach — both in financial terms and in the irreplaceable currency of patient trust.")

# References
add_heading(doc, "References", level=1)
refs = [
    "National Information Technology Development Agency (NITDA). (2019). Nigeria Data Protection Regulation (NDPR). https://nitda.gov.ng/ndpr/",
    "NIST. (2018). Framework for Improving Critical Infrastructure Cybersecurity. https://www.nist.gov/cyberframework",
    "NIST. (2017). SP 800-63B: Digital Identity Guidelines. https://pages.nist.gov/800-63-3/sp800-63b.html",
    "International Organisation for Standardisation. (2022). ISO/IEC 27001:2022 — Information Security Management Systems. https://www.iso.org/standard/27001",
    "PCI Security Standards Council. (2022). PCI DSS v4.0. https://www.pcisecuritystandards.org/",
    "CISA. (2023). Healthcare and Public Health Sector Cybersecurity. https://www.cisa.gov/topics/critical-infrastructure-security-and-resilience/critical-infrastructure-sectors/healthcare-and-public-health-sector",
    "World Health Organisation (WHO). (2021). Cybersecurity in Health. https://www.who.int/publications/i/item/9789240033955",
    "ngCERT. (2023). Nigeria Computer Emergency Response Team Advisories. https://ngcert.gov.ng/",
    "SANS Institute. (2021). Incident Handler's Handbook. https://www.sans.org/white-papers/33901/",
    "Microsoft Security. (2023). Zero Trust Security Model. https://www.microsoft.com/en-us/security/business/zero-trust"
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    set_run(r, size=11)

# Save
output_path = os.path.expanduser("~/Kali-Linux-Security-Labs/INT304/Lab2/INT304_Lab2_Report.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print("Report saved successfully!")
