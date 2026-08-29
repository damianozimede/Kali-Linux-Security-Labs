from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('INT310: Web Application Source Code Vulnerability Analysis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Lab 1: Understanding Web Application Security Mechanisms', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

name = doc.add_paragraph('Patrick Damian Ozimede')
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.runs[0].bold = True

doc.add_page_break()

doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'This report documents research conducted for Lab 1, focusing on two foundational '
    'web application security mechanisms: Authentication and Session Management. Both '
    'mechanisms were selected because they build directly on hands-on work completed in '
    'earlier courses (session fixation, hijacking, and CSRF testing against DVWA), and '
    'this research extends that practical experience with a deeper look at how these '
    'mechanisms fail in real-world production systems.'
)

doc.add_heading('Security Mechanism 1: Authentication', level=1)
doc.add_paragraph(
    'Authentication is the process by which a system verifies that a user is who they '
    'claim to be, before granting any access to resources. Common methods include '
    'password-based login, multi-factor authentication (MFA), and federated approaches '
    'such as OAuth. Password-based authentication alone is widely considered weak: it '
    'depends entirely on secrets that are frequently reused, guessed, or leaked in '
    'unrelated breaches, and offers no protection once a credential is stolen. MFA adds '
    'a second verification factor (something the user has, such as a phone, or something '
    'they are, such as a fingerprint) on top of the password, meaningfully raising the '
    'cost of a successful attack. OAuth and similar federated protocols shift '
    'authentication to a trusted third party, reducing how many places a password must '
    'be stored, but introducing new risks around token handling and third-party trust.'
)
doc.add_paragraph(
    'Despite these protections, authentication remains one of the most commonly '
    'exploited areas of web application security. Common vulnerabilities include weak or '
    'reused passwords, credential stuffing (automated reuse of previously breached '
    'username and password pairs across other services), and social engineering attacks '
    'that bypass MFA entirely by manipulating a human into approving a fraudulent login '
    'request rather than breaking the technical control itself. Best practices to '
    'mitigate these risks include enforcing MFA using number-matching or app-based '
    'approval rather than simple push notifications, rate-limiting and locking accounts '
    'after repeated failed or repeated MFA prompts, and monitoring for credential reuse '
    'against known breach databases.'
)

doc.add_heading('Real-World Example: The 2022 Uber Breach (MFA Fatigue)', level=2)
doc.add_paragraph(
    'In September 2022, Uber suffered a major internal breach that illustrated how even '
    'a properly deployed MFA system can be defeated through social engineering rather '
    'than a technical flaw. An attacker had obtained valid VPN credentials belonging to a '
    'third-party contractor, likely through a credential-stuffing attack using previously '
    'leaked passwords. Because the contractor\'s VPN access was protected by push-based '
    'MFA, the attacker could not log in directly. Instead, the attacker repeatedly '
    'attempted to log in with the stolen credentials, triggering a flood of MFA approval '
    'requests on the contractor\'s phone, a technique known as MFA fatigue or MFA '
    'bombing. When the contractor did not respond, the attacker escalated by contacting '
    'them directly over WhatsApp, impersonating Uber IT support and claiming that '
    'approving one more notification would make the flood of requests stop. Under this '
    'pressure, the contractor approved the request, granting the attacker full VPN access '
    'to Uber\'s internal network. From there, the attacker moved laterally, ultimately '
    'gaining access to internal tools, source code repositories, and communication '
    'platforms including Slack.'
)
doc.add_paragraph(
    'What went wrong was not the presence of MFA itself but its implementation: a simple '
    'accept/deny push notification gives an attacker unlimited free attempts to wear down '
    'a single human being, and provides that person no context about who or where the '
    'request is actually coming from. This could have been avoided through number-'
    'matching MFA (requiring the user to enter a code shown on the login screen, rather '
    'than a single tap), automatic account lockout after a small number of denied or '
    'unanswered MFA prompts, and security awareness training that specifically covers '
    'MFA fatigue and unsolicited "IT support" contact as a red flag rather than a routine '
    'inconvenience.'
)

doc.add_heading('Security Mechanism 2: Session Management', level=1)
doc.add_paragraph(
    'Session management governs how a web application recognizes a user as "logged in" '
    'across multiple requests, typically through a session identifier stored in a cookie '
    'after the user successfully authenticates. Because HTTP itself is stateless, nearly '
    'every modern web application depends on session tokens to maintain a continuous, '
    'authenticated experience without requiring the user to log in on every single page '
    'load. Best practices for secure session management include marking cookies as '
    'Secure and HttpOnly, generating session identifiers that are long and '
    'cryptographically random so they cannot be guessed, rotating the session identifier '
    'after login to prevent fixation, and enforcing both an idle timeout and an absolute '
    'session lifetime.'
)
doc.add_paragraph(
    'When session management is implemented poorly, the consequences are severe, because '
    'a session token is functionally equivalent to a password for the duration it '
    'remains valid; whoever holds it is treated as the authenticated user, with no '
    'further verification required. Session hijacking occurs when an attacker steals a '
    'valid session token, whether through network interception, cross-site scripting, or '
    'a server-side memory disclosure vulnerability, and replays it to impersonate the '
    'victim without ever needing their password. Session fixation occurs when an '
    'application fails to issue a new session identifier after a user logs in, allowing '
    'an attacker to pre-set a known session ID for the victim and then use that same ID '
    'to access the account once the victim authenticates. Both attack classes were '
    'demonstrated firsthand in earlier coursework against a deliberately vulnerable DVWA '
    'instance.'
)

doc.add_heading('Real-World Example: CitrixBleed 2 (CVE-2025-5777)', level=2)
doc.add_paragraph(
    'CitrixBleed 2 is a critical memory-disclosure vulnerability disclosed in mid-2025, '
    'affecting Citrix NetScaler ADC and Gateway appliances widely used for enterprise VPN '
    'and remote access. The flaw allowed an unauthenticated attacker to send specially '
    'crafted, malformed requests to the appliance\'s login endpoint, causing it to leak '
    'small fragments of its memory in the response. Because active session tokens and '
    'authentication cookies were sometimes present in that leaked memory, attackers could '
    'harvest valid session identifiers belonging to already-authenticated users, then '
    'replay those tokens from an entirely different device and location to hijack the '
    'session outright. Because the victim had already completed authentication, including '
    'MFA, before their token was stolen, this technique bypassed multi-factor '
    'authentication entirely, the attacker never needed to pass through the login process '
    'themselves at all, they simply reused a session that had already been granted. Once '
    'inside, attackers were observed escalating privileges, creating rogue administrator '
    'accounts, and in confirmed cases deploying ransomware such as DragonForce and '
    'LockBit.'
)
doc.add_paragraph(
    'This case demonstrates a critical lesson in session management: authentication and '
    'session validity are not the same guarantee, and a system is only as strong as the '
    'protection given to a session token after login succeeds. What went wrong was a '
    'server-side memory-handling defect that exposed session data that should never have '
    'been retrievable by an unauthenticated party, combined with session tokens that '
    'remained valid regardless of the requesting device or location. This could have been '
    'reduced through binding sessions to properties of the original login (such as IP '
    'address range or device fingerprint) so that a token replayed from an unrelated '
    'location is automatically rejected, shorter session lifetimes that reduce the window '
    'in which a leaked token remains useful, and mandatory session invalidation '
    'immediately after any related vulnerability patch, since simply patching the flaw '
    'without revoking already-issued tokens left some organizations exposed to continued '
    'abuse of previously stolen sessions.'
)

doc.add_heading('Exercise 2: Mapping Security Mechanisms to Common Vulnerabilities', level=1)
doc.add_paragraph(
    'Authentication and session management map directly onto some of the most persistent '
    'entries in the OWASP Top Ten. Weak authentication is the root cause behind '
    '"Identification and Authentication Failures," and directly enables credential '
    'stuffing and account takeover, the Uber case is a clear real-world instance of this '
    'category, where a valid credential combined with a weak MFA implementation was '
    'enough to defeat the control entirely. Poor session management similarly maps to '
    'the same OWASP category, since session fixation and hijacking are explicitly listed '
    'failure modes; the CitrixBleed 2 case shows this playing out at internet scale, '
    'where a single memory-disclosure bug undermined session integrity across every '
    'authenticated user on an affected appliance. Both mechanisms also intersect with '
    '"Security Misconfiguration": push-based MFA without number matching, and session '
    'tokens without device or location binding, are configuration choices rather than '
    'unavoidable flaws in the underlying technology.'
)

doc.add_heading('Reflection', level=1)
doc.add_paragraph(
    'This research reinforced that the most damaging security failures rarely stem from '
    'a single missing control, they stem from a working control being deployed without '
    'the surrounding safeguards that make it resilient under real attack conditions. Uber '
    'had MFA in place, and it was still defeated, not because MFA is ineffective but '
    'because a simple accept-or-deny push notification places the entire burden of '
    'resisting a persistent, well-disguised social engineering attempt on one tired '
    'human being. CitrixBleed 2 similarly shows that a session token, once issued, '
    'carries nearly the full weight of a user\'s authenticated identity, and if that '
    'token can be extracted through a completely unrelated memory-handling bug, then '
    'authentication strength upstream becomes almost irrelevant. The most critical '
    'security practices identified through this research are: MFA implementations that '
    'require active user input rather than passive approval, session tokens that are '
    'bound to more than just their own secrecy (through device or location context), '
    'and short, enforced session lifetimes that limit how long any single stolen token '
    'remains useful to an attacker. Comparing these findings against earlier coursework, '
    'the hands-on session fixation and hijacking exercises performed against DVWA now '
    'read less like isolated lab exercises and more like small-scale versions of exactly '
    'the failure modes that led to real breaches at major organizations.'
)

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab explored authentication and session management as two of the most '
    'consequential web application security mechanisms, and grounded that research in '
    'documented real-world breaches. Both case studies, Uber\'s 2022 MFA fatigue '
    'incident and the 2025 CitrixBleed 2 vulnerability, demonstrate that well-known '
    'security mechanisms can still fail catastrophically when their implementation '
    'overlooks how a determined attacker will actually try to defeat them. This '
    'reinforces the value of source code-level analysis, the subject of the next lab in '
    'this course, since many of these implementation gaps are only visible by examining '
    'exactly how a mechanism is coded, rather than assuming it works correctly simply '
    'because it exists.'
)

doc.add_heading('References', level=1)
references = [
    'centrexIT. "How Uber Was Breached Through MFA Fatigue: A Security Wake-Up Call." 2026.',
    'TeamPassword. "What the Uber Breach Taught Us About MFA Fatigue (and How to Prevent It)." 2026.',
    'Galink. "Uber Breach and MFA Fatigue." 2025.',
    'Ascella Group. "How Uber Was Breached Through Social Engineering." 2026.',
    'Splunk Threat Research Team. "CitrixBleed 2: When Memory Leaks Become Session Hijacks." 2025.',
    'Cyberwarzone. "What is CitrixBleed 2 (CVE-2025-5777)?" 2025.',
    'GBHackers. "Hackers Exploit CitrixBleed 2 to Hijack MFA-Protected Sessions and Deploy DragonForce Ransomware." 2026.',
    'CyberPress. "Hackers Abuse CitrixBleed 2 to Steal Session Tokens and Bypass MFA Protections." 2026.',
]
for ref in references:
    doc.add_paragraph(ref, style='List Bullet')

doc.save('INT310_Lab1_Report.docx')
print('Report generated: INT310_Lab1_Report.docx')
