from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT308: Session Management and Web Security', level=0)
subtitle = doc.add_heading('Lab 1: Session Management Vulnerabilities', level=1)

doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Environment: Kali Linux (legion@pastor, 192.168.5.139) targeting DVWA v1.8 on OWASP BWA VM (192.168.5.130)')

doc.add_heading('Overview', level=1)
doc.add_paragraph(
    "This lab explored session management vulnerabilities in DVWA v1.8, focusing on how "
    "insecure session cookie configuration and lack of session regeneration can lead to "
    "session fixation and session hijacking attacks. Three exercises were conducted: "
    "analysis of session cookie attributes, a session fixation attack via manual cookie "
    "manipulation, and a session hijacking demonstration using Burp Suite as an "
    "intercepting proxy."
)

# ---------------- Exercise 1 ----------------
doc.add_heading('Exercise 1: Analyzing Session Management Practices', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "Logged into DVWA using valid credentials (admin/admin) and used browser DevTools "
    "(Storage > Cookies) to inspect the PHPSESSID session cookie and its attributes "
    "(HttpOnly, Secure, SameSite).",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: PHPSESSID cookie attributes after login (DevTools > Storage > Cookies)]')

doc.add_heading('Findings', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Attribute'
hdr[1].text = 'Observed Value'
rows_data = [
    ('PHPSESSID', 'mcm3ucsdcvge7553f68t63fnb1'),
    ('HttpOnly', 'false'),
    ('Secure', 'false'),
    ('SameSite', 'None'),
]
for attr, val in rows_data:
    row = table.add_row().cells
    row[0].text = attr
    row[1].text = val

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "Because HttpOnly is set to false, the PHPSESSID cookie is accessible to client-side "
    "JavaScript through document.cookie. This means that if the application has any XSS "
    "vulnerability, an attacker could inject a script to steal the session cookie directly "
    "and impersonate the logged-in user without ever knowing their password. Similarly, "
    "because Secure is set to false, the cookie can be transmitted over plain, unencrypted "
    "HTTP, so anyone able to intercept network traffic between the user and the server "
    "(e.g. on a shared Wi-Fi network) could capture the session ID and hijack the session. "
    "Together, these missing flags mean the session token has no protection against either "
    "script-based theft or network-based interception -- the two most common vectors for "
    "session hijacking."
)

# ---------------- Exercise 2 ----------------
doc.add_heading('Exercise 2: Session Fixation Attack', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "Before authenticating, the PHPSESSID cookie value was manually set to a known, "
    "attacker-chosen value (fixedsession12345) via browser DevTools. The application was "
    "then logged into (admin/admin) using this same browser session, and the PHPSESSID "
    "value was checked again post-login to determine whether the server regenerated it.",
    style='List Bullet'
)
doc.add_paragraph(
    "A second, independent browser session (private/incognito window) then presented the "
    "same known session ID and attempted to access an authenticated page directly, without "
    "submitting any login credentials.",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: PHPSESSID manually set to a known value before login]')
doc.add_paragraph('[INSERT SCREENSHOT: Session ID unchanged after authentication -- session fixation confirmed]')
doc.add_paragraph('[INSERT SCREENSHOT: Unauthenticated browser accesses protected page using fixed session ID -- session hijacking successful]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "The PHPSESSID value (fixedsession12345) remained identical before and after login, "
    "confirming that DVWA does not regenerate the session identifier upon successful "
    "authentication. A second browser session that never submitted any credentials was "
    "then able to access the authenticated DVWA security settings page simply by "
    "presenting this same session ID."
)

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "This exercise demonstrated a session fixation vulnerability in DVWA. By manually "
    "setting the PHPSESSID cookie to a known value before authenticating, and observing "
    "that the same value remained active after login, it was confirmed that the "
    "application does not regenerate session identifiers upon successful authentication. "
    "This allowed a second browser session, which never submitted any login credentials, "
    "to access an authenticated page simply by presenting the same session ID. This "
    "confirms that an attacker who can plant a known session ID on a victim's browser "
    "(for example, via a crafted link) could hijack that victim's session once they log "
    "in, without ever needing to know their password. The core mitigation is to "
    "regenerate the session ID immediately after a successful login (e.g. using "
    "session_regenerate_id() in PHP), which invalidates any pre-set or attacker-known "
    "session values."
)

# ---------------- Exercise 3 ----------------
doc.add_heading('Exercise 3: Session Hijacking Using Burp Suite', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "Burp Suite's built-in Chromium browser was used to route traffic through Burp's "
    "intercepting proxy. The DVWA login request was captured with Intercept enabled, and "
    "the PHPSESSID value in the Cookie header was modified to an arbitrary, attacker-chosen "
    "string (hijacktest99999) before forwarding the request to the server.",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: Intercepted login POST request in Burp showing Cookie header with PHPSESSID]')
doc.add_paragraph('[INSERT SCREENSHOT: Authenticated access achieved after modifying PHPSESSID mid-request via Burp Suite]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "Despite substituting an arbitrary, attacker-chosen PHPSESSID value mid-request, the "
    "server granted full authenticated access to the DVWA dashboard. This confirms DVWA "
    "does not bind sessions to any additional validation (such as IP address, user-agent "
    "fingerprinting, or a server-side session token check) -- it trusts whatever PHPSESSID "
    "value is presented."
)

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "Using Burp Suite as an intercepting proxy, the login request was captured and the "
    "PHPSESSID cookie value was modified to an arbitrary, attacker-chosen string before "
    "forwarding the request. Despite this tampering, the application granted full "
    "authenticated access, confirming that DVWA does not validate or bind session IDs to "
    "any server-side state beyond the string itself. This demonstrates that session "
    "management in this application relies entirely on the secrecy of the session ID with "
    "no additional integrity checks (such as IP address binding or user-agent "
    "verification). An attacker who intercepts or guesses a valid session ID -- whether "
    "via network sniffing, XSS, or a fixation attack as shown in Exercise 2 -- can fully "
    "hijack a user's session. Mitigations include regenerating session IDs after login, "
    "binding sessions to additional client attributes, enforcing short session timeouts, "
    "and always transmitting session cookies over HTTPS with the Secure and HttpOnly flags "
    "set."
)

# ---------------- Conclusion & Recommendations ----------------
doc.add_heading('Additional Insights and Recommendations', level=1)
recs = [
    "Set the HttpOnly flag on all session cookies to prevent JavaScript access and mitigate XSS-based cookie theft.",
    "Set the Secure flag to ensure cookies are only transmitted over HTTPS.",
    "Set SameSite=Strict or Lax to reduce cross-site request exposure.",
    "Call session_regenerate_id(true) immediately after successful authentication to invalidate any pre-existing or attacker-known session IDs.",
    "Enforce session binding to additional client attributes (e.g. IP address or user-agent) where feasible, and implement short session timeouts with server-side invalidation on logout.",
    "Migrate the application to HTTPS to protect session tokens from network-level interception.",
]
for r in recs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This lab confirmed that DVWA v1.8 exhibits multiple session management weaknesses: "
    "session cookies lack protective flags (HttpOnly, Secure, SameSite), the application "
    "does not regenerate session IDs after login, and it does not bind sessions to any "
    "additional client-side validation. Together, these weaknesses make DVWA fully "
    "susceptible to session fixation and session hijacking attacks, underscoring the "
    "importance of secure session management practices in real-world web applications."
)

doc.save('INT308_Lab1_Session_Management.docx')
print("Report generated: INT308_Lab1_Session_Management.docx")
