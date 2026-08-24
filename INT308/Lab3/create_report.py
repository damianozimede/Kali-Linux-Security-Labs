from docx import Document

doc = Document()

doc.add_heading('INT308: Session Management and Web Security', level=0)
doc.add_heading('Lab 3: Cross-Site Request Forgery (CSRF) Protection', level=1)

doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Environment: Kali Linux (legion@pastor, 192.168.5.139) targeting DVWA v1.8 on OWASP BWA VM (192.168.5.130) via root SSH access')

doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "Cross-Site Request Forgery (CSRF) is an attack that exploits the trust a web "
    "application places in an authenticated user's browser, tricking that browser into "
    "submitting unwanted requests on the user's behalf. This lab demonstrated a working "
    "CSRF attack against DVWA's password-change function, implemented server-side CSRF "
    "token protection directly in DVWA's source code, and verified that the same attack "
    "was blocked once the protection was in place."
)

# Exercise 1
doc.add_heading('Exercise 1: Understanding CSRF Attacks', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "DVWA's CSRF page (vulnerabilities/csrf/) allows an authenticated user to change the "
    "admin password via a GET request with password_new, password_conf, and Change "
    "parameters. A standalone HTML file was created outside of DVWA, containing a form "
    "with hidden inputs pre-filled with an attacker-chosen password, disguised as a "
    "prize-claim page. This file was opened in the same browser session already "
    "authenticated to DVWA, and the submit button was clicked.",
    style='List Bullet'
)
doc.add_paragraph(
    "An initial attempt appeared to succeed based on the resulting URL parameters, but "
    "logging in with the new password failed, while the original admin/admin credentials "
    "still worked. Reviewing DVWA's server-side source (source/low.php) revealed the "
    "handler only processes the request if isset($_GET['Change']) is true -- the initial "
    "attack form omitted this hidden field. Adding a hidden Change=Change input to the "
    "attack form resolved this, and the subsequent attack successfully changed the "
    "account password without any credentials or direct interaction with DVWA.",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: CSRF attack page disguised as a prize claim, hosted outside DVWA]')
doc.add_paragraph('[INSERT SCREENSHOT: CSRF attack succeeded -- password change request submitted via external page using authenticated session]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "The unprotected CSRF endpoint allowed a fully external, attacker-controlled page to "
    "change the admin account password using only the victim's existing authenticated "
    "session -- no credentials, confirmation, or additional user interaction beyond a "
    "single click were required."
)

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "This exercise demonstrated a working CSRF attack against DVWA's unprotected "
    "password-change function. A form hosted entirely outside DVWA, containing hidden "
    "fields matching the target's expected parameters, was able to submit a "
    "password-change request using the victim's existing authenticated session, without "
    "any user interaction beyond a single click. Notably, an initial attack attempt "
    "appeared to fail because the form omitted the submit button's name/value (Change), "
    "which DVWA's backend checks via isset($_GET['Change']). This silently prevented the "
    "update with no error shown, illustrating that even minor mismatches in required "
    "parameters can cause an attack to fail undetected -- but also that an attacker who "
    "studies the target form closely can trivially replicate every required field."
)

# Exercise 2
doc.add_heading('Exercise 2: Implementing CSRF Protection', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "CSRF token protection was added directly to DVWA's low-security CSRF handler "
    "(source/low.php): a per-session token is generated if one does not already exist, "
    "and any password-change request is now rejected unless it includes a csrf_token "
    "parameter matching the value stored in $_SESSION. Since this environment runs PHP "
    "5.3.2, the modern random_bytes() function used in the lab's example was unavailable "
    "and md5(uniqid(mt_rand(), true)) was used instead to generate the token. The token "
    "is regenerated after each use to prevent reuse. The legitimate form in index.php "
    "was updated to include a hidden csrf_token field populated from the session.",
    style='List Bullet'
)
doc.add_paragraph(
    "While implementing this, an unrelated bug from Lab 2's session-expiration code "
    "surfaced: the redirect header('Location: logout.php') used a relative path, which "
    "resolved incorrectly (to a non-existent nested path) when triggered from within the "
    "csrf/ subdirectory. This was fixed by changing the redirect to an absolute path "
    "(/dvwa/logout.php).",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: Legitimate password change succeeds with valid CSRF token present]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "The legitimate DVWA form, now including the correct csrf_token, successfully "
    "changed the password and displayed \"Password Changed\", confirming the token "
    "generation and validation logic did not break normal application functionality."
)

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "Implementing a CSRF token involved generating a random token per session, embedding "
    "it as a hidden field in the legitimate form, and validating it server-side before "
    "processing any password change. Since this environment runs PHP 5.3.2, the modern "
    "random_bytes() function was unavailable and md5(uniqid(mt_rand(), true)) was used "
    "as a legacy-compatible alternative for token generation. This exercise also "
    "surfaced an unrelated but real bug from the previous lab's session-expiration code: "
    "a relative redirect path (logout.php) failed when triggered from a nested URL, "
    "requiring a fix to use an absolute path (/dvwa/logout.php). This is a good reminder "
    "that redirect paths in web applications should generally be absolute to avoid "
    "breaking depending on which page triggers them."
)

# Exercise 3
doc.add_heading('Exercise 3: Testing CSRF Protection', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "The original attack HTML file, which does not include a csrf_token field, was "
    "re-opened and submitted again while authenticated to DVWA, to test whether the new "
    "protection would block the same attack that succeeded in Exercise 1.",
    style='List Bullet'
)
doc.add_paragraph(
    "The password was then verified to remain unchanged by logging out and logging back "
    "in with the original admin/admin credentials.",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: CSRF attack blocked after token protection implemented]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "The attack request was rejected with a \"CSRF token validation failed!\" message, "
    "and the account password remained unchanged, confirming the token-based protection "
    "successfully blocked the same attack that previously succeeded."
)

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "Re-running the original attack form, which did not include a csrf_token field, "
    "against the now-protected endpoint resulted in a clear \"CSRF token validation "
    "failed!\" message, and the password was confirmed unchanged. This confirms the "
    "token-based protection is effective against a naive replay of the original attack. "
    "However, this protection has limitations: if the token were leaked via another "
    "vulnerability (such as XSS, explored in INT307), an attacker could still forge a "
    "valid request. CSRF protection should therefore be treated as one layer of defense, "
    "complemented by measures such as the SameSite cookie attribute (limited by this "
    "environment's legacy PHP, as seen in Lab 2), re-authentication for sensitive "
    "actions, and general output encoding to prevent token theft via XSS."
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This lab demonstrated the full lifecycle of a CSRF vulnerability: successful "
    "exploitation against an unprotected endpoint, implementation of a server-side "
    "token-based defense compatible with a legacy PHP 5.3.2 environment, and verification "
    "that the same attack was blocked once the defense was in place. Understanding both "
    "how to exploit and how to properly defend against CSRF is essential for building and "
    "maintaining secure web applications."
)

doc.save('INT308_Lab3_CSRF_Protection.docx')
print("Report generated: INT308_Lab3_CSRF_Protection.docx")
