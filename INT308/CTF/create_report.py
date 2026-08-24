from docx import Document

doc = Document()

doc.add_heading('INT308 Session Management and Web Security', level=0)
doc.add_heading('CTF Challenge: Testing for Session Hijacking on OWASP Juice Shop', level=1)

doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Environment: Kali Linux (legion@pastor, 192.168.5.139) targeting a locally deployed OWASP Juice Shop instance (Docker, 127.0.0.1:3000)')

doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "Session hijacking occurs when an attacker gains unauthorized access to a "
    "legitimate user's session, typically by stealing, predicting, fixating, or "
    "replaying a session token. Because a valid session token often grants the same "
    "level of access as a username and password, weaknesses in session management can "
    "completely undermine an otherwise secure authentication system. This report "
    "documents a session hijacking assessment performed against a locally deployed "
    "instance of OWASP Juice Shop, covering session token predictability, session "
    "fixation, session expiration, and transport security."
)

doc.add_heading('Methodology', level=1)
doc.add_paragraph(
    "OWASP Juice Shop was deployed locally via Docker (bkimminich/juice-shop image), "
    "exposed on http://127.0.0.1:3000. Testing was performed using the Kali Linux "
    "browser's built-in Developer Tools (Storage/Cookies panel) to inspect the "
    "application's authentication cookie, and jwt.io was used to decode the JSON Web "
    "Token (JWT) issued by the application. Each of the four focus areas specified in "
    "the challenge -- session ID predictability, session fixation, session expiration, "
    "and transport security -- was tested directly against this instance using only the "
    "tester's own account."
)

# Findings
doc.add_heading('Findings', level=1)

doc.add_heading('1. Session ID Predictability', level=2)
doc.add_paragraph(
    "Unlike a simple sequential or random session ID, Juice Shop issues a signed JSON "
    "Web Token (JWT) as its authentication cookie ('token'). The token is not "
    "predictable in the traditional sense, since it is cryptographically signed by the "
    "server. However, decoding the token's payload (which is base64-encoded, not "
    "encrypted) revealed the full account record, including the account ID, email "
    "address, role, and -- notably -- the account's password hash (MD5). This means "
    "anyone who obtains the token, whether via network interception or client-side "
    "script access, can read this sensitive data without needing the server's signing "
    "key."
)
doc.add_paragraph('[INSERT SCREENSHOT: Juice Shop session token and cookie attributes after login]')
doc.add_paragraph('[INSERT SCREENSHOT: Decoded JWT payload revealing account data via jwt.io]')

doc.add_heading('2. Session Fixation', level=2)
doc.add_paragraph(
    "A token value ('fixedtoken12345') was manually set in the browser's cookie store "
    "before logging in, replicating the fixation technique used successfully against "
    "DVWA in Lab 1. Upon logging in, Juice Shop issued a brand-new, properly signed JWT, "
    "completely overwriting the fixed value. This confirms Juice Shop is not vulnerable "
    "to session fixation: because a valid token must be cryptographically signed with "
    "the server's private key, a client cannot pre-plant or forge an acceptable token "
    "value, unlike DVWA's server-generated-but-unvalidated PHPSESSID."
)
doc.add_paragraph('[INSERT SCREENSHOT: Session token removed from cookies after logout]')
doc.add_paragraph('[INSERT SCREENSHOT: Token cookie manually set to a known value before login]')
doc.add_paragraph('[INSERT SCREENSHOT: New JWT issued after login, overwriting the fixed token value]')

doc.add_heading('3. Session Expiration', level=2)
doc.add_paragraph(
    "Decoding an active JWT via jwt.io showed no 'exp' (expiration) claim in the "
    "payload, meaning the token contains no built-in, self-enforced expiry. Logging out "
    "correctly cleared the token cookie client-side, and re-adding the same (old) token "
    "value afterward did not restore the logged-in state in the application UI. This "
    "suggests session termination is enforced at the application/front-end state level "
    "rather than by the JWT itself becoming cryptographically invalid, since a JWT with "
    "no exp claim would otherwise remain valid indefinitely if replayed directly against "
    "a backend API endpoint. The absence of a token blocklist or expiration mechanism "
    "represents a real gap, even though the front-end behavior appeared secure in this "
    "test."
)
doc.add_paragraph('[INSERT SCREENSHOT: Decoded JWT payload showing no "exp" (expiration) claim]')

doc.add_heading('4. Transport Security', level=2)
doc.add_paragraph(
    "The Juice Shop instance was served over plain HTTP (http://127.0.0.1:3000), and "
    "the 'token' cookie's Secure attribute was set to false, meaning the session token "
    "could be transmitted over an unencrypted connection and intercepted by anyone able "
    "to observe network traffic between the client and server, such as on a shared or "
    "untrusted network."
)

doc.add_heading('Countermeasures', level=1)
recs = [
    "Add an exp (expiration) claim to issued JWTs with a reasonable, short expiration "
    "window, and implement refresh-token rotation for maintaining longer sessions.",
    "Avoid embedding sensitive fields, such as password hashes, inside JWT payloads, "
    "since JWT payloads are base64-encoded and readable by anyone who obtains the token, "
    "not encrypted.",
    "Implement a server-side token revocation or blocklist mechanism for logout events, "
    "since JWTs cannot otherwise be invalidated before their natural expiry once issued.",
    "Enforce HTTPS in all environments and set the Secure flag on authentication "
    "cookies to prevent transmission over unencrypted connections.",
    "Set the HttpOnly flag on the token cookie to reduce its exposure to theft via "
    "client-side script injection (XSS).",
]
for r in recs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This assessment of OWASP Juice Shop found that while its JWT-based authentication "
    "architecture inherently resists session fixation, since tokens must be "
    "cryptographically signed by the server, it introduces different risks compared to "
    "the traditional session-ID model tested against DVWA in earlier labs. Most notably, "
    "the JWT payload exposed sensitive account data, including a password hash, in "
    "plainly decodable form, and the token itself carried no expiration claim. Securing "
    "web sessions requires understanding the specific architecture in use: a defense "
    "that closes one class of vulnerability, such as fixation, does not guarantee "
    "protection against others, such as data exposure or indefinite token validity. "
    "Comparing DVWA's legacy session-ID model against Juice Shop's modern JWT model "
    "over the course of this course's labs demonstrated that both older and newer "
    "session management approaches carry their own distinct risks if not implemented "
    "carefully."
)

doc.save('INT308_CTF_Session_Hijacking_JuiceShop.docx')
print("Report generated: INT308_CTF_Session_Hijacking_JuiceShop.docx")
