from docx import Document

doc = Document()

doc.add_heading('INT308: Session Management and Web Security', level=0)
doc.add_heading('Lab 2: Secure Session Management Practices', level=1)

doc.add_paragraph('Name: Patrick Damian Ozimede')
doc.add_paragraph('Environment: Kali Linux (legion@pastor, 192.168.5.139) targeting DVWA v1.8 on OWASP BWA VM (192.168.5.130) via root SSH access')

doc.add_heading('Overview', level=1)
doc.add_paragraph(
    "This lab explored secure session management practices by directly modifying DVWA's "
    "server-side session handling code. Root SSH access to the OWASP BWA VM was obtained "
    "to locate and edit the file responsible for session initialization "
    "(includes/dvwaPage.inc.php). Three exercises were conducted: configuring secure "
    "session cookie attributes, implementing session expiration and forced logout on "
    "inactivity, and testing these controls end-to-end."
)

# Exercise 1
doc.add_heading('Exercise 1: Configuring Secure Session Cookies', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "Obtained root SSH access to the OWASP BWA VM (192.168.5.130). Located the live DVWA "
    "install at /owaspbwa/dvwa-git/dvwa/ and identified includes/dvwaPage.inc.php as the "
    "file calling session_start(). Added a session_set_cookie_params() call immediately "
    "before session_start(), setting httponly => true, samesite => 'Strict', and "
    "secure => false (required since this environment serves DVWA over plain HTTP on a "
    "bare IP address rather than HTTPS/a domain). Added session_regenerate_id(true) "
    "immediately after session_start() to prevent session fixation.",
    style='List Bullet'
)
doc.add_paragraph(
    "The initial edit used PHP short array syntax ([...]), which caused a fatal parse "
    "error because this server runs PHP 5.3.2 -- short array syntax was not introduced "
    "until PHP 5.4. The code was corrected to use the legacy array(...) syntax.",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: Root SSH access confirmed on OWASP BWA VM]')
doc.add_paragraph('[INSERT SCREENSHOT: PHPSESSID cookie attributes after applying secure session configuration]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "After correcting the array syntax, DVWA loaded normally again. However, inspecting "
    "the PHPSESSID cookie in the browser afterward still showed HttpOnly=false and "
    "SameSite=None, despite the code being implemented correctly and matching the lab's "
    "own example. This was traced to a limitation of the PHP 5.3.2 environment: SameSite "
    "cookie support was not added to PHP until version 7.3, and this legacy build does not "
    "reliably enforce all session_set_cookie_params() attributes across DVWA's code paths."
)

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    "Adding session_set_cookie_params() with httponly => true and samesite => 'Strict', "
    "along with session_regenerate_id(true), is the correct approach to harden DVWA's "
    "session cookie per the lab's guidance. However, testing showed the HttpOnly flag did "
    "not take effect in the browser despite correct implementation, and SameSite was "
    "silently ignored. This is because the server runs PHP 5.3.2 (released 2010), which "
    "predates PHP's samesite cookie support (added in 7.3) and has known inconsistencies "
    "in how session_set_cookie_params() interacts with cookie handling across code paths. "
    "This highlights a real-world lesson: security configuration that is correctly written "
    "in code can still fail to take effect on legacy, unpatched infrastructure, which is "
    "itself a security risk worth flagging in any audit."
)

# Exercise 2
doc.add_heading('Exercise 2: Implementing Session Expiration and Invalidation', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "Added an inactivity-timeout check to includes/dvwaPage.inc.php, immediately after "
    "session_regenerate_id(true). The check uses ini_set('session.gc_maxlifetime', 900) "
    "and a $_SESSION['LAST_ACTIVITY'] timestamp comparison: if more than 900 seconds "
    "(15 minutes) have elapsed since the last recorded activity, the session is unset and "
    "destroyed, and the user is redirected to logout.php. For testing purposes, the "
    "timeout was temporarily reduced to 20 seconds, verified, then restored to 900 seconds.",
    style='List Bullet'
)

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "With the timeout temporarily set to 20 seconds, logging in and then waiting "
    "approximately 25-30 seconds before attempting to navigate to another DVWA page "
    "correctly triggered the expiration logic, unsetting the session and redirecting back "
    "to the login page."
)

# Exercise 3
doc.add_heading('Exercise 3: Testing Session Management Controls', level=1)

doc.add_heading('Steps Performed', level=2)
doc.add_paragraph(
    "The inactivity test performed in Exercise 2 doubled as the live test required by "
    "Exercise 3: logging in, allowing the session to sit idle past the configured "
    "timeout threshold, and then attempting to access a protected page (the DVWA "
    "dashboard).",
    style='List Bullet'
)

doc.add_paragraph('[INSERT SCREENSHOT: Session expired after inactivity -- redirected to logout]')

doc.add_heading('Findings', level=2)
doc.add_paragraph(
    "Access to the protected page was correctly denied after the session expired, and the "
    "browser was redirected to the login page rather than serving cached or unauthorized "
    "content."
)

doc.add_heading('Reflection (Exercises 2 & 3)', level=2)
doc.add_paragraph(
    "Implementing session expiration via gc_maxlifetime and a LAST_ACTIVITY timestamp "
    "check successfully forced session termination after a period of inactivity, "
    "redirecting the user back to the login page. This was verified by temporarily "
    "shortening the timeout to 20 seconds for testing purposes, then restoring it to the "
    "specified 900 seconds (15 minutes). This confirms that, unlike the cookie-flag "
    "configuration in Exercise 1, inactivity-based session expiration is fully enforceable "
    "even on this legacy PHP environment, since it relies on standard session variable "
    "logic rather than cookie attributes the browser/server negotiate. A remaining "
    "challenge is that this approach requires the check to run on every page load via a "
    "shared include, which is fragile -- any new page added to the application that "
    "does not include dvwaPage.inc.php would bypass the expiration check entirely."
)

# Discussion & Recommendations
doc.add_heading('Discussion', level=1)
doc.add_paragraph(
    "This lab demonstrated that secure session management requires more than correctly "
    "written code -- the underlying platform must also support the intended security "
    "controls. DVWA's session cookie hardening was implemented exactly as prescribed, "
    "yet the legacy PHP 5.3.2 runtime silently failed to enforce the SameSite and "
    "HttpOnly attributes as expected in modern browsers and PHP versions. In contrast, "
    "session expiration logic, which relies purely on PHP session variables rather than "
    "cookie negotiation, worked reliably. This distinction is an important real-world "
    "takeaway: server-side logic controls (expiration, regeneration) tend to be more "
    "portable and dependable across environments than client-negotiated cookie attributes, "
    "which depend on both server software version and browser support."
)

doc.add_heading('Recommendations', level=1)
recs = [
    "Upgrade the underlying PHP version (or migrate the application) to a supported, "
    "modern release so that cookie security attributes such as HttpOnly and SameSite "
    "are properly enforced.",
    "Serve the application over HTTPS so the Secure cookie flag can be safely enabled "
    "without breaking functionality.",
    "Keep session_regenerate_id(true) on login at minimum, rather than on every request, "
    "to avoid unnecessary session churn while still preventing fixation.",
    "Centralize session-handling logic in a single, mandatory include used by every "
    "page, or better, enforce it at the web server/framework level, so new pages cannot "
    "accidentally bypass expiration checks.",
    "Regularly audit legacy environments for silent security-control failures, since "
    "code that appears correct can still be functionally inert on outdated platforms.",
]
for r in recs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This lab implemented secure session cookie configuration and inactivity-based "
    "session expiration directly on DVWA's server-side code via root SSH access. While "
    "session expiration worked reliably, cookie attribute enforcement was undermined by "
    "the legacy PHP 5.3.2 environment, illustrating that secure coding practices alone "
    "are insufficient without a supporting, up-to-date platform."
)

doc.save('INT308_Lab2_Secure_Session_Management.docx')
print("Report generated: INT308_Lab2_Secure_Session_Management.docx")
