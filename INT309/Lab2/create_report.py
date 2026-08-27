from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('INT309: Web Technologies and Database Security', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Lab 2: Web Development Components and Security Threats', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

name = doc.add_paragraph('Patrick Damian Ozimede')
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.runs[0].bold = True

doc.add_page_break()

doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'This report documents the completion of Lab 2, which explored front-end and '
    'back-end web development components and the security threats commonly associated '
    'with each. A small custom web application (an HTML/CSS/JavaScript front end backed '
    'by a PHP and MariaDB back end) was built on Kali Linux to demonstrate SQL injection '
    'and cross-site scripting vulnerabilities firsthand, followed by a cross-site request '
    'forgery test against a previously hardened DVWA instance.'
)

doc.add_heading('Environment Setup', level=1)
doc.add_paragraph(
    'PHP 8.4 and MariaDB were installed locally on Kali Linux to host a standalone '
    'vulnerable application, separate from the OWASP BWA virtual machine used in prior '
    'labs. A database (test1) and a dedicated application user (webapp) were created, and '
    'PHP\'s built-in development server was used to serve the application at '
    'http://127.0.0.1:8080. One environment note worth recording: PHP 8.1 and later throw '
    'exceptions on mysqli errors by default, rather than failing silently. This meant that '
    'malformed injection payloads produced visible 500 errors with full stack traces in '
    'the server log, which turned out to be useful evidence of the injection reaching the '
    'database layer, rather than an obstacle.'
)

doc.add_heading('Exercise 1: Exploring Front-End Development Components', level=1)
doc.add_paragraph(
    'Common front-end technologies (HTML, CSS, JavaScript, and frameworks such as React, '
    'Angular, and Vue.js) were reviewed. A simple front-end application was then built, '
    'consisting of a header with a navigation menu, a contact form (First Name, Last '
    'Name, School), and an interactive button that toggles displayed text using '
    'JavaScript.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 1 - index.php, front-end form and interactive button]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'Front-end technologies work together to create a cohesive user experience by '
    'dividing responsibilities: HTML provides structure and content, CSS handles '
    'presentation and layout, and JavaScript adds interactivity and dynamic behavior. In '
    'this exercise, this division was visible in a simple form: HTML defined the header, '
    'navigation, and input fields, CSS styled the layout and spacing, and JavaScript '
    'powered the button that changed the displayed text on click without requiring a page '
    'reload. This separation of concerns is what makes modern front-end frameworks like '
    'React, Angular, and Vue.js effective, they build on the same fundamental model but '
    'add structure for managing state and reusable components at scale. However, this '
    'exercise also demonstrated how easily front-end code becomes an attack surface. '
    'Because the form submits directly to a back-end script with no client-side or '
    'server-side sanitization, an attacker can inject malicious JavaScript into an input '
    'field, and if that input is later rendered back to a page without escaping, the '
    'browser executes it as if it were legitimate code. This is exactly how cross-site '
    'scripting occurs, and it highlights that front-end trust boundaries are not enough. '
    'Similarly, a form that performs a state-changing action without a CSRF token can be '
    'triggered from an entirely different site the user happens to have open, since the '
    'browser will still send the user\'s session cookies along with the forged request. '
    'Both vulnerabilities stem from the same root issue: treating user input, or a '
    'browser request, as inherently trustworthy.'
)

doc.add_heading('Exercise 2: Understanding Back-End Development Components', level=1)
doc.add_paragraph(
    'Common back-end technologies were reviewed, including server-side languages (PHP, '
    'Python, Java, Node.js), databases (MySQL/MariaDB, MongoDB, PostgreSQL), and web '
    'frameworks (Express, Django, Laravel). A basic back-end script (process.php) was '
    'built using PHP and the mysqli extension to connect to a MariaDB database, insert '
    'submitted form data, and display stored records. The script was written to be '
    'intentionally vulnerable, mirroring the lab\'s example: user input was concatenated '
    'directly into the SQL query with no sanitization or prepared statement, and '
    'retrieved data was echoed back to the page without output escaping.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 2 - process.php, legitimate form submission result]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'Back-end components form the trust boundary of a web application: they process '
    'user input, communicate with databases, and enforce (or fail to enforce) business '
    'logic and access rules. In this exercise, the back-end consisted of a PHP script '
    'using the mysqli extension to insert form data directly into a MariaDB database. '
    'Because the query was built through direct string concatenation rather than a '
    'prepared statement, user input was able to escape its intended context as a quoted '
    'string value and be interpreted as executable SQL syntax, which is the definition of '
    'a SQL injection vulnerability. Attempting this firsthand made the mechanism concrete '
    'rather than theoretical: a single unescaped quote character was enough to break the '
    'query, and in a differently structured query such as a login check, could be used to '
    'bypass authentication entirely. Beyond injection, the back-end code also exposed a '
    'broader issue: the application echoed raw database contents back to the page, '
    'including a full "database dump" with no access control, meaning any visitor could '
    'view every record ever submitted by any user. This illustrates that back-end '
    'vulnerabilities tend to have more severe consequences than front-end ones, since the '
    'back-end is where data persistence, authentication, and authorization actually '
    'happen. A compromised or poorly written back-end can expose an entire database, not '
    'just a single user\'s session, which is why input validation, prepared statements, '
    'and the principle of least privilege for database accounts are considered '
    'foundational rather than optional security practices.'
)

doc.add_heading('Exercise 3: Identifying Security Threats', level=1)
doc.add_paragraph(
    'Five common OWASP Top Ten threats were reviewed: SQL Injection, Cross-Site '
    'Scripting (XSS), Cross-Site Request Forgery (CSRF), Security Misconfiguration, and '
    'Insecure Deserialization. Three of these (SQLi, XSS, and CSRF) were then tested '
    'directly.'
)

doc.add_heading('SQL Injection Test', level=2)
doc.add_paragraph(
    'The payload \' OR \'1\'=\'1 was submitted in the First Name field of the custom '
    'application. Because the field value is concatenated directly into an INSERT query '
    'with no sanitization, the injected quote terminated the intended string early and '
    'caused the remaining OR \'1\'=\'1 fragment to be parsed as SQL syntax rather than '
    'text. This produced an unhandled mysqli exception (a 500 Internal Server Error), '
    'with the server log confirming a "Truncated incorrect DOUBLE value" error, direct '
    'evidence that the injected input reached and altered the structure of the SQL query '
    'sent to the database.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3a - SQL injection attempt, 500 error]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3b - server log, mysqli exception confirming injection]')

doc.add_heading('Cross-Site Scripting (XSS) Test', level=2)
doc.add_paragraph(
    'A script payload was submitted in the School field. An initial attempt using single '
    'quotes inside the payload also broke the SQL query (the same unsanitized '
    'concatenation issue affecting every field), producing another 500 error. A revised '
    'payload using double quotes instead, <script>alert("XSS")</script>, avoided breaking '
    'the SQL syntax and was successfully stored and later executed by the browser when '
    'the page rendered the stored data back without output escaping, confirming a stored '
    'XSS vulnerability. The alert fired twice, once for the "Submitted Data" table and '
    'once for the "Database Dump" section, since both display the same unsanitized '
    'field. The affected School column rendered as empty in the page, not because no data '
    'was stored, but because the browser executed the payload as a script rather than '
    'displaying it as visible text.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 4a - XSS payload with single quote, 500 error]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 4b - stored XSS, alert box triggered]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 4c - stored XSS firing a second time]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 4d - rendered page, School column empty after XSS execution]')

doc.add_heading('Cross-Site Request Forgery (CSRF) Test', level=2)
doc.add_paragraph(
    'CSRF was tested against the DVWA CSRF module, which had previously been patched '
    'with a custom per-session token system in an earlier lab. A classic CSRF attack URL '
    'was crafted (a GET request with password_new and password_conf parameters but no '
    'csrf_token), and submitted while logged in as admin to simulate a forged request '
    'from an attacker-controlled page. The application correctly rejected the request '
    'with a "CSRF token validation failed!" message, confirming the token-based '
    'protection was functioning. A legitimate password change was then submitted through '
    'the actual form, which succeeded normally, confirming the protection blocks forged '
    'requests without interfering with genuine use. The password was reverted to its '
    'original value afterward.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 5 - CSRF attack blocked, no token supplied]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 6 - legitimate password change succeeds]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'This exercise brought together SQL injection, cross-site scripting, and cross-site '
    'request forgery, three of the OWASP Top Ten\'s most consistently ranked threats, and '
    'demonstrated each one against a real application rather than as abstract concepts. '
    'The SQL injection test showed how directly concatenated user input can escape its '
    'intended string context and be parsed as executable SQL, in this case triggering a '
    'database-level type conversion error, but in a differently structured query (such as '
    'a login check) the identical technique is capable of bypassing authentication '
    'entirely. The stored XSS test showed how unsanitized output allows attacker-supplied '
    'script tags to execute in the context of any user who later views that data, with '
    'the added danger that the payload persists in the database and re-executes on every '
    'page load, not just at the moment of submission. This is a meaningfully more serious '
    'risk than reflected XSS, since it requires no ongoing interaction from the attacker '
    'once the payload is stored. The CSRF test, run against a previously hardened DVWA '
    'instance from earlier coursework, demonstrated the opposite side of the equation: a '
    'properly implemented per-session token successfully rejected a forged request while '
    'still allowing the legitimate, intentional request to succeed. Together, these three '
    'tests reinforce a common theme across all of the OWASP Top Ten: nearly every major '
    'web application vulnerability traces back to a failure to treat user-controllable '
    'input, whether from a form field, a URL parameter, or a cross-origin request, as '
    'untrusted until it has been validated, sanitized, or cryptographically verified. The '
    'practical mitigations demonstrated or discussed here, prepared statements for SQL '
    'injection, output encoding and escaping for XSS, and per-session CSRF tokens, are '
    'not exotic defenses; they are foundational practices that, when omitted, account for '
    'a disproportionate share of real-world application breaches.'
)

doc.add_heading('Recommendations for Secure Coding Practices', level=1)
doc.add_paragraph('Based on the findings above, the following practices are recommended:', style='List Bullet')
recommendations = [
    'Use parameterized queries or prepared statements for all database interactions, never concatenate user input directly into SQL strings.',
    'Escape or encode all output rendered back to a page (e.g., htmlspecialchars() in PHP) to prevent stored and reflected XSS.',
    'Implement per-session, unpredictable CSRF tokens on all state-changing forms, and validate them server-side before processing the request.',
    'Apply the principle of least privilege to database accounts; avoid using a root or administrative account for routine application queries.',
    'Avoid exposing full database contents or debug information (such as unhandled exception stack traces) to end users in production environments.',
    'Validate and sanitize all user input on the server side, client-side validation alone is not a security control.',
]
for rec in recommendations:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab provided hands-on experience building both sides of a simple web '
    'application and identifying how vulnerabilities are introduced through common '
    'coding shortcuts. Constructing a custom vulnerable application, rather than only '
    'exploiting a pre-built one, clarified exactly how SQL injection and stored XSS arise '
    'from unsanitized input and output handling. Revisiting DVWA\'s CSRF module after an '
    'earlier hardening exercise also demonstrated that the same principles apply in '
    'reverse: a correctly implemented defense measurably blocks the same class of attack '
    'without disrupting legitimate functionality. Together, these exercises reinforce '
    'that secure coding is not an afterthought but a set of concrete, well-understood '
    'practices that must be applied consistently across both front-end and back-end '
    'components.'
)

doc.save('INT309_Lab2_Report.docx')
print('Report generated: INT309_Lab2_Report.docx')
