from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INT309: Web Technologies and Database Security', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Lab 1: Introduction to Web Technologies and Understanding HTTP and HTTPS', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

name = doc.add_paragraph('Patrick Damian Ozimede')
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.runs[0].bold = True

doc.add_page_break()

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'This report documents the completion of Lab 1, covering the foundational concepts '
    'of web technologies, with a focus on the HTTP and HTTPS protocols. The lab explored '
    'the structure of a basic web page, the anatomy of HTTP requests and responses using '
    'browser developer tools, and the security guarantees provided by HTTPS through '
    'certificate inspection and TLS analysis.'
)

# Exercise 1
doc.add_heading('Exercise 1: Creating a Simple Web Page', level=1)
doc.add_paragraph(
    'A basic HTML page (index.html) was created, including a header, a paragraph, a '
    'placeholder image, and a link to an external resource. Note: the lab\'s original '
    'placeholder image URL (via.placeholder.com) has been discontinued; it was replaced '
    'with placehold.co, a functionally equivalent, currently active service.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 1 — index.html rendered in browser]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'This exercise reinforced how HTML uses semantic structure to organize content in a '
    'way that\'s both machine-readable and human-usable. The header element visually and '
    'structurally separates introductory content from the body, giving users an immediate '
    'sense of what the page is about. Paragraph tags break content into readable chunks '
    'rather than one dense block of text, which improves readability. The image tag '
    'demonstrates how visual media is embedded and referenced via a src attribute rather '
    'than stored inline, meaning a page\'s appearance depends on external resources being '
    'available — a point reinforced firsthand when the lab\'s original placeholder image '
    'URL failed to load because the service had been discontinued, requiring a working '
    'alternative to be substituted. Finally, the anchor tag shows how HTML enables '
    'navigation between resources, forming the basis of the interconnected web. Together, '
    'these elements illustrate that HTML isn\'t just about displaying text — it\'s about '
    'structuring content in a way that conveys meaning, hierarchy, and relationships to '
    'both browsers and users.'
)

# Exercise 2
doc.add_heading('Exercise 2: Analyzing HTTP Requests and Responses', level=1)
doc.add_paragraph(
    'The page was first inspected via the file:// protocol (opened directly from disk), '
    'and the Network tab in Firefox DevTools was used to examine the request. This showed '
    'no headers, since no real HTTP transaction occurs when loading a file locally.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 2 — Network tab showing HTTP requests]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3a — index.html via file:// — no headers]')

doc.add_paragraph(
    'To properly capture an HTTP transaction, a local server was started using '
    '"python3 -m http.server 8000" and the page was reloaded at http://127.0.0.1:8000. '
    'This produced a genuine request/response cycle, including a 200 OK on first load, '
    'a subsequent 304 Not Modified (demonstrating browser cache validation), and a benign '
    '404 for the automatically-requested favicon.ico.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3b-1 — index.html via HTTP — Response Headers]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3b-2 — index.html via HTTP — Request Headers]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3b-3 and 3b-4 — additional header detail]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'This exercise highlighted a distinction that isn\'t obvious until you see it '
    'firsthand: opening a file directly via the file:// protocol is not the same as '
    'loading it over HTTP, even though the rendered page looks identical. When index.html '
    'was opened locally, the Network tab reported no headers for the request, because no '
    'actual HTTP transaction occurred; the browser simply read bytes off disk. To properly '
    'observe HTTP behavior, a local server was started and the page reloaded over HTTP. '
    'This surfaced a genuine HTTP request/response cycle, complete with a status code, '
    'response headers, and caching behavior — including a 304 Not Modified response on '
    'the second load, which demonstrated the browser\'s cache validation mechanism in '
    'action. The 404 for favicon.ico was also observed and is benign, expected browser '
    'behavior rather than a page defect. Reviewing the Request and Response headers '
    '(Content-Type, Content-Length, Date, Server, etc.) reinforced why these matter for '
    'security: headers like Content-Type prevent MIME-sniffing attacks when set correctly, '
    'Cache-Control governs whether sensitive data might be inappropriately cached, and '
    'improperly configured or missing headers (e.g. absent X-Content-Type-Options or '
    'Content-Security-Policy) are a common source of web application vulnerabilities. This '
    'exercise made clear that understanding the full HTTP request/response cycle — not '
    'just the rendered page — is foundational to identifying and mitigating such risks.'
)

# Exercise 3
doc.add_heading('Exercise 3: Understanding HTTPS and Its Importance', level=1)
doc.add_paragraph(
    'A secure website (https://www.google.com) was accessed and its network traffic '
    'inspected. All 35 observed requests, including POST requests, carried the padlock '
    'indicator confirming HTTPS encryption.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 4 — HTTPS requests to google.com]')

doc.add_paragraph(
    'The site\'s certificate was inspected via the padlock icon, showing a certificate '
    'for *.google.com issued by Google Trust Services (WR2), chaining to the GTS Root R1 '
    'root certificate, valid from 10 Aug 2026 to 2 Nov 2026.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 5 — google.com certificate details]')

doc.add_paragraph(
    'The Page Info Security tab confirmed the connection was secured using TLS 1.3 with '
    'the TLS_AES_128_GCM_SHA256 cipher suite.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 6 — google.com Page Info Security tab, TLS version and cipher suite]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'This exercise demonstrated concretely what HTTPS provides over plain HTTP. Every '
    'request to google.com carried a padlock icon in the Network tab, confirming that all '
    'requests observed — including POST requests carrying data — were encrypted in '
    'transit. Inspecting the certificate showed a chain of trust: the leaf certificate for '
    '*.google.com was issued by Google Trust Services (WR2), which itself chains up to the '
    'GTS Root R1 root certificate. This chain is what allows a browser to verify the '
    'site\'s identity without having to trust every server individually — trust is '
    'inherited from a small set of root certificate authorities. The certificate\'s short '
    'validity window (roughly three months) reflects a broader industry shift toward '
    'short-lived certificates, which reduces the impact of key compromise and encourages '
    'automated renewal. The Page Info panel further confirmed the connection was secured '
    'using TLS 1.3 with the TLS_AES_128_GCM_SHA256 cipher suite — an authenticated '
    'encryption mode that protects both confidentiality and integrity of the data in '
    'transit. Together, these findings illustrate why HTTPS is essential: it prevents '
    'eavesdropping, protects against man-in-the-middle attacks through certificate '
    'validation, and preserves data integrity. By contrast, the HTTP request examined in '
    'Exercise 2 had none of these protections — data was transmitted in plaintext, with no '
    'verification of server identity and no protection against tampering.'
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab provided hands-on exposure to the foundational mechanics of web '
    'communication. Building a simple page reinforced HTML structure and semantics; '
    'inspecting HTTP traffic — both locally via file:// and over an actual HTTP server — '
    'clarified the real difference between a rendered page and a genuine network '
    'transaction, along with headers, status codes, and caching behavior; and analyzing '
    'HTTPS traffic to a live production site demonstrated encryption, certificate chains '
    'of trust, and modern TLS in practice. Together, these exercises establish the '
    'groundwork needed for the more advanced web and database security topics ahead in '
    'this course.'
)

doc.save('INT309_Lab1_Report.docx')
print('Report generated: INT309_Lab1_Report.docx')
