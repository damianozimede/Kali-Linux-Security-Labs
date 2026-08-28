from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('INT309: Web Technologies and Database Security', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Lab 3: Introduction to Databases and SQL', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

name = doc.add_paragraph('Patrick Damian Ozimede')
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.runs[0].bold = True

doc.add_page_break()

doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'This report documents the completion of Lab 3, which covered the fundamentals of '
    'relational databases using MySQL (MariaDB). A student_management_system database '
    'was designed and built, consisting of students, courses, and enrollments tables '
    'linked by foreign keys, followed by SQL operations to insert, retrieve, update, and '
    'delete data across those related tables.'
)

doc.add_heading('Environment Setup', level=1)
doc.add_paragraph(
    'phpMyAdmin was installed and configured to work with the MariaDB instance already '
    'running on Kali Linux from prior labs, rather than installing a separate XAMPP '
    'stack, since XAMPP bundles its own Apache, MySQL, and PHP versions that would have '
    'conflicted with the environment already in place. Two environment notes are worth '
    'recording: phpMyAdmin blocks blank-password logins by default regardless of the '
    'underlying database configuration, so the existing webapp application user '
    '(created in Lab 2) was used to log in instead of root. That user initially lacked '
    'privileges to create new databases, since it had only been granted access to a '
    'single database previously, this was resolved by granting broader privileges at the '
    'MySQL level and re-authenticating in phpMyAdmin to refresh its cached privilege '
    'information. Additionally, the students table could not be created successfully '
    'through phpMyAdmin\'s graphical table designer (it returned a "#1075 - Incorrect '
    'table definition" error despite the form appearing correctly filled in), so all '
    'table creation was instead performed directly through the SQL tab, which proved '
    'reliable and is documented as the approach used throughout this lab.'
)

doc.add_heading('Database Design Summary', level=1)
doc.add_paragraph(
    'The student_management_system database consists of three related tables. The '
    'students table stores individual student records (student_id, first_name, '
    'last_name, email, date_of_birth), with student_id as an auto-incrementing primary '
    'key and email enforced as unique. The courses table stores course offerings '
    '(course_id, course_name, credits), also with an auto-incrementing primary key. The '
    'enrollments table serves as a junction table modeling the many-to-many relationship '
    'between students and courses: each row links one student_id and one course_id (both '
    'enforced as foreign keys referencing their respective parent tables), along with an '
    'enrollment_date. This design allows a single student to enroll in multiple courses '
    'and a single course to have multiple students, while the foreign key constraints '
    'ensure that no enrollment can reference a student or course that does not actually '
    'exist.'
)

doc.add_heading('Exercise 1: Setting Up MySQL with phpMyAdmin', level=1)
doc.add_paragraph(
    'The student_management_system database was created via phpMyAdmin\'s Databases tab.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 1 - student_management_system database created in phpMyAdmin]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'A web application like a Student Management System needs a relational database '
    'because the data it manages, students, courses, and enrollments, is inherently '
    'interconnected rather than standalone. A single student can enroll in multiple '
    'courses, and a single course can have many students, this many-to-many relationship '
    'is exactly what relational databases are designed to model cleanly through separate '
    'tables linked by keys, rather than duplicating data or storing it in unstructured '
    'files. MySQL specifically offers several benefits for this kind of application: it '
    'enforces data integrity through constraints (primary keys, foreign keys, unique '
    'constraints), supports efficient querying across related tables through JOINs, and '
    'scales well for the kind of structured, relationship-heavy data that administrative '
    'systems typically handle. Additionally, MySQL\'s widespread adoption means extensive '
    'tooling, documentation, and community support are available, phpMyAdmin itself being '
    'one example of the ecosystem built around it. Using a relational database instead of, '
    'for example, storing everything in a single flat table, avoids data duplication (a '
    'student\'s information is stored once and referenced by ID elsewhere) and makes '
    'updates safer and more consistent.'
)

doc.add_heading('Exercise 2: Creating Tables and Defining Columns', level=1)
doc.add_paragraph(
    'The students, courses, and enrollments tables were created using SQL statements '
    'executed directly through phpMyAdmin\'s SQL tab, after the graphical table designer '
    'produced an error on the students table\'s primary key and auto-increment '
    'configuration.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 2 - students table created via SQL]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 3 - courses table created via SQL]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 4 - enrollments table created via SQL with foreign keys]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'Relationships between tables in a relational database are established through '
    'foreign keys, columns in one table that reference the primary key of another table. '
    'In this exercise, the enrollments table doesn\'t duplicate student or course '
    'details, instead it stores only student_id and course_id, which point back to the '
    'corresponding rows in students and courses. This design is what allows a single '
    'student to enroll in multiple courses and a single course to have multiple '
    'students, a many-to-many relationship modeled through this intermediate "junction" '
    'table. Foreign keys are important for maintaining data integrity because the '
    'database engine enforces them at the storage layer, it becomes impossible to insert '
    'an enrollment record referencing a student_id or course_id that doesn\'t actually '
    'exist in the respective table. This prevents orphaned or inconsistent data that '
    'could otherwise arise from application-level bugs or oversight. Foreign keys also '
    'communicate the intended structure of the data model itself, anyone examining the '
    'schema can immediately see how the tables relate to one another, which supports '
    'both correctness and long-term maintainability of the database.'
)

doc.add_heading('Exercise 3: Inserting Data into the Tables', level=1)
doc.add_paragraph(
    'Sample records were inserted into all three tables: three students, three courses, '
    'and four enrollment records linking them together.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 5 - sample data inserted into students table]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 6 - sample data inserted into courses table]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 7 - sample data inserted into enrollments table]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'Inserting data is what transforms an empty schema into a functioning database, the '
    'table structure defines the rules, but the data itself is what makes the system '
    'useful for actually managing students, courses, and enrollments. In this exercise, '
    'populating the three tables demonstrated how a well-designed schema pays off '
    'immediately: the enrollments table\'s foreign key constraints meant that only valid '
    'student_id and course_id values (already present in the students and courses '
    'tables) could be inserted, any attempt to reference a nonexistent student or course '
    'would have been rejected automatically. Data consistency is crucial when inserting '
    'related data across tables because a single inconsistency, such as an enrollment '
    'record pointing to a student or course that doesn\'t exist, can silently corrupt '
    'reports, break application logic, or produce misleading results in later queries. '
    'This becomes increasingly important as a database grows: manually verifying every '
    'reference isn\'t feasible at scale, which is why enforcing consistency through '
    'database-level constraints, rather than relying solely on application code, is a '
    'foundational best practice.'
)

doc.add_heading('Exercise 4: Querying the Database', level=1)
doc.add_paragraph(
    'Four SQL operations were performed to demonstrate the full range of basic query '
    'types: a SELECT retrieving all students, a multi-table JOIN retrieving students '
    'enrolled in a specific course, an UPDATE modifying a student\'s email address, and '
    'a DELETE removing a specific enrollment record.'
)
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 8 - SELECT * FROM students]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 9 - JOIN query, students enrolled in Web Development]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 10 - UPDATE query, John Doe email changed]')
doc.add_paragraph('[INSERT SCREENSHOT: Screenshot 11 - DELETE query, Jane Smith enrollment removed]')

doc.add_heading('Reflection', level=2)
doc.add_paragraph(
    'SQL allows querying and manipulating data in a relational database through a small '
    'set of core commands, SELECT for retrieval, INSERT for adding, UPDATE for '
    'modifying, and DELETE for removing, each operating declaratively, meaning the query '
    'describes the desired result rather than the exact steps to obtain it. This exercise '
    'demonstrated all four in a realistic context: retrieving all students, updating a '
    'specific student\'s contact information, and removing a specific enrollment record. '
    'JOINs are what make relational databases genuinely useful beyond simple single-table '
    'storage, they facilitate retrieving related information from multiple tables by '
    'matching rows based on a shared key, in this case connecting students and courses '
    'through the enrollments junction table. Without the JOIN, answering a question like '
    '"which students are enrolled in Web Development" would require manually '
    'cross-referencing three separate result sets in application code, an approach that '
    'is both inefficient and error-prone. By letting the database engine perform this '
    'matching internally, JOINs push the relational logic down to where it belongs, the '
    'data layer, keeping application code simpler and ensuring the relationships defined '
    'by foreign keys are actually put to use.'
)

doc.add_heading('Potential Improvements', level=1)
improvements = [
    'Add NOT NULL constraints to fields such as first_name, last_name, and course_name to prevent incomplete records.',
    'Add an index on enrollment_date if the application will frequently query enrollments by date range.',
    'Consider a composite unique constraint on (student_id, course_id) in the enrollments table to prevent a student being enrolled in the same course more than once.',
    'Introduce a status column on enrollments (e.g., active, completed, dropped) rather than deleting records outright, preserving historical enrollment data.',
    'Apply the principle of least privilege to database accounts used by the application, rather than granting broad administrative privileges as was done here for lab convenience.',
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This lab provided hands-on experience designing and populating a relational '
    'database from the ground up, and demonstrated the four fundamental SQL operations '
    'in a realistic multi-table scenario. Building the student_management_system '
    'database reinforced how foreign keys and normalized table design work together to '
    'maintain data integrity, while the querying exercises, particularly the JOIN, '
    'illustrated why relational databases remain the standard choice for applications '
    'with interconnected data. Understanding these fundamentals is essential groundwork '
    'for the more advanced database security topics ahead in this course.'
)

doc.save('INT309_Lab3_Report.docx')
print('Report generated: INT309_Lab3_Report.docx')
