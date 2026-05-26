from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Web Application Security Testing Report")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
info = [
    ("Target:", "http://zero.webappsecurity.com"),
    ("Prepared by:", "Damian Patrick Ozimede"),
    ("Course:", "INT302: Kali Linux Tools and System Security"),
    ("Date:", "25 May 2026"),
    ("Tools Used:", "Burp Suite Community Edition v2026.3.2, OWASP ZAP 2.17.0"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + value)

doc.add_paragraph()
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph("A web application security assessment was conducted against zero.webappsecurity.com, a deliberately vulnerable banking application used for security testing. Using Burp Suite and OWASP ZAP, a total of 15 vulnerabilities were identified ranging from critical SQL injection to missing security headers. Manual fuzzing confirmed that the search functionality is vulnerable to SQL injection with payloads being reflected back without sanitisation.")

doc.add_heading("2. Methodology", level=1)
for item in [
    "Traffic Interception — Burp Suite and OWASP ZAP configured as proxies to capture all HTTP traffic",
    "Site Mapping — Manual browsing and automated spidering to discover all available endpoints",
    "Automated Scanning — OWASP ZAP automated scan to identify vulnerabilities",
    "Manual Fuzzing — jbrofuzz SQL injection payload library used against the search parameter",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3. Identified Vulnerabilities", level=1)
vuln_table = doc.add_table(rows=1, cols=4)
vuln_table.style = "Table Grid"
hdr = vuln_table.rows[0].cells
hdr[0].text = "#"
hdr[1].text = "Vulnerability"
hdr[2].text = "Severity"
hdr[3].text = "Location"
vulns = [
    ("1", "SQL Injection - SQLite Time Based", "High", "/search?searchTerm="),
    ("2", "Absence of Anti-CSRF Tokens", "Medium", "Multiple forms"),
    ("3", "Content Security Policy Not Set", "Medium", "All pages"),
    ("4", "Cross-Domain Misconfiguration", "Medium", "Site-wide"),
    ("5", "Hidden File Found", "Medium", "Server"),
    ("6", "Missing Anti-Clickjacking Header", "Medium", "All pages"),
    ("7", "Vulnerable JS Library", "Medium", "jQuery 1.7.2, 1.8.2"),
    ("8", "Cookie without SameSite Attribute", "Medium", "Session cookie"),
    ("9", "In Page Banner Information Leak", "Medium", "Multiple pages"),
    ("10", "Server Version Information Leak", "Medium", "Server header"),
    ("11", "X-Content-Type-Options Missing", "Medium", "All pages"),
    ("12", "Authentication Request Identified", "Informational", "/login.html"),
    ("13", "Information Disclosure - Comments", "Informational", "Source code"),
    ("14", "Modern Web Application", "Informational", "Site-wide"),
    ("15", "User Agent Fuzzer", "Informational", "Site-wide"),
]
for v in vulns:
    row = vuln_table.add_row().cells
    row[0].text = v[0]
    row[1].text = v[1]
    row[2].text = v[2]
    row[3].text = v[3]

doc.add_heading("4. Key Vulnerability Details", level=1)

doc.add_heading("SQL Injection", level=2)
doc.add_paragraph("The /search?searchTerm= parameter accepts unsanitised user input. Fuzzing with jbrofuzz confirmed multiple SQL injection payloads were reflected including 'a' or 1=1; --' and '' and 1=0) union'. An attacker could extract the entire database contents, modify data, or bypass authentication.")

doc.add_heading("Missing Security Headers", level=2)
doc.add_paragraph("The application lacks Content-Security-Policy, X-Frame-Options, and X-Content-Type-Options headers, leaving it vulnerable to XSS, clickjacking, and MIME-type sniffing attacks.")

doc.add_heading("Vulnerable JavaScript Libraries", level=2)
doc.add_paragraph("jQuery versions 1.7.2 and 1.8.2 are in use. Both contain known XSS vulnerabilities that have been publicly disclosed and patched in later versions.")

doc.add_heading("Server Information Disclosure", level=2)
doc.add_paragraph("The Server header reveals Apache-Coyote/1.1, giving attackers precise information about the server technology to target known exploits.")

doc.add_heading("5. Remediation Recommendations", level=1)
rem_table = doc.add_table(rows=1, cols=2)
rem_table.style = "Table Grid"
rem_table.rows[0].cells[0].text = "Vulnerability"
rem_table.rows[0].cells[1].text = "Recommendation"
rems = [
    ("SQL Injection", "Implement parameterised queries and input validation"),
    ("Missing CSP", "Add Content-Security-Policy header to all responses"),
    ("CSRF Tokens", "Implement anti-CSRF tokens on all forms"),
    ("Clickjacking", "Add X-Frame-Options: DENY header"),
    ("Vulnerable jQuery", "Upgrade to latest jQuery version"),
    ("Cookie SameSite", "Set SameSite=Strict on all session cookies"),
    ("Server Header", "Remove or obscure the Server response header"),
    ("Hidden Files", "Audit and remove unnecessary files from the server"),
]
for r in rems:
    row = rem_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]

doc.add_heading("6. Conclusion", level=1)
doc.add_paragraph("The assessment revealed significant security weaknesses in zero.webappsecurity.com, most critically a confirmed SQL injection vulnerability in the search functionality. The combination of Burp Suite for traffic interception and manual analysis, and OWASP ZAP for automated scanning and fuzzing, provided comprehensive coverage of the application's attack surface. Immediate remediation of the SQL injection vulnerability is recommended as a priority.")

doc.save("/home/legion/Kali-Linux-Security-Labs/Lab8/Lab8_Security_Testing_Report.docx")
print("Report saved!")
