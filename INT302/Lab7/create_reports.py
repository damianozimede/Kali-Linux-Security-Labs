from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── FORENSIC REPORT ──────────────────────────────────────────────────────────
doc = Document()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BookWorld Network Breach - Forensic Investigation Report")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
info = [
    ("Prepared by:", "Damian Patrick Ozimede"),
    ("Course:", "INT302: Kali Linux Tools and System Security"),
    ("Date:", "15 March 2024"),
    ("Classification:", "Confidential"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + value)

doc.add_paragraph()
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph("BookWorld suffered a targeted cyberattack on 15 March 2024 between 12:01:36 and 12:24:18 UTC. A single external threat actor operating from IP address 111.224.250.131 conducted a systematic attack against BookWorld's web application hosted at 73.124.22.98. The attacker exploited a SQL injection vulnerability in the site's search functionality, successfully enumerated the MySQL database, brute-forced the admin panel using weak credentials, and planted a PHP web shell for persistent access. Customer data stored in the database was likely compromised during the exploitation phase.")

doc.add_heading("2. Attack Timeline", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Time (UTC)"
hdr[1].text = "Phase"
hdr[2].text = "Activity"
rows = [
    ("11:39:21", "Capture Begins", "Normal traffic observed"),
    ("12:00:58", "Legitimate Activity", "IP 170.40.150.126 performs normal book searches"),
    ("12:01:36", "Initial Access", "Attacker IP 111.224.250.131 connects"),
    ("12:01:36-12:02:00", "Reconnaissance", "Site enumeration using GoBuster"),
    ("12:02:23", "Vulnerability Discovery", "SQL injection vulnerability identified in /search.php"),
    ("12:04:04-12:07:34", "Exploitation", "Automated SQLMap SQL injection attack"),
    ("12:12:20", "DB Enumeration", "MySQL file enumeration using GoBuster"),
    ("12:17:35-12:17:36", "Credential Attack", "Admin panel brute-forced — admin/admin123!"),
    ("12:24:18", "Persistence", "PHP web shell uploaded to /admin/index.php"),
]
for r in rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_heading("3. Entry Point and Exploited Services", level=1)
doc.add_paragraph("Primary Entry Point: /search.php — SQL injection vulnerability in the search parameter allowed the attacker to interact directly with the backend MySQL database without authentication.")
doc.add_paragraph("Secondary Entry Point: /admin/login.php — weak administrator credentials (admin/admin123!) allowed the attacker to gain full admin panel access through credential guessing.")
doc.add_paragraph("Persistence Mechanism: /admin/index.php — a PHP web shell was uploaded via the admin panel, giving the attacker persistent remote control of the server.")

doc.add_heading("4. Affected Systems and Data", level=1)
doc.add_paragraph("Affected Server: 73.124.22.98 (Apache/2.4.52 on Ubuntu Linux)")
for item in ["Web application (/search.php)", "Admin panel (/admin/login.php, /admin/index.php)", "MySQL database backend"]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph("Data at Risk:")
for item in ["Customer personal information stored in the MySQL database", "Administrator credentials", "Internal server configuration and file structure"]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("5. Indicators of Compromise (IoCs)", level=1)
ioc_table = doc.add_table(rows=1, cols=2)
ioc_table.style = "Table Grid"
ioc_table.rows[0].cells[0].text = "Category"
ioc_table.rows[0].cells[1].text = "Detail"
iocs = [
    ("Malicious IP", "111.224.250.131 — primary attacker IP"),
    ("Suspicious IP", "170.40.150.126 — secondary IP, possible reconnaissance"),
    ("Targeted Endpoints", "/search.php?search=, /admin/login.php, /admin/index.php, /.mysql_history"),
    ("Tools Identified", "SQLMap — automated SQL injection; GoBuster 3.6 — directory enumeration"),
    ("Compromised Credentials", "Username: admin | Password: admin123!"),
    ("Attack Protocol", "HTTP (port 80) — all traffic unencrypted"),
]
for i in iocs:
    row = ioc_table.add_row().cells
    row[0].text = i[0]
    row[1].text = i[1]

doc.add_heading("6. Supporting Evidence", level=1)
evidence = [
    ("88,484 packets from attacker IP", "Wireshark Conversations — IPv4 tab"),
    ("SQL injection payloads", "Packets 357-681, /search.php?search= parameter"),
    ("GoBuster enumeration", "Packets 2008-2026, User-Agent: gobuster/3.6"),
    ("Admin login attempts", "Packets 655, 658, 659, 661 — POST /admin/login.php"),
    ("Credentials in plaintext", "Packet 661 — username=admin&password=admin123!"),
    ("PHP web shell upload", "Packet 88757 — POST /admin/index.php (application/x-php)"),
    ("MySQL targeting", "63 packets containing MySQL references"),
    ("No DNS traffic", "0 DNS packets — attacker had prior knowledge of target IP"),
]
ev_table = doc.add_table(rows=1, cols=2)
ev_table.style = "Table Grid"
ev_table.rows[0].cells[0].text = "Evidence"
ev_table.rows[0].cells[1].text = "Reference"
for e in evidence:
    row = ev_table.add_row().cells
    row[0].text = e[0]
    row[1].text = e[1]

doc.add_heading("7. Remediation Recommendations", level=1)
recs = [
    "Patch SQL Injection — implement parameterised queries and input validation on all user-supplied fields",
    "Enforce Strong Passwords — replace weak admin credentials and implement password complexity policy",
    "Enable Account Lockout — block accounts after 3 failed login attempts",
    "Implement HTTPS — encrypt all web traffic",
    "Remove Web Shell — investigate and remove /admin/index.php and audit all server files",
    "Block Attacker IP — blacklist 111.224.250.131 at the firewall level",
    "Deploy WAF — implement a Web Application Firewall to block SQL injection and enumeration",
    "Enable DNS Logging — implement DNS monitoring to detect future C2 activity",
    "Audit Admin Panel — restrict /admin/ directory access by IP whitelist",
    "Incident Response — notify affected customers of potential data breach per data protection regulations",
]
for r in recs:
    doc.add_paragraph(r, style="List Bullet")

doc.save("/home/legion/Kali-Linux-Security-Labs/Lab7/Lab7_Forensic_Report.docx")
print("Report saved!")

# ── EXECUTIVE PRESENTATION ────────────────────────────────────────────────────
doc2 = Document()

title2 = doc2.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = title2.add_run("BookWorld Security Incident Briefing")
run2.bold = True
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

for line in ["Presented to: BookWorld Executive Team", "Date: 15 March 2024", "Presenter: Damian Patrick Ozimede, Digital Forensics Team"]:
    p = doc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line)

doc2.add_heading("Slide 1: What Happened?", level=1)
doc2.add_paragraph("On the morning of 15 March 2024, BookWorld was targeted by a deliberate and sophisticated cyberattack. An external attacker gained unauthorised access to your web application, extracted data from your customer database, obtained full administrative control of your web server, and planted a backdoor for future access.")
doc2.add_paragraph("The attack lasted approximately 23 minutes and went undetected until forensic analysis was conducted.")

doc2.add_heading("Slide 2: How Did They Get In?", level=1)
doc2.add_paragraph("The attacker exploited two weaknesses:")
p = doc2.add_paragraph()
p.add_run("Weakness 1: Vulnerable Search Box").bold = True
doc2.add_paragraph("Your website's book search feature was not properly secured. The attacker used it as a doorway to access your customer database directly.")
p = doc2.add_paragraph()
p.add_run("Weakness 2: Weak Admin Password").bold = True
doc2.add_paragraph("Your administrator account was protected by a simple, guessable password. The attacker gained full admin access within four attempts.")
p = doc2.add_paragraph()
p.add_run("NOTE: Neither of these required advanced hacking skills — both are entirely preventable.").bold = True

doc2.add_heading("Slide 3: What Was Affected?", level=1)
affected_table = doc2.add_table(rows=1, cols=2)
affected_table.style = "Table Grid"
affected_table.rows[0].cells[0].text = "Area"
affected_table.rows[0].cells[1].text = "Status"
affected = [
    ("Customer database", "Likely compromised"),
    ("Admin panel", "Fully breached"),
    ("Web server", "Backdoor planted"),
    ("Customer passwords/emails", "At risk"),
    ("Payment data", "Under investigation"),
]
for a in affected:
    row = affected_table.add_row().cells
    row[0].text = a[0]
    row[1].text = a[1]

doc2.add_heading("Slide 4: How Bad Is It?", level=1)
for item in [
    "The attacker had full administrative control of your web server by 12:24 AM",
    "A backdoor was planted: They can return at any time without needing to repeat the attack",
    "All traffic was unencrypted: Everything transmitted was readable",
    "Customer personal data stored in your database was directly accessible to the attacker",
]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_heading("Slide 5: What Are We Doing About It?", level=1)
doc2.add_paragraph("Immediate Actions (This Week):").runs[0].bold = True
for item in ["Remove the attacker's backdoor from the server", "Block the attacker's IP address", "Reset all administrator credentials", "Take the vulnerable search feature offline until patched"]:
    doc2.add_paragraph(item, style="List Bullet")
doc2.add_paragraph("Short Term (This Month):").runs[0].bold = True
for item in ["Fix the search vulnerability permanently", "Encrypt all website traffic with HTTPS", "Implement account lockout after failed logins", "Deploy a Web Application Firewall"]:
    doc2.add_paragraph(item, style="List Bullet")
doc2.add_paragraph("Long Term (This Quarter):").runs[0].bold = True
for item in ["Full security audit of all web applications", "Staff security awareness training", "Regular penetration testing", "Customer breach notification process"]:
    doc2.add_paragraph(item, style="List Bullet")

doc2.add_heading("Slide 6: Business Impact and Risk", level=1)
risk_table = doc2.add_table(rows=1, cols=2)
risk_table.style = "Table Grid"
risk_table.rows[0].cells[0].text = "Risk"
risk_table.rows[0].cells[1].text = "Detail"
risks = [
    ("Regulatory", "Potential GDPR/data protection violation if customer data was exposed"),
    ("Reputational", "Customer trust at risk if breach becomes public"),
    ("Financial", "Potential fines, litigation, and remediation costs"),
    ("Operational", "Server integrity compromised — backdoor still active"),
]
for r in risks:
    row = risk_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
p = doc2.add_paragraph()
p.add_run("The longer you delay remediation, the greater your exposure.").bold = True

doc2.add_heading("Slide 7: Key Message to Executives", level=1)
doc2.add_paragraph("This attack succeeded not because of sophisticated hacking but because of basic security gaps that are well known and entirely fixable. The good news is that forensic analysis has given the organisation a complete picture of what happened. We know exactly how they got in, what they did, and how to stop it from happening again.")
p = doc2.add_paragraph()
p.add_run("We recommend immediate approval of the remediation plan to protect your customers, reputation, and business.").bold = True

doc2.save("/home/legion/Kali-Linux-Security-Labs/Lab7/Lab7_Executive_Presentation.docx")
print("Presentation saved!")
