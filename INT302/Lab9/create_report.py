from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Reverse Shell via Netcat Using DVWA Command Execution")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
info = [
    ("Prepared by:", "Damian Patrick Ozimede"),
    ("Course:", "INT302: Kali Linux Tools and System Security"),
    ("Lab:", "Lab 9"),
    ("Date:", "25 May 2026"),
    ("Tools Used:", "DVWA, Netcat, Kali Linux, OWASP BWA VM"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + value)

doc.add_paragraph()

# Exercise 1
doc.add_heading("Exercise 1: DVWA Target IP", level=1)
doc.add_paragraph("DVWA Target IP Address: 192.168.5.130")
doc.add_paragraph("DVWA URL: http://192.168.5.130/dvwa")
doc.add_paragraph("Kali Linux (Attacker) IP: 192.168.5.139")

# Exercise 2
doc.add_heading("Exercise 2: Why Does whoami Execute Successfully?", level=1)
doc.add_paragraph("The whoami command executes successfully because DVWA's Command Execution page passes user input directly to the operating system's ping command without any sanitisation or validation. The semicolon (;) is a Linux command separator that tells the shell to execute the next command regardless of the previous result. So '127.0.0.1; whoami' becomes two separate commands: ping 127.0.0.1 and whoami. This is a classic OS Command Injection vulnerability — the application blindly trusts user input and passes it to a system shell, allowing any command to be appended and executed.")

# Exercise 3
doc.add_heading("Exercise 3: Purpose of Netcat in Listening Mode", level=1)
doc.add_paragraph("Netcat in listening mode (nc -lvnp 4444) acts as a receiver waiting for an incoming connection on port 4444. When the victim machine connects back, Netcat establishes a communication channel giving the attacker an interactive shell on the victim machine.")
doc.add_paragraph("The flags used:")
for item in [
    "-l — listen for incoming connections",
    "-v — verbose output to show connection details",
    "-n — no DNS resolution, use IP addresses only",
    "-p 4444 — listen on port 4444",
]:
    doc.add_paragraph(item, style="List Bullet")

# Exercise 4
doc.add_heading("Exercise 4: Reverse Shell User Account", level=1)
doc.add_paragraph("The reverse shell connected successfully with the following output:")
doc.add_paragraph("connect to [192.168.5.139] from (UNKNOWN) [192.168.5.130] 41174")
doc.add_paragraph("Commands executed and results:")
result_table = doc.add_table(rows=1, cols=2)
result_table.style = "Table Grid"
result_table.rows[0].cells[0].text = "Command"
result_table.rows[0].cells[1].text = "Output"
results = [
    ("whoami", "www-data"),
    ("id", "uid=33(www-data) gid=33(www-data) groups=33(www-data)"),
    ("uname -a", "Linux owaspbwa 2.6.32-25-generic-pae #44-Ubuntu SMP i686 GNU/Linux"),
]
for r in results:
    row = result_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("The reverse shell is running as www-data (uid=33, gid=33). This is the Apache web server user account, confirming the shell was spawned through the web application.")

# Exercise 5
doc.add_heading("Exercise 5: Post-Exploitation Validation", level=1)
doc.add_paragraph("Post-exploitation commands executed:")
post_table = doc.add_table(rows=1, cols=2)
post_table.style = "Table Grid"
post_table.rows[0].cells[0].text = "Command"
post_table.rows[0].cells[1].text = "Output"
post_results = [
    ("pwd", "/owaspbwa/dvwa-git/vulnerabilities/exec"),
    ("ls", "help, index.php, source"),
    ("ip addr", "192.168.5.130 on eth0"),
    ("cat /etc/issue", "OWASP Broken Web Applications VM Version 1.2"),
    ("id", "uid=33(www-data) gid=33(www-data)"),
]
for r in post_results:
    row = post_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Why a reverse shell is more powerful than simple command execution:")
doc.add_paragraph("Simple command execution only returns output in the web browser one command at a time. A reverse shell provides persistent interactive access with full terminal session on the victim machine, full command capability without going through the browser, complete file system access, network information, and stealth since the connection originates from victim to attacker bypassing inbound firewall rules.")

# Exercise 6
doc.add_heading("Exercise 6: Security Impact and Risk Analysis", level=1)
p = doc.add_paragraph()
p.add_run("What could an attacker do with persistent shell access?").bold = True
for item in [
    "Read sensitive configuration files containing database credentials",
    "Access and exfiltrate customer data from the web application",
    "Plant backdoors or web shells for persistent future access",
    "Use the compromised server to attack other internal systems",
    "Modify web application files to serve malware to visitors",
    "Create new user accounts for persistent access",
]:
    doc.add_paragraph(item, style="List Bullet")

p = doc.add_paragraph()
p.add_run("How could this lead to full system compromise?").bold = True
doc.add_paragraph("The attacker currently has www-data access but can escalate privileges by exploiting local kernel vulnerabilities (kernel 2.6.32 is outdated), reading configuration files for database credentials, or exploiting SUID binaries. Full root access would give complete control of the entire system including all hosted applications.")

p = doc.add_paragraph()
p.add_run("Why is command execution considered a critical vulnerability?").bold = True
doc.add_paragraph("Command execution is classified as critical because it bypasses all application-level security controls and gives the attacker direct access to the underlying operating system. Unlike SQL injection which is limited to the database, command execution exposes the entire server, its file system, network interfaces, and all running services. It is listed in the OWASP Top 10 as one of the most dangerous web application vulnerabilities.")

# Exercise 7
doc.add_heading("Exercise 7: Defensive Controls", level=1)
p = doc.add_paragraph()
p.add_run("Input Validation").bold = True
doc.add_paragraph("Always check and clean user input before using it. Only allow expected characters like numbers and dots for an IP address field. Reject anything containing special characters like semicolons, pipes, or ampersands.")

p = doc.add_paragraph()
p.add_run("Command Whitelisting").bold = True
doc.add_paragraph("Implement a whitelist of allowed commands and inputs. Validate that input matches a strict IP address regex pattern and reject everything else. This ensures only legitimate IP addresses are accepted and no command chaining is possible.")

p = doc.add_paragraph()
p.add_run("Use of Safe APIs").bold = True
doc.add_paragraph("Instead of passing user input to system shell commands, use programming libraries that perform the same function safely. For example, use a Python socket library to ping instead of calling the OS ping command directly.")

p = doc.add_paragraph()
p.add_run("Web Application Firewall (WAF)").bold = True
doc.add_paragraph("A WAF sits between the user and the web application and automatically blocks requests containing known attack patterns like command injection payloads before they reach the server.")

doc.save("/home/legion/Kali-Linux-Security-Labs/INT302/Lab9/Lab9_Reverse_Shell_Report.docx")
print("Report saved!")
