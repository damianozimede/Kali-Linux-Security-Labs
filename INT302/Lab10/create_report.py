from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("DNS Query Tools and SMB Enumeration Report")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
info = [
    ("Prepared by:", "Damian Patrick Ozimede"),
    ("Course:", "INT302: Kali Linux Tools and System Security"),
    ("Lab:", "Lab 10"),
    ("Date:", "27 May 2026"),
    ("Tools Used:", "nslookup, host, dig, enum4linux"),
    ("Target Domain:", "example.com"),
    ("Target IP:", "192.168.5.130 (OWASP BWA VM)"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + value)

doc.add_paragraph()

# Exercise 1
doc.add_heading("Exercise 1: nslookup Results", level=1)
doc.add_paragraph("Command used: nslookup example.com")
doc.add_paragraph("The following information was obtained:")
nslookup_table = doc.add_table(rows=1, cols=2)
nslookup_table.style = "Table Grid"
nslookup_table.rows[0].cells[0].text = "Field"
nslookup_table.rows[0].cells[1].text = "Value"
nslookup_rows = [
    ("DNS Server Used", "192.168.5.2 (port 53)"),
    ("Query Type", "Non-authoritative answer"),
    ("IPv4 Address 1", "172.66.147.243"),
    ("IPv4 Address 2", "104.20.23.154"),
    ("IPv6 Address 1", "2606:4700:10::6814:179a"),
    ("IPv6 Address 2", "2606:4700:10::ac42:93f3"),
]
for r in nslookup_rows:
    row = nslookup_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Key Observations: The response came from DNS server 192.168.5.2 (local resolver). The non-authoritative answer means it came from cache, not the domain's own DNS server. Two IPv4 and two IPv6 addresses suggest load balancing or CDN usage, likely Cloudflare based on the IP ranges.")

# Exercise 2
doc.add_heading("Exercise 2: host vs nslookup Comparison", level=1)
doc.add_paragraph("Command used: host example.com")
host_table = doc.add_table(rows=1, cols=3)
host_table.style = "Table Grid"
host_table.rows[0].cells[0].text = "Feature"
host_table.rows[0].cells[1].text = "nslookup"
host_table.rows[0].cells[2].text = "host"
host_rows = [
    ("IPv4 addresses", "Yes", "Yes"),
    ("IPv6 addresses", "Yes", "Yes"),
    ("DNS server shown", "Yes", "No"),
    ("Mail (MX) record", "No", "Yes"),
    ("HTTP service bindings", "No", "Yes"),
    ("Output format", "Verbose", "Cleaner/simpler"),
]
for r in host_rows:
    row = host_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
doc.add_paragraph("Additional findings from host: Mail record showed 'example.com mail is handled by 0 .' meaning no mail server is configured. HTTP service bindings revealed the site supports HTTP/2 (h2) protocol. The host command provides a cleaner output and automatically shows MX and HTTP service records without extra flags, making it quicker for initial reconnaissance.")

# Exercise 3
doc.add_heading("Exercise 3: dig Analysis", level=1)
doc.add_paragraph("Command used: dig example.com")
dig_table = doc.add_table(rows=1, cols=3)
dig_table.style = "Table Grid"
dig_table.rows[0].cells[0].text = "Field"
dig_table.rows[0].cells[1].text = "Value"
dig_table.rows[0].cells[2].text = "Significance"
dig_rows = [
    ("DiG version", "9.20.22-1-Debian", "Tool version confirmed"),
    ("Status", "NOERROR", "Query successful"),
    ("Flags", "qr rd ra", "Query Response, Recursion Desired, Recursion Available"),
    ("TTL", "5 seconds", "Records expire quickly — typical of CDN"),
    ("Query time", "43 msec", "DNS response speed"),
    ("Protocol", "UDP", "DNS uses UDP by default"),
    ("Timestamp", "Wed May 27 09:15:52 WAT 2026", "Exact query time recorded"),
    ("Message size", "72 bytes", "Size of DNS response"),
]
for r in dig_rows:
    row = dig_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
doc.add_paragraph("Key advantages of dig: Shows TTL values confirming Cloudflare CDN, shows exact query timestamps important for forensic reporting, shows DNS flags revealing server capabilities, and is the most detailed and scriptable tool preferred by professional penetration testers.")

# Exercise 4
doc.add_heading("Exercise 4: Advanced DNS Record Types", level=1)
doc.add_paragraph("Commands used: dig example.com MX and dig example.com TXT")
p = doc.add_paragraph()
p.add_run("MX Record Results:").bold = True
doc.add_paragraph("MX Record: 0 . (No mail server configured). TTL: 5 seconds. Query time: 2111 msec.")
p = doc.add_paragraph()
p.add_run("TXT Record Results:").bold = True
txt_table = doc.add_table(rows=1, cols=2)
txt_table.style = "Table Grid"
txt_table.rows[0].cells[0].text = "TXT Record"
txt_table.rows[0].cells[1].text = "Significance"
txt_rows = [
    ("_k2n1y4vw3qtb4skdx9e7dxt97qrmmq9", "Domain verification token — reveals third-party service registration"),
    ("v=spf1 -all", "SPF hardfail — no server authorised to send email for this domain"),
]
for r in txt_rows:
    row = txt_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Penetration Testing Value: MX records reveal email infrastructure for phishing or spoofing opportunities. SPF records reveal email security posture. TXT verification tokens expose third-party services in use. In a real penetration test, TXT records can reveal cloud services, email providers, and technology stack details that inform attack strategies.")

# Exercise 5
doc.add_heading("Exercise 5: enum4linux Full Enumeration Results", level=1)
doc.add_paragraph("Command used: enum4linux -a 192.168.5.130")
p = doc.add_paragraph()
p.add_run("Users Found:").bold = True
users_table = doc.add_table(rows=1, cols=3)
users_table.style = "Table Grid"
users_table.rows[0].cells[0].text = "Account"
users_table.rows[0].cells[1].text = "RID"
users_table.rows[0].cells[2].text = "Description"
users_rows = [
    ("nobody", "0x1f5", "Default nobody account"),
    ("user", "0x3e8", "Standard user account"),
    ("root", "0x3e9", "Root/admin account"),
]
for r in users_rows:
    row = users_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Shares Found:").bold = True
shares_table = doc.add_table(rows=1, cols=3)
shares_table.style = "Table Grid"
shares_table.rows[0].cells[0].text = "Share"
shares_table.rows[0].cells[1].text = "Type"
shares_table.rows[0].cells[2].text = "Comment"
shares_rows = [
    ("print$", "Disk", "Printer Drivers"),
    ("apache", "Disk", "Apache Web Server Root"),
    ("tomcat", "Disk", "Tomcat6 Root"),
    ("var", "Disk", "/var directory"),
    ("etc", "Disk", "/etc directory"),
    ("usr", "Disk", "/usr directory"),
    ("owaspbwa", "Disk", "/owaspbwa directory"),
    ("IPC$", "IPC", "IPC Service"),
]
for r in shares_rows:
    row = shares_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("OS Information:").bold = True
for item in [
    "Server: OWASPBWA (Samba, Ubuntu)",
    "OS Version: 4.9",
    "Platform ID: 500",
    "Workgroup: WORKGROUP",
    "Anonymous sessions allowed — no username or password required",
    "MAC Address: 00-00-00-00-00-00",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph("Critical Finding: The server allows sessions with empty username and password — a serious security misconfiguration that allows any attacker to enumerate the system without credentials.")

# Exercise 6
doc.add_heading("Exercise 6: DNS vs SMB Enumeration Comparison", level=1)
comp_table = doc.add_table(rows=1, cols=3)
comp_table.style = "Table Grid"
comp_table.rows[0].cells[0].text = "Information"
comp_table.rows[0].cells[1].text = "DNS Queries"
comp_table.rows[0].cells[2].text = "SMB Enumeration"
comp_rows = [
    ("IP Addresses", "Yes", "No"),
    ("Mail servers", "Yes", "No"),
    ("Domain verification", "Yes", "No"),
    ("OS information", "No", "Yes"),
    ("User accounts", "No", "Yes"),
    ("Shared directories", "No", "Yes"),
    ("Server software", "No", "Yes"),
]
for r in comp_rows:
    row = comp_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
doc.add_paragraph("Key Insights: DNS queries reveal the external-facing infrastructure including IP addresses, mail servers, and third-party services. SMB enumeration reveals internal system details including user accounts, shared directories, and OS information. Combined, both techniques provide a complete picture of the target. The discovery of root and user accounts via SMB is particularly dangerous as these can be targeted for brute-force attacks. The etc and var shares being exposed suggests sensitive configuration files may be accessible.")

# Exercise 7
doc.add_heading("Exercise 7: Methodology, Tools, and Key Insights", level=1)
p = doc.add_paragraph()
p.add_run("Methodology:").bold = True
for item in [
    "Phase 1 — DNS Reconnaissance: Used nslookup, host, and dig to gather domain information including IP addresses, mail servers, SPF records, and TXT verification tokens",
    "Phase 2 — SMB Enumeration: Used enum4linux with the -a flag to perform full enumeration of the target system including users, shares, OS information, and password policy",
    "Phase 3 — Analysis: Combined findings from both phases to build a complete picture of the target network",
]:
    doc.add_paragraph(item, style="List Bullet")

p = doc.add_paragraph()
p.add_run("Tools Employed:").bold = True
for item in [
    "nslookup — basic DNS queries, good for quick IP lookups",
    "host — cleaner DNS output with automatic MX and HTTP service records",
    "dig — most detailed DNS tool with timing, flags, and TTL information",
    "enum4linux — comprehensive SMB enumeration tool for users, shares, and OS details",
]:
    doc.add_paragraph(item, style="List Bullet")

p = doc.add_paragraph()
p.add_run("Key Insights for Penetration Testing:").bold = True
for item in [
    "DNS enumeration revealed example.com uses Cloudflare CDN with no mail server configured and strict SPF policy",
    "SMB enumeration revealed three user accounts (nobody, user, root) on the OWASP BWA VM",
    "Anonymous SMB access is enabled — a critical misconfiguration allowing unauthenticated enumeration",
    "Seven SMB shares exposed including sensitive directories like /etc, /var, and /usr",
    "The root account being enumerable via SMB makes it a prime target for brute-force attacks",
    "Combining DNS and SMB data gives attackers a complete map of both external and internal attack surfaces",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.save("/home/legion/Kali-Linux-Security-Labs/INT302/Lab10/Lab10_DNS_SMB_Report.docx")
print("Report saved!")
