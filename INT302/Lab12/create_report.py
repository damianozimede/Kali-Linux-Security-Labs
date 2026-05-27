from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("John the Ripper - Password Cracking Lab Report")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
info = [
    ("Prepared by:", "Damian Patrick Ozimede"),
    ("Course:", "INT302: Kali Linux Tools and System Security"),
    ("Lab:", "Lab 12"),
    ("Date:", "27 May 2026"),
    ("Tools Used:", "John the Ripper 1.9.0-jumbo-1, rockyou.txt, Python hashlib"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + value)

doc.add_paragraph()

# Exercise 1
doc.add_heading("Exercise 1: John the Ripper Version", level=1)
doc.add_paragraph("Command used: john --help")
doc.add_paragraph("Version: John the Ripper 1.9.0-jumbo-1+bleeding-aec1328d6c (2021-11-02). This is the Jumbo community-enhanced version which supports a wider range of hash formats than the standard version, including over 416 hash formats.")

# Exercise 2
doc.add_heading("Exercise 2: Identifying Hash Types", level=1)
doc.add_paragraph("Command used: john --format=raw-md5 ~/hash.txt and john --list=formats | grep -i md5")
doc.add_paragraph("John the Ripper identifies hash types using the --format flag. Running john --list=formats shows all 416 supported formats including MD5 variants such as raw-md5, md5crypt, asa-md5, and net-md5. John can also auto-detect hash types when run without specifying a format by analysing the hash structure and length.")
doc.add_paragraph("The /etc/shadow file stores Linux password hashes in the format $y$ indicating yescrypt hashing algorithm used by modern Kali Linux systems.")

# Exercise 3
doc.add_heading("Exercise 3: Cracking with rockyou.txt Wordlist", level=1)
doc.add_paragraph("Command used: john --wordlist=/usr/share/wordlists/rockyou.txt --format=raw-md5 ~/hash.txt")
ex3_table = doc.add_table(rows=1, cols=2)
ex3_table.style = "Table Grid"
ex3_table.rows[0].cells[0].text = "Field"
ex3_table.rows[0].cells[1].text = "Value"
ex3_rows = [
    ("Hash", "482c811da5d5b4bc6d497ffa98491e38"),
    ("Hash Type", "MD5"),
    ("Password Cracked", "password123"),
    ("Time Taken", "Under 1 second"),
    ("Speed", "153,600 passwords per second"),
    ("Result", "Successful"),
]
for r in ex3_rows:
    row = ex3_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("The password password123 was found almost instantly in rockyou.txt confirming it is a commonly used weak password. This demonstrates the effectiveness of dictionary attacks against common passwords.")

# Exercise 4
doc.add_heading("Exercise 4: Custom Wordlist Attack", level=1)
doc.add_paragraph("Command used: john --wordlist=~/custom_wordlist.txt --format=raw-md5 ~/hash2.txt")
ex4_table = doc.add_table(rows=1, cols=2)
ex4_table.style = "Table Grid"
ex4_table.rows[0].cells[0].text = "Field"
ex4_table.rows[0].cells[1].text = "Value"
ex4_rows = [
    ("Hash", "0192023a7bbd73250516f069df18b500"),
    ("Hash Type", "MD5"),
    ("Custom Wordlist", "password1, letmein, admin123"),
    ("Password Cracked", "admin123"),
    ("Time Taken", "Under 1 second"),
    ("Result", "Successful"),
]
for r in ex4_rows:
    row = ex4_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Yes, the custom wordlist was successful. This demonstrates that even a small targeted wordlist can be effective if it contains the right passwords. Custom wordlists are useful in penetration testing when you have knowledge of the target password patterns or naming conventions.")

# Exercise 5
doc.add_heading("Exercise 5: Brute Force Attack", level=1)
doc.add_paragraph("Command used: john --incremental --format=raw-md5 ~/hash3.txt")
ex5_table = doc.add_table(rows=1, cols=2)
ex5_table.style = "Table Grid"
ex5_table.rows[0].cells[0].text = "Field"
ex5_table.rows[0].cells[1].text = "Value"
ex5_rows = [
    ("Hash", "MD5 of abc"),
    ("Password Cracked", "abc"),
    ("Time Taken", "Under 1 second"),
    ("Speed", "116,438 passwords per second"),
    ("Result", "Successful"),
]
for r in ex5_rows:
    row = ex5_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("The 3-character password abc was cracked almost instantly. This demonstrates that short passwords are extremely vulnerable to brute force attacks. Longer and more complex passwords exponentially increase the time needed — an 8-character password with mixed case, numbers and symbols could take years to crack by brute force.")

# Exercise 6
doc.add_heading("Exercise 6: NTLM Hash Cracking", level=1)
doc.add_paragraph("Command used: john --format=nt --wordlist=/usr/share/wordlists/rockyou.txt ~/ntlm_hash.txt")
ex6_table = doc.add_table(rows=1, cols=2)
ex6_table.style = "Table Grid"
ex6_table.rows[0].cells[0].text = "Field"
ex6_table.rows[0].cells[1].text = "Value"
ex6_rows = [
    ("Hash", "64f12cddaa88057e06a81b54e73b949b"),
    ("Hash Type", "NTLM"),
    ("Password Cracked", "Password1"),
    ("Time Taken", "Under 1 second"),
    ("Speed", "364,800 passwords per second"),
    ("Result", "Successful"),
]
for r in ex6_rows:
    row = ex6_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Despite having a capital letter and number, Password1 is a commonly used password pattern found in rockyou.txt. This demonstrates that simple complexity rules like capitalising the first letter and adding a number at the end are not sufficient protection against dictionary attacks.")

# Exercise 7
doc.add_heading("Exercise 7: Rules-Based Attack", level=1)
doc.add_paragraph("Command used: john --wordlist=/usr/share/wordlists/rockyou.txt --rules --format=raw-md5 ~/hash4.txt")
ex7_table = doc.add_table(rows=1, cols=2)
ex7_table.style = "Table Grid"
ex7_table.rows[0].cells[0].text = "Field"
ex7_table.rows[0].cells[1].text = "Value"
ex7_rows = [
    ("Hash", "MD5 of password1!"),
    ("Password Cracked", "password1!"),
    ("Time Taken", "2 seconds"),
    ("Speed", "43,978 passwords per second"),
    ("Result", "Successful"),
]
for r in ex7_rows:
    row = ex7_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("The rules engine modified wordlist entries by appending special characters like ! to existing passwords. Without rules, password1! might not be in the wordlist directly. Rules make John significantly more powerful by generating variations including capitalisation, number appending, special character substitution, and letter replacement.")

# Analysis and Conclusion
doc.add_heading("Analysis and Conclusion", level=1)
p = doc.add_paragraph()
p.add_run("Success Rates of Different Cracking Techniques:").bold = True
analysis_table = doc.add_table(rows=1, cols=4)
analysis_table.style = "Table Grid"
analysis_table.rows[0].cells[0].text = "Technique"
analysis_table.rows[0].cells[1].text = "Success Rate"
analysis_table.rows[0].cells[2].text = "Speed"
analysis_table.rows[0].cells[3].text = "Best Used For"
analysis_rows = [
    ("Wordlist (rockyou.txt)", "Very High", "Fastest", "Common passwords"),
    ("Custom Wordlist", "High if targeted", "Fast", "Known password patterns"),
    ("Rules-based", "High", "Moderate", "Password variations"),
    ("Brute Force", "100% eventually", "Slowest", "Short passwords only"),
]
for r in analysis_rows:
    row = analysis_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
    row[3].text = r[3]

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Why Password Length, Complexity, and Salts are Essential:").bold = True
for item in [
    "Length — every extra character exponentially increases cracking time. An 8-character password has billions more combinations than a 6-character one",
    "Complexity — mixing uppercase, lowercase, numbers, and symbols dramatically reduces the chance of the password appearing in any wordlist",
    "Salts — random data added to a password before hashing ensures that even if two users have the same password their hashes will be different, preventing attackers from cracking multiple accounts simultaneously",
]:
    doc.add_paragraph(item, style="List Bullet")

# Additional Exercises
doc.add_heading("Additional Exercise 1: Custom Hash with Python", level=1)
doc.add_paragraph("Python code used: import hashlib; print(hashlib.md5(b'password').hexdigest())")
doc.add_paragraph("MD5 hash generated: 5f4dcc3b5aa765d61d8327deb882cf99")
doc.add_paragraph("Password cracked by John: password — cracked in under 1 second using rockyou.txt. This confirms that even programmatically generated hashes are vulnerable if the underlying password is weak.")

doc.add_heading("Additional Exercise 2: Benchmark Performance", level=1)
doc.add_paragraph("Command used: john --test")
bench_table = doc.add_table(rows=1, cols=2)
bench_table.style = "Table Grid"
bench_table.rows[0].cells[0].text = "Hash Type"
bench_table.rows[0].cells[1].text = "Speed"
bench_rows = [
    ("DES", "9,461K c/s"),
    ("MD5crypt", "87,264 c/s"),
    ("bcrypt (32 iterations)", "1,352 c/s"),
]
for r in bench_rows:
    row = bench_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("DES is the fastest and least secure. bcrypt is deliberately slow making it the most resistant to cracking. At 1,352 attempts per second a complex bcrypt password would take years to crack, demonstrating why modern systems should use bcrypt, scrypt, or Argon2 for password hashing.")

doc.add_heading("Additional Exercise 3: Hashcat Comparison (Optional)", level=1)
doc.add_paragraph("Hashcat is GPU-accelerated making it significantly faster than John the Ripper which uses CPU. For MD5 hashes Hashcat can achieve hundreds of millions of attempts per second on a modern GPU compared to John's hundreds of thousands per second on CPU. However John is more versatile with automatic hash detection and broader format support making it better for general forensic work while Hashcat excels at raw cracking speed.")

doc.save("/home/legion/Kali-Linux-Security-Labs/INT302/Lab12/Lab12_John_The_Ripper_Report.docx")
print("Report saved!")
