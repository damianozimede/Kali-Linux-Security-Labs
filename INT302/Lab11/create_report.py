from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Anonymity Testing with Tor and Proxychains Report")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
info = [
    ("Prepared by:", "Damian Patrick Ozimede"),
    ("Course:", "INT302: Kali Linux Tools and System Security"),
    ("Lab:", "Lab 11"),
    ("Date:", "27 May 2026"),
    ("Tools Used:", "Tor, Proxychains-ng, Curl, Firefox"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" " + value)

doc.add_paragraph()

# Exercise 1
doc.add_heading("Exercise 1: Tor Service Status", level=1)
doc.add_paragraph("Command used: sudo service tor@default start && systemctl status tor@default")
doc.add_paragraph("Tor service status output:")
for item in [
    "Service: tor@default.service — Active (running)",
    "Bootstrap: 100% complete",
    "Memory usage: 130.8MB",
    "Tor successfully loaded relay descriptors and established circuits",
    "Exit nodes confirmed available",
    "Workgroup: WORKGROUP",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph("Yes, the Tor service is running successfully. The bootstrap reached 100% confirming Tor has fully connected to the network and is ready to route traffic anonymously.")

# Exercise 2
doc.add_heading("Exercise 2: Proxychains Chain Types", level=1)
chain_table = doc.add_table(rows=1, cols=2)
chain_table.style = "Table Grid"
chain_table.rows[0].cells[0].text = "Chain Type"
chain_table.rows[0].cells[1].text = "Description"
chain_rows = [
    ("strict_chain", "Traffic must pass through every proxy in the list in order. If any proxy is down the connection fails. Most secure and predictable — used in this lab."),
    ("dynamic_chain", "Traffic passes through proxies in order but skips dead ones. More reliable but less predictable than strict_chain."),
    ("random_chain", "Traffic passes through a random proxy from the list each time. Good for evading IDS/IPS detection during penetration testing."),
]
for r in chain_rows:
    row = chain_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]

# Exercise 3
doc.add_heading("Exercise 3: Baseline Real IP Address", level=1)
doc.add_paragraph("Command used: curl https://httpbin.org/ip")
doc.add_paragraph("Real IP Address: 102.89.47.96")
doc.add_paragraph("Yes, this IP address is associated with the ISP and local network. This is the real public IP address that websites and servers see when connecting without any anonymisation tools. Exposing this IP during penetration testing could reveal identity and location.")

# Exercise 4
doc.add_heading("Exercise 4: Anonymity Verification", level=1)
doc.add_paragraph("Command used: proxychains curl https://api.ipify.org")
ip_table = doc.add_table(rows=1, cols=2)
ip_table.style = "Table Grid"
ip_table.rows[0].cells[0].text = "Connection Type"
ip_table.rows[0].cells[1].text = "IP Address"
ip_rows = [
    ("Real IP (without Tor)", "102.89.47.96"),
    ("Tor Exit IP (with Proxychains)", "185.220.100.248"),
]
for r in ip_rows:
    row = ip_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Yes, the IP is different. This confirms that traffic is being routed through the Tor network and exiting through a Tor exit node. Any server receiving this connection sees the exit node IP instead of the real IP, successfully masking identity and location.")

# Exercise 5
doc.add_heading("Exercise 5: Tor Circuit Persistence", level=1)
doc.add_paragraph("The IP address remained the same (192.76.153.253) across all three consecutive requests. This indicates that Tor reuses the same circuit for multiple connections within a session. Tor builds a circuit through three relays (entry, middle, exit) and keeps it active for approximately 10 minutes by default to improve performance. A new circuit is only built when the current one expires or is explicitly requested, not on every new connection.")

# Exercise 6
doc.add_heading("Exercise 6: Tor Exit IP After Restart", level=1)
restart_table = doc.add_table(rows=1, cols=2)
restart_table.style = "Table Grid"
restart_table.rows[0].cells[0].text = "Test"
restart_table.rows[0].cells[1].text = "IP Address"
restart_rows = [
    ("Before restart", "192.76.153.253"),
    ("After restart", "192.42.116.145"),
]
for r in restart_rows:
    row = restart_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("Yes, the IP changed after restarting Tor. Restarting the Tor service forces all circuits to be torn down and rebuilt with potentially different relays. IP rotation is not guaranteed on every request because Tor reuses existing circuits for performance efficiency. A new exit IP is only assigned when a new circuit is built. Even after restart, the same exit node could theoretically be selected again since Tor chooses from a pool of available exit nodes.")

# Exercise 7
doc.add_heading("Exercise 7: Browser-Based Anonymity Test", level=1)
doc.add_paragraph("Command used: proxychains firefox — visited https://www.whatismyip.com/")
browser_table = doc.add_table(rows=1, cols=2)
browser_table.style = "Table Grid"
browser_table.rows[0].cells[0].text = "Field"
browser_table.rows[0].cells[1].text = "Value"
browser_rows = [
    ("Shown IPv4 Address", "192.42.116.108"),
    ("ISP", "TOR EXIT AND MORE"),
    ("City", "Soest"),
    ("Country", "Netherlands"),
    ("Real IP", "102.89.47.96 (Nigeria)"),
]
for r in browser_rows:
    row = browser_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
doc.add_paragraph("The browser shows IP 192.42.116.108 which is different from the real IP 102.89.47.96. The ISP field explicitly shows TOR EXIT AND MORE confirming this is a Tor exit node located in the Netherlands. The browser IP differs slightly from the curl IP as Firefox may use a different Tor circuit.")
doc.add_paragraph("Why consistency is important: During a penetration test, all tools must route through the same anonymisation layer. If some traffic goes through Tor and other traffic does not, the real IP can be leaked and the tester's identity exposed. Consistency ensures complete anonymisation across all tools and connections.")

# Exercise 8
doc.add_heading("Exercise 8: Limitations of Routing Scans Through Tor", level=1)
doc.add_paragraph("Command used: proxychains nmap -sT -Pn 192.168.5.130")
for item in [
    "Local network traffic bypasses Tor — Tor only anonymises internet-bound traffic. Scanning local/private IP ranges routes directly without going through Tor",
    "Speed degradation — Tor adds significant latency for internet targets as traffic bounces through three relays globally",
    "TCP connect scans only — Tor only supports TCP connections. SYN scans, UDP scans, and ICMP ping sweeps are not supported through Tor",
    "Unreliable connections — Tor circuits can drop mid-scan causing incomplete results",
    "Exit node blocking — many servers block known Tor exit nodes preventing scans from completing",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph("Why Tor is unsuitable for active scanning: Tor is designed for anonymous browsing, not high-speed network scanning. Active scanning generates large volumes of traffic that overwhelm Tor circuits, violates Tor acceptable use policy, and is ineffective due to timeouts and blocked exit nodes.")

# Exercise 9
doc.add_heading("Exercise 9: Risk and Limitation Analysis", level=1)
p = doc.add_paragraph()
p.add_run("What are the risks of Tor exit nodes?").bold = True
doc.add_paragraph("Tor exit nodes are the final relay between the Tor network and the destination server. The operator of an exit node can see all unencrypted traffic passing through it. A malicious exit node operator could intercept, modify, or log unencrypted data such as HTTP traffic, credentials, and session cookies. Additionally, law enforcement monitors known exit nodes making them a surveillance point.")

p = doc.add_paragraph()
p.add_run("How can DNS leaks compromise anonymity?").bold = True
doc.add_paragraph("A DNS leak occurs when DNS queries are sent outside the Tor network directly to the ISP DNS server instead of being routed through Tor. This reveals the websites being visited even if the actual connection goes through Tor. The ISP DNS server still receives and logs the domain name lookup, exposing browsing activity. This is why proxy_dns is enabled in the Proxychains configuration — it forces DNS queries through Tor preventing leaks.")

p = doc.add_paragraph()
p.add_run("Why should Tor not be used for authenticated personal accounts?").bold = True
doc.add_paragraph("Using Tor with personal accounts completely defeats the purpose of anonymity. When logging into a personal account such as email or social media, the user voluntarily identifies themselves to that service regardless of IP address. The service knows who the user is from their credentials and can correlate Tor-based activity with their real identity. Additionally, malicious exit node operators can intercept unencrypted login credentials.")

# Summary
doc.add_heading("Summary: Was Anonymity Successfully Achieved?", level=1)
doc.add_paragraph("Yes, anonymity was successfully achieved in this lab. The real IP address 102.89.47.96 was completely masked and replaced with Tor exit node IPs across both curl and browser tests. The ISP field on whatismyip.com explicitly confirmed the connection was identified as coming from a Tor exit node located in the Netherlands rather than Nigeria. However, the lab also demonstrated that anonymity has real limitations — Tor circuits persist and reuse the same exit IP, local network traffic bypasses Tor entirely, and active scanning tools like Nmap are not fully compatible with Tor routing. True operational anonymity requires consistent use of Tor across all tools combined with encrypted protocols and strict avoidance of identity-revealing actions such as logging into personal accounts.")

doc.save("/home/legion/Kali-Linux-Security-Labs/INT302/Lab11/Lab11_Tor_Proxychains_Report.docx")
print("Report saved!")
